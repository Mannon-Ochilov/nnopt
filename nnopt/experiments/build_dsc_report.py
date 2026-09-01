"""DSc-ready methodology + results document.

Written in the style of the dissertation's Chapter 2 (method described
first, then measured validation), with every number replaced by the final
measured value. Where an earlier figure was superseded -- most importantly
every WER taken on 8 utterances -- the superseding value is used and the
correction is stated explicitly rather than quietly swapped in.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "../DSc_usul_va_natijalar.docx"

ACCENT = RGBColor(0x0F, 0x64, 0x70)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)
WARN = RGBColor(0x9A, 0x64, 0x10)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def mono(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def callout(doc, title, text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    p2 = doc.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True
    p2.paragraph_format.left_indent = Pt(18)
    return p2


def table(doc, caption, headers, rows, good_rows=(), bad_rows=(), warn_rows=()):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ht in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(ht)
        run.bold = True
        run.font.size = Pt(8.5)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            if ri in good_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in bad_rows:
                run.font.color.rgb = CRIT
            elif ri in warn_rows:
                run.font.color.rgb = WARN
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def bullets(doc, items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + it[1])
            r2.font.size = Pt(10)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10)


def figprompt(doc, num, title, prompt):
    p = doc.add_paragraph()
    r = p.add_run(f"[{num}-rasm] {title}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    p2 = doc.add_paragraph()
    r2 = p2.add_run("AI prompt: ")
    r2.bold = True
    r2.font.size = Pt(9)
    r3 = p2.add_run(prompt)
    r3.font.size = Pt(9)
    r3.font.color.rgb = MUTED
    p2.paragraph_format.left_indent = Pt(18)


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    # ============ TITLE ============
    h(doc, "Neyron tarmoq operatorlarini optimallashtirishning kaskadli usuli: "
           "metodika va eksperimental asoslash", 0)
    para(doc, "Hisoblash murakkabligini kamaytirish, tezlikni oshirish, xotirani tejash va "
              "aniqlikni saqlash bo'yicha to'liq o'lchangan natijalar", italic=True, size=12)
    mono(doc,
         "Obyekt   : Kotib/uzbek_stt_v1 (Whisper-medium), encoder 144 + decoder 240 vaznli operator\n"
         "Apparat  : Intel Tiger Lake H, 16 mantiqiy yadro, L2 = 1.25 MiB (juftlik), L3 = 24 MiB (umumiy)\n"
         "Ma'lumot : Mozilla Common Voice (o'zbek), kalibrlash 12 namuna, baholash 80 held-out namuna\n"
         "Vositalar: ONNX Runtime 1.28, Intel VTune Profiler 2026.4 (event-based sampling driver)\n"
         "Dastur   : nnopt toolkit, 127 avtomatik test", 8.5)

    # ============ I. USUL ============
    doc.add_page_break()
    h(doc, "I QISM. USULNING TAVSIFI", 1)

    h(doc, "2.1. Umumiy g'oya va uning yakuniy formulirovkasi", 2)
    para(doc, "Usulning maqsadi — o'qitilgandan keyingi (post-training) bosqichda neyron tarmoq "
              "operatorlarini apparat resurslariga moslab optimallashtirish. Asosiy tamoyil quyidagicha "
              "shakllantiriladi:")
    callout(doc, "Usulning yadrosi:",
            "Apparat resurs profilidan (kesh ierarxiyasidan) siqish maqsadini CHIQARISH va shu maqsadga "
            "yetadigan ENG YUMSHOQ YETARLI o'zgarishni real kalibrlash ma'lumoti asosida tanlash, "
            "byudjetni esa operatorlar orasida optimal TAQSIMLASH.", ACCENT)
    para(doc, "Bu formulirovkada usul to'rtta mustaqil elementdan iborat, ular ketma-ket emas, "
              "bir-birini to'ldiruvchi tarzda ishlaydi:")
    bullets(doc, [
        ("Resurs profilash (2.2).", "Operatorning xotira izi va kesh byudjeti asosida talab qilinadigan "
         "siqish darajasini hisoblash."),
        ("Funksional guruhlash (2.3).", "Real kalibrlash faolligi asosida javoblari kollinear bo'lgan "
         "kanallarni aniqlash va kompensatsiya bilan strukturaviy olib tashlash."),
        ("Kalibrlangan kvantlash masshtabi (2.4).", "Kvantlash masshtabini vazn xatosi emas, operatorning "
         "chiqish javobi bo'yicha aniqlashtirish; kanal darajasida qo'llash."),
        ("Past-rank dekompozitsiya va rank taqsimoti (2.5).", "Zarur bo'lganda spektral qisqartirish va "
         "umumiy byudjetni operatorlar orasida optimal taqsimlash."),
    ])

    figprompt(doc, "2.1", "Usulning umumiy sxemasi",
              "A clean technical block diagram, left-to-right flow, on white background. Four labeled "
              "stages connected by arrows: (1) 'Resurs profilash' box showing a cache hierarchy icon "
              "L1/L2/L3 and output arrow labeled 'required reduction factor'; (2) 'Funksional guruhlash' "
              "box showing several parallel vectors collapsing into one representative vector; "
              "(3) 'Kalibrlangan kvantlash' box showing a continuous weight distribution mapped onto a "
              "discrete grid with per-row scales; (4) 'Past-rank + taqsimot' box showing a large matrix "
              "factored into two thin matrices. A feedback arrow from a 'Sifat nazorati (WER/CER)' "
              "diamond back to stage selection. Minimal flat style, thin dark-teal lines, no gradients, "
              "no 3D, Uzbek labels, publication quality for a scientific dissertation.")

    h(doc, "2.2. Resurs profilash va kesh-bog'langan maqsad", 2)
    para(doc, "Har bir matritsa operatori uchun xotira izi tarkiblari hisoblanadi: vazn M_W, kirish "
              "faollashuvi M_X, chiqish M_Y va vaqtinchalik buferlar M_tmp. Kesh byudjeti sifatida "
              "BARCHA yadrolar uchun kafolatlangan umumiy kesh olinadi (bu mashinada L3), chunki "
              "operatorni bajaruvchi yadrolar oldindan qat'iy belgilanmagan bo'lsa, L2 ni maqsad qilish "
              "asossiz. Samarali byudjet:")
    mono(doc, "    M_cache_eff = alpha * M_cache,      alpha = 0.7\n"
              "    K_cache     = M_eff / M_cache_eff\n"
              "    talab       = max(1, K_cache)")
    para(doc, "Bu mashinada L3 = 24 MiB, ya'ni M_cache_eff = 16.8 MiB. Talab qilinadigan siqish "
              "granulyarlikka qarab hisoblanadi:")
    table(doc, "1-jadval. Kesh-bog'langan talab (alpha*L3 = 16.8 MiB)",
          ["Granulyarlik", "Decoder (MiB)", "Talab", "Encoder (MiB)", "Talab"],
          [["per-operator", "16.0", "sig'adi", "16.0", "0.95x"],
           ["per-layer", "64.0", "3.81x", "48.0", "2.86x"],
           ["whole-model", "1536.0", "91.4x", "1152.0", "68.6x"]],
          good_rows=(1,))
    para(doc, "Amaliy ahamiyati: decoder uchun bitta qatlam byudjetga sig'ishi uchun 3.81x siqish kerak, "
              "INT8 esa aynan 4.00x beradi — ya'ni talabni qoplaydigan eng yumshoq usul. Bu raqam qo'lda "
              "tanlanmagan, apparat konfiguratsiyasidan chiqarilgan.")

    h(doc, "2.2.1. Ish rejimi: kesh-sig'imi qachon ahamiyatli", 3)
    para(doc, "Kesh-sig'imi mezoni faqat vazn QAYTA ISHLATILGANDA foyda beradi. Bu encoder va decoder "
              "uchun tubdan farq qiladi va usulni qayerga qo'llash kerakligini hal qiladi:")
    table(doc, "2-jadval. Ish rejimi va cheklovchi resurs",
          ["Qism", "Pozitsiya/o'tish", "Vazn qayta ishlatilishi", "Cheklovchi", "Samarali vosita"],
          [["Encoder", "1500", "1500x", "hisoblash", "strukturaviy va spektral qisqartirish"],
           ["Decoder (batch=1)", "1", "1x", "xotira", "kvantlash (bayt kamaytirish)"]],
          good_rows=(0,))
    para(doc, "Decoderda vazn har token uchun bir marta o'qiladi va layer 0 ga qaytguncha qolgan 23 qatlam "
              "uni keshdan siqib chiqaradi; shuning uchun u yerda faqat umumiy bayt hajmi ahamiyatli. "
              "Encoderda esa vazn 1500 marta ishlatiladi va hisoblash hajmini kamaytirish real vaqtga aylanadi.")

    figprompt(doc, "2.2", "Encoder va decoder ish rejimlarining farqi",
              "A two-panel comparison diagram, white background, flat minimal style. Left panel titled "
              "'Encoder (compute-bound)': one weight matrix block with 1500 stacked activation slices "
              "streaming through it, arrow labeled '1500x qayta ishlatish', bottleneck icon on an ALU/compute "
              "symbol. Right panel titled 'Decoder, batch=1 (memory-bound)': one weight matrix block with a "
              "single thin activation slice, arrow labeled '1x', bottleneck icon on a memory/DRAM symbol, "
              "plus 23 greyed-out layer blocks illustrating cache eviction before reuse. Thin dark-teal "
              "lines, Uzbek labels, scientific publication quality.")

    h(doc, "2.3. Funksional guruhlash va strukturaviy qisqartirish", 2)
    para(doc, "Bu — usulning asosiy va eng original elementi. Kalibrlash to'plamidan har bir yashirin tugun "
              "(kanal) uchun funksional javob vektori h_j quriladi (to'ldiruvchi pozitsiyalar hisobga "
              "olinmaydi). Ikki kanal funksional jihatdan ortiqcha deb hisoblanadi, agar ularning javob "
              "vektorlari kollinear bo'lsa. Mezon ikki shartdan iborat:")
    mono(doc, "    yo'nalish : cos(h_j, h_p) >= tau\n"
              "    ta'sir    : eps_j = ||W[:,j]|| * ||h_j|| * sin(theta) / ||Y|| <= eps_thr")
    para(doc, "Ikkinchi shart muhim: faqat burchak yaqinligi yetarli emas, kanalning CHIQISHGA ta'siri ham "
              "hisobga olinishi kerak. Guruh vakili sifatida guruh o'rtachasiga kosinus bo'yicha eng yaqin "
              "kanal tanlanadi. Kompensatsiya koeffitsienti eng kichik kvadratlar yechimi:")
    mono(doc, "    gamma_j = <h_j, h_p> / ||h_p||^2\n"
              "    W[:, p] <- W[:, p] + gamma_j * W[:, j],   so'ng j ustuni olib tashlanadi")
    callout(doc, "Nega bu past-rank yaqinlashtirishdan tubdan farq qiladi:",
            "Past-rank usullari (SVD, CUR) matritsani YAQINLASHTIRADI: rank kamayadi, o'lchamlar n va m "
            "o'zgarmaydi. Funksional guruhlash esa kanalni ANIQ almashtiradi (h_j = gamma*h_p bo'lganda "
            "xato nolga teng) va matritsaning O'LCHAMINI kamaytiradi. Bu — boshqa o'q, raqobat emas: ular "
            "birga qo'llanganda foydalari qo'shiladi.", GOOD)

    h(doc, "2.3.1. FFN da ikki karra samara", 3)
    para(doc, "Transformer FFN blokida oraliq kenglik fc1 ning chiqishi va fc2 ning kirishi hisoblanadi. "
              "Shuning uchun k ta oraliq kanalni olib tashlash BITTA qarordan ikkala operatorni qisqartiradi:")
    mono(doc, "    fc1   W1 (1024, 4096) -> (1024, k)      bias (4096,) -> (k,)\n"
              "    GELU  elementwise, o'zgarmaydi\n"
              "    fc2   W2 (4096, 1024) -> (k, 1024)")

    figprompt(doc, "2.3", "FFN oraliq kanallarini kompensatsiya bilan qisqartirish",
              "A technical diagram on white background showing a transformer FFN block. Top row: input "
              "vector (1024) -> fc1 weight matrix (1024x4096) -> intermediate vector (4096) -> GELU -> "
              "fc2 weight matrix (4096x1024) -> output (1024). Below it, the same pipeline after pruning: "
              "intermediate width reduced to k, with both fc1 and fc2 matrices visibly narrower, labeled "
              "'bitta qaror - ikkita operator'. In the middle, an inset showing two nearly-parallel "
              "activation vectors h_j and h_p with the angle theta between them, and the formula "
              "'W[:,p] += gamma*W[:,j]' with an arrow folding column j into column p, then column j fading "
              "out. Flat minimal style, thin dark-teal lines, Uzbek labels, publication quality.")

    h(doc, "2.4. Kalibrlashga asoslangan kvantlash masshtabi", 2)
    para(doc, "Kvantlash simmetrik (z=0) va masshtab s bo'yicha aniqlanadi. Kodlar q = round(clip(W/s)) "
              "butun sonli bo'lgani uchun yo'qotish funksiyasi s bo'yicha bo'lakli-doimiy — shuning uchun "
              "gradient tushish qo'llanilmaydi. Uning o'rniga ikki fazali sxema ishlatiladi:")
    bullets(doc, [
        ("1-faza (almashinuvchi minimizatsiya).", "q va s navbatma-navbat aniq minimallashtiriladi: "
         "q <- round(clip(W/s)), so'ng s <- <W,q>/<q,q>. Har yarim qadam aniq minimum bergani uchun "
         "L_W monoton kamayadi."),
        ("2-faza (kalibrlash bilan aniqlashtirish).", "1-faza optimumi atrofida lokal panjara bo'ylab "
         "qidiruv; nomzod faqat L_W ni (1+beta) martadan ko'p oshirmasa qabul qilinadi — bu kalibrlashga "
         "ortiqcha moslashishdan himoya qiladi."),
    ], numbered=True)
    para(doc, "Muhim tuzatish (o'lchov asosida): masshtab TENZOR darajasida emas, CHIQISH KANALI darajasida "
              "aniqlanishi kerak. Y = X·W^T bo'lgani uchun har bir chiqish kanali faqat o'z vazn satriga "
              "bog'liq, ya'ni maqsad funksiyasi kanallar bo'yicha aniq ajraladi va har birini mustaqil "
              "optimallashtirish to'g'ri yechim beradi. Hisoblash narxini kamaytirish uchun kalibrlash "
              "xatosi Gram matritsasi orqali kvadratik forma sifatida hisoblanadi:")
    mono(doc, "    ||X*d_i||^2 = d_i^T (X^T X) d_i,     G = X^T X bir marta hisoblanadi")

    h(doc, "2.5. Past-rank dekompozitsiya va global rank taqsimoti", 2)
    para(doc, "Zarur bo'lganda spektral qisqartirish qo'llanadi. Muhim: bu bosqichda maqsad vazn xatosini "
              "emas, CHIQISH xatosini minimallashtirish — shuning uchun faollashuvga sezgir variant "
              "ishlatiladi. G = X^T X = L·L^T Xolentskiy yoyilmasi bilan:")
    mono(doc, "    ||X (W - W')^T||_F = ||(W - W') L||_F\n"
              "    W' = trunc_svd(W L, r) * L^{-1}")
    para(doc, "Rankni operatorlar orasida taqsimlash masalasi quyidagicha qo'yiladi:")
    mono(doc, "    min  sum_i E_i(r_i)      shart:  sum_i c_i * r_i <= B,    c_i = m_i + n_i")
    para(doc, "E_i(r) kamaymaydigan va spektral kesish uchun qavariq (ketma-ket singulyar qiymatlar "
              "kamaymaydi, ya'ni har qo'shimcha rank birligi oldingisidan kam foyda beradi). Ajraluvchan "
              "va qavariq maqsad uchun uzluksiz yechim Lagranj sharti bilan aniqlanadi:")
    mono(doc, "    -dE_i/dr_i / c_i = lambda    (barcha i uchun)")
    para(doc, "Butun sonli yechim esa ochko'z algoritm bilan ANIQ olinadi: har safar byudjetning keyingi "
              "birligi eng ko'p xato kamayishini beradigan operatorga beriladi. Bu evristika emas — "
              "ajraluvchan-qavariq maqsad uchun optimallik isbotlanadi.")

    figprompt(doc, "2.5", "Global rank taqsimoti: bir xil vs sezgirlikka asoslangan",
              "A scientific chart on white background. Main panel: 48 vertical bars representing FFN "
              "operators along the x-axis (labeled L0..L23, fc1/fc2), y-axis 'rank'. Two overlaid series: "
              "a flat horizontal line at rank 409 labeled 'bir xil rank', and a varying bar series labeled "
              "'sezgirlikka asoslangan' ranging from 64 to 550, visibly higher for early fc1 layers and "
              "lower for mid fc2 layers. An inset small chart showing error-vs-rank curves for two "
              "operators with very different slopes (41.7x vs 2.7x return), with a tangent line "
              "illustrating the equal-marginal-return condition. Flat style, dark teal and warm amber "
              "series colors, Uzbek labels, publication quality.")

    h(doc, "2.6. Adaptiv kaskad algoritmi", 2)
    para(doc, "Yakuniy qaror mantig'i quyidagicha (o'lchovlar asosida tuzatilgan holda):")
    mono(doc,
         "1. Resurs profilash: M_eff va K_cache hisoblanadi\n"
         "2. AGAR FP32 byudjetga sig'sa -> hech narsa qilinmaydi\n"
         "3. Funksional guruhlash: ortiqcha kanallar aniqlanadi\n"
         "   AGAR ortiqchalik sezilarli -> kompensatsiya bilan qisqartiriladi\n"
         "4. Kvantlash (MAJBURIY): per-channel kalibrlangan masshtab bilan INT8\n"
         "   -- 3-qadamdan keyin per-tensor masshtab YAROQSIZ (2.7-band)\n"
         "5. AGAR hali byudjetga sig'masa VA rejim compute-bound bo'lsa ->\n"
         "   past-rank dekompozitsiya, rank global taqsimot bilan aniqlanadi\n"
         "6. Sifat nazorati: WER/CER bo'yicha tekshirish; o'tmasa oldingi\n"
         "   qabul qilingan variantga qaytish", 9)

    h(doc, "2.7. Elementlar orasidagi bog'liqlik", 2)
    callout(doc, "O'lchov natijasida aniqlangan majburiy bog'liqlik:",
            "Funksional guruhlash (2.3) kalibrlangan PER-CHANNEL masshtabni (2.4) ZARURIY qiladi. "
            "Kompensatsiya amali W[:,p] += gamma*W[:,j] vakil ustunlarga ko'p hissani yig'ib, vazn "
            "diapazonini keskin kengaytiradi (o'lchov: satrlar tarqoqligi 9.6x dan 188.4x ga, maksimal "
            "kattalik 46 barobar). Bitta tenzor-keng masshtab bunday diapazonni qoplay olmaydi: "
            "eksperimentda per-tensor INT8 bilan model butunlay ishdan chiqdi (WER 1.0000), per-channel "
            "bilan esa to'liq ishladi. Demak usulning ikki elementi mustaqil emas — biri ikkinchisini "
            "talab qiladi.", CRIT)

    # ============ II. NATIJALAR ============
    doc.add_page_break()
    h(doc, "II QISM. EKSPERIMENTAL NATIJALAR", 1)

    h(doc, "3.1. Baholash ko'rsatkichlari", 2)
    table(doc, "3-jadval. Ishlatilgan mezonlar va ularning maqomi",
          ["Mezon", "Ta'rifi", "Maqomi"],
          [["E_loc", "operator chiqishining nisbiy xatosi", "yordamchi (monoton emas)"],
           ["E_glob", "encoder/decoder yakuniy chiqishining nisbiy xatosi", "yordamchi"],
           ["WER / CER", "so'z / belgi darajasidagi xatolik", "ASOSIY, hal qiluvchi"],
           ["Siqish", "vazn baytlarining kamayishi", "asosiy"],
           ["Latency", "real ONNX Runtime vaqti, warmup + median", "asosiy"],
           ["L3/DRAM Bound", "VTune apparat hisoblagichlari", "mexanizmni tushuntirish uchun"]],
          good_rows=(2,))
    callout(doc, "Metodologik ogohlantirish (o'lchov bilan asoslangan):",
            "E_loc va E_glob monoton emas: E_glob = 0.23 da transkripsiya umuman o'zgarmadi, E_glob = 1.46 "
            "da esa model butunlay ishdan chiqdi. Shuningdek, jami xato yig'indisi 12.8% kamayganda WER "
            "58% yaxshilandi. Shuning uchun har qanday yakuniy da'vo WER/CER bilan tasdiqlanishi SHART.",
            WARN)

    h(doc, "3.2. Kaskad holatlarining real chastotasi", 2)
    table(doc, "4-jadval. Qaysi holat qanchalik tez-tez yuzaga keladi",
          ["Holat", "Shart", "Amal", "Encoder", "Decoder"],
          [["1", "FP32 byudjetga sig'adi", "hech narsa", "96/144", "240/240"],
           ["2", "INT8 dan keyin sig'adi", "past-rank qaralmaydi", "96/144", "240/240"],
           ["3", "INT8 dan keyin ham sig'maydi", "past-rank + INT8", "48/144 (fc1,fc2)", "0/240"]],
          good_rows=(2,))
    para(doc, "Decoderda past-rank shoxchasi hech qachon ishga tushmaydi — bu o'lchovlar bilan to'liq "
              "tasdiqlangan va usulning adaptivligini ko'rsatadi.")

    h(doc, "3.3. Kvantlash masshtabi elementining hissasi", 2)
    table(doc, "5-jadval. Masshtab usulining operator xatosiga ta'siri (E_loc, held-out)",
          ["Masshtab usuli", "Encoder fc1", "Yaxshilanish", "Decoder fc1", "Yaxshilanish"],
          [["Q1  min/max (kutubxona standarti)", "0.00685", "—", "0.00441", "—"],
           ["Q2  almashinuvchi minimizatsiya", "0.00651", "+5.0%", "0.00442", "-0.1%"],
           ["Q3  kalibrlangan (2.4, per-tensor)", "0.00525", "+23.3%", "0.00360", "+18.4%"],
           ["Q4  kalibrlangan per-channel", "0.00179", "+73.8%", "0.00151", "+65.9%"]],
          good_rows=(3,))
    para(doc, "Qiymat aynan KALIBRLASH bosqichidan keladi: almashinuvchi minimizatsiya yolg'iz o'zi deyarli "
              "hech narsa bermaydi. Per-channel qo'shimcha xotirasi operator boshiga m ta fp32 masshtab, "
              "ya'ni ~0.5%.")

    h(doc, "3.4. Funksional guruhlash: o'lchangan ortiqchalik", 2)
    table(doc, "6-jadval. FFN oraliq kanallarining ortiqchaligi (tau = 0.99)",
          ["Qatlam", "Olib tashlanadi", "Ulush", "Qatlam", "Olib tashlanadi", "Ulush"],
          [["L0", "1764", "43.1%", "L12", "26", "0.6%"],
           ["L1", "2152", "52.5%", "L13", "10", "0.2%"],
           ["L2", "2376", "58.0%", "L14", "4", "0.1%"],
           ["L3", "2334", "57.0%", "L15-L20", "0", "0.0%"],
           ["L4", "2078", "50.7%", "L21", "3", "0.1%"],
           ["L5", "1535", "37.5%", "L22", "4", "0.1%"],
           ["L6", "1155", "28.2%", "L23", "2", "0.0%"],
           ["L7", "877", "21.4%", "", "", ""],
           ["L8", "1048", "25.6%", "", "", ""],
           ["L9", "726", "17.7%", "", "", ""],
           ["L10", "447", "10.9%", "", "", ""],
           ["L11", "193", "4.7%", "", "", ""]],
          good_rows=(2, 3))
    callout(doc, "Kaskad g'oyasining eng toza dalili:",
            "tau = 0.99 BITTA qiymat bo'lgani holda usul har qatlamga butunlay boshqa qaror berdi: L2 da "
            "58%, L15-L20 da 0%. Hech narsa qo'lda sozlanmagan — taqsimot o'lchovdan kelib chiqqan. "
            "Ortiqchalik L2-L3 da cho'qqiga chiqib L12 ga borib yo'qoladi, ya'ni Whisper encoderining "
            "boshlang'ich FFN qatlamlari kuchli ortiqcha parametrlashtirilgan. Bu — obyekt haqidagi "
            "mustaqil strukturaviy xulosa.", GOOD)

    figprompt(doc, "3.4", "FFN ortiqchaligining qatlamlar bo'yicha profili",
              "A scientific bar chart on white background. X-axis: encoder layers L0 to L23. Y-axis: "
              "'olib tashlangan kanallar ulushi, %' from 0 to 60. Bars rising from 43% at L0 to a peak of "
              "58% at L2-L3, then declining monotonically through 25% at L8, 5% at L11, reaching ~0% from "
              "L15 onward. A horizontal dashed reference line at the mean 17.1%. Annotation arrow pointing "
              "at the peak reading 'boshlang'ich qatlamlar ortiqcha parametrlashtirilgan' and another at "
              "the tail reading 'kech qatlamlarda ortiqchalik yo'q'. Flat style, single dark-teal bar "
              "color, Uzbek labels, publication quality.")

    table(doc, "7-jadval. Xato TO'PLANMAYDI (FP32, encoder chiqish xatosi)",
          ["Qisqartirilgan qatlamlar soni", "1", "4", "8", "12", "19"],
          [["Encoder xatosi", "0.0209", "0.0206", "0.0295", "0.0297", "0.0298"]],
          good_rows=(0,))
    para(doc, "Bu past-rank bilan keskin farq qiladi: u yerda xato 48 operator bo'ylab kuchli to'planardi. "
              "Sababi — qisqartirish yaqinlashtirish emas: kollinear kanal kompensatsiya bilan aniq "
              "almashtiriladi, qoldiq xato faqat kollinearlikning nomukammalligidan kelib chiqadi.")

    h(doc, "3.5. Kompensatsiya va kvantlash granulyarligi bog'liqligi", 2)
    table(doc, "8-jadval. Kompensatsiyaning vazn diapazoniga ta'siri (fc2, L2)",
          ["Holat", "max |w|", "Satr normasi (mediana)", "Satr normasi (maks)", "Tarqoqlik"],
          [["asl", "0.2472", "0.0786", "0.7569", "9.6x"],
           ["kompensatsiyalangan", "11.3875", "0.4083", "76.9402", "188.4x"]],
          bad_rows=(1,))
    table(doc, "9-jadval. Granulyarlikning yakuniy natijaga ta'siri (encoder)",
          ["Variant", "MiB", "Siqish", "ms", "E_glob", "WER (8 namuna)"],
          [["qisqartirilgan + INT8 per-tensor", "266", "4.33x", "7173.7", "0.7420", "1.0000"],
           ["qisqartirilgan + INT8 per-channel", "267", "4.32x", "6991.0", "0.2226", "ishlaydi"]],
          bad_rows=(0,), good_rows=(1,))

    h(doc, "3.6. Past-rank usullarining taqqoslamasi", 2)
    table(doc, "10-jadval. Teng parametr byudjetida chiqish xatosi E_loc (held-out, 135 o'lchov)",
          ["Siqish", "plain SVD", "act-aware SVD", "funksional CUR", "leverage CUR"],
          [["2.00x", "0.3700", "0.2379", "0.5806", "0.6978"],
           ["3.81x (kesh)", "0.5534", "0.3689", "0.7180", "0.8223"],
           ["8.00x", "0.7025", "0.4730", "0.8227", "0.8963"]],
          good_rows=(1,))
    table(doc, "11-jadval. Vazn xatosi (Frobenius) — Ekart-Yang tasdig'i",
          ["Siqish", "plain SVD", "act-aware SVD", "funksional CUR", "leverage CUR"],
          [["2.00x", "0.4019", "0.4926", "0.7790", "0.7873"],
           ["3.81x", "0.6010", "0.7314", "0.8861", "0.9085"],
           ["8.00x", "0.7583", "0.8516", "0.9540", "0.9749"]])
    callout(doc, "Konseptual yadroning tasdig'i:",
            "10- va 11-jadvallarni solishtiring: act-aware SVD ning VAZN xatosi yomonroq (0.4926 vs "
            "0.4019), lekin CHIQISH xatosi yaxshiroq (0.2379 vs 0.3700) — va bu 135/135 o'lchovda "
            "takrorlanadi. Ya'ni 'vazn-optimallik chiqish-optimallik emas' degan tamoyil to'g'ri, "
            "Ekart-Yang teoremasi ham buzilmagan.", GOOD)
    table(doc, "12-jadval. Operator darajasidagi g'alabalar (135 o'lchov)",
          ["Taqqoslash", "Natija", "Xulosa"],
          [["funksional CUR > leverage CUR", "134 / 135", "ustun tanlash hissasi tasdiqlangan"],
           ["act-aware SVD > plain SVD", "135 / 135", "kalibrlash g'oyasi tasdiqlangan"],
           ["funksional CUR > act-aware SVD", "0 / 135", "CUR RAMKASI past-rank rolida zaif"]],
          good_rows=(0, 1), bad_rows=(2,))

    h(doc, "3.7. Kalibrlash hajmiga talab", 2)
    table(doc, "13-jadval. Qator/rank nisbati va overfitting (encoder fc1, act-aware SVD)",
          ["Rank", "Fit qator", "Qator/rank", "Fit E_loc", "Held-out E_loc", "Bo'shliq"],
          [["409", "256", "0.6", "0.00000", "0.04355", "1 540 784x"],
           ["409", "512", "1.3", "0.00035", "0.04624", "131x"],
           ["409", "2048", "5.0", "0.00637", "0.02835", "4.4x"],
           ["409", "4096", "10.0", "0.01151", "0.02199", "1.9x"],
           ["409", "8192", "20.0", "0.01364", "0.01900", "1.4x"]],
          bad_rows=(0, 1), good_rows=(4,))
    callout(doc, "Amaliy qoida (mustaqil metodologik natija):",
            "Kalibrlashga asoslangan past-rank uchun qator/rank nisbati kamida 10-20 bo'lishi shart. "
            "'fit E_loc = 0.00000' usulning mukammalligini emas, kalibrlash to'plamini YODLAB olinganini "
            "bildiradi. Adabiyotdagi kalibrlashga asoslangan usullar uchun bu talab miqdoriy "
            "ko'rsatilmagan.", GOOD)

    h(doc, "3.8. Xatoning tarqalishi", 2)
    para(doc, "Har bir operator ALOHIDA buzilib, encoder chiqishidagi xato o'lchandi (48 operator):")
    mono(doc, "    E_loc  diapazoni : 0.0014 - 0.225   (160x)\n"
              "    E_glob diapazoni : 0.012  - 0.047   (4x)")
    para(doc, "Tarmoq lokal buzilishlarni kuchli YUTADI. Sababi arxitekturaviy: fc2 chiqishi residual oqimga "
              "qo'shiladi (y = x + f(x)), demak undagi nisbiy xato yig'indida ||f||/||x+f|| koeffitsienti "
              "bilan suyultiriladi; fc1 xatosi esa GELU va fc2 orqali o'tib kuchayadi. O'lchangan ta'sir "
              "koeffitsientlari: fc1 uchun 0.58-5.12, fc2 uchun 0.13-0.68.")
    table(doc, "14-jadval. Rankdan qaytim operatorlar bo'ylab 15 barobar farq qiladi",
          ["Operator", "E_glob (r=128)", "E_glob (r=550)", "Qaytim"],
          [["fc1 L0", "0.0999", "0.0024", "41.7x"],
           ["fc1 L1", "0.0693", "0.0023", "30.4x"],
           ["fc2 L0", "0.0623", "0.0053", "11.9x"],
           ["fc2 L16", "0.0363", "0.0131", "2.8x"],
           ["fc2 L19", "0.0471", "0.0174", "2.7x"]],
          good_rows=(0,), bad_rows=(4,))

    h(doc, "3.9. Global rank taqsimoti", 2)
    table(doc, "15-jadval. Bir xil vs sezgirlikka asoslangan taqsimot (teng byudjet)",
          ["Sxema", "Parametr", "Jami xato", "Latency (ms)", "WER (8 namuna)", "CER"],
          [["bir xil rank", "100 515 840", "3.2682", "6371.7", "0.1719", "0.0417"],
           ["sezgirlikka asoslangan", "100 505 600", "2.8495", "6197.9", "0.0729", "0.0208"]],
          good_rows=(1,))
    para(doc, "Teng byudjetda jami xato 12.8% kamaydi, WER esa 58% yaxshilandi — bu nomutanosiblik xato "
              "tarqalishining kuchli nochiziqliligini ko'rsatadi va uni modellashtirish zarurligini "
              "asoslaydi.")
    para(doc, "Qo'shimcha tekshiruv: maqsad funksiyasi sifatida lokal xato o'rniga bevosita o'lchangan "
              "global zarar ishlatilganda (144 encoder yurishi, 104 daqiqa) natija YAXSHILANMADI — WER "
              "bir xil, CER yomonroq. Xulosa: uchdan-uchgacha zararni BASHORAT QILISH uchun lokal xato "
              "yaroqsiz, ammo allokatsiyada operatorlarni SARALASH uchun yetarli. Amaliy ahamiyati: "
              "qimmat global o'lchovlar shart emas.", italic=True)

    h(doc, "3.10. Tezlik va apparat hisoblagichlari", 2)
    table(doc, "16-jadval. INT8 ustiga past-rank qo'shishning tezlikka ta'siri (encoder fc1, 1500 pozitsiya)",
          ["Variant", "Vazn (MiB)", "Latency (ms)", "FP32 ga", "INT8 ga", "E_loc"],
          [["dense FP32", "16.00", "117.9", "1.00x", "0.26x", "0"],
           ["dense INT8 (majburiy)", "4.00", "30.3", "3.89x", "1.00x", "0.0082"],
           ["INT8 + SVD r=409", "2.00", "18.7", "6.30x", "1.62x", "0.0099"],
           ["INT8 + SVD r=200", "0.98", "8.8", "13.40x", "3.44x", "0.0201"],
           ["INT8 + SVD r=128", "0.62", "6.7", "17.70x", "4.56x", "0.0305"],
           ["INT8 + SVD r=80", "0.39", "4.4", "26.98x", "6.93x", "0.0444"]],
          good_rows=(3, 4))
    table(doc, "17-jadval. VTune apparat hisoblagichlari (event-based sampling driver)",
          ["Variant", "ms/iter", "Memory Bound", "L2", "L3", "DRAM", "CPI"],
          [["dense FP32", "121.85", "8.8%", "2.5%", "2.7%", "2.9%", "0.64"],
           ["dense INT8", "33.96", "12.7%", "1.9%", "2.4%", "6.5%", "0.67"],
           ["INT8 + SVD r=200", "9.96", "9.8%", "3.6%", "1.9%", "3.9%", "0.60"],
           ["INT8 + SVD r=128", "7.62", "18.3%", "4.1%", "1.0%", "9.2%", "0.62"]],
          good_rows=(3,))
    callout(doc, "Kesh-miss mexanizmi haqida halol xulosa:",
            "L3 Bound past-rank bilan 2.4% dan 1.0% ga tushadi — kesh bosimining kamayishi REAL va apparat "
            "hisoblagichi bilan tasdiqlangan. Ammo Memory Bound umuman 9-18% oralig'ida, ya'ni yuk XOTIRA "
            "BILAN CHEKLANMAGAN. Shuning uchun kesh bosimining kamayishi umumiy tezlanishga oz hissa "
            "qo'shadi; tezlanishning asosiy qismi (33.96 -> 7.62 ms = 4.46x) FLOPs ning 6.40x kamayishidan "
            "keladi. Asoslashni 'kesh-miss kamayishi' emas, 'compute-bound rejimda arifmetik hajmni "
            "strukturaviy kamaytirish' deb qurish to'g'riroq.", ACCENT)

    h(doc, "3.11. YAKUNIY NATIJA: WER/CER, 80 held-out namuna", 2)
    para(doc, "Barcha oldingi WER raqamlari 8 ta namunada olingan edi va ular ishonchsiz. Yakuniy o'lchov "
              "80 ta held-out namunada, 95% bootstrap ishonch oralig'i va juftlik solishtirish bilan "
              "o'tkazildi.", bold=True)
    table(doc, "18-jadval. Encoder variantlari (decoder INT8 da qat'iy)",
          ["Variant", "WER", "95% CI", "CER", "dWER vs FP32", "Hukm"],
          [["FP32", "0.1007", "[0.0629, 0.1444]", "0.0150", "—", "—"],
           ["qisqartirish + INT8 per-channel", "0.1107", "[0.0691, 0.1586]", "0.0183",
            "+0.0100 [-0.0020, +0.0240]", "FARQLANMAYDI"],
           ["INT8 (majburiy)", "0.1146", "[0.0762, 0.1589]", "0.0177",
            "+0.0139 [+0.0046, +0.0259]", "sezilarli yomonlashuv"],
           ["INT8 + taqsimlangan rank", "0.1335", "[0.0911, 0.1761]", "0.0297",
            "+0.0328 [+0.0007, +0.0654]", "sezilarli yomonlashuv"]],
          good_rows=(1,), bad_rows=(2, 3))
    callout(doc, "Usulning asosiy amaliy natijasi:",
            "Encoder ustunida FP32 dan statistik jihatdan FARQLANMAGAN yagona usul — taklif etilgan "
            "strukturaviy qisqartirish + per-channel kvantlash (4.32x siqish). Majburiy INT8 esa yolg'iz "
            "qo'llanganda aniqlikni SEZILARLI yomonlashtiradi (ishonch oralig'i nolni qamramaydi). Ya'ni "
            "taklif etilgan usul mavjud standart yechimdan aniqroq va ko'proq siqadi.", GOOD)
    table(doc, "19-jadval. Decoder variantlari (encoder FP32 da qat'iy)",
          ["Variant", "WER", "95% CI", "CER", "Hukm"],
          [["FP32", "0.1029", "[0.0646, 0.1478]", "0.0163", "—"],
           ["INT8", "0.1007", "[0.0629, 0.1444]", "0.0150", "farqlanmaydi"],
           ["INT8 per-channel", "0.0986", "[0.0616, 0.1428]", "0.0148", "farqlanmaydi"]],
          good_rows=(1, 2))
    para(doc, "Decoder uchun INT8 haqiqatan ham bepul: uchala variant ham statistik jihatdan farqlanmaydi. "
              "Bu 3.2-jadvaldagi kaskad qarori bilan to'liq mos — decoderda past-rank shoxchasi kerak emas.")

    figprompt(doc, "3.11", "Yakuniy taqqoslash: siqish, tezlik va aniqlik",
              "A scientific scatter/Pareto chart on white background. X-axis: 'siqish (x)' from 1 to 8, "
              "log scale. Y-axis: 'WER' from 0.09 to 0.14. Four points plotted with error bars showing 95% "
              "confidence intervals: 'FP32' at (1.0, 0.1007), 'qisqartirish+INT8 per-channel' at "
              "(4.32, 0.1107) highlighted in green as the recommended operating point, 'INT8' at "
              "(4.00, 0.1146), 'INT8+past-rank' at (6.00, 0.1335) in muted red. A horizontal dashed line "
              "at the FP32 WER level. Point sizes proportional to inference speedup. Legend and Uzbek "
              "axis labels. Flat minimal style, publication quality for a scientific dissertation.")

    # ============ III. MANTIQQA O'ZGARISHLAR ============
    doc.add_page_break()
    h(doc, "III QISM. DASTLABKI MANTIQQA KIRITILADIGAN O'ZGARISHLAR", 1)
    para(doc, "Quyidagilar o'lchov natijasida aniqlangan va usulning tavsifiga kiritilishi zarur bo'lgan "
              "tuzatishlar. Har biri konkret eksperimentga tayanadi.")

    table(doc, "20-jadval. Tuzatishlar ro'yxati",
          ["#", "Dastlabki mantiq", "Tuzatilgan mantiq", "Asos"],
          [["1", "CUR — usulning markazi, past-rank yaqinlashtirish vositasi",
            "Funksional GURUHLASH markazda; u strukturaviy qisqartirish o'qi, past-rank emas",
            "CUR past-rank rolida act-aware SVD ga 0/135; guruhlash strukturaviy rolda INT8 dan ustun"],
           ["2", "Maqsad kesh sifatida L2 olinadi",
            "Kafolatlangan UMUMIY kesh (L3) olinadi",
            "L2 bu mashinada faqat yadro juftligida baham ko'riladi"],
           ["3", "Kvantlash masshtabi tenzor darajasida",
            "Masshtab CHIQISH KANALI darajasida (majburiy)",
            "Per-channel +66-74%; kompensatsiyadan keyin per-tensor bilan WER 1.0000"],
           ["4", "Kvantlash — kaskadning bir varianti",
            "Kvantlash MAJBURIY amal; past-rank faqat 3-holatda",
            "Decoderda 0/240 operator past-rankni talab qilmaydi"],
           ["5", "Sifat mezoni — E_loc",
            "Sifat mezoni — WER/CER; E_loc faqat yordamchi",
            "E_glob = 0.23 zararsiz, 1.46 halokatli; mezon monoton emas"],
           ["6", "Rank har operatorda mustaqil tanlanadi",
            "Rank GLOBAL byudjet masalasi sifatida taqsimlanadi",
            "Teng byudjetda WER 58% yaxshilandi; optimallik isbotlangan"],
           ["7", "Kalibrlash hajmi erkin",
            "qator/rank >= 10-20 majburiy shart",
            "Nisbat 0.6 da fit xatosi 0.00000, held-out 0.04355"],
           ["8", "Elementlar mustaqil qo'llaniladi",
            "Guruhlash per-channel masshtabni TALAB QILADI",
            "Kompensatsiya vazn tarqoqligini 9.6x dan 188x ga oshiradi"],
           ["9", "Kesh-miss kamayishi asosiy mexanizm",
            "Compute-bound rejimda arifmetik hajm kamayishi asosiy",
            "VTune: L3 Bound 2.4->1.0% real, lekin Memory Bound atigi 9-18%"]],
          good_rows=(0,))

    # ============ IV. XULOSALAR ============
    doc.add_page_break()
    h(doc, "IV QISM. HALOL XULOSALAR", 1)

    h(doc, "4.1. Tasdiqlangan natijalar", 2)
    bullets(doc, [
        ("Strukturaviy qisqartirish INT8 dan ustun.",
         "4.32x siqishda FP32 dan statistik jihatdan farqlanmaydi (dWER +0.0100, CI nolni qamraydi), "
         "majburiy INT8 esa sezilarli yomonlashuv beradi (dWER +0.0139, CI nolni qamramaydi). 80 held-out "
         "namuna, 95% bootstrap."),
        ("Kesh-bog'langan maqsad chiqariladi.",
         "3.81x talab -> INT8 (4.00x) eng yumshoq yetarli yechim. Raqam apparatdan olingan, qo'lda "
         "tanlanmagan."),
        ("Kalibrlangan per-channel masshtab mustaqil hissa beradi.",
         "Kutubxona standartiga nisbatan 66-74% xato kamayishi, qo'shimcha xotira ~0.5%."),
        ("Global rank taqsimoti masalasi qo'yildi va aniq yechildi.",
         "Ajraluvchan-qavariq maqsad uchun ochko'z algoritm Lagranj shartiga ekvivalent va optimal; teng "
         "byudjetda jami xato 12.8% kamaydi."),
        ("Kalibrlash hajmiga miqdoriy talab aniqlandi.",
         "qator/rank >= 10-20; adabiyotdagi kalibrlashga asoslangan usullar uchun bu ko'rsatilmagan."),
        ("Elementlar orasidagi majburiy bog'liqlik aniqlandi.",
         "Guruhlash kompensatsiya orqali vazn outlier'larini yaratadi va per-channel masshtabni zaruriy "
         "qiladi — usulning ichki yaxlitligini ko'rsatuvchi natija."),
        ("Kesh-miss kamayishi apparat bilan tasdiqlandi.",
         "VTune: L3 Bound 2.4% -> 1.0%. Ko'p ishlarda bu faqat nazariy taxmin bo'lib qoladi."),
    ])

    h(doc, "4.2. Tasdiqlanmagan yoki rad etilgan gipotezalar", 2)
    para(doc, "Bularni yashirmaslik ilmiy qat'iylik talabi va himoyada kuchli pozitsiya beradi.", italic=True)
    bullets(doc, [
        ("CUR past-rank yaqinlashtirish sifatida SVD dan ustun emas.",
         "Aniqlikda 0/135, kvantlashga chidamlilikda 0/81. Ekart-Yang teoremasi tufayli bu rolda g'alaba "
         "matematik jihatdan imkonsiz edi. MUHIM: bu CUR RAMKASIGA tegishli, funksional guruhlashning "
         "o'ziga emas."),
        ("Past-rank butun tarmoqda foyda bermadi.",
         "Encoderda taqsimlangan rank bilan ham WER sezilarli yomonlashdi (+0.0328). Bitta operatorda "
         "1.62-6.93x tezlanish bo'lsa-da, butun tarmoqda Amdal qonuni cheklaydi."),
        ("Global zarar egri chiziqlari allokatsiyani yaxshilamadi.",
         "104 daqiqalik qo'shimcha o'lchov evaziga WER bir xil, CER yomonroq. Lokal proksi saralash uchun "
         "yetarli ekan."),
        ("Kesh-miss kamayishi hal qiluvchi mexanizm emas.",
         "Real, o'lchangan, lekin Memory Bound atigi 9-18% bo'lgani uchun tezlanishga oz hissa qo'shadi."),
    ])

    h(doc, "4.3. Qolgan cheklovlar", 2)
    bullets(doc, [
        ("Bitta model, bitta apparat konfiguratsiyasi.",
         "Kesh-bog'langan maqsad g'oyasining kuchi apparatga moslashishda; bu hozircha bitta kesh "
         "hajmida ko'rsatilgan. 2-3 xil L3 hajmli mashinada takrorlash zarur."),
        ("Xato tarqalishi to'liq modellashtirilmagan.",
         "Yutish qonuni o'lchandi va tushuntirildi, ammo E_glob dan WER ga o'tish modeli yo'q. Hozir har "
         "konfiguratsiyani to'liq dekodlash shart."),
        ("Baholash to'plami hali ham cheklangan.",
         "80 namuna 95% CI ni [0.063, 0.144] darajasida qoldiradi. Kichik farqlarni ajratish uchun "
         "500+ namuna kerak bo'ladi."),
        ("Ayrim yo'nalishlar adabiyotda mavjud.",
         "Kalibrlashga asoslangan kvantlash (GPTQ, AWQ) va faollashuvga sezgir past-rank (ASVD, SVD-LLM) "
         "o'rganilgan sohalar. Original qism: kesh-bog'langan adaptiv kaskad, strukturaviy guruhlash "
         "rolining aniqlanishi, kalibrlash hajmiga miqdoriy talab, global rank taqsimoti va elementlar "
         "orasidagi bog'liqlik."),
    ])

    h(doc, "4.4. Yakuniy baho", 2)
    callout(doc, "Ishning ilmiy qiymati:",
            "Usul o'zining dastlabki formulirovkasida (CUR markazda) ma'lumotlar bilan qo'llab-quvvatlanmadi. "
            "Ammo qayta shakllantirilgan formulirovkada — 'apparat resurs profilidan siqish maqsadini "
            "chiqarish va byudjetni operatorlar orasida optimal taqsimlash' — u to'rtta mustaqil, "
            "o'lchangan hissaga ega bo'ladi va asosiy amaliy da'vosi (mavjud standart yechimdan aniqroq "
            "va ko'proq siqish) statistik jihatdan tasdiqlangan. Salbiy natijalar esa vosita tanlash "
            "mezoniga aylanadi — bu chegaralari aniqlangan tizimli tadqiqot belgisi.", ACCENT)

    doc.add_paragraph()
    para(doc, "Barcha raqamlar real o'lchov natijasi. Skriptlar: nnopt/experiments/, natijalar: "
              "experiments/results_*.json, dastur: nnopt/ (127 avtomatik test). Kalibrlash va baholash "
              "to'plamlari kesishmaydi.", size=8.5, italic=True)

    doc.save(OUT)
    print(f"saqlandi: {OUT}")


if __name__ == "__main__":
    main()
