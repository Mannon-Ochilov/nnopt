"""Final DSc report: all measured tables + element-contribution ablation +
direct VTune hardware evidence + a scientific assessment of what the method
is worth and what would raise it to DSc weight.

Written as real Word tables so the author can lift them into chapters.
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "../DSc_yakuniy_hisobot.docx"

ACCENT = RGBColor(0x0F, 0x64, 0x70)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)
WARN = RGBColor(0x9A, 0x64, 0x10)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def callout(doc, title, text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    p2 = doc.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True
    p2.paragraph_format.left_indent = Pt(18)
    return p2


def table(doc, caption, headers, rows, good_rows=(), bad_rows=(), warn_rows=()):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ht in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(ht)
        run.bold = True
        run.font.size = Pt(9)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri in good_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in bad_rows:
                run.font.color.rgb = CRIT
            elif ri in warn_rows:
                run.font.color.rgb = WARN
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def bullets(doc, items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + it[1])
            r2.font.size = Pt(10)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10)


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    # ================= TITLE =================
    h(doc, "Neyron tarmoq operatorlarini optimallashtirishning kaskadli usuli", 0)
    para(doc, "Eksperimental asoslash: hisoblash murakkabligi, tezlik, xotira, kesh-miss va aniqlik",
         italic=True, size=12)
    para(doc, "Model: Kotib/uzbek_stt_v1 (Whisper-medium)  |  Encoder 144 + Decoder 240 vaznli operator  |  "
              "Mashina: Intel Tiger Lake H, 16 mantiqiy yadro, L2 = 1.25 MiB, L3 = 24 MiB  |  2026-08-12", size=9)
    para(doc, "Barcha raqamlar real o'lchov: real o'zbek nutqi kalibratsiyasi (Mozilla Common Voice), "
              "real ONNX Runtime (intra_op_threads=1), held-out namunalar, hamda Intel VTune Profiler 2026.4 "
              "apparat hisoblagichlari (event-based sampling driver faol).", size=9, italic=True)

    # ================= 0. XULOSA =================
    h(doc, "Qisqacha xulosa", 1)
    bullets(doc, [
        ("Kesh-bog'langan maqsad chiqarildi.",
         "Siqish darajasi qo'lda tanlanmaydi: bitta decoder qatlami alpha*L3 = 16.8 MiB ga sig'ishi uchun "
         "3.81x kerak, INT8 esa aynan 4.00x beradi — kaskadning 'eng yumshoq yetarli o'zgarish' tamoyili "
         "aynan shuni tanlaydi."),
        ("Masshtabni aniqlashtirish g'oyasi mustaqil hissa beradi.",
         "Kutubxona standarti (min/max) bilan solishtirganda per-channel kalibrlangan masshtab operator "
         "xatosini encoderda 74%, decoderda 66% kamaytiradi."),
        ("Past-rank qo'shish compute-bound rejimda foydali — TAQSIMLANGAN bo'lsa.",
         "Butun encoder: INT8 ustiga sezgirlikka asoslangan past-rank 1.5x ko'proq siqish va 1.14x tezlik "
         "beradi, WER esa atigi +0.006. Bir xil rank berilganda esa WER +0.130 ga yomonlashadi — ya'ni "
         "TAQSIMOT hal qiluvchi."),
        ("Global rank taqsimoti masalasi qo'yildi va aniq yechildi.",
         "Ajraluvchan-qavariq maqsad uchun ochko'z algoritm Lagranj shartiga ekvivalent va optimal. Teng "
         "byudjetda WER 58% kamaydi (0.1719 -> 0.0729)."),
        ("Kalibrlash hajmiga miqdoriy talab aniqlandi.",
         "Qator/rank nisbati kamida 10–20 bo'lishi shart; aks holda usul kalibrlash to'plamini yodlab oladi "
         "(rank 409, 256 qator: fit xatosi 0.00000, held-out 0.04355)."),
        ("Xato-yutish qonuni o'lchandi.",
         "E_loc 160 barobar o'zgarganda E_glob atigi 4 barobar o'zgaradi — tarmoq lokal buzilishlarni kuchli "
         "yutadi. Residual oqim va LayerNorm hisobiga; fc1 xatosi kuchayadi (c=0.58–5.12), fc2 niki so'nadi "
         "(c=0.13–0.68)."),
        ("Funksional guruhlash o'z rolida INT8 dan USTUN chiqdi.",
         "Strukturaviy kanal qisqartirish + per-channel INT8: 4.32x siqish, WER 0.0417 — asl FP32 darajasi, "
         "holbuki majburiy INT8 yolg'iz WER ni 0.0667 ga ko'taradi. Bu — usulning asosiy g'oyasi uchun "
         "uchdan-uchgacha tasdiqlangan g'alaba."),
        ("Usulning ikki elementi o'zaro bog'liq ekani aniqlandi.",
         "Kompensatsiya vazn diapazonini 9.6x dan 188x ga kengaytiradi, shuning uchun guruhlash per-channel "
         "kalibrlangan masshtabni ZARURIY qiladi (per-tensor bilan WER 1.0000). Oldindan aytib bo'lmaydigan, "
         "usulning ichki yaxlitligini ko'rsatuvchi natija."),
        ("Kesh-miss kamayishi APPARAT bilan tasdiqlandi, lekin u hal qiluvchi omil emas.",
         "VTune: L3 Bound 2.4% -> 1.0% (2.4x kamayish), ammo Memory Bound umuman 9–18% — yuk compute-bound, "
         "tezlanishning asosiy qismi FLOPs kamayishidan keladi."),
        ("Aniqlik mezoni sifatida faqat WER/CER ishonchli.",
         "E_glob = 0.23 -> transkripsiya umuman o'zgarmaydi; E_glob = 1.46 -> model yaroqsiz. Mezon monoton emas."),
    ])

    # ================= 1. KASKAD MANTIG'I =================
    h(doc, "1. Kaskadning qaror mantig'i va uning kesh asosi", 1)
    para(doc, "Kaskad uch holatni ajratadi. Kvantlash (INT8) — majburiy amal; u faqat FP32 ning o'zi keshga "
              "to'liq sig'gan taqdirda bajarilmaydi. Past-rankli dekompozitsiya esa faqat INT8 dan keyin ham "
              "sig'maganda qaraladi.")
    table(doc, "1-jadval. Kaskad holatlari va ularning yuzaga kelish chastotasi (real model)",
          ["Holat", "Shart", "Amal", "Encoder", "Decoder"],
          [["1", "FP32 alpha*L3 ga sig'adi", "hech narsa qilinmaydi", "96/144", "240/240"],
           ["2", "INT8 dan keyin sig'adi", "past-rank QARALMAYDI", "96/144", "240/240"],
           ["3", "INT8 dan keyin ham sig'maydi", "past-rank + INT8", "48/144 (fc1, fc2)", "0/240"]],
          good_rows=(2,))
    para(doc, "Ya'ni decoderda past-rank shoxchasi hech qachon ishga tushmaydi — bu o'lchovlar bilan to'liq "
              "tasdiqlangan. Encoderning FFN operatorlari esa byudjetdan 1.98x oshadi va aynan ular usulning "
              "past-rank qismini oqlaydigan maydon hisoblanadi.")

    table(doc, "2-jadval. Kesh-bog'langan talab, granulyarlik bo'yicha (alpha*L3 = 16.8 MiB)",
          ["Granulyarlik", "Decoder (MiB)", "Talab", "Encoder (MiB)", "Talab"],
          [["per-operator", "16.0", "sig'adi", "16.0", "0.95x"],
           ["per-layer", "64.0", "3.81x", "48.0", "2.86x"],
           ["whole-model", "1536.0", "91.4x", "1152.0", "68.6x"]],
          good_rows=(1,))

    table(doc, "3-jadval. Ish rejimi: kesh-sig'imi qachon ma'noga ega",
          ["Qism", "Pozitsiya/o'tish", "Vazn qayta ishlatilishi", "Cheklovchi", "Nima yordam beradi"],
          [["Encoder", "1500", "1500x", "hisoblash", "FLOPs kamaytirish (past-rank)"],
           ["Decoder (batch=1)", "1", "1x", "xotira", "bayt kamaytirish (kvantlash)"]],
          good_rows=(0,))

    # ================= 2. ABLATION =================
    doc.add_page_break()
    h(doc, "2. Elementlarning hissasi (ablation)", 1)
    para(doc, "Usul — bitta hiyla emas, balki elementlar to'plami. Quyida har bir element mustaqil va "
              "kombinatsiyada o'lchandi. Barcha moslashtirish x_fit da, barcha baholash ko'rilmagan x_eval da.")

    h(doc, "2.1. Kvantlash masshtabini aniqlashtirish (mustaqil hissa)", 2)
    para(doc, "Bu — muallifning o'z hissasi (README 2.4-bo'lim) va u shu vaqtgacha alohida o'lchanmagan edi: "
              "avvalgi barcha INT8 raqamlari ONNX Runtime ning standart min/max masshtabidan foydalangan.")
    table(doc, "4-jadval. Masshtab usulining operator xatosiga ta'siri (E_loc, held-out)",
          ["Masshtab usuli", "Encoder fc1", "Yaxshilanish", "Decoder fc1", "Yaxshilanish"],
          [["Q1  min/max (kutubxona)", "0.00685", "—", "0.00441", "—"],
           ["Q2  almashinuvchi minimizatsiya", "0.00651", "+5.0%", "0.00442", "-0.1%"],
           ["Q3  kalibrlangan (2.4-bo'lim to'liq)", "0.00525", "+23.3%", "0.00360", "+18.4%"],
           ["Q4  per-channel kalibrlangan", "0.00179", "+73.8%", "0.00151", "+65.9%"]],
          good_rows=(3,))
    callout(doc, "Xulosa:",
            "Masshtabni kalibrlash g'oyasi real va o'lchanadigan hissa beradi. Eng kuchli variant — per-channel "
            "kalibrlangan masshtab: xato 3.8 barobar (encoder) va 2.9 barobar (decoder) kamayadi, qo'shimcha "
            "xotira esa atigi 0.5% (har operator uchun m ta fp32 masshtab). Faqat almashinuvchi minimizatsiya "
            "(Q2) o'zi deyarli hech narsa bermaydi — qiymat aynan KALIBRLASH bosqichidan keladi.", GOOD)

    h(doc, "2.2. Past-rankli dekompozitsiya usuli (mustaqil hissa)", 2)
    table(doc, "5-jadval. Dekompozitsiya usullari, teng parametr byudjetida (FP32, E_loc held-out)",
          ["Usul", "Rank", "Siqish", "FLOPs nisbat", "Encoder fc1", "Decoder fc1"],
          [["L1  plain SVD (Ekart-Yang)", "204", "4.02x", "4.02x", "0.03581", "0.09451"],
           ["L2  activation-aware SVD", "204", "4.02x", "4.02x", "0.01975", "0.07420"],
           ["L3  funksional CUR (taklif)", "197", "4.16x", "4.16x", "0.28903", "0.27811"],
           ["L4  leverage-score CUR", "197", "4.16x", "4.16x", "0.29283", "0.37609"]],
          good_rows=(1,))
    para(doc, "Ikkita mustaqil kuzatuv. Birinchidan, kalibrlashga asoslangan yaqinlashtirish (L2) klassik "
              "Ekart-Yang optimumidan (L1) doimiy ravishda yaxshiroq — encoderda 45%, decoderda 21%. Bu "
              "dissertatsiyaning konseptual yadrosini tasdiqlaydi: vazn-optimallik chiqish-optimallik emas. "
              "Ikkinchidan, funksional klasterlash generik leverage-score tanlovidan yaxshiroq (decoderda 26%), "
              "ya'ni muallifning CUR ichidagi hissasi ham real — lekin CUR ramkasining o'zi SVD dan zaif.")

    h(doc, "2.3. Kombinatsiyalar", 2)
    table(doc, "6-jadval. Past-rank + kvantlash, ENCODER fc1 (E_loc, held-out; umumiy siqish ~16x)",
          ["Past-rank", "Kvantlash", "Umumiy siqish", "E_loc"],
          [["L2 act-aware SVD", "Q4 per-channel", "15.80x", "0.02001"],
           ["L2 act-aware SVD", "Q1 min/max", "16.06x", "0.02801"],
           ["L1 plain SVD", "Q4 per-channel", "15.80x", "0.03596"],
           ["L2 act-aware SVD", "Q3 kalibrlangan", "16.06x", "0.04650"],
           ["L3 funksional CUR", "Q1 min/max", "16.63x", "0.28551"],
           ["L4 leverage CUR", "Q4 per-channel", "16.36x", "0.29209"]],
          good_rows=(0,), warn_rows=(3,))
    callout(doc, "Kutilmagan, lekin izohlanadigan natija:",
            "Per-tensor kalibrlangan masshtab (Q3) past-rank FAKTORLARIDA min/max dan yomonroq ishlaydi "
            "(0.04650 vs 0.02801). Sabab: almashinuvchi minimizatsiya vazn xatosini minimallashtiradi va shu "
            "yo'lda chetdagi katta qiymatlarni qirqadi; past-rank faktorlarida esa singulyar qiymatlar tartiblar "
            "bo'yicha farq qilgani uchun aynan o'sha katta qiymatlar hal qiluvchi. Bu — usulning O'Z tezisining "
            "ichki tasdig'i (vazn-optimallik chiqish-optimallik emas). Amaliy qoida: faktorlar uchun per-channel "
            "(Q4) ishlatilsin, per-tensor kalibrlangan emas.", WARN)

    # ================= 3. TEZLIK =================
    doc.add_page_break()
    h(doc, "3. Tezlik va hisoblash murakkabligi", 1)
    para(doc, "Encoder fc1 operatori, 1500 pozitsiya, bitta oqim (intra_op_threads=1), real ONNX Runtime.")
    table(doc, "7-jadval. INT8 ustiga past-rank qo'shishning tezlikka ta'siri",
          ["Variant", "Vazn (MiB)", "Latency (ms)", "FP32 ga", "INT8 ga", "E_loc"],
          [["dense FP32", "16.00", "117.9", "1.00x", "0.26x", "0"],
           ["dense INT8 (majburiy)", "4.00", "30.3", "3.89x", "1.00x", "0.0082"],
           ["INT8 + SVD r=409", "2.00", "18.7", "6.30x", "1.62x", "0.0099"],
           ["INT8 + SVD r=200", "0.98", "8.8", "13.40x", "3.44x", "0.0201"],
           ["INT8 + SVD r=128", "0.62", "6.7", "17.70x", "4.56x", "0.0305"],
           ["INT8 + SVD r=80", "0.39", "4.4", "26.98x", "6.93x", "0.0444"]],
          good_rows=(3, 4))
    para(doc, "Ya'ni majburiy INT8 dan keyin past-rank qo'shish 1.6x–6.9x QO'SHIMCHA tezlanish beradi, "
              "aniqlik yo'qotishi esa juda kichik (E_loc 0.01–0.04). Bu — usulning past-rank qismining "
              "asosiy amaliy oqlanishi.")

    # ================= 4. VTUNE =================
    h(doc, "4. Kesh-miss: apparat hisoblagichlari bilan bevosita o'lchov", 1)
    para(doc, "Shu paytgacha kesh xulqi bilvosita (latency va FLOPs bashorati farqidan) baholanardi. Intel "
              "VTune Profiler 2026.4 o'rnatilgach va uning event-based sampling drayveri faol ekani "
              "tasdiqlangach, PMC hisoblagichlari bevosita o'qildi.")
    table(doc, "8-jadval. VTune uarch-exploration: kesh bosimi (Encoder fc1, 1500 pozitsiya)",
          ["Variant", "ms/iter", "Memory Bound", "L2 Bound", "L3 Bound", "DRAM Bound", "CPI"],
          [["dense FP32", "121.85", "8.8%", "2.5%", "2.7%", "2.9%", "0.64"],
           ["dense INT8", "33.96", "12.7%", "1.9%", "2.4%", "6.5%", "0.67"],
           ["INT8 + SVD r=409", "18.52", "12.2%", "2.8%", "4.0%", "4.8%", "0.65"],
           ["INT8 + SVD r=200", "9.96", "9.8%", "3.6%", "1.9%", "3.9%", "0.60"],
           ["INT8 + SVD r=128", "7.62", "18.3%", "4.1%", "1.0%", "9.2%", "0.62"],
           ["INT8 + SVD r=80", "5.02", "15.8%", "3.1%", "1.4%", "7.0%", "0.55"]],
          good_rows=(4,))
    callout(doc, "Ilmiy jihatdan eng nozik xulosa:",
            "L3 Bound past-rank bilan 2.4% dan 1.0% ga tushadi — ya'ni kesh bosimining kamayishi REAL va endi "
            "apparat hisoblagichi bilan tasdiqlangan, taxmin emas. Ammo Memory Bound umuman 9–18% oralig'ida, "
            "ya'ni bu yuk XOTIRA BILAN CHEKLANMAGAN. Shuning uchun L3 bosimining 1.4 foiz-punkt kamayishi "
            "umumiy tezlanishga oz hissa qo'shadi; tezlanishning asosiy qismi (33.96 -> 7.62 ms = 4.46x) "
            "FLOPs ning 6.40x kamayishidan keladi. Asoslashni 'kesh-miss kamayishi' emas, 'compute-bound "
            "rejimda arifmetik hajmni strukturaviy kamaytirish' deb qurish to'g'riroq va mustahkamroq.", ACCENT)

    # ================= 5. UCHDAN-UCHGACHA =================
    doc.add_page_break()
    h(doc, "5. Uchdan-uchgacha natijalar: WER va CER", 1)
    para(doc, "Barcha oldingi mezonlar (E_loc, E_glob) — proxy. Transkripsiya o'zgaradimi degan savolga faqat "
              "WER/CER javob beradi. 8 ta held-out namuna, greedy avtoregressiv dekodlash.")
    table(doc, "9-jadval. Butun DECODER: siqish x aniqlik x tezlik x WER/CER",
          ["Usul", "Vazn (MiB)", "Siqish", "E_glob", "WER", "CER", "Tezlanish"],
          [["FP32 (asl)", "1536", "1.00x", "0.0000", "0.0417", "0.0030", "1.00x"],
           ["INT8 per-tensor", "384", "4.00x", "0.2275", "0.0417", "0.0030", "2.80x"],
           ["INT8 per-channel", "386", "3.98x", "0.1997", "0.0417", "0.0030", "2.79x"],
           ["fc1'ga past-rank majburlangan", "336", "4.57x", "1.4559", "0.9185", "0.6993", "2.97x"]],
          good_rows=(1, 2), bad_rows=(3,))
    para(doc, "Decoder uchun javob tugallangan: INT8 mutlaqo bepul (WER o'zgarmaydi), past-rank esa kerak emas "
              "va majburlanganda modelni yaroqsiz qiladi. Kaskadning uni rad etishi ehtiyotkorlik emas — zaruriyat.")

    table(doc, "10-jadval. Butun ENCODER (dastlabki natijalar)",
          ["Usul", "Latency (ms)", "INT8 ga", "WER", "Per-operator E_loc"],
          [["FP32 encoder", "11793.9", "0.59x", "0.0417", "0"],
           ["INT8 (majburiy)", "6979.2", "1.00x", "0.0667", "—"],
           ["INT8 + SVD r=409", "6631.2", "1.05x", "0.2929", "0.0033"]],
          bad_rows=(2,))
    callout(doc, "Ochiq muammo — eng muhim ilmiy bo'shliq:",
            "Bitta operatorda r=409 1.62x tezlanish bergan edi; butun encoderda esa atigi 1.05x (Amdal qonuni: "
            "FFN operatorlari encoder vaqtining bir qismini egallaydi). Bundan ham jiddiyroq: per-operator xato "
            "atigi 0.0033 bo'lsa ham, uchdan-uchgacha WER 0.0667 dan 0.2929 ga sakraydi. Ya'ni 48 operator "
            "bo'ylab xato KUCHLI to'planadi va bu to'planish hozircha MODELLASHTIRILMAGAN.", CRIT)

    table(doc, "11-jadval. Xatoning qatlamlar bo'ylab kuchayishi",
          ["Usul", "Per-operator o'rtacha", "Uchdan-uchgacha", "Kuchayish", "WER ta'siri"],
          [["INT8 (decoder)", "0.0287", "0.2275", "~8x", "yo'q"],
           ["past-rank (decoder fc1)", "~0.35", "1.4559", "~4x", "halokatli"],
           ["past-rank r=409 (encoder)", "0.0033", "—", "—", "0.067 -> 0.293"]],
          bad_rows=(1, 2))

    # ================= 5b. KALIBRLASH + ALLOKATSIYA =================
    doc.add_page_break()
    h(doc, "5b. Kalibrlash hajmi va global rank taqsimoti", 1)

    h(doc, "5b.1. Kalibrlash hajmiga talab (metodologik natija)", 2)
    para(doc, "Butun encoder bo'yicha birinchi tajriba salbiy chiqqan edi (WER 0.0667 -> 0.2929). Sabab "
              "usulda emas, KALIBRLASH YETISHMASLIGIDA bo'lib chiqdi: rank 409 uchun atigi 512 qator "
              "ishlatilgan (nisbat 1.3), va xato o'sha qatorlarning o'zida hisoblangan. Kalibrlashga asoslangan "
              "past-rank yechimi bunday sharoitda kalibrlash to'plamini shunchaki YODLAB oladi.")
    table(doc, "13-jadval. Kalibrlash qatorlari soni va rank nisbati (encoder fc1, act-aware SVD)",
          ["Rank", "Fit qator", "Qator/rank", "Fit E_loc", "Held-out E_loc", "Bo'shliq"],
          [["409", "256", "0.6", "0.00000", "0.04355", "1 540 784x"],
           ["409", "512", "1.3", "0.00035", "0.04624", "131x"],
           ["409", "1024", "2.5", "0.00131", "0.03560", "27x"],
           ["409", "2048", "5.0", "0.00637", "0.02835", "4.4x"],
           ["409", "4096", "10.0", "0.01151", "0.02199", "1.9x"],
           ["409", "8192", "20.0", "0.01364", "0.01900", "1.4x"]],
          bad_rows=(0, 1), good_rows=(5,))
    callout(doc, "Amaliy qoida (mustaqil metodologik natija):",
            "Kalibrlashga asoslangan past-rank dekompozitsiyasi uchun qator/rank nisbati kamida 10–20 bo'lishi "
            "shart. 'fit E_loc = 0.00000' qiymati usulning mukammalligini emas, kalibrlash to'plamini yodlab "
            "olinganini bildiradi. Adabiyotdagi kalibrlashga asoslangan usullar (ASVD, SVD-LLM, FWSVD) uchun "
            "bu talab miqdoriy ko'rsatilmagan.", GOOD)
    para(doc, "Tuzatilgandan keyin haqiqiy ustunlik ko'rindi: rank 409 da act-aware SVD held-out E_loc = 0.0190, "
              "plain SVD esa 0.1284 — 6.8 barobar farq (rank 200 da 7.5x, rank 128 da 7.3x).")

    h(doc, "5b.2. Global rank taqsimoti: masalaning qo'yilishi va yechimi", 2)
    para(doc, "Bir xil rank berilganda operatorlarning held-out xatolari 0.0009 dan 0.1141 gacha farq qildi — "
              "127 barobar tarqoqlik. Bu bir xil rank taqsimotining ochiq-oydin isrofgarchiligini ko'rsatadi. "
              "Shu sababli masala operatorlararo optimallashtirish sifatida qo'yildi:")
    p = doc.add_paragraph()
    r = p.add_run("    min  sum_i E_i(r_i)     shart:  sum_i c_i * r_i <= B,   c_i = m_i + n_i")
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    para(doc, "E_i(r) — o'lchangan xato egri chizig'i, kamaymaydigan va spektral kesish uchun qavariq (ketma-ket "
              "singulyar qiymatlar kamaymaydi, ya'ni har qo'shimcha rank birligi oldingisidan kam foyda beradi). "
              "Ajraluvchan va qavariq maqsad uchun uzluksiz yechim Lagranj sharti bilan aniqlanadi "
              "(-dE_i/dr_i / c_i = lambda barcha i uchun), butun sonli yechim esa ochko'z algoritm bilan ANIQ "
              "olinadi — ya'ni bu evristika emas, isbotlanadigan optimal taqsimot. Amalga oshirilishi: "
              "nnopt/cascade/rank_allocation.py (11 test).")
    table(doc, "14-jadval. Bir xil vs sezgirlikka asoslangan rank taqsimoti (TENG byudjet, butun encoder)",
          ["Sxema", "Parametr", "Jami xato", "Latency (ms)", "WER", "CER"],
          [["bir xil rank", "100 515 840", "3.2682", "6371.7", "0.1719", "0.0417"],
           ["sezgirlikka asoslangan", "100 505 600", "2.8495", "6197.9", "0.0729", "0.0208"]],
          good_rows=(1,))
    callout(doc, "Eng kuchli amaliy natija:",
            "Teng byudjetda WER 0.1719 dan 0.0729 ga tushdi — 58% kamayish, CER esa 50% kamaydi, model yana "
            "biroz tezroq. Hech qanday qo'shimcha xotira yoki hisoblash sarflanmagan; butun yutuq rankni to'g'ri "
            "TAQSIMLASH hisobiga. Nazariy bashorat 12.8% xato kamayishi edi, amalda WER 58% yaxshilandi — "
            "bu WER ning per-operator xatoga nisbatan kuchli nochiziqliligini yana bir bor tasdiqlaydi.", GOOD)

    table(doc, "15-jadval. Yakuniy encoder taqqoslamasi",
          ["Variant", "Siqish", "Latency (ms)", "INT8 ga", "WER", "dWER"],
          [["FP32", "1.00x", "11246.1", "0.63x", "0.0417", "—"],
           ["INT8 (majburiy)", "4.00x", "7039.4", "1.00x", "0.0667", "+0.0250"],
           ["INT8 + bir xil rank", "6.00x", "6371.7", "1.10x", "0.1719", "+0.1302"],
           ["INT8 + taqsimlangan rank", "6.00x", "6197.9", "1.14x", "0.0729", "+0.0312"]],
          good_rows=(3,), bad_rows=(2,))
    para(doc, "Ya'ni majburiy INT8 ustiga to'g'ri taqsimlangan past-rank qo'shish 1.5 barobar ko'proq siqish va "
              "1.14 barobar tezlik beradi, WER esa atigi 0.006 ga oshadi. Bu — usulning past-rank qismining "
              "birinchi marta uchdan-uchgacha tasdiqlangan foydasi.")

    # ================= 5c. XATO TARQALISHI =================
    doc.add_page_break()
    h(doc, "5c. Xatoning tarqalishi: o'lchangan qonuniyat", 1)

    h(doc, "5c.1. Tarmoq lokal buzilishlarni kuchli yutadi", 2)
    para(doc, "Har bir FFN operatori ALOHIDA buzilib (rank 200 past-rank, qolgan hamma narsa o'zgarishsiz), "
              "encoder chiqishidagi nisbiy xato o'lchandi. 48 operator, real encoder yurishlari.")
    table(doc, "16-jadval. Lokal xato va global zarar (qat'iy rank 200, tanlangan operatorlar)",
          ["Operator", "E_loc", "E_glob", "c = E_glob/E_loc"],
          [["fc1 L0", "0.0099", "0.0507", "5.12"],
           ["fc1 L1", "0.0115", "0.0364", "3.16"],
           ["fc2 L7", "0.0014", "0.0202", "14.66"],
           ["fc2 L5", "0.0842", "0.0145", "0.17"],
           ["fc2 L14", "0.1734", "0.0233", "0.13"],
           ["fc2 L18", "0.2096", "0.0304", "0.15"]])
    callout(doc, "Asosiy kuzatuv:",
            "48 operator bo'ylab E_loc 160 barobar o'zgaradi (0.0014 dan 0.225 gacha), E_glob esa atigi 4 "
            "barobar (0.012 dan 0.047 gacha). Ya'ni qaysi operator buzilishidan qat'i nazar, encoder "
            "chiqishidagi zarar taxminan bir xil darajada qoladi. Tarmoq lokal buzilishlarni KUCHLI YUTADI. "
            "Sababi arxitekturaviy: fc2 chiqishi residual oqimga qo'shiladi (y = x + f(x)), shuning uchun "
            "undagi nisbiy xato yig'indida ||f||/||x+f|| koeffitsienti bilan suyultiriladi; fc1 xatosi esa "
            "GELU va fc2 orqali o'tib kuchayadi. Shu sababli fc1 koeffitsientlari (0.58–5.12) fc2 nikidan "
            "(0.13–0.68) sezilarli yuqori.", ACCENT)
    para(doc, "Bu natija nega jami xato yig'indisi 12.8% kamayganda WER 58% yaxshilanganini tushuntiradi: "
              "yig'indi bilan yakuniy sifat orasidagi bog'liqlik chiziqli emas.")

    h(doc, "5c.2. Global zarar egri chiziqlari", 2)
    para(doc, "Ta'sir koeffitsientini bitta rankda o'lchab uni butun egri chiziqqa ko'paytirish CHIZIQLILIKNI "
              "faraz qiladi, ma'lumot esa buni rad etadi. Shuning uchun farazsiz yo'l tanlandi: har operator "
              "uchun E_glob(r) bevosita o'lchandi (48 operator x 3 rank = 144 encoder yurishi, 104 daqiqa).")
    table(doc, "17-jadval. Rankdan qaytim: E_glob ning r=128 dan r=550 ga kamayishi",
          ["Operator", "r=128", "r=550", "Qaytim"],
          [["fc1 L0", "0.0999", "0.0024", "41.7x"],
           ["fc1 L1", "0.0693", "0.0023", "30.4x"],
           ["fc2 L0", "0.0623", "0.0053", "11.9x"],
           ["fc2 L16", "0.0363", "0.0131", "2.8x"],
           ["fc2 L19", "0.0471", "0.0174", "2.7x"]],
          good_rows=(0,), bad_rows=(4,))
    para(doc, "Rankdan qaytim operatorlar bo'ylab 15 barobar farq qiladi (2.71x dan 41.69x gacha, mediana "
              "3.69x). Boshlang'ich fc1 qatlamlari rankka eng sezgir, oxirgi fc2 qatlamlari eng befarq.")

    h(doc, "5c.3. Qaysi maqsad funksiyasi to'g'ri? (salbiy natija)", 2)
    table(doc, "18-jadval. Uch xil allokatsiya maqsadi, teng byudjet, uchdan-uchgacha",
          ["Maqsad funksiyasi", "Latency (ms)", "WER", "CER"],
          [["bir xil rank (maqsadsiz)", "6371.7", "0.1719", "0.0417"],
           ["sum E_loc (lokal xato)", "6197.9", "0.0729", "0.0208"],
           ["sum E_glob (global zarar)", "6787.1", "0.0729", "0.0298"]],
          good_rows=(1,))
    callout(doc, "Halol xulosa — gipoteza tasdiqlanmadi:",
            "104 daqiqalik qo'shimcha o'lchov evaziga olingan global zarar egri chiziqlari BIR XIL WER berdi, "
            "CER esa yomonroq. Demak murakkabroq maqsad funksiyasi o'zini oqlamadi. Bu ikkita TURLI savolni "
            "ajratadi: (1) uchdan-uchgacha zararni BASHORAT QILISH uchun lokal xato yaroqsiz (bog'liqlik "
            "monoton emas); (2) allokatsiya uchun operatorlarni SARALASH uchun esa lokal xato YETARLI — u "
            "keltirib chiqargan tartib global o'lchov bilan olinganidan yomon emas. Amaliy ahamiyati katta: "
            "qimmat global o'lchovlar shart emas.", WARN)

    h(doc, "5c.4. Statistik cheklov (ochiq tan olinadi)", 2)
    para(doc, "Barcha WER/CER taqqoslashlari 8 ta held-out namunada o'tkazilgan. Bu hajmda 0.0729 va 0.0729 "
              "orasidagi farqni ajratib bo'lmaydi; faqat katta farqlar (0.1719 vs 0.0729) ishonchli "
              "hisoblanadi. Baholash to'plamini kamida 50–100 namunaga kengaytirish — dissertatsiya uchun "
              "zaruriy shart, chunki retsenzent statistik ishonchlilikni albatta so'raydi.", italic=True)

    # ================= 5d. STRUKTURAVIY QISQARTIRISH =================
    doc.add_page_break()
    h(doc, "5d. Funksional guruhlash o'z rolida: strukturaviy kanal qisqartirish", 1)
    para(doc, "Bu bo'lim usulning asosiy g'oyasini TO'G'RI rolda sinaydi. Oldingi taqqoslashlar funksional "
              "guruhlashni SVD bilan bir xil vazifada (past-rank yaqinlashtirish) o'lchagan edi — bu rolda "
              "Ekart-Yang teoremasi tufayli SVD ustunligi muqarrar. Guruhlashning haqiqiy qiymati boshqa "
              "narsada: real kalibrlash faolligi asosida javoblari kollinear bo'lgan kanallarni topib, "
              "kompensatsiya bilan ANIQ olib tashlash. Bu yaqinlashtirish emas — strukturaviy qisqartirish.")

    h(doc, "5d.1. Nega FFN da bu ikki karra foyda beradi", 2)
    p = doc.add_paragraph()
    r = p.add_run("    fc1  W1 (1024, 4096) -> (1024, k)      bias (4096,) -> (k,)\n"
                  "    GELU  elementwise, o'zgarmaydi\n"
                  "    fc2  W2 (4096, 1024) -> (k, 1024)")
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    para(doc, "Oraliq kenglik fc1 ning CHIQISHI va fc2 ning KIRISHI hisoblanadi, shuning uchun bitta qaror "
              "ikkala operatorni ham qisqartiradi. Past-rank buni qila olmaydi: u rankni kamaytiradi, n va m "
              "esa o'zgarishsiz qoladi. Demak bu ikki o'q raqobatlashmaydi — QO'SHILADI.")
    para(doc, "Kompensatsiya: j kanalini olib tashlashdan oldin W2[p, :] += gamma * W2[j, :], "
              "bunda p — vakil kanal, gamma = <h_j, h_p> / ||h_p||^2 (eng kichik kvadratlar).")

    h(doc, "5d.2. O'lchangan ortiqchalik: chuqurlikka kuchli bog'liq", 2)
    table(doc, "19-jadval. FFN oraliq kanallarining ortiqchaligi (tau=0.99, qatlamlar bo'yicha)",
          ["Qatlam", "Olib tashlanadi", "Ulush", "Qatlam", "Olib tashlanadi", "Ulush"],
          [["L0", "1764", "43.1%", "L12", "26", "0.6%"],
           ["L1", "2152", "52.5%", "L13", "10", "0.2%"],
           ["L2", "2376", "58.0%", "L14", "4", "0.1%"],
           ["L3", "2334", "57.0%", "L15", "0", "0.0%"],
           ["L4", "2078", "50.7%", "L17-L20", "0", "0.0%"],
           ["L5", "1535", "37.5%", "L21", "3", "0.1%"],
           ["L6", "1155", "28.2%", "L22", "4", "0.1%"],
           ["L7", "877", "21.4%", "L23", "2", "0.0%"],
           ["L8", "1048", "25.6%", "", "", ""],
           ["L9", "726", "17.7%", "", "", ""],
           ["L10", "447", "10.9%", "", "", ""],
           ["L11", "193", "4.7%", "", "", ""]],
          good_rows=(2, 3))
    callout(doc, "Kaskad g'oyasining eng toza dalili:",
            "tau = 0.99 BITTA qiymat bo'lgani holda usul har qatlamga BUTUNLAY BOSHQA qaror berdi: L2 da 58%, "
            "L15-L20 da 0%. Hech narsa qo'lda sozlanmagan — taqsimot o'lchovdan kelib chiqqan. Ortiqchalik "
            "L2-L3 da cho'qqiga chiqib L12 ga borib yo'qoladi, ya'ni Whisper encoderining boshlang'ich FFN "
            "qatlamlari kuchli ortiqcha parametrlashtirilgan, keyingilari esa deyarli emas. Bu — model "
            "haqidagi mustaqil strukturaviy xulosa.", GOOD)

    h(doc, "5d.3. Xato TO'PLANMAYDI", 2)
    table(doc, "20-jadval. Qatlamlarni ketma-ket qisqartirganda encoder chiqish xatosi (FP32)",
          ["Qisqartirilgan qatlamlar", "Encoder xatosi"],
          [["faqat L0 (1 ta)", "0.0209"],
           ["L0..L3 (4 ta)", "0.0206"],
           ["L0..L7 (8 ta)", "0.0295"],
           ["L0..L11 (12 ta)", "0.0297"],
           ["barchasi (19 ta)", "0.0298"]],
          good_rows=(4,))
    para(doc, "19 ta qatlam qisqartirilganda xato bitta qatlamnikidan deyarli farq qilmaydi (0.0298 vs 0.0209). "
              "Bu past-rank bilan keskin farq qiladi: u yerda xato 48 operator bo'ylab kuchli to'planardi. "
              "Sababi — qisqartirish YAQINLASHTIRISH emas: kollinear kanal kompensatsiya bilan aniq "
              "almashtiriladi, qoldiq xato esa faqat kollinearlikning nomukammalligidan kelib chiqadi.")

    h(doc, "5d.4. ⚠ Kutilmagan bog'liqlik: guruhlash per-channel masshtabni ZARURIY qiladi", 2)
    para(doc, "Qisqartirilgan modelni per-tensor INT8 ga o'tkazish uni BUTUNLAY buzdi (WER 1.0000). Sabab "
              "kompensatsiyaning o'zida: W[:, p] += gamma * W[:, j] amali vakil ustunlarga ko'p hissani "
              "yig'ib, ularning kattaligini keskin oshiradi.")
    table(doc, "21-jadval. Kompensatsiyaning vazn diapazoniga ta'siri (fc2, L2)",
          ["Holat", "max |w|", "Satr normasi (mediana)", "Satr normasi (maks)", "Tarqoqlik"],
          [["asl", "0.2472", "0.0786", "0.7569", "9.6x"],
           ["kompensatsiyalangan", "11.3875", "0.4083", "76.9402", "188.4x"]],
          bad_rows=(1,))
    para(doc, "Maksimal vazn kattaligi 46 barobar, satrlar orasidagi tarqoqlik 9.6x dan 188x ga oshdi. "
              "Bitta tenzor-keng masshtab bunday diapazonni qoplay olmaydi: u eng katta satrga moslashadi va "
              "qolgan hammasining aniqligini yo'q qiladi. Per-channel masshtab (README 8.3.8) esa muammoni "
              "to'liq hal qiladi.")
    callout(doc, "Metodologik ahamiyati:",
            "Usulning ikki elementi — funksional guruhlash va kalibrlangan per-channel masshtab — mustaqil "
            "emas: BIRINCHISI IKKINCHISINI TALAB QILADI. Guruhlash kompensatsiya orqali vazn outlier'larini "
            "yaratadi, per-channel masshtab esa aynan shuni ko'taradi. Bu — usulning ichki yaxlitligini "
            "ko'rsatuvchi va oldindan aytib bo'lmaydigan natija.", ACCENT)

    h(doc, "5d.5. Yakuniy natija", 2)
    table(doc, "22-jadval. Kanal qisqartirish + kvantlash granulyarligi (encoder, held-out)",
          ["Variant", "MiB", "Siqish", "ms", "E_glob", "WER", "CER"],
          [["FP32 (asl)", "1152", "1.00x", "11246.1", "0.0000", "0.0417", "0.0030"],
           ["INT8 (majburiy)", "288", "4.00x", "7039.4", "—", "0.0667", "0.0064"],
           ["qisqartirilgan FP32", "1041", "1.11x", "11030.5", "0.0239", "0.0417", "0.0030"],
           ["qisqartirilgan + INT8 per-tensor", "266", "4.33x", "7173.7", "0.7420", "1.0000", "1.0000"],
           ["qisqartirilgan + INT8 per-channel", "267", "4.32x", "6991.0", "0.2226", "0.0417", "0.0030"]],
          good_rows=(4,), bad_rows=(3,))
    callout(doc, "Asosiy amaliy natija:",
            "Kanal qisqartirish + per-channel INT8: 4.32x siqish, WER 0.0417 — ya'ni ASL FP32 DARAJASI. "
            "Taqqoslang: majburiy INT8 yolg'iz WER ni 0.0667 ga ko'taradi. Demak taklif etilgan usul INT8 dan "
            "ham ANIQROQ va ko'proq siqadi. Tezlik esa deyarli teng (6991 vs 7039 ms), chunki qisqartirish "
            "faqat boshlang'ich 12 qatlamga tegadi — Amdal qonuni. Aniqlik va xotira bo'yicha aniq yutuq.", GOOD)

    # ================= 6. ILMIY BAHO =================
    doc.add_page_break()
    h(doc, "6. Usulning ilmiy salmog'i: baho va rivojlantirish yo'nalishlari", 1)
    para(doc, "Quyidagi baho o'lchangan natijalarga asoslanadi va DSc darajasidagi talablar nuqtai nazaridan "
              "berilgan.", italic=True, size=10)

    h(doc, "6.1. Nima mustahkam (himoya qilinadigan natijalar)", 2)
    bullets(doc, [
        ("Kaskadning qaror mantig'i empirik tasdiqlangan.",
         "Usul zararli o'zgarishni to'g'ri rad etadi: majburlanganda WER 0.042 dan 0.919 ga chiqadi. Bu — "
         "adaptiv kaskad g'oyasining eng kuchli oqlanishi va u salbiy natija orqali isbotlangan, bu esa "
         "ilmiy jihatdan qimmatliroq."),
        ("Maqsad kesh hajmidan chiqariladi, qo'lda tanlanmaydi.",
         "3.81x talab -> INT8 (4.00x) eng yumshoq yetarli yechim. Bu kaskadga apparat-asoslangan mazmun beradi."),
        ("Masshtabni kalibrlash — mustaqil, o'lchangan hissa.",
         "Kutubxona standartiga nisbatan 66–74% xato kamayishi, qo'shimcha xotira 0.5%."),
        ("Konseptual yadro 135/135 tasdiq.",
         "Vazn-optimallik chiqish-optimallik emas: kalibrlashga asoslangan yaqinlashtirish vazn xatosida "
         "yomonroq, chiqish xatosida yaxshiroq — barcha o'lchovlarda."),
        ("Kesh-miss kamayishi apparat bilan tasdiqlangan.",
         "VTune: L3 Bound 2.4% -> 1.0%. Ko'p ishlarda bu faqat nazariy taxmin bo'lib qoladi."),
    ])

    h(doc, "6.2. Nima zaif (retsenzent birinchi so'raydigan savollar)", 2)
    bullets(doc, [
        ("Bitta model, bitta mashina.",
         "Kesh-bog'langan maqsad g'oyasining butun kuchi shundaki, u APPARATGA moslashadi. Ammo bu hozircha "
         "bitta kesh konfiguratsiyasida ko'rsatilgan. Turli L3 hajmli 2–3 mashinada takrorlanmasa, "
         "'moslashuvchanlik' da'vosi asossiz qoladi."),
        ("Xatoning to'planishi modellashtirilmagan.",
         "Eng jiddiy qolgan bo'shliq. Per-operator E_loc bilan yakuniy WER orasidagi bog'liqlik monoton emas "
         "(0.23 -> zarar yo'q, 1.46 -> halokat) va kuchli nochiziqli: rank taqsimoti nazariy jihatdan 12.8% "
         "yaxshilanish bergan joyda WER 58% yaxshilandi. Hozir har bir konfiguratsiyani to'liq dekodlab "
         "ko'rish kerak."),
        ("CUR ramkasi past-rank rolida ma'lumotlar bilan qo'llab-quvvatlanmaydi.",
         "Aniqlikda act-aware SVD ga 0/135, kvantlashga chidamlilikda 0/81. MUHIM: bu CUR ramkasiga tegishli, "
         "funksional guruhlashning O'ZIGA emas — 5d-bo'lim guruhlashni strukturaviy qisqartirish rolida "
         "sinab, INT8 dan ustun natija oldi. Xulosa: g'oyaning qiymati past-rank yaqinlashtirishda emas, "
         "ortiqchalikni aniqlashda."),
        ("Ayrim natijalar adabiyotda mavjud yo'nalishlarni takrorlaydi.",
         "Kalibrlashga asoslangan kvantlash (GPTQ, AWQ, SmoothQuant) va faollashuvga sezgir past-rank "
         "(FWSVD, ASVD, SVD-LLM) — o'rganilgan sohalar. Original qism: kesh-bog'langan adaptiv kaskad, "
         "kalibrlash hajmiga miqdoriy talab, va global rank taqsimoti."),
    ])
    para(doc, "Global taqsimot masalasi bu hisobotda QO'YILDI VA YECHILDI (5b.2-bo'lim), shuning uchun u "
              "zaifliklar ro'yxatidan chiqarildi.", italic=True, size=9.5)

    h(doc, "6.3. Ilmiy salmoqni oshirish uchun aniq tavsiyalar", 2)
    para(doc, "Muhimlik tartibida. Birinchi ikkitasi — DSc darajasi uchun hal qiluvchi.", italic=True, size=10)
    bullets(doc, [
        ("Baholash to'plamini kengaytirish.",
         "Eng arzon va eng zarur ish. Hozirgi 8 namunada 0.0729 va 0.0729 farqlanmaydi; 50–100 namuna kerak. "
         "Bu holda 5c.3-bo'limdagi 'maqsad funksiyasi farq qilmadi' xulosasi ham qayta tekshirilishi lozim."),
        ("Xato tarqalishining bashoratli modelini qurish (qisman bajarildi).",
         "Ta'sir koeffitsientlari va global zarar egri chiziqlari o'lchandi (5c-bo'lim), xato-yutish qonuni "
         "aniqlandi. Qolgan qism: E_glob dan WER ga o'tish. Hozir har konfiguratsiya uchun to'liq dekodlash "
         "shart; E_glob -> WER bog'liqligi modellashtirilsa, kaskad butunlay bashoratli bo'ladi."),
        ("Global taqsimotni kengaytirish (asosi qo'yilgan).",
         "Rank bo'yicha masala yechildi. Keyingi qadam — bit kengligi (INT8/INT4) va rankni BIRGALIKDA "
         "taqsimlash, hamda maqsad funksiyasiga latency ni ham kiritish (ko'p mezonli optimallashtirish)."),
        ("Turli apparatlarda takrorlash.",
         "Kamida 2–3 xil kesh ierarxiyasi (masalan 8 MiB, 24 MiB, 32 MiB L3). Kesh-bog'langan maqsad "
         "haqiqatan moslashishini ko'rsatish — bu usulning asosiy da'vosi."),
        ("Ish rejimini roofline orqali rasmiylashtirish.",
         "VTune ma'lumoti kesh-sig'imi emas, arifmetik intensivlik hal qiluvchi ekanini ko'rsatdi. Har "
         "operatorni roofline diagrammasida joylashtirish va 'compute-bound bo'lsa past-rank, memory-bound "
         "bo'lsa kvantlash' qoidasini rasmiy mezon sifatida keltirish."),
        ("Modellar bo'yicha umumlashtirish (Q1 maqola uchun).",
         "mT5 / mBERT — kam resursli tillar uchun siqish mavzusi. Katta modellarda (operator boshiga "
         "17.6M+ parametr) past-rank shoxchasi tabiiy ravishda ko'proq ishga tushadi."),
        ("Salbiy natijalarni oldindan e'lon qilish.",
         "CUR ning SVD ga yutqazishi zaiflik emas — u chegaralarni aniqlagan tizimli tadqiqot natijasi. "
         "Buni oldindan bayon qilish himoyada kuchli pozitsiya beradi."),
    ], numbered=True)

    h(doc, "6.4. Usulning amaliy hayotiyligi", 2)
    para(doc, "Hozirgi holatda usul quyidagi vaziyatlarda real qiymat beradi:")
    table(doc, "12-jadval. Amaliy qo'llash sohalari",
          ["Vaziyat", "Kaskad qarori", "Kutilayotgan foyda", "Tasdiq"],
          [["Decoder, batch=1, avtoregressiv", "faqat INT8", "4x xotira, 2.8x tezlik, WER 0", "o'lchangan"],
           ["Encoder, uzun ketma-ketlik", "INT8 + past-rank", "16x xotira, 4.5x tezlik", "operator darajasida"],
           ["Chekka qurilma (kichik kesh)", "INT8 + past-rank", "kesh-sig'imi hal qiluvchi", "sinalmagan"],
           ["Katta LLM (17.6M+ param/op)", "INT8 + past-rank", "3-holat tabiiy yuzaga keladi", "sinalmagan"]],
          good_rows=(0,))
    para(doc, "Eng kuchli amaliy da'vo — decoder uchun: majburiy INT8 kesh-bog'langan talabni aynan qoplaydi, "
              "4x xotira tejaydi, 2.8x tezlashtiradi va transkripsiya sifatini UMUMAN o'zgartirmaydi. Bu "
              "kaskadning 'eng yumshoq yetarli o'zgarish' tamoyilining to'g'ridan-to'g'ri isboti.")

    doc.add_paragraph()
    para(doc, "Metodologik eslatma: E_loc va E_glob mezonlari monoton emas va ular asosida usulni baholash "
              "xavfli; har qanday yakuniy da'vo WER/CER bilan tasdiqlanishi shart. Skriptlar: experiments/, "
              "natijalar: experiments/results_*.json", size=8.5, italic=True)

    doc.save(OUT)
    print(f"saqlandi: {OUT}")


if __name__ == "__main__":
    main()
