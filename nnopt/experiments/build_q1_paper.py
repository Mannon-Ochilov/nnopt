"""Q1 journal manuscript (MDPI Sensors format) from all measured results.

Structure follows the Sensors template: Abstract, Keywords, Introduction,
Related Work, Materials and Methods, Results, Discussion, Conclusions, and
the required back matter. Equations are written in plain notation that
converts cleanly to LaTeX/MathType. Figure placeholders carry generation
prompts so the author can render them.

Every number is taken from the measured JSON outputs in experiments/.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "../Q1_Sensors_manuscript.docx"
ACCENT = RGBColor(0x0F, 0x64, 0x70)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def eq(doc, text, number=None):
    """Display equation, numbered on the right as MDPI requires."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if number:
        r2 = p.add_run(f"\t\t({number})")
        r2.font.size = Pt(10.5)
    return p


def mono(doc, text, size=8.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def table(doc, caption, headers, rows, good_rows=(), bad_rows=()):
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ht in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(ht)
        run.bold = True
        run.font.size = Pt(8)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            if ri in good_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in bad_rows:
                run.font.color.rgb = CRIT
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def bullets(doc, items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + it[1])
            r2.font.size = Pt(10.5)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10.5)


def figure(doc, num, caption, prompt):
    p = doc.add_paragraph()
    r = p.add_run(f"Figure {num}. {caption}")
    r.bold = True
    r.font.size = Pt(9)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("[AI generation prompt] ")
    r2.bold = True
    r2.font.size = Pt(8)
    r2.font.color.rgb = ACCENT
    r3 = p2.add_run(prompt)
    r3.font.size = Pt(8)
    r3.font.color.rgb = MUTED
    p2.paragraph_format.left_indent = Pt(14)
    doc.add_paragraph()


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Palatino Linotype"
    st.font.size = Pt(10.5)

    # ===================== FRONT MATTER =====================
    h(doc, "Cache-Anchored Cascade Optimization of Transformer Operators: "
           "Structural Redundancy Removal, Calibrated Quantization, and "
           "Budget-Optimal Rank Allocation", 0)
    para(doc, "Firstname Lastname 1,*, Co-Author 2", italic=True, size=10)
    mono(doc,
         "1  Department, University, City, Country; author@email\n"
         "2  Department, University, City, Country; coauthor@email\n"
         "*  Correspondence: author@email", 8.5)

    h(doc, "Abstract", 1)
    para(doc,
         "Post-training compression of transformer models is usually driven by a "
         "target compression ratio chosen by hand, and evaluated by weight- or "
         "layer-level error proxies. We show that both choices are unsound, and "
         "propose a cascade in which the compression target is DERIVED from the "
         "hardware cache topology and the transformation is selected per operator "
         "by measurement. The method has four components: (i) a cache-anchored "
         "target obtained from the guaranteed shared cache; (ii) functional "
         "grouping, which identifies channels whose calibration responses are "
         "collinear and removes them exactly by compensation, reducing operator "
         "DIMENSIONS rather than rank; (iii) calibrated per-output-channel "
         "quantization; and (iv) budget-optimal rank allocation formulated as a "
         "separable convex program whose greedy solution is exact. On a "
         "Whisper-medium Uzbek ASR encoder, structural removal of 17.1% of "
         "feed-forward channels with per-channel INT8 yields 4.32x compression at "
         "a word error rate statistically indistinguishable from FP32 (WER 0.1107 "
         "vs 0.1007, 95% CI of the difference [-0.0020, +0.0240]), whereas plain "
         "INT8 degrades WER significantly (+0.0139, CI [+0.0046, +0.0259]). We "
         "further report three findings of independent interest: compensation "
         "widens the per-row weight dynamic range from 9.6x to 188.4x and thereby "
         "MAKES per-channel quantization mandatory; calibration-based low-rank "
         "factorization requires a rows-to-rank ratio of at least 10-20, below "
         "which it memorizes the calibration set (fit error 0.00000 versus "
         "held-out 0.04355); and transformer blocks absorb local perturbations "
         "strongly, with per-operator error spanning 160x while the network output "
         "error spans only 4x. Cross-architecture evaluation on mBERT and "
         "open_llama_3b shows that the cache-anchored diagnostic transfers while "
         "the specific tool does not: feed-forward redundancy is abundant in the "
         "audio encoder (58% peak), modest in Llama (25% at a relaxed threshold), "
         "and absent in mBERT. All measurements use real hardware counters (Intel "
         "VTune), real runtimes, and held-out data.",
         size=10)

    h(doc, "Keywords", 1)
    para(doc, "post-training compression; cache-aware optimization; structural pruning; "
              "low-rank factorization; per-channel quantization; automatic speech "
              "recognition; low-resource languages; edge inference", italic=True, size=10)

    # ===================== 1. INTRODUCTION =====================
    doc.add_page_break()
    h(doc, "1. Introduction", 1)
    para(doc,
         "Deploying transformer models on CPU-class hardware is limited less by "
         "parameter count than by the interaction between an operator's memory "
         "footprint and the cache hierarchy that must hold it. Existing "
         "post-training methods - quantization [GPTQ, AWQ, SmoothQuant] and "
         "low-rank factorization [FWSVD, ASVD, SVD-LLM] - compress to a ratio the "
         "practitioner selects, typically 4x or 8x, without reference to the "
         "machine that will execute the result. Two consequences follow. First, "
         "the chosen ratio may be unnecessary (the operator already fits) or "
         "unreachable (no ratio helps because activations, not weights, dominate). "
         "Second, because the ratio is fixed a priori, the same transformation is "
         "applied to operators with very different tolerance.")
    para(doc,
         "This paper makes the target a derived quantity. Given the guaranteed "
         "shared cache of the execution platform, each operator's required "
         "reduction follows from its measured footprint, and the cascade then "
         "applies the SOFTEST SUFFICIENT transformation. On the machine used here "
         "(24 MiB L3 shared by 16 logical cores) the derived per-layer requirement "
         "for the Whisper decoder is 3.81x, and INT8 supplies exactly 4.00x - the "
         "cascade therefore selects quantization and rejects further compression, "
         "a decision we verify is correct by measuring what happens when it is "
         "overridden.")

    h(doc, "1.1. Contributions", 2)
    bullets(doc, [
        ("Cache-anchored target derivation.",
         "The compression requirement is computed from cache topology and operator "
         "footprint rather than chosen. We show the derived target explains which "
         "transformation is appropriate on three different architectures."),
        ("Functional grouping as structural removal.",
         "We identify channels whose calibration responses are collinear and remove "
         "them exactly via least-squares compensation. Unlike low-rank "
         "approximation this reduces operator dimensions, so in a feed-forward "
         "block one decision shrinks two matrices (three in gated architectures)."),
        ("A coupling between two components.",
         "Compensation concentrates mass into representative columns, widening the "
         "per-row dynamic range from 9.6x to 188.4x. Per-tensor INT8 then destroys "
         "the model (WER 1.0000) while per-channel INT8 preserves it. The two "
         "components are therefore not independent: grouping REQUIRES calibrated "
         "per-channel scales."),
        ("A quantitative calibration-size requirement.",
         "Calibration-based low-rank factorization needs a rows-to-rank ratio of at "
         "least 10-20. Below it the solution interpolates the calibration set: at "
         "ratio 0.6 the fit error is 0.00000 while the held-out error is 0.04355."),
        ("Budget-optimal rank allocation.",
         "Cross-operator allocation is posed as a separable convex program; the "
         "greedy solution satisfies the Lagrangian equal-marginal-return condition "
         "and is exact. At equal budget it reduces WER from 0.1719 to 0.0729."),
        ("Measured error-absorption law.",
         "Perturbing single operators shows per-operator error spanning 160x while "
         "network output error spans only 4x, explained by the residual stream "
         "diluting relative error. Confirmed on both an audio encoder and a "
         "decoder-only LLM."),
        ("Delimiting negative results.",
         "We report where the method does not help: CUR-style assembly loses to "
         "activation-aware SVD in 135/135 comparisons, and feed-forward redundancy "
         "is architecture-dependent."),
    ], numbered=True)

    figure(doc, 1,
           "Overview of the cache-anchored cascade. The compression target is derived "
           "from the cache hierarchy; two orthogonal reduction axes are applied "
           "conditionally; calibrated per-channel quantization is mandatory.",
           "A clean black-and-white technical block diagram for an IEEE/MDPI paper, "
           "white background, thin black rectangular boxes, orthogonal connector "
           "arrows, no color fill, no gradients, no shadows, sans-serif labels. "
           "Left: a box 'Operator footprint M_W + M_X + M_Y' feeding a box "
           "'Cache-anchored target: required = M_eff / (alpha x L3)'. Center: two "
           "parallel horizontal branches. Upper branch titled 'Axis 1: structural "
           "reduction' with boxes 'Calibration responses h_j' -> 'Representative "
           "nodes J_rep (cos >= tau)' -> 'Compensation W[:,p] += gamma_j W[:,j]' -> "
           "'Channel removal n->n-k'. Lower branch titled 'Axis 2: spectral "
           "reduction (case 3 only)' with boxes 'Gram matrix G = X^T X = L L^T' -> "
           "'Activation-aware SVD' -> 'Rank from cache budget r <= alpha L3/(m+n)'. "
           "Both branches converge into one wide box 'Mandatory: calibrated "
           "per-channel INT8'. A curved dashed arrow from the Compensation box to "
           "this box labeled 'dynamic range 9.6x -> 188.4x makes per-channel "
           "mandatory'. Right: box 'Optimized operator' then a diamond 'WER/CER "
           "gate' with a 'no' arrow looping back labeled 'rollback'. Bottom: three "
           "small legend boxes reading 'Case 1: FP32 fits - no change', 'Case 2: "
           "INT8 suffices - no low-rank', 'Case 3: still over budget - add "
           "low-rank'.")

    # ===================== 2. RELATED WORK =====================
    h(doc, "2. Related Work", 1)
    para(doc,
         "Post-training quantization. GPTQ and AWQ refine quantization scales using "
         "calibration activations, and SmoothQuant redistributes activation outliers "
         "into weights. These establish that calibration-aware scale selection "
         "outperforms min/max, which our measurements confirm (Section 4.2). Our "
         "addition is not the calibrated scale itself but the finding that "
         "structural compensation makes per-channel granularity mandatory rather "
         "than merely beneficial.")
    para(doc,
         "Low-rank factorization. FWSVD, ASVD and SVD-LLM weight the factorization "
         "by activation statistics. We reproduce their central claim - the "
         "activation-aware solution beats the Eckart-Young optimum on output error "
         "in 135/135 measurements while being worse on weight error - and add a "
         "quantitative calibration-size requirement that, to our knowledge, has not "
         "been reported.")
    para(doc,
         "Structured pruning. Channel and neuron pruning for transformers is well "
         "studied, generally with magnitude or gradient criteria and followed by "
         "fine-tuning. Our variant is purely post-training, uses a collinearity "
         "criterion on calibration responses, and compensates exactly, so no "
         "retraining is required.")
    para(doc,
         "Cache- and roofline-aware optimization. Hardware-aware scheduling is "
         "standard in kernel libraries but is rarely used to SELECT a compression "
         "method. The cache-anchored target of Section 3.1 is the component we "
         "believe is new.")

    # ===================== 3. METHODS =====================
    doc.add_page_break()
    h(doc, "3. Materials and Methods", 1)

    h(doc, "3.1. Cache-Anchored Compression Target", 2)
    para(doc,
         "For a matrix operator with weights W of shape (m, n) and calibration "
         "activations X of shape (B, n), the footprint components are the weight "
         "bytes M_W, input M_X, output M_Y and workspace M_tmp. Because the set of "
         "cores executing an operator is not fixed a priori, the target cache must "
         "be the level guaranteed shared by ALL logical processors (L3 on our "
         "platform). With utilization coefficient alpha the effective budget and "
         "the cache pressure are")
    eq(doc, "M_cache_eff = alpha · M_cache ,    K_cache = M_eff / M_cache_eff", 1)
    para(doc, "and the required reduction factor is")
    eq(doc, "rho = max(1, K_cache) .", 2)
    para(doc,
         "M_eff admits two bracketing definitions: the upper bound M_total = M_W + "
         "M_X + M_Y + M_tmp, and a blocked-GEMM estimate. We report both and note "
         "in Section 5 which is physically relevant for each execution regime.")
    para(doc,
         "The cascade distinguishes three cases: (1) FP32 already fits, no change; "
         "(2) mandatory INT8 makes it fit, low-rank is NOT considered; (3) still "
         "over budget after INT8, low-rank is added. Case frequencies are reported "
         "in Section 4.1.")

    figure(doc, 2,
           "Execution-regime dependence of the cache criterion. Weight reuse decides "
           "whether cache residency or byte volume is the binding constraint.",
           "A two-panel black-and-white scientific diagram, white background, flat "
           "line art, no color. Left panel titled 'Encoder (compute-bound)': a "
           "weight matrix block with 1500 stacked thin activation slices streaming "
           "through it, an arrow labeled '1500x weight reuse', and a bottleneck "
           "symbol placed on an ALU/compute icon. Right panel titled 'Decoder, "
           "batch = 1 (memory-bound)': the same weight matrix block with a single "
           "thin activation slice, arrow labeled '1x reuse', bottleneck symbol on a "
           "DRAM icon, and 23 greyed-out layer blocks illustrating cache eviction "
           "before the weight is reused. Thin black lines, sans-serif labels, "
           "publication quality.")

    h(doc, "3.2. Functional Grouping and Structural Removal", 2)
    para(doc,
         "For hidden node j let h_j denote its functional response vector, formed by "
         "concatenating that channel's activations over all non-padding calibration "
         "positions. Two nodes are functionally redundant when their responses are "
         "collinear. We require two conditions simultaneously:")
    eq(doc, "cos(h_j, h_p) = <h_j, h_p> / (||h_j|| ||h_p||) >= tau ,", 3)
    eq(doc, "eps_j = ||W[:, j]|| · ||h_j|| · sin(theta_jp) / (||Y|| + xi) <= eps_thr .", 4)
    para(doc,
         "The second condition is essential: angular proximity alone ignores how "
         "much the channel contributes to the operator output. Equation (4) weights "
         "the residual by the column norm of W and normalizes by the calibration "
         "output magnitude, so eps_j is the channel's relative contribution to "
         "output error if it is merged.")
    para(doc,
         "Within a group anchored at p, the least-squares optimal compensation "
         "coefficient for member j and the resulting weight update are")
    eq(doc, "gamma_j = <h_j, h_p> / ||h_p||^2 ,    W[:, p] <- W[:, p] + gamma_j W[:, j] ,", 5)
    para(doc,
         "after which column j is deleted. If h_j = gamma_j h_p exactly the "
         "substitution is lossless; the residual error is governed only by the "
         "departure from collinearity. The group representative is the member "
         "maximizing cosine similarity to the group mean.")
    para(doc,
         "In a feed-forward block the intermediate width is the OUTPUT dimension of "
         "the first projection and the INPUT dimension of the second, so removing k "
         "intermediate channels shrinks both matrices from one decision:")
    mono(doc, "    W1 (d, F) -> (d, F-k)     bias (F,) -> (F-k,)\n"
              "    activation: elementwise, unchanged\n"
              "    W2 (F, d) -> (F-k, d)")
    para(doc,
         "In gated architectures (Llama-style, h = SiLU(W_gate x) * (W_up x)) the "
         "same decision shrinks three matrices. This is the structural axis; it "
         "reduces n and m, whereas low-rank factorization reduces rank while "
         "leaving n and m unchanged. The two axes are therefore complementary: "
         "removing k channels frees budget for a higher rank at equal parameter "
         "count, since r'(m + n - k) = r(m + n) implies r' > r.")

    figure(doc, 3,
           "Structural channel removal with compensation in a feed-forward block. One "
           "decision shrinks both projections.",
           "A black-and-white technical diagram, white background, thin line art. Top "
           "row: input vector block (width d) -> matrix W1 (d x F) -> intermediate "
           "vector (width F) -> activation symbol -> matrix W2 (F x d) -> output "
           "vector (width d). Bottom row: the same pipeline after removal, with the "
           "intermediate width visibly narrowed to F-k and BOTH matrices drawn "
           "narrower, annotated 'one decision, two matrices'. A circular inset in "
           "the middle shows two nearly-parallel vectors h_j and h_p with the angle "
           "theta between them, the formula 'gamma_j = <h_j,h_p>/||h_p||^2', and an "
           "arrow folding column j into column p while column j fades out. "
           "Sans-serif labels, publication quality, no color.")

    h(doc, "3.3. Calibrated Per-Channel Quantization", 2)
    para(doc,
         "Symmetric quantization with scale s maps weights to integer codes q = "
         "round(clip(W/s, -q_max, q_max)). Because the codes are integers the "
         "reconstruction loss is piecewise constant in s and gradient descent is "
         "not well founded. We therefore use alternating minimization, in which "
         "each half-step is an exact minimizer:")
    eq(doc, "q_t = round(clip(W / s_t)) ,    s_{t+1} = <W, q_t> / <q_t, q_t> ,", 6)
    para(doc,
         "so the weight reconstruction loss L_W is non-increasing. A second phase "
         "searches a local grid around the phase-one optimum using the calibration "
         "objective, accepting a candidate only if it does not inflate L_W beyond a "
         "tolerance beta:")
    eq(doc, "s* = argmin_s [ L_W(s) + lambda L_calib(s) ]   s.t.  L_W(s) <= (1+beta) L_W(s_1) .", 7)
    para(doc,
         "Granularity matters more than the refinement itself. Since Y = X W^T, "
         "output channel i depends only on weight row i, so the calibration "
         "objective separates exactly across output channels and each row's scale "
         "can be optimized independently without approximation. To keep this "
         "affordable the per-channel calibration error is evaluated as a quadratic "
         "form using a Gram matrix computed once:")
    eq(doc, "|| X d_i ||^2 = d_i^T G d_i ,   G = X^T X ,   d_i = W_deq[i,:] - W[i,:] .", 8)

    h(doc, "3.4. Activation-Aware Low-Rank Factorization", 2)
    para(doc,
         "When case 3 applies, the objective is output error rather than weight "
         "error. With the Cholesky factor of the Gram matrix, G = L L^T,")
    eq(doc, "|| X (W - W')^T ||_F = || (W - W') L ||_F ,", 9)
    para(doc, "so the optimal rank-r solution is obtained by truncating in the "
              "transformed space and mapping back:")
    eq(doc, "W' = trunc_svd(W L, r) L^{-1} .", 10)
    para(doc,
         "The rank is not chosen but derived: for a two-factor INT8 representation "
         "costing r(m+n) bytes, fitting the cache budget requires")
    eq(doc, "r <= alpha · M_cache / (m + n) .", 11)

    h(doc, "3.5. Budget-Optimal Rank Allocation", 2)
    para(doc,
         "Allocating one shared budget across operators is posed as")
    eq(doc, "min  sum_i E_i(r_i)    s.t.  sum_i c_i r_i <= B ,   c_i = m_i + n_i ,", 12)
    para(doc,
         "where E_i is operator i's measured error curve. E_i is non-increasing and, "
         "for spectral truncation, convex in r: successive singular values are "
         "non-increasing, so each additional rank unit contributes less than the "
         "previous one. For a separable convex objective the continuous relaxation "
         "is solved by equalizing marginal return per parameter,")
    eq(doc, "- (dE_i / dr_i) / c_i = lambda   for all i ,", 13)
    para(doc,
         "and the integer solution follows from the greedy that repeatedly spends "
         "the next budget unit where it buys the largest error reduction per "
         "parameter. That greedy is EXACT for separable convex objectives, so the "
         "allocator is optimal rather than heuristic.")

    h(doc, "3.6. Experimental Setup", 2)
    para(doc,
         "Platform: Intel Tiger Lake H, 16 logical cores, L2 = 1.25 MiB per core "
         "pair, L3 = 24 MiB shared; alpha = 0.7 gives a 16.8 MiB budget. Runtime: "
         "ONNX Runtime 1.28, single intra-op thread for latency, warmup plus median "
         "of repeated runs. Hardware counters: Intel VTune Profiler 2026.4 with the "
         "event-based sampling driver active.")
    para(doc,
         "Models: (i) Whisper-medium Uzbek ASR (encoder 144 and decoder 240 weighted "
         "matrix operators); (ii) mBERT; (iii) open_llama_3b_v2. Data: Mozilla "
         "Common Voice Uzbek; 12 utterances for calibration and 80 held-out "
         "utterances for evaluation, with an 8000-sentence text corpus for the "
         "language-model experiments. Calibration and evaluation sets are disjoint "
         "throughout, and padding positions are excluded from response vectors.")
    para(doc,
         "Metrics: word and character error rate (primary), relative output error at "
         "operator level (E_loc) and network level (E_glob) as diagnostics, weight "
         "bytes, measured latency, and VTune memory-bound metrics. Confidence "
         "intervals are percentile bootstrap over utterances with 2000 resamples; "
         "paired comparisons use the paired bootstrap of per-utterance differences.")

    # ===================== 4. RESULTS =====================
    doc.add_page_break()
    h(doc, "4. Results", 1)

    h(doc, "4.1. Derived Targets and Case Frequencies", 2)
    table(doc, "Table 1. Cache-anchored requirement by granularity (budget alpha·L3 = 16.8 MiB).",
          ["Granularity", "Decoder (MiB)", "Required", "Encoder (MiB)", "Required"],
          [["per-operator", "16.0", "fits", "16.0", "0.95x"],
           ["per-layer", "64.0", "3.81x", "48.0", "2.86x"],
           ["whole model", "1536.0", "91.4x", "1152.0", "68.6x"]],
          good_rows=(1,))
    table(doc, "Table 2. Frequency of each cascade case across three architectures.",
          ["Case", "Condition", "Action", "Whisper enc.", "Whisper dec.", "Llama"],
          [["1", "FP32 fits", "no change", "96/144", "240/240", "attn only"],
           ["2", "INT8 suffices", "no low-rank", "96/144", "240/240", "q/k/v/o 0.58x"],
           ["3", "still over budget", "add low-rank", "48/144", "0/240", "FFN 1.57x, head 5.81x"]],
          good_rows=(2,))
    para(doc,
         "The derived per-layer requirement for the decoder is 3.81x and INT8 supplies "
         "4.00x, so the cascade selects quantization. Case 3 never arises for the "
         "decoder, arises for encoder feed-forward operators, and arises broadly in "
         "Llama - the first model in our set where the low-rank branch is genuinely "
         "required.")

    h(doc, "4.2. Contribution of the Quantization Scale", 2)
    table(doc, "Table 3. Effect of scale estimation method on operator output error "
               "(E_loc, held-out).",
          ["Scale method", "Encoder fc1", "Improvement", "Decoder fc1", "Improvement"],
          [["Q1 min/max (library default)", "0.00685", "-", "0.00441", "-"],
           ["Q2 alternating minimization", "0.00651", "+5.0%", "0.00442", "-0.1%"],
           ["Q3 calibrated, per-tensor", "0.00525", "+23.3%", "0.00360", "+18.4%"],
           ["Q4 calibrated, per-channel", "0.00179", "+73.8%", "0.00151", "+65.9%"]],
          good_rows=(3,))
    para(doc,
         "The benefit comes specifically from the calibration stage; alternating "
         "minimization alone contributes almost nothing. Per-channel scales cost m "
         "additional FP32 values per operator, about 0.5% of the weight bytes.")

    h(doc, "4.3. Structural Redundancy in Feed-Forward Blocks", 2)
    table(doc, "Table 4. Removable feed-forward channels at tau = 0.99 (Whisper encoder).",
          ["Layer", "Removed", "Share", "Layer", "Removed", "Share"],
          [["L0", "1764", "43.1%", "L12", "26", "0.6%"],
           ["L1", "2152", "52.5%", "L13", "10", "0.2%"],
           ["L2", "2376", "58.0%", "L14", "4", "0.1%"],
           ["L3", "2334", "57.0%", "L15-L20", "0", "0.0%"],
           ["L4", "2078", "50.7%", "L21", "3", "0.1%"],
           ["L5", "1535", "37.5%", "L22", "4", "0.1%"],
           ["L6", "1155", "28.2%", "L23", "2", "0.0%"],
           ["L8", "1048", "25.6%", "mean", "-", "17.1%"]],
          good_rows=(2,))
    para(doc,
         "A single threshold produces a strongly layer-dependent decision: 58% of "
         "channels are removable at L2-L3 and none from L15 onward. Nothing is "
         "hand-tuned per layer; the profile is a property of the model recovered by "
         "measurement.")
    table(doc, "Table 5. Error does not compound across pruned layers (FP32, encoder "
               "output error).",
          ["Pruned layers", "1", "4", "8", "12", "19"],
          [["Encoder output error", "0.0209", "0.0206", "0.0295", "0.0297", "0.0298"]],
          good_rows=(0,))
    para(doc,
         "Removing 19 layers costs barely more than removing one, in sharp contrast "
         "to low-rank approximation where error accumulates across operators. The "
         "reason is that a collinear channel is replaced exactly rather than "
         "approximated.")

    figure(doc, 4,
           "Layer-wise feed-forward redundancy across three architectures. Redundancy "
           "is a model property, not a universal one.",
           "A scientific line chart, white background, publication style, no 3D. "
           "X-axis 'layer index (normalized depth 0-1)', y-axis 'removable channels "
           "(%)' from 0 to 60. Three series with distinct markers and line styles in "
           "grayscale: 'Whisper encoder' rising to a peak of 58% at normalized depth "
           "0.1 then decaying to 0 by depth 0.6; 'open_llama_3b' starting at 3.4% "
           "and decaying to 0 by depth 0.3; 'mBERT' flat near 0 throughout. A "
           "horizontal dashed line marks the Whisper mean of 17.1%. Legend inside "
           "the plot area, thin axis lines, sans-serif labels.")

    h(doc, "4.4. Coupling Between Compensation and Quantization Granularity", 2)
    table(doc, "Table 6. Compensation widens the weight dynamic range (fc2, layer 2).",
          ["State", "max |w|", "Row norm (median)", "Row norm (max)", "Spread"],
          [["original", "0.2472", "0.0786", "0.7569", "9.6x"],
           ["after compensation", "11.3875", "0.4083", "76.9402", "188.4x"]],
          bad_rows=(1,))
    table(doc, "Table 7. Consequence for quantization granularity (Whisper encoder).",
          ["Variant", "Size (MiB)", "Compression", "E_glob", "Outcome"],
          [["pruned, INT8 per-tensor", "266", "4.33x", "0.7420", "model destroyed"],
           ["pruned, INT8 per-channel", "267", "4.32x", "0.2226", "preserved"]],
          bad_rows=(0,), good_rows=(1,))
    para(doc,
         "Folding gamma_j W[:, j] into the representative column concentrates mass, "
         "so a single tensor-wide scale must span outliers and loses precision "
         "everywhere else. This is a dependency between two components of the method "
         "rather than an independent design choice.")

    h(doc, "4.5. Comparison of Low-Rank Constructions", 2)
    table(doc, "Table 8. Output error at equal parameter budget (held-out, 135 "
               "measurements).",
          ["Compression", "plain SVD", "activation-aware SVD", "functional CUR", "leverage CUR"],
          [["2.00x", "0.3700", "0.2379", "0.5806", "0.6978"],
           ["3.81x (cache)", "0.5534", "0.3689", "0.7180", "0.8223"],
           ["8.00x", "0.7025", "0.4730", "0.8227", "0.8963"]],
          good_rows=(1,))
    table(doc, "Table 9. Weight error at the same budgets, confirming Eckart-Young.",
          ["Compression", "plain SVD", "activation-aware SVD", "functional CUR", "leverage CUR"],
          [["2.00x", "0.4019", "0.4926", "0.7790", "0.7873"],
           ["3.81x", "0.6010", "0.7314", "0.8861", "0.9085"],
           ["8.00x", "0.7583", "0.8516", "0.9540", "0.9749"]])
    table(doc, "Table 10. Pairwise outcomes over 135 operator-level comparisons.",
          ["Comparison", "Wins", "Interpretation"],
          [["functional CUR > leverage CUR", "134 / 135", "calibration-guided column ranking works"],
           ["activation-aware SVD > plain SVD", "135 / 135", "output-optimality beats weight-optimality"],
           ["functional CUR > activation-aware SVD", "0 / 135", "CUR assembly is not competitive"]],
          good_rows=(0, 1), bad_rows=(2,))
    para(doc,
         "Tables 8 and 9 together isolate the central principle: the "
         "activation-aware solution is WORSE on weight error yet BETTER on output "
         "error in every measurement. Eckart-Young is not violated - it simply "
         "optimizes the wrong objective for this purpose. The CUR assembly, however, "
         "carries an additional r^2 block (r(m+n) + r^2 versus r(m+n)) and therefore "
         "affords a lower rank at equal budget; it loses uniformly.")

    h(doc, "4.6. Calibration-Size Requirement", 2)
    table(doc, "Table 11. Overfitting versus rows-to-rank ratio (encoder fc1, rank 409).",
          ["Fit rows", "Rows/rank", "Fit error", "Held-out error", "Gap"],
          [["256", "0.6", "0.00000", "0.04355", "1 540 784x"],
           ["512", "1.3", "0.00035", "0.04624", "131x"],
           ["2048", "5.0", "0.00637", "0.02835", "4.4x"],
           ["4096", "10.0", "0.01151", "0.02199", "1.9x"],
           ["8192", "20.0", "0.01364", "0.01900", "1.4x"]],
          bad_rows=(0, 1), good_rows=(4,))
    para(doc,
         "A fit error of exactly zero indicates memorization of the calibration set, "
         "not method quality. We recommend a rows-to-rank ratio of at least 10, and "
         "20 where the data permit.")

    h(doc, "4.7. Error Propagation", 2)
    para(doc,
         "Perturbing one operator at a time and measuring the network output error "
         "gives per-operator influence coefficients c_i = E_glob(i alone) / E_loc(i). "
         "Across 48 encoder operators E_loc spans 160x (0.0014 to 0.225) while "
         "E_glob spans only 4x (0.012 to 0.047). Coefficients are 0.58-5.12 for the "
         "expanding projection and 0.13-0.68 for the contracting one. The asymmetry "
         "follows from the residual stream: for y = x + f(x), a relative error in f "
         "is diluted by the factor ||f|| / ||x + f||, whereas an error in the "
         "expanding projection is amplified through the nonlinearity before reaching "
         "it.")

    figure(doc, 5,
           "Measured error absorption. Local operator error varies by two orders of "
           "magnitude while network output error remains within a narrow band.",
           "A scientific scatter plot, white background, publication style. X-axis "
           "'per-operator error E_loc' on a logarithmic scale from 0.001 to 0.3, "
           "y-axis 'network output error E_glob' on a linear scale from 0 to 0.06. "
           "48 points using two distinct marker shapes: circles labeled 'expanding "
           "projection (fc1)' clustered at lower E_loc, triangles labeled "
           "'contracting projection (fc2)' clustered at higher E_loc. Despite the "
           "160x horizontal spread, all points lie within a narrow horizontal band. "
           "A shaded horizontal band marks the 4x E_glob range. An inset diagram in "
           "a corner shows a residual connection y = x + f(x) with the annotation "
           "'relative error diluted by ||f||/||x+f||'. Grayscale, thin axis lines.")

    h(doc, "4.8. Rank Allocation", 2)
    table(doc, "Table 12. Uniform versus budget-optimal allocation at equal budget "
               "(Whisper encoder).",
          ["Scheme", "Parameters", "Summed error", "Latency (ms)", "WER", "CER"],
          [["uniform rank", "100 515 840", "3.2682", "6371.7", "0.1719", "0.0417"],
           ["sensitivity-based", "100 505 600", "2.8495", "6197.9", "0.0729", "0.0208"]],
          good_rows=(1,))
    para(doc,
         "The summed objective improves by 12.8% while WER improves by 58%, a "
         "disproportion consistent with the non-linear propagation of Section 4.7. "
         "We additionally measured allocation against directly observed global damage "
         "curves (144 network evaluations); it produced the same WER and slightly "
         "worse CER, indicating that the inexpensive local proxy is adequate for "
         "RANKING operators even though it is inadequate for PREDICTING end-to-end "
         "damage.")

    h(doc, "4.9. End-to-End Quality", 2)
    table(doc, "Table 13. Encoder variants, 80 held-out utterances, 95% bootstrap CI.",
          ["Variant", "WER", "95% CI", "CER", "dWER vs FP32", "Significance"],
          [["FP32", "0.1007", "[0.0629, 0.1444]", "0.0150", "-", "-"],
           ["structural removal + per-channel INT8", "0.1107", "[0.0691, 0.1586]", "0.0183",
            "+0.0100 [-0.0020, +0.0240]", "not significant"],
           ["INT8 (mandatory step only)", "0.1146", "[0.0762, 0.1589]", "0.0177",
            "+0.0139 [+0.0046, +0.0259]", "significant"],
           ["INT8 + allocated low-rank", "0.1335", "[0.0911, 0.1761]", "0.0297",
            "+0.0328 [+0.0007, +0.0654]", "significant"]],
          good_rows=(1,), bad_rows=(2, 3))
    para(doc,
         "At 4.32x compression the proposed structural variant is the only encoder "
         "configuration whose difference from FP32 is not statistically resolvable, "
         "while plain INT8 degrades significantly. Decoder variants (FP32 0.1029, "
         "INT8 0.1007, per-channel INT8 0.0986) are mutually indistinguishable, "
         "consistent with the cascade's decision that no structural change is needed "
         "there.")

    h(doc, "4.10. Speed and Hardware Counters", 2)
    table(doc, "Table 14. Latency of adding low-rank on top of INT8 (encoder fc1, 1500 positions).",
          ["Variant", "Weights (MiB)", "Latency (ms)", "vs FP32", "vs INT8", "E_loc"],
          [["dense FP32", "16.00", "117.9", "1.00x", "0.26x", "0"],
           ["dense INT8", "4.00", "30.3", "3.89x", "1.00x", "0.0082"],
           ["INT8 + SVD r=200", "0.98", "8.8", "13.40x", "3.44x", "0.0201"],
           ["INT8 + SVD r=128", "0.62", "6.7", "17.70x", "4.56x", "0.0305"]],
          good_rows=(3,))
    table(doc, "Table 15. VTune hardware counters (encoder fc1).",
          ["Variant", "ms/iter", "Memory bound", "L2", "L3", "DRAM", "CPI"],
          [["dense FP32", "121.85", "8.8%", "2.5%", "2.7%", "2.9%", "0.64"],
           ["dense INT8", "33.96", "12.7%", "1.9%", "2.4%", "6.5%", "0.67"],
           ["INT8 + SVD r=128", "7.62", "18.3%", "4.1%", "1.0%", "9.2%", "0.62"]],
          good_rows=(2,))
    para(doc,
         "L3 pressure falls from 2.4% to 1.0%, so cache benefit is real and directly "
         "observed rather than inferred. However memory-bound stalls account for only "
         "9-18% of pipeline slots, so the dominant mechanism of the 4.46x speedup is "
         "the 6.40x arithmetic reduction, not cache residency. We therefore describe "
         "the effect as structural reduction of arithmetic volume in a compute-bound "
         "regime.")

    h(doc, "4.11. Cross-Architecture Transfer", 2)
    table(doc, "Table 16. Redundancy and cache verdict across three architectures.",
          ["Model", "Redundancy (tau=0.99)", "Peak", "Case-3 operator"],
          [["Whisper encoder", "17.1%", "58.0%", "none (FFN fits after INT8)"],
           ["open_llama_3b", "0.6%", "3.4%", "gate/up/down 1.57x, head 5.81x"],
           ["mBERT", "0.1%", "0.7%", "vocabulary matrix 5.21x"]])
    table(doc, "Table 17. First end-to-end validation of the low-rank branch "
               "(open_llama_3b, perplexity on held-out Uzbek text).",
          ["Variant", "Perplexity", "vs FP32"],
          [["FP32", "230.122", "1.000x"],
           ["per-channel INT8 (mandatory)", "230.900", "1.003x"],
           ["cache-anchored low-rank + INT8", "231.642", "1.007x"]],
          good_rows=(2,))
    para(doc,
         "At the derived rank of 1487 (46% of full rank) the mean per-operator error "
         "is 0.148 with a maximum of 0.308, yet perplexity degrades by only 0.7%. "
         "This confirms the absorption law of Section 4.7 on a decoder-only "
         "architecture, and is the first end-to-end validation of the low-rank branch "
         "in our study, since neither Whisper nor mBERT entered case 3 for their "
         "feed-forward operators.")

    figure(doc, 6,
           "Compression-quality operating points for the encoder with 95% confidence "
           "intervals.",
           "A scientific scatter plot with error bars, white background, publication "
           "style, grayscale with one highlighted marker. X-axis 'compression factor' "
           "logarithmic from 1 to 8, y-axis 'WER' from 0.09 to 0.15. Four points with "
           "vertical 95% CI bars: 'FP32' at (1.0, 0.1007); 'structural removal + "
           "per-channel INT8' at (4.32, 0.1107) drawn as a filled highlighted marker "
           "and labeled 'proposed'; 'INT8' at (4.00, 0.1146); 'INT8 + low-rank' at "
           "(6.00, 0.1335). A horizontal dashed line at the FP32 WER. Marker area "
           "proportional to measured speedup. Legend inside the plot, sans-serif "
           "labels, thin axis lines.")

    # ===================== 5. DISCUSSION =====================
    doc.add_page_break()
    h(doc, "5. Discussion", 1)
    para(doc,
         "Three results generalize beyond the specific models studied. First, the "
         "cache-anchored target converts an arbitrary hyperparameter into a derived "
         "quantity, and the derived value proved informative on all three "
         "architectures - it identified feed-forward operators in the audio encoder, "
         "the vocabulary matrix in mBERT, and both feed-forward and head operators in "
         "Llama. Second, the absorption law explains why layer-level error proxies "
         "mislead: because the residual stream dilutes relative error, a large local "
         "perturbation can be harmless while a moderate one distributed differently "
         "is not. Third, the coupling between compensation and quantization "
         "granularity is a design constraint that is invisible if the two components "
         "are evaluated separately.")
    para(doc,
         "The negative results are equally informative. The CUR construction, which "
         "motivated the initial design, is not competitive as a low-rank "
         "approximator: it carries an r^2 penalty and Eckart-Young bounds the "
         "achievable weight error. Its column-selection criterion nevertheless "
         "outperforms leverage scores in 134/135 comparisons, and the underlying "
         "grouping mechanism succeeds when redirected from approximation to exact "
         "structural removal. We therefore recommend interpreting calibration-guided "
         "column selection as a redundancy detector rather than a factorization "
         "method.")
    para(doc,
         "Limitations. The evaluation set of 80 utterances leaves confidence "
         "intervals wide enough that differences below roughly 0.03 WER are not "
         "resolvable; larger corpora would sharpen the comparisons of Section 4.9. "
         "The cache-anchored target was validated on a single cache configuration, "
         "and reproducing the analysis on platforms with different L3 capacities is "
         "the most direct test of the central claim. Quantization in the "
         "language-model experiments was simulated on our per-channel grid, so those "
         "results characterize quality rather than wall-clock speed. Finally, "
         "mapping E_glob to task metrics remains empirical: predicting WER from "
         "operator-level error would make the cascade fully predictive, which it is "
         "not yet.")

    # ===================== 6. CONCLUSIONS =====================
    h(doc, "6. Conclusions", 1)
    para(doc,
         "We presented a cascade in which the compression target is derived from "
         "cache topology and the transformation is selected per operator by "
         "measurement. On a Whisper-medium Uzbek ASR encoder, structural removal of "
         "functionally redundant feed-forward channels combined with calibrated "
         "per-channel INT8 achieves 4.32x compression with a word error rate "
         "statistically indistinguishable from FP32, while plain INT8 at 4.00x "
         "degrades quality significantly. The method requires no retraining. Three "
         "supporting findings are of independent interest: compensation makes "
         "per-channel quantization mandatory rather than optional; calibration-based "
         "low-rank factorization requires a rows-to-rank ratio of at least 10-20; and "
         "transformer blocks absorb local perturbations strongly, which explains why "
         "layer-level error proxies are unreliable predictors of task quality. "
         "Cross-architecture evaluation indicates that the diagnostic transfers while "
         "the specific tool does not, supporting the case for measurement-driven "
         "rather than fixed compression recipes.")

    # ===================== BACK MATTER =====================
    h(doc, "Author Contributions", 1)
    para(doc, "Conceptualization, X.Y.; methodology, X.Y.; software, X.Y.; validation, "
              "X.Y. and Z.W.; formal analysis, X.Y.; investigation, X.Y.; data curation, "
              "X.Y.; writing-original draft preparation, X.Y.; writing-review and "
              "editing, Z.W.; visualization, X.Y.; supervision, Z.W. All authors have "
              "read and agreed to the published version of the manuscript.", size=9.5)
    h(doc, "Funding", 1)
    para(doc, "This research received no external funding.", size=9.5)
    h(doc, "Data Availability Statement", 1)
    para(doc, "The Common Voice Uzbek corpus is publicly available. Implementation code, "
              "experiment scripts and all measured result files are available from the "
              "authors on request.", size=9.5)
    h(doc, "Conflicts of Interest", 1)
    para(doc, "The authors declare no conflict of interest.", size=9.5)

    h(doc, "References (to be completed)", 1)
    mono(doc,
         "1.  Frantar, E.; et al. GPTQ: Accurate Post-Training Quantization for\n"
         "    Generative Pre-trained Transformers. ICLR 2023.\n"
         "2.  Lin, J.; et al. AWQ: Activation-aware Weight Quantization for LLM\n"
         "    Compression and Acceleration. MLSys 2024.\n"
         "3.  Xiao, G.; et al. SmoothQuant: Accurate and Efficient Post-Training\n"
         "    Quantization for Large Language Models. ICML 2023.\n"
         "4.  Hsu, Y.-C.; et al. Language Model Compression with Weighted Low-Rank\n"
         "    Factorization (FWSVD). ICLR 2022.\n"
         "5.  Yuan, Z.; et al. ASVD: Activation-aware Singular Value Decomposition\n"
         "    for Compressing Large Language Models. arXiv 2023.\n"
         "6.  Wang, X.; et al. SVD-LLM: Truncation-aware Singular Value\n"
         "    Decomposition for Large Language Model Compression. 2024.\n"
         "7.  Eckart, C.; Young, G. The Approximation of One Matrix by Another of\n"
         "    Lower Rank. Psychometrika 1936, 1, 211-218.\n"
         "8.  Drineas, P.; Mahoney, M.W. On the Nystrom Method for Approximating a\n"
         "    Gram Matrix. JMLR 2005.\n"
         "9.  Radford, A.; et al. Robust Speech Recognition via Large-Scale Weak\n"
         "    Supervision (Whisper). ICML 2023.\n"
         "10. Devlin, J.; et al. BERT: Pre-training of Deep Bidirectional\n"
         "    Transformers for Language Understanding. NAACL 2019.\n"
         "11. Williams, S.; et al. Roofline: An Insightful Visual Performance Model\n"
         "    for Multicore Architectures. CACM 2009.\n"
         "[Add 15-25 more references per Sensors norms, including recent edge-ASR\n"
         " and low-resource-language compression work.]", 8.5)

    doc.save(OUT)
    print(f"saqlandi: {OUT}")


if __name__ == "__main__":
    main()
