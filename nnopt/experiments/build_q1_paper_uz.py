"""Q1 (MDPI Sensors) uslubidagi maqola — o'zbek tilida.

Tuzilishi Sensors shabloniga mos: Annotatsiya, Kalit so'zlar, Kirish,
Tegishli ishlar, Materiallar va usullar, Natijalar, Muhokama, Xulosalar va
majburiy yakuniy bo'limlar. Formulalar LaTeX/MathType ga oson o'tadigan
oddiy yozuvda. Rasm promtlari ingliz tilida qoldirilgan, chunki tasvir
generatorlari ingliz tilida ancha barqaror ishlaydi.

Barcha raqamlar experiments/ dagi o'lchangan JSON natijalaridan olingan.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = "../Q1_Sensors_maqola_uz.docx"
ACCENT = RGBColor(0x0F, 0x64, 0x70)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)
WARN = RGBColor(0x9A, 0x64, 0x10)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def eq(doc, text, number=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if number:
        r2 = p.add_run(f"\t\t({number})")
        r2.font.size = Pt(10.5)
    return p


def mono(doc, text, size=8.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def table(doc, caption, headers, rows, good_rows=(), bad_rows=()):
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ht in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(ht)
        run.bold = True
        run.font.size = Pt(8)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            if ri in good_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in bad_rows:
                run.font.color.rgb = CRIT
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def bullets(doc, items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + it[1])
            r2.font.size = Pt(10.5)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10.5)


FIGURE_DIR = "figures"


def figure(doc, num, caption, prompt):
    """Embed the rendered figure if it exists; otherwise carry the prompt.

    Figures that display measurements are drawn from the result files by
    experiments/make_figures.py and embedded here. The prompt is kept only for
    the schematics, which have no data behind them. It must never be used to
    generate a chart: an image model asked for a scatter plot invents the
    points, and invented data has no place in a paper.
    """
    path = os.path.join(FIGURE_DIR, f"fig{num}.png")
    if os.path.exists(path):
        pimg = doc.add_paragraph()
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.add_run().add_picture(path, width=Inches(5.6))

    p = doc.add_paragraph()
    r = p.add_run(f"{num}-rasm. {caption}")
    r.bold = True
    r.font.size = Pt(9)

    if not os.path.exists(path):
        p2 = doc.add_paragraph()
        r2 = p2.add_run("[AI prompt — ingliz tilida; SXEMA uchun, "
                        "grafik uchun EMAS] ")
        r2.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = ACCENT
        r3 = p2.add_run(prompt)
        r3.font.size = Pt(8)
        r3.font.color.rgb = MUTED
        p2.paragraph_format.left_indent = Pt(14)
    doc.add_paragraph()


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Palatino Linotype"
    st.font.size = Pt(10.5)

    # ===================== SARLAVHA =====================
    h(doc, "A Cache-Aware Adaptive Compression Framework for Improving the "
           "Inference Efficiency of Transformer Neural Networks", 0)
    para(doc, "Transformer neyron tarmoqlari inferens samaradorligini "
              "oshirish uchun keshni hisobga oluvchi adaptiv siqish "
              "freymvorki", italic=True, size=10)
    para(doc, "Ism Familiya 1,*, Hammuallif 2", italic=True, size=10)
    mono(doc,
         "1  Kafedra, Universitet, Shahar, Mamlakat; muallif@email\n"
         "2  Kafedra, Universitet, Shahar, Mamlakat; hammuallif@email\n"
         "*  Bog'lanish uchun: muallif@email", 8.5)

    h(doc, "Annotatsiya", 1)
    para(doc,
         "Transformer modellarini o'qitilgandan keyin siqishda siqish darajasi odatda "
         "qo'lda tanlanadi, sifat esa vazn yoki qatlam darajasidagi bilvosita "
         "mezonlar bilan baholanadi. Biz ikkala yondashuv ham asossiz ekanini "
         "ko'rsatamiz va siqish maqsadi apparatning kesh topologiyasidan CHIQARILADIGAN, "
         "o'zgartirish esa har operator uchun o'lchov asosida tanlanadigan kaskadni "
         "taklif qilamiz. Usul to'rt qismdan iborat: (i) kafolatlangan umumiy keshdan "
         "olinadigan kesh-bog'langan maqsad; (ii) kalibrlashdagi javoblari kollinear "
         "bo'lgan kanallarni aniqlab, ularni kompensatsiya orqali aniq olib "
         "tashlaydigan funksional guruhlash — bu rankni emas, operator O'LCHAMINI "
         "kamaytiradi; (iii) kvantlash bosqichi, unda o'lchov asosida eng yaxshi "
         "mavjud usul (GPTQ) tanlanadi; "
         "(iv) ajraluvchan qavariq masala sifatida qo'yilgan va ochko'z algoritmi aniq "
         "yechim beradigan rank taqsimoti. Nashr etilgan post-training usullari "
         "bilan taqqoslash kvantlash bosqichining tanlovini belgilaydi: GPTQ [1] "
         "ning Hessian orqali xato kompensatsiyasi bizning kalibrlangan "
         "masshtabimizdan ustun (60 operatordan 47 tasida g'alaba, xatoni 54.3% va "
         "12.9% kamaytiradi), shuning uchun kaskad kvantlashni unga topshiradi va "
         "faqat ORTOGONAL strukturaviy o'qni da'vo qiladi. Asosiy natija shu o'qning "
         "tekin ekanini ko'rsatadi: Whisper-medium o'zbek ASR enkoderida oldinga "
         "uzatuvchi (FFN) kanallarning 17.1% ini olib tashlash GPTQ ustiga "
         "qo'shilganda xotirani 300 dan 267 MiB gacha (-11%) kamaytiradi va 3% "
         "tezlik beradi, so'z xatoligini esa o'zgartirmaydi (dWER = -0.0014, 95% "
         "ishonch oralig'i [-0.0111, +0.0096]); ikkala variant ham FP32 dan "
         "statistik jihatdan farqlanmaydi. To'liq 2x2 taqqoslash bu xossaning "
         "SHARTLI ekanini ko'rsatadi: kvantlagich oddiy yaxlitlashga "
         "almashtirilsa, qisqartirishning narxi -0.0014 dan +0.0084 ga ko'tariladi "
         "va birikma FP32 dan sezilarli yomon bo'lib qoladi (+0.0150), holbuki "
         "qisqartirishsiz ikkala kvantlagich deyarli teng (farq -0.0011). Sabab "
         "4.4-bo'limda o'lchangan: kompensatsiya satr diapazonini 188x "
         "kengaytiradi va GPTQ ning Hessian orqali xato tarqatishi aynan shu "
         "qoldiqni yutadi. Butun model darajasida kaskadning qiymati "
         "assimetrik bo'lib chiqadi: eng sodda 'hamma joyda INT8' bazasidan u "
         "atigi 4.5% xotira yutadi (738 -> 705 MiB, aniqlik o'zgarmaydi), "
         "ammo qarorni agressiv tomonga bekor qilish — 5.34x ga siqish, ya'ni "
         "odatda qo'lda tanlanadigan 4x va 8x oralig'idagi butunlay oqilona "
         "ko'rinadigan daraja — WER ni 0.1833 dan 0.6101 ga ko'taradi. "
         "Shuning uchun usulning hissasi ko'proq siqishda emas, QAYERDA "
         "TO'XTASHNI oldindan aytishda. Baholash Common Voice o'zbek TEST "
         "splitining 300 namunasida, kalibrlash esa VALIDATION splitida "
         "o'tkazilgan, ya'ni ular turli taqsimotlardan. Bu tanlov muhim: xuddi shu "
         "taqqoslash kalibrlash bilan bir splitda o'lchanganda kalibrlashga "
         "tayanadigan variantlar sun'iy ravishda ustun ko'ringan. Mustaqil "
         "ahamiyatga ega yana uchta natija "
         "keltiriladi: kompensatsiya satrlar bo'yicha vazn diapazonini 9.6x dan 188.4x "
         "gacha kengaytiradi va shu sababli per-channel kvantlashni MAJBURIY qiladi; "
         "kalibrlashga asoslangan past-rank yoyilma qator/rank nisbati kamida 10-20 "
         "bo'lishini talab qiladi, aks holda u kalibrlash to'plamini yodlab oladi "
         "(moslash xatosi 0.00000, held-out xatosi 0.04355); transformer bloklari lokal "
         "buzilishlarni kuchli yutadi — operator darajasidagi xato 160 barobar "
         "o'zgarganda tarmoq chiqishidagi xato atigi 4 barobar o'zgaradi. Standart "
         "WikiText-2 benchmarkida yutishning chegarasi ham aniqlanadi: INT4 da "
         "masshtabni vazn xatosi bo'yicha optimallashtirish har bir operatorda "
         "26% aniqroq bo'lsa-da, perplexity ni 1.696x ga yomonlashtiradi, chunki "
         "bunday masshtab har bir kanalni bir xil yo'nalishda kichraytiradi va "
         "siljish 78 ta operator bo'ylab ko'payadi (kuchaytirish 0.9896, "
         "0.9896^78 = 0.44); masshtabni chiqish domenida qayta tanlash butun "
         "sonli kodlarni va bit kengligini o'zgartirmasdan perplexity ni 8.246 "
         "ga tushiradi. Bundan umumiy tavsiya kelib chiqadi: chuqur tarmoqlarda "
         "operator darajasidagi maqsad xato kattaligini emas, SILJIMAGANLIKNI "
         "muhofaza qilishi kerak. Ish nuqtasi bo'lgan INT8 da siljish 100 "
         "barobar kichik (gain 0.9999) va usulimiz uchdan-uchgacha ozgina "
         "ustun turadi (7.549 va RTN ning 7.550 i, FP32 7.547). mBERT va "
         "open_llama_3b da o'tkazilgan tekshiruv kesh-bog'langan diagnostika "
         "ko'chishini, aniq vosita esa ko'chmasligini ko'rsatadi: FFN ortiqchaligi "
         "audio enkoderda ko'p (cho'qqida 58%), Llama da kam (yumshoq chegarada 25%), "
         "mBERT da esa yo'q. Kesh hajmi ushbu ishda QATTIQ CHEKLOV emas — YUMSHOQ "
         "maqsad sifatida qo'yiladi: byudjetdan (masalan, L3 = 24 MiB) chiqib "
         "ketish kesh o'tkazib yuborishlarini (miss) proportsional "
         "oshiradigan, ammo sig'ishga majburlamaydigan uzluksiz funksiya "
         "orqali jarimalanadi, shunday qilib har bir apparat uchun eng kam "
         "kesh o'tkazib yuborishga olib keladigan model tanlanadi, sig'ish "
         "shart emas. Bu maqsad model-mustaqil FREYMVORKKA joylashtirilgan: "
         "har bir arxitektura (Whisper, mBERT, open_llama_3b) uch metodni "
         "(parts, structural_ladder, evaluate) taqdim etadigan PROFIL orqali "
         "ulanadi, freymvork esa berilgan model, L3 hajmi va sifat "
         "byudjeti asosida optimal rejimni avtomatik qidiradi va uni "
         "KO'R-KO'RONA (blind) INT8/pruning bazalariga nisbatan tekshiradi. "
         "Olti haqiqiy apparat konfiguratsiyasida (Raspberry Pi 5 dan EPYC "
         "gacha) sinov shuni ko'rsatadiki, kesh hajmidan kelib chiqadigan "
         "QARORLAR (qaysi bosqichda to'xtash) haqiqatan foydali bo'lsa-da, "
         "kesh sig'imining o'zi kechikishga sezilarli ta'sir qilishi to'g'ridan "
         "-to'g'ri o'lchovda TASDIQLANMADI (interleaved dizaynda nisbat "
         "0.98-1.02x) — ya'ni kaskadning ustunligi FLOP-hajmiga bog'liq, kesh "
         "mexanizmiga emas, va bu farq ishda ochiq tan olinadi. Uchala "
         "arxitekturada strukturaviy kanal kesishni majburlash xarajati "
         "o'lchandi: Whisper da 0.7% (17.1% ortiqchalik bilan), mBERT da 6% "
         "(3.5% ortiqchalik), Llama da 18-26% (0.6% ortiqchalik) — "
         "ortiqchalik kamayishi bilan xarajat monoton o'sadi, bu kaskadning "
         "kvantlashni birinchi qo'yish qoidasini uch modelda mustaqil "
         "asoslaydi. Barcha o'lchovlar real apparat hisoblagichlari (Intel "
         "VTune), real ishlash vaqti va held-out ma'lumotlarga asoslanadi.",
         size=10)

    h(doc, "Kalit so'zlar", 1)
    para(doc, "o'qitilgandan keyingi siqish; kesh-hisobga oluvchi optimallashtirish; "
              "strukturaviy qisqartirish; past-rankli yoyilma; per-channel kvantlash; "
              "avtomatik nutq tanish; kam resursli tillar; chekka qurilmada inferens",
         italic=True, size=10)

    # ===================== 1. KIRISH =====================
    doc.add_page_break()
    h(doc, "1. Kirish", 1)
    para(doc,
         "Transformer modellarini CPU sinfidagi apparatda joylashtirish parametrlar "
         "soni bilan emas, operatorning xotira izi va uni saqlashi kerak bo'lgan kesh "
         "ierarxiyasi o'rtasidagi munosabat bilan cheklanadi. Mavjud o'qitilgandan "
         "keyingi usullar — kvantlash [1-6] va past-rankli yoyilma [14-17] — "
         "tadqiqotchi tanlagan darajaga, odatda 4x yoki 8x "
         "ga siqadi va natijani bajaradigan mashinani hisobga olmaydi. Bundan ikkita "
         "oqibat kelib chiqadi. Birinchidan, tanlangan daraja keraksiz bo'lishi "
         "(operator allaqachon sig'adi) yoki erishib bo'lmasligi mumkin (hech qanday "
         "daraja yordam bermaydi, chunki vazn emas, faollashuvlar hal qiluvchi). "
         "Ikkinchidan, daraja oldindan qat'iy belgilangani uchun chidamliligi juda "
         "har xil operatorlarga bir xil o'zgartirish qo'llanadi.")
    para(doc,
         "Buning zamirida kompyuter arxitekturasining fundamental "
         "assimetriyasi — XOTIRA DEVORI — yotadi. Zamonaviy CPU yadrosi "
         "bitta taktda o'nlab arifmetik amalni bajaradi, DRAM ga bitta "
         "murojaat esa yuzlab taktga tushadi; kesh ierarxiyasi shu "
         "jarlikni yashirish uchun mavjud va L3 dan o'qish DRAM dan "
         "o'qishdan taxminan tartibga arzon. Inferensda vaznlar har "
         "o'tishda to'liq oqiziladigan yagona yirik ma'lumot oqimi "
         "bo'lgani uchun bajarish vaqtining xotira qismi bevosita KESH "
         "O'TKAZIB YUBORISHLAR (miss) hajmi bilan belgilanadi: keshga "
         "sig'magan har bayt DRAM dan qayta keladi. Demak siqishning "
         "vazifasi shunchaki diskdagi hajmni emas, HAR O'TISHDAGI MISS "
         "HAJMINI kamaytirishdir — va bu ishning maqsad funksiyasi aynan "
         "shu kattalik ustida quriladi. Ushbu bog'lanish 4.10-bo'limda "
         "olti mustaqil o'lchov bilan tasdiqlanadi: ko'chirilgan baytlar "
         "va vaqt orasidagi korrelyatsiyadan (r = +0.974) apparat "
         "hisoblagichlaridagi xotira to'xtashlarining 2.41 barobar "
         "qisqarishigacha.")
    para(doc,
         "Ushbu ishda maqsad chiqariladigan kattalikka aylantiriladi. Bajarish "
         "platformasining kafolatlangan umumiy keshi berilganda, har bir operator "
         "uchun talab qilinadigan qisqartirish uning o'lchangan izidan kelib chiqadi "
         "va kaskad ENG YUMSHOQ YETARLI o'zgartirishni qo'llaydi. Ishlatilgan "
         "mashinada (16 mantiqiy yadro baham ko'radigan 24 MiB L3) Whisper dekoderi "
         "uchun qatlam darajasidagi chiqarilgan talab 3.81x ni, INT8 esa aynan 4.00x "
         "ni beradi — shuning uchun kaskad kvantlashni tanlaydi va keyingi siqishni "
         "rad etadi. Bu qarorning to'g'riligini uni majburan bekor qilganda nima "
         "bo'lishini o'lchash orqali tasdiqlaymiz.")
    para(doc,
         "Yondashuvning qamrovini boshidanoq aniq belgilaymiz. Maqsad "
         "kafolatlangan umumiy keshdan chiqarilgani uchun usul CPU sinfidagi "
         "apparatga TA'RIFAN bog'langan: grafik protsessorlarda hukmron cheklov "
         "kesh sig'imi emas, xotira o'tkazuvchanligi bo'lib, alpha x L3 "
         "byudjetining u yerda analogi yo'q. Shu sababli ushbu ish GPU "
         "xizmatiga muqobil sifatida emas, apparat qat'iy va kamtar bo'lgan "
         "joylashtirishlar uchun taklif etiladi — kam resursli tillar uchun "
         "nutqni tanish, kodlovchi modellarda klassifikatsiya va qidiruv, "
         "hamda chekka qurilmalardagi inferens. Qamrovning batafsil asoslanishi "
         "5.1-bo'limda keltirilgan.")

    h(doc, "1.1. Ishning hissasi (contributions)", 2)
    bullets(doc, [
        ("Kesh-bog'langan maqsadni chiqarish.",
         "Siqish talabi tanlanmaydi, balki kesh topologiyasi va operator izidan "
         "hisoblanadi. Chiqarilgan maqsad uchta turli arxitekturada qaysi "
         "o'zgartirish o'rinli ekanini tushuntirib berishi ko'rsatiladi."),
        ("Funksional guruhlash — strukturaviy olib tashlash sifatida.",
         "Kalibrlashdagi javoblari kollinear bo'lgan kanallar aniqlanadi va eng kichik "
         "kvadratlar kompensatsiyasi orqali aniq olib tashlanadi. Past-rankli "
         "yaqinlashtirishdan farqli o'laroq bu operator o'lchamlarini kamaytiradi, "
         "shuning uchun FFN blokida bitta qaror ikkita matritsani (gated "
         "arxitekturada uchtasini) qisqartiradi."),
        ("Ikki komponent orasidagi bog'liqlik.",
         "Kompensatsiya massani vakil ustunlarga to'playdi va satrlar bo'yicha "
         "diapazonni 9.6x dan 188.4x ga kengaytiradi. Shundan keyin per-tensor INT8 "
         "modelni butunlay buzadi (WER 1.0000), per-channel INT8 esa saqlaydi. Demak "
         "komponentlar mustaqil emas: guruhlash kalibrlangan per-channel masshtabni "
         "TALAB QILADI."),
        ("Strukturaviy mezonning o'lchangan qamrovi.",
         "Nashr etilgan kanal-qisqartirish mezonlari (magnitude, Wanda) bilan "
         "to'rtta teng byudjetda taqqoslash shuni ko'rsatadiki, mezonlar WER "
         "bo'yicha ajralmaydi — barcha juftlik oraliqlari nolni qamraydi va "
         "tartib monoton emas. Ular DEGRADATSIYA SHAKLI bo'yicha ajraladi: "
         "eng katta bir qadamli yomonlashish taklif etilgan mezonda +0.0173, "
         "Wanda da +0.0351, magnitude da esa +3.0612 — oxirgisi bir qadamda "
         "modelni ish holatidan chiqaradi. Bu aniqlik byudjetini boshqarish "
         "uchun zaruriy shart, chunki monoton bo'lmagan egri chiziqda "
         "tolerans ushlab turilmaydi. Taqsimot ta'sirini ajratganda ham "
         "(29-jadval) magnitude sezilarli darajada ortda qoladi, garchi "
         "qulash o'z taqsimoti bilan ancha yumshoqroq bo'lsa ham. "
         "Ablation esa kompensatsiyaning TARKIBIY ekanini ko'rsatadi: xuddi "
         "shu kanallar kompensatsiyasiz olib tashlanganda WER 0.2006 dan "
         "1.3393 ga chiqadi, chunki bizning mezon ortiqcha — ammo kattaligi "
         "mumkin bo'lgan — kanallarni tanlaydi va ularning hissasi faqat "
         "vakilga qo'shilgandagina saqlanadi."),
        ("Kalibrlash hajmiga miqdoriy talab.",
         "Kalibrlashga asoslangan past-rank yoyilma uchun qator/rank nisbati kamida "
         "10-20 bo'lishi kerak. Undan past bo'lsa yechim kalibrlash to'plamini "
         "interpolatsiya qiladi: nisbat 0.6 da moslash xatosi 0.00000, held-out xatosi "
         "esa 0.04355."),
        ("Byudjet bo'yicha optimal rank taqsimoti.",
         "Operatorlararo taqsimot ajraluvchan qavariq masala sifatida qo'yiladi; "
         "ochko'z yechim Lagranjning teng marginal qaytim shartini qanoatlantiradi va "
         "aniq hisoblanadi. Teng byudjetda WER 0.1719 dan 0.0729 ga tushadi."),
        ("O'lchangan xato-yutish qonuni.",
         "Operatorlarni alohida buzish per-operator xato 160 barobar, tarmoq chiqish "
         "xatosi esa atigi 4 barobar o'zgarishini ko'rsatadi; buni residual oqimning "
         "nisbiy xatoni suyultirishi tushuntiradi. Audio enkoderda ham, dekoder-only "
         "LLM da ham tasdiqlangan."),
        ("Strukturaviy o'qning kvantlashga ORTOGONALLIGI.",
         "Nashr etilgan post-training usullari bilan taqqoslash: GPTQ kvantlashda "
         "bizdan ustun (54.3% va 12.9%), lekin taklif etilgan qisqartirish uning "
         "ustiga qo'shiladi — 11% kamroq xotira, 3% tezlik, aniqlik o'zgarmaydi "
         "(dWER +0.0003, IO nolni qamraydi). Demak usul kvantlash bilan "
         "raqobatlashmaydi, uni to'ldiradi."),
        ("Siljish chuqurlik bo'ylab to'planadi — yutish qonunining chegarasi.",
         "Standart WikiText-2 da ko'rsatiladiki, kvantlash masshtabini har bir "
         "operatorda alohida optimallashtirish INT4 da tarmoqni buzadi (perplexity "
         "1.696x), garchi u operator xatosini RTN ga nisbatan 26% kamaytirsa ham. "
         "Sabab o'lchandi: vazn xatosini minimallashtiruvchi masshtab har bir "
         "kanalni bir xil yo'nalishda ~1% kichraytiradi, shuning uchun siljish 78 "
         "ta operator bo'ylab ko'payadi (0.9896^78 = 0.44), tasodifiy xato esa "
         "so'nadi. Masshtabni chiqish domenida qayta tanlash — butun sonli kodlar, "
         "bit kengligi va xotira formati o'zgarmaydi — perplexity ni 8.246 ga "
         "tushiradi va usulni AWQ bilan raqobatbardosh qiladi. Xulosa ushbu "
         "kaskaddan tashqarida, ixtiyoriy post-training kvantlash sxemasiga "
         "tegishli: chuqur tarmoqlarda maqsad funksiyasi xato kattaligini emas, "
         "siljimaganlikni muhofaza qilishi kerak. Siljish bit kengligiga "
         "keskin bog'liq: INT8 da gain 0.9999 bo'lib qoladi va masshtabimiz "
         "uchdan-uchgacha RTN dan ozgina ustun turadi (7.549 va 7.550), "
         "shuning uchun kaskadning ish nuqtasi buzilmaydi."),
        ("Kvantlash rejimining aniqlanishi.",
         "Taklif etilgan sxema vazn-only ekani o'lchov bilan asoslanadi: xuddi shu FFN "
         "operatorlarida faollashuvlarni ham kvantlash perplexity ni 15.3x "
         "yomonlashtiradi (230.9 va 3521.5), ayb esa chiqish proyeksiyasida emas — "
         "faqat lm_head ni kvantlash atigi 1.023x zarar keltiradi."),
        ("Chegaralarni belgilovchi salbiy natijalar.",
         "Usul yordam bermaydigan joylar ham keltiriladi: CUR uslubidagi yig'ish "
         "135/135 taqqoslashda faollashuvga sezgir SVD ga yutqazadi va FFN "
         "ortiqchaligi arxitekturaga bog'liq."),
    ], numbered=True)

    figure(doc, 1,
           "Kesh-bog'langan kaskadning umumiy sxemasi. Siqish maqsadi kesh "
           "ierarxiyasidan chiqariladi; ikkita ortogonal qisqartirish o'qi shartli "
           "qo'llanadi; kalibrlangan per-channel kvantlash majburiy.",
           "A clean black-and-white technical block diagram for an IEEE/MDPI paper, "
           "white background, thin black rectangular boxes, orthogonal connector "
           "arrows, no color fill, no gradients, no shadows, sans-serif labels. "
           "Left: a box 'Operator footprint M_W + M_X + M_Y' feeding a box "
           "'Cache-anchored target: required = M_eff / (alpha x L3)'. Center: two "
           "parallel horizontal branches. Upper branch titled 'Axis 1: structural "
           "reduction' with boxes 'Calibration responses h_j' -> 'Representative "
           "nodes J_rep (cos >= tau)' -> 'Compensation W[:,p] += gamma_j W[:,j]' -> "
           "'Channel removal n->n-k'. Lower branch titled 'Axis 2: spectral reduction "
           "(case 3 only)' with boxes 'Gram matrix G = X^T X = L L^T' -> "
           "'Activation-aware SVD' -> 'Rank from cache budget r <= alpha L3/(m+n)'. "
           "Both branches converge into one wide box 'Mandatory: calibrated "
           "per-channel INT8'. A curved dashed arrow from the Compensation box to "
           "this box labeled 'dynamic range 9.6x -> 188.4x makes per-channel "
           "mandatory'. Right: box 'Optimized operator' then a diamond 'WER/CER gate' "
           "with a 'no' arrow looping back labeled 'rollback'. Bottom: three small "
           "legend boxes reading 'Case 1: FP32 fits - no change', 'Case 2: INT8 "
           "suffices - no low-rank', 'Case 3: still over budget - add low-rank'.")

    # ===================== 2. TEGISHLI ISHLAR =====================
    h(doc, "2. Tegishli ishlar", 1)
    para(doc,
         "O'qitilgandan keyingi kvantlash. GPTQ va AWQ kvantlash masshtablarini "
         "kalibrlash faollashuvlari yordamida aniqlashtiradi, SmoothQuant esa "
         "faollashuvdagi chetlanishlarni vaznlarga qayta taqsimlaydi. Bu ishlar "
         "kalibrlashga asoslangan masshtab tanlash min/max dan ustun ekanini "
         "o'rnatgan va bizning o'lchovlarimiz buni tasdiqlaydi (4.2-bo'lim). Bizning "
         "qo'shimchamiz kalibrlangan masshtabning o'zi emas, balki strukturaviy "
         "kompensatsiya per-channel granulyarlikni shunchaki foydali emas, MAJBURIY "
         "qilishi haqidagi topilma.")
    para(doc,
         "Past-rankli yoyilma. FWSVD, ASVD va SVD-LLM yoyilmani faollashuv "
         "statistikasi bilan tortadi. Biz ularning markaziy da'vosini takrorlaymiz — "
         "faollashuvga sezgir yechim 135/135 o'lchovda Ekart-Yang optimumini chiqish "
         "xatosida yutadi, vazn xatosida esa unga yutqazadi — va bilishimizcha "
         "adabiyotda keltirilmagan kalibrlash hajmiga miqdoriy talabni qo'shamiz.")
    para(doc,
         "Strukturaviy qisqartirish. Transformerlar uchun kanal va neyron "
         "qisqartirish yaxshi o'rganilgan, odatda amplituda yoki gradient mezoni "
         "bilan va keyin qo'shimcha o'qitish talab qiladi. Bizning variantimiz "
         "butunlay o'qitilgandan keyingi, kalibrlash javoblaridagi kollinearlik "
         "mezonidan foydalanadi va aniq kompensatsiya qiladi, shuning uchun qayta "
         "o'qitish kerak emas.")
    para(doc,
         "Kesh va roofline asosidagi optimallashtirish. Apparatni hisobga oluvchi "
         "rejalashtirish yadro kutubxonalarida standart, ammo u siqish usulini "
         "TANLASH uchun kamdan-kam ishlatiladi. HAQ va AMC bit kengligi "
         "yoki kesish nisbatini reinforcement learning bilan qidiradi — "
         "apparatni MUKOFOT orqali ko'radi, maqsadni undan chiqarib "
         "olmaydi. 3.1-bo'limdagi kesh-bog'langan "
         "maqsad — bizning fikrimizcha yangi bo'lgan komponent.")
    para(doc,
         "Birgalikda kesish va kvantlash. JPQD kesish, kvantlash va "
         "distillashni transfer-o'qitish davomida parallel yuritadi [51], "
         "GETA kvantlash-ogoh bog'liqlik grafi ustida qo'shma qidiruv "
         "quradi [52], LLM lar uchun strukturaviy kesish bilan "
         "aralash-aniqlikdagi PTQ ni birlashtiradigan freymvorklar ham "
         "mavjud [53]. Bu ishlar yo QAYTA O'QITISH ichida, yo bit/nisbat "
         "QIDIRUVIDA ishlaydi. Ushbu ish ataylab o'qitishsiz rejimda "
         "qoladi, va bu texnik tanlov emas, masala qo'yilishining o'zi: "
         "kam resursli til sharoitida joylashtiruvchining qo'lida odatda "
         "na yorliqlangan o'qitish ma'lumoti, na GPU byudjeti bo'ladi — "
         "unda tayyor modelni bor kalibrlash namunasi bilan moslashtirish "
         "yagona amaliy yo'l. Bizning hissamiz boshqa qatlamda — "
         "o'qitishsiz, qidiruvsiz rejimda ikki bosqichning O'LCHANGAN "
         "o'zaro ta'siri: kompensatsiya keltiradigan satr-diapazon "
         "inflyatsiyasi, uning granulyarlikni majburiy qilishi va "
         "Hessian-kompensatsiyali kvantlagichning aynan shu qoldiqni "
         "yutishi (4.4, 4.9b-bo'limlar). Bilishimizcha bu bog'lanish "
         "zanjiri hujjatlashtirilmagan.")

    # ===================== 3. USUL =====================
    doc.add_page_break()
    h(doc, "3. Materiallar va usullar", 1)

    h(doc, "3.1. Kesh-bog'langan siqish maqsadi", 2)
    para(doc,
         "Vazni W (m, n) shaklda va kalibrlash faollashuvlari X (B, n) shaklda bo'lgan "
         "matritsa operatori uchun iz tarkiblari quyidagilar: vazn baytlari M_W, "
         "kirish M_X, chiqish M_Y va ishchi bufer M_tmp. Operatorni bajaradigan "
         "yadrolar to'plami oldindan qat'iy belgilanmagani uchun maqsad kesh sifatida "
         "BARCHA mantiqiy protsessorlar baham ko'radigan daraja olinishi kerak (bizning "
         "platformada L3). Foydalanish koeffitsienti alpha bilan samarali byudjet va "
         "kesh bosimi:")
    eq(doc, "M_cache_eff = alpha В· M_cache ,    K_cache = M_eff / M_cache_eff", 1)
    para(doc, "talab qilinadigan qisqartirish koeffitsienti esa")
    eq(doc, "rho = max(1, K_cache) .", 2)
    para(doc,
         "M_eff ikkita chegaraviy ta'rifga ega: yuqori chegara M_total = M_W + M_X + "
         "M_Y + M_tmp va blokli GEMM bahosi. Biz ikkalasini ham keltiramiz va "
         "5-bo'limda har bir bajarish rejimi uchun qaysi biri fizik jihatdan "
         "ahamiyatli ekanini ko'rsatamiz.")
    para(doc,
         "Kaskad uchta holatni ajratadi: (1) FP32 allaqachon sig'adi, o'zgartirish "
         "yo'q; (2) majburiy INT8 sig'dirР°Рґi, past-rank QARALMAYDI; (3) INT8 dan "
         "keyin ham byudjetdan oshadi, past-rank qo'shiladi. Holatlar chastotasi "
         "4.1-bo'limda keltirilgan.")

    figure(doc, 2,
           "Kesh mezonining bajarish rejimiga bog'liqligi. Vaznning qayta ishlatilishi "
           "keshda saqlanish yoki bayt hajmi qaysi biri cheklovchi ekanini hal qiladi.",
           "A two-panel black-and-white scientific diagram, white background, flat "
           "line art, no color. Left panel titled 'Encoder (compute-bound)': a weight "
           "matrix block with 1500 stacked thin activation slices streaming through "
           "it, an arrow labeled '1500x weight reuse', and a bottleneck symbol placed "
           "on an ALU/compute icon. Right panel titled 'Decoder, batch = 1 "
           "(memory-bound)': the same weight matrix block with a single thin "
           "activation slice, arrow labeled '1x reuse', bottleneck symbol on a DRAM "
           "icon, and 23 greyed-out layer blocks illustrating cache eviction before "
           "the weight is reused. Thin black lines, sans-serif labels, publication "
           "quality.")

    h(doc, "3.2. Funksional guruhlash va strukturaviy olib tashlash", 2)
    para(doc,
         "j yashirin tuguni uchun h_j orqali uning funksional javob vektorini "
         "belgilaymiz — u shu kanalning barcha to'ldiruvchi bo'lmagan kalibrlash "
         "pozitsiyalaridagi faollashuvlarini birlashtirish orqali quriladi. Ikki tugun "
         "javob vektorlari kollinear bo'lganda funksional jihatdan ortiqcha "
         "hisoblanadi. Biz bir vaqtda ikkita shartni talab qilamiz:")
    eq(doc, "cos(h_j, h_p) = <h_j, h_p> / (||h_j|| ||h_p||) >= tau ,", 3)
    eq(doc, "eps_j = ||W[:, j]|| В· ||h_j|| В· sin(theta_jp) / (||Y|| + xi) <= eps_thr .", 4)
    para(doc,
         "Ikkinchi shart hal qiluvchi ahamiyatga ega: faqat burchak yaqinligi "
         "kanalning operator CHIQISHIGA qanchalik hissa qo'shishini hisobga olmaydi. "
         "(4) tenglama qoldiqni W ning ustun normasi bilan tortadi va kalibrlash "
         "chiqishi kattaligiga normallashtiradi, ya'ni eps_j — kanal birlashtirilganda "
         "chiqish xatosiga qo'shiladigan nisbiy hissa.")
    para(doc,
         "p ga bog'langan guruh ichida j a'zosi uchun eng kichik kvadratlar bo'yicha "
         "optimal kompensatsiya koeffitsienti va vaznning yangilanishi:")
    eq(doc, "gamma_j = <h_j, h_p> / ||h_p||^2 ,    W[:, p] <- W[:, p] + gamma_j W[:, j] ,", 5)
    para(doc,
         "shundan so'ng j ustuni o'chiriladi. Agar h_j = gamma_j h_p aniq bajarilsa, "
         "almashtirish yo'qotishsiz bo'ladi; qoldiq xato faqat kollinearlikdan "
         "chetlanish bilan belgilanadi. Guruh vakili sifatida guruh o'rtachasiga "
         "kosinus o'xshashligi eng katta bo'lgan a'zo tanlanadi.")
    para(doc,
         "FFN blokida oraliq kenglik birinchi proyeksiyaning CHIQISH o'lchami va "
         "ikkinchisining KIRISH o'lchami hisoblanadi, shuning uchun k ta oraliq "
         "kanalni olib tashlash bitta qarordan ikkala matritsani qisqartiradi:")
    mono(doc, "    W1 (d, F) -> (d, F-k)     bias (F,) -> (F-k,)\n"
              "    faollashuv: elementwise, o'zgarmaydi\n"
              "    W2 (F, d) -> (F-k, d)")
    para(doc,
         "Gated arxitekturalarda (Llama uslubi, h = SiLU(W_gate x) * (W_up x)) xuddi "
         "shu qaror uchta matritsani qisqartiradi. Bu — strukturaviy o'q; u n va m ni "
         "kamaytiradi, past-rankli yoyilma esa rankni kamaytirib n va m ni "
         "o'zgarishsiz qoldiradi. Demak ikki o'q bir-birini to'ldiradi: k ta kanalni "
         "olib tashlash teng parametr sonida yuqoriroq rank uchun joy bo'shatadi, "
         "chunki r'(m + n - k) = r(m + n) dan r' > r kelib chiqadi.")

    figure(doc, 3,
           "FFN blokida kompensatsiya bilan strukturaviy kanal olib tashlash. Bitta "
           "qaror ikkala proyeksiyani qisqartiradi.",
           "A black-and-white technical diagram, white background, thin line art. Top "
           "row: input vector block (width d) -> matrix W1 (d x F) -> intermediate "
           "vector (width F) -> activation symbol -> matrix W2 (F x d) -> output "
           "vector (width d). Bottom row: the same pipeline after removal, with the "
           "intermediate width visibly narrowed to F-k and BOTH matrices drawn "
           "narrower, annotated 'one decision, two matrices'. A circular inset in the "
           "middle shows two nearly-parallel vectors h_j and h_p with the angle theta "
           "between them, the formula 'gamma_j = <h_j,h_p>/||h_p||^2', and an arrow "
           "folding column j into column p while column j fades out. Sans-serif "
           "labels, publication quality, no color.")

    h(doc, "3.3. Kalibrlangan per-channel kvantlash", 2)
    para(doc,
         "Masshtab s bilan simmetrik kvantlash vaznlarni q = round(clip(W/s, -q_max, "
         "q_max)) butun kodlariga akslantiradi. Kodlar butun bo'lgani uchun tiklash "
         "yo'qotishi s bo'yicha bo'lakli-doimiy va gradient tushish asoslanmagan. "
         "Shuning uchun har yarim qadami aniq minimum beradigan almashinuvchi "
         "minimizatsiya ishlatiladi:")
    eq(doc, "q_t = round(clip(W / s_t)) ,    s_{t+1} = <W, q_t> / <q_t, q_t> ,", 6)
    para(doc,
         "natijada vazn tiklash yo'qotishi L_W kamaymaydi. Ikkinchi faza birinchi faza "
         "optimumi atrofidagi lokal panjarani kalibrlash maqsadi bo'yicha qidiradi va "
         "nomzodni faqat L_W ni beta chegarasidan ortiq oshirmasa qabul qiladi:")
    eq(doc, "s* = argmin_s [ L_W(s) + lambda L_calib(s) ]   shart:  L_W(s) <= (1+beta) L_W(s_1) .", 7)
    para(doc,
         "Granulyarlik aniqlashtirishning o'zidan ko'ra muhimroq. Y = X W^T bo'lgani "
         "uchun i chiqish kanali faqat i vazn satriga bog'liq, ya'ni kalibrlash "
         "maqsadi chiqish kanallari bo'yicha aniq ajraladi va har bir satr masshtabini "
         "mustaqil optimallashtirish yaqinlashtirishsiz to'g'ri yechim beradi. Buni "
         "arzon qilish uchun per-channel kalibrlash xatosi bir marta hisoblangan Gram "
         "matritsasi orqali kvadratik forma sifatida baholanadi:")
    eq(doc, "|| X d_i ||^2 = d_i^T G d_i ,   G = X^T X ,   d_i = W_deq[i,:] - W[i,:] .", 8)

    h(doc, "3.4. Faollashuvga sezgir past-rankli yoyilma", 2)
    para(doc,
         "3-holat qo'llanganda maqsad vazn xatosi emas, chiqish xatosi hisoblanadi. "
         "Gram matritsasining Xolentskiy ko'paytuvchisi bilan, G = L L^T:")
    eq(doc, "|| X (W - W')^T ||_F = || (W - W') L ||_F ,", 9)
    para(doc, "shuning uchun optimal rank-r yechim o'zgartirilgan fazoda kesish va "
              "orqaga akslantirish orqali olinadi:")
    eq(doc, "W' = trunc_svd(W L, r) L^{-1} .", 10)
    para(doc,
         "Rank tanlanmaydi, chiqariladi: r(m+n) bayt talab qiladigan ikki faktorli "
         "INT8 tasvir uchun kesh byudjetiga sig'ish sharti")
    eq(doc, "r <= alpha В· M_cache / (m + n) .", 11)
    para(doc,
         "Ushbu chegara ULAR UCHUN faol bo'ladigan operatorlarda ma'noga ega, "
         "ammo o'rganilgan audio enkoderda u ishlamaydi va buni ochiq qayd "
         "etamiz. Uning fc1 operatori uchun m = 1024, n = 4096 bo'lgani va "
         "byudjet alphaВ·L3 = 16.8 MiB ni tashkil qilgani uchun (11) "
         "r <= 3440 ni beradi, matritsaning maksimal ranki esa 1024. Ya'ni "
         "chegara hech qachon bog'lamaydi va bu model uchun rankni (11) emas, "
         "3.5-bo'limdagi byudjet taqsimoti belgilaydi. Chegaraning o'zi "
         "operatori keshdan sezilarli katta bo'lgan modellarda — masalan "
         "19-jadvaldagi Llama lug'at va head matritsalarida — kuchga kiradi.",
         italic=True, size=10)

    h(doc, "3.5. Byudjet bo'yicha optimal rank taqsimoti", 2)
    para(doc, "Bitta umumiy byudjetni operatorlar orasida taqsimlash masalasi:")
    eq(doc, "min  sum_i E_i(r_i)    shart:  sum_i c_i r_i <= B ,   c_i = m_i + n_i ,", 12)
    para(doc,
         "bunda E_i — i operatorining o'lchangan xato egri chizig'i. E_i kamaymaydigan "
         "va spektral kesish uchun r bo'yicha qavariq: ketma-ket singulyar qiymatlar "
         "kamaymaydi, ya'ni har qo'shimcha rank birligi oldingisidan kam hissa "
         "qo'shadi. Ajraluvchan qavariq maqsad uchun uzluksiz yechim marginal qaytimni "
         "parametr birligiga tenglashtirish bilan aniqlanadi,")
    eq(doc, "- (dE_i / dr_i) / c_i = lambda   (barcha i uchun) ,", 13)
    para(doc,
         "butun sonli yechim esa byudjetning keyingi birligini eng katta xato "
         "kamayishini beradigan operatorga sarflaydigan ochko'z algoritmdan kelib "
         "chiqadi. Bu ochko'z ajraluvchan qavariq maqsadlar uchun ANIQ, ya'ni "
         "taqsimlagich evristika emas, optimal.")

    h(doc, "3.6. Freymvork: kirishlar, maqsad funksiyasi va qidiruv tartibi", 2)
    para(doc,
         "Yuqoridagi bosqichlar usulni tashkil qiladi; ularni ishlatiladigan "
         "vositaga aylantirish uchun uchta narsa aniqlanishi kerak — nima "
         "beriladi, nima optimallashtiriladi va nomzodlar qanday tartibda "
         "sinaladi. Freymvork uchta kirish oladi: MODEL (ONNX grafigi), "
         "MAQSADLI KESH HAJMI va ANIQLIK BYUDJETI (ruxsat etilgan eng yuqori "
         "WER). Kesh hajmi apparatdan o'qilmasdan, argument sifatida "
         "beriladi: vositaning asosiy qiymati oldinda turmagan apparat uchun "
         "javob bera olishida, mahalliy L3 esa xususiy hol.")
    para(doc, "Keshga sig'ish — maqsad, darvoza emas.", bold=True, size=10)
    para(doc,
         "Ish davomida kesh talabini QAT'IY cheklov sifatida qo'yish ikki "
         "joyda mo'rtlik keltirib chiqardi. Birinchidan, maqsad ikkilik "
         "bo'lgani uchun qaror pichoq tig'ida turadi: alfa bo'yicha chegara "
         "0.033 uzoqlikda, L3 bo'yicha esa 1.1 MiB uzoqlikda (3- va "
         "4-jadval). Ikkinchidan, maqsad erishib bo'lmaydigan bo'lsa, qat'iy "
         "cheklov kaskadni o'z mezoni tasdiqlamaydigan joyda ham kesishda "
         "davom etishga majbur qiladi. Shuning uchun maqsad funksiyasi "
         "sig'ish emas, KESHDAN O'TKAZIB YUBORISH (miss) hajmi:")
    eq(doc, "Miss(P, t) = L * [ b(P, t) + max(0, b(P, t) - B) * (R - 1) ] ,", 14)
    para(doc,
         "bunda b(P, t) — P qismining t ishlovidan keyingi qatlam hajmi, "
         "B = alpha*L3 byudjet, L — qatlamlar soni, R — vaznning bir "
         "o'tishdagi qayta ishlatilishi. Mantiq oddiy: byudjetga sig'gan "
         "vaznlar bir marta o'qiladi, oshgan qismi esa har qayta "
         "ishlatishda qayta o'qiladi. Bu funksiya uzluksiz (Lipschitz "
         "konstantasi R ga teng), sig'ishni TALAB QILMAY rag'batlantiradi va "
         "sig'ish imkonsiz bo'lganda ham ma'noli tartib beradi.")
    para(doc,
         "Ifodaning BIRINCHI hadi — L * b(P, t) — xotira devori "
         "mantig'ining bevosita ifodasi: vaznlar keshda turmasa, har "
         "o'tishda ularning har bayti DRAM dan keladi, ya'ni bu had "
         "o'tish boshiga MAJBURIY miss hajmining o'zi. Iz kamaytirilsa "
         "miss hajmi proportsional kamayadi, va 4.10-bo'limda "
         "ko'rsatilganidek vaqt shunga ergashadi. Shuning uchun maqsad "
         "funksiyasining tayanchi shu had; ikkinchi had (overflow "
         "jarimasi) esa qo'shimcha, yadro sinfiga bog'liq effektni "
         "modellashtiradi.", italic=True, size=10)
    para(doc,
         "Ifodaning ikkinchi hadi — overflow jarimasi — KEYINCHALIK "
         "TEKSHIRILDI VA TASDIQLANMADI. 4.10a-bo'limda byudjetni ikki "
         "tomondan qamrab oluvchi operatorlar o'lchanganda MAC boshiga vaqt "
         "tekis chiqdi: fp32 da 0.95x, modellar aslida ishlatadigan INT8 "
         "yadrosida 1.02x — ya'ni byudjetni kesib o'tishning narxi 2% dan "
         "oshmaydi, ifoda esa o'sha baytlarga R - 1 ni (bu yerda 1499) "
         "ko'paytiradi. Sabab bloklangan GEMM da: keshda turishi kerak "
         "bo'lgan narsa vazn PLITKASI, butun matritsa emas. Shuning uchun "
         "quyida miss faqat BIRINCHI hadi — umumiy bayt hajmi — bo'yicha "
         "ishonchli tartiblovchi hisoblanadi, va o'sha tartib o'lchangan "
         "vaqt bilan mos keladi (r = +0.974, 30-jadval); overflow hadiga "
         "tayanadigan bashoratlar esa tasdiqlanmagan deb belgilanadi.",
         italic=True, size=10)
    para(doc, "Bajarilishni oldindan tekshirish.", bold=True, size=10)
    para(doc,
         "Strukturaviy bosqich faqat FFN ga tegadi, attention qo'zg'almaydi. "
         "Shuning uchun INT8 dan keyingi qoldiq faqat FFN ulushidan "
         "chiqarilishi kerak va talab qilinadigan olib tashlash ulushi "
         "quyidagicha aniqlanadi:")
    eq(doc, "f = [ M_layer * (1 - 1/rho_resid) ] / M_FFN ,   rho_resid = rho / 4 ,", 15)
    para(doc,
         "f > 1 bo'lsa maqsad bu model uchun umuman erishib bo'lmaydi. Bu "
         "tekshiruv soniyalar oladi va qurishdan OLDIN yuritiladi, chunki u "
         "'kaskad yomon model berdi' degan xulosani 'maqsad erishib "
         "bo'lmas edi' degan aniq bayonotga aylantiradi.")
    para(doc, "Nomzodlar zinapoyasi va yalqov baholash.", bold=True, size=10)
    para(doc,
         "Har bir nomzod qurish uchun ~1 soat va baholash uchun ~0.5 soat "
         "talab qiladi, ya'ni to'liq skanerlash imkonsiz. Shuning uchun "
         "reja YUMSHOQDAN QATTIQQA to'liq tartiblangan zinapoya beradi va "
         "baholash yalqov: byudjet buzilgan birinchi pog'onada to'xtaydi, "
         "javob esa undan pastdagi pog'ona bo'ladi. To'xtash asosli, chunki "
         "zinapoya monoton — buzilgan pog'onadan yuqoridagilar yo'qotilgan "
         "aniqlikni tiklay olmaydi.")
    para(doc,
         "Zinapoyaning tartibi ikki qoidadan iborat. Avval hamma qismda "
         "kvantlash, keyin strukturaviy qisqartirish — bu 1.1-bo'limdagi "
         "orqaga rejalashtirish tamoyilining o'zi. Bu qoida OG'IRLIK "
         "sifatida emas, FAZA CHEGARASI sifatida qo'yiladi, chunki og'irlik "
         "uni ifodalay olmaydi: har qadamni 'miss daromadi / xavf' bo'yicha "
         "saralash birinchi urinish edi va u tuzilmaviy ravishda buziladi — "
         "enkoderning R = 1500 i uning overflow ini shu qadar "
         "kattalashtiradiki, L3 = 12 MiB da rejalashtiruvchi dekoderni "
         "umuman kvantlashdan oldin enkoderni 50% gacha qisqartirardi. "
         "Faza ichida esa eng ko'p miss olib tashlaydigan qadam tanlanadi.")
    para(doc, "Arxitekturaga bog'liq bo'lmagan blok topish.", bold=True, size=10)
    para(doc,
         "Qisqartiriladigan blok operator NOMI bo'yicha emas, usul haqiqatan "
         "talab qiladigan xossa bo'yicha topiladi: ikkita matritsa operatori "
         "orasidagi KENGAYADIGAN umumiy o'q — A ning chiqish kengligi B ning "
         "kirish kengligiga teng va undan katta, orasida esa faqat "
         "koordinata bo'yicha ishlaydigan tugunlar (aktivatsiya, bias) "
         "turadi. Bu shart kanalni ANIQ olib tashlash mumkinligini "
         "kafolatlaydi. Kengayish sharti attention ni ataylab chetda "
         "qoldiradi: v va o proyeksiyalari kenglikni bo'lishadi, lekin uni "
         "kengaytirmaydi, ortiqchalik esa aynan kengaytirilgan o'qda "
         "yashaydi. Nom bo'yicha qidiruv Whisper ning fc1/fc2 iga "
         "moslashgan va boshqa ASR arxitekturalarida HECH NARSA topmaydi; "
         "jimgina hech narsa topish esa eng xavfli nosozlik, chunki "
         "rejalashtiruvchi butun FFN steki mavjud bo'lgani holda 'maqsad "
         "erishilmaydi' degan xulosaga kelardi. Bloki topilmagan model nol "
         "qiymat bilan qayd etiladi, taxmin bilan emas.")
    para(doc, "Umumiylik da'vosining aniq chegarasi.", bold=True, size=10)
    para(doc,
         "Freymvorkning qaysi qismi arxitekturadan mustaqil ekanini aniq "
         "ajratamiz, chunki bu da'voni kengroq o'qish oson. Kesh maqsadi, "
         "miss funksiyasi, bajarilish tekshiruvi, nomzodlar zinapoyasi va "
         "blok topish — bularning hammasi baytlar va graf topologiyasi "
         "bilan ishlaydi, ya'ni ilgari ko'rilmagan transformer ASR modeli "
         "uchun ham REJA to'g'ri chiqadi. Umumiylik faqat bitta joyda "
         "to'xtaydi: audioni gipotezaga aylantirish uchun belgi ajratgich, "
         "chiqish lug'ati va dekodlash sikli kerak, va bular grafning emas, "
         "modelning xossalari. Shuning uchun BAHOLASH tor interfeys orqali "
         "o'tadi — uchta metod (enkoder kirishini tayyorlash, matn olish, "
         "normallashtirish) — va yangi arxitekturani qo'shish narxi shu "
         "uchtasidan iborat. Ushbu ishda interfeysning yagona bajarilishi "
         "Whisper uchun, chunki o'lchovlar shu model bo'yicha; adapteri "
         "yo'q model rejalashtiriladi, lekin baholanmaydi va vosita buni "
         "ochiq aytadi. Bu ataylab tanlangan chegara: dalili bo'lmagan "
         "arxitektura uchun yarim yozilgan adapter tajribalar "
         "tasdiqlamaydigan qamrovni nazarda tutgan bo'lardi.")
    para(doc, "Kalibrlash to'plami — oshkor kirish.", bold=True, size=10)
    para(doc,
         "Ushbu ishdagi har bir siqish qarori kalibrlash ma'lumotining "
         "funksiyasi: qaysi kanallar kollinear ko'rinadi, qaysi masshtab "
         "xatoni minimallashtiradi, qaysi rank saqlashga arziydi. Shuning "
         "uchun kalibrlash to'plami oshkor argument bo'lib, uning nomi "
         "artefakt fayl nomiga kiradi — bir xil byudjetni turli "
         "namunalardan so'ragan ikki yugurish turli kanal xaritalarini "
         "beradi. Bundan tashqari kalibrlash va baholash namunalarining "
         "kesishmasligi OGOHLANTIRISH emas, XATO sifatida tekshiriladi: "
         "4.9-bo'limda ko'rsatilganidek, variantlar kalibrlash bilan bir "
         "xil ma'lumotda tanlanganda tartib mustaqil splitda saqlanmagan.")

    h(doc, "3.7. Eksperiment sharoiti", 2)
    para(doc,
         "Platforma: Intel Tiger Lake H, 16 mantiqiy yadro, L2 = 1.25 MiB har yadro "
         "juftligiga, L3 = 24 MiB umumiy; alpha = 0.7 bilan 16.8 MiB byudjet. Ishga "
         "tushirish muhiti: ONNX Runtime 1.28, latency uchun bitta intra-op oqim, "
         "qizdirish va takroriy yurishlar medianasi. Apparat hisoblagichlari: Intel "
         "VTune Profiler 2026.4, event-based sampling drayveri faol.")
    para(doc,
         "Modellar: (i) Whisper-medium o'zbek ASR (enkoder 144 va dekoder 240 vaznli "
         "matritsa operatori); (ii) mBERT; (iii) open_llama_3b_v2. Ma'lumot: Mozilla "
         "Common Voice o'zbek; kalibrlash uchun 12 namuna, baholash uchun 80 held-out "
         "namuna, til modeli tajribalari uchun 8000 jumlalik matn korpusi. Kalibrlash "
         "va baholash to'plamlari hech qayerda kesishmaydi, javob vektorlaridan "
         "to'ldiruvchi pozitsiyalar chiqarib tashlanadi.")
    para(doc,
         "Mezonlar: so'z va belgi xatoligi (asosiy), operator (E_loc) va tarmoq "
         "(E_glob) darajasidagi nisbiy chiqish xatosi (diagnostik), vazn baytlari, "
         "o'lchangan latency va VTune xotira mezonlari. Ishonch oraliqlari — namunalar "
         "bo'yicha 2000 qayta tanlashli persentil bootstrap; juftlik taqqoslashlar "
         "namuna bo'yicha farqlarning juftlik bootstrapi.")

    # ===================== 4. NATIJALAR =====================
    doc.add_page_break()
    h(doc, "4. Natijalar", 1)

    h(doc, "4.1. Chiqarilgan maqsadlar va holatlar chastotasi", 2)
    table(doc, "1-jadval. Granulyarlik bo'yicha kesh-bog'langan talab (byudjet alphaВ·L3 = 16.8 MiB).",
          ["Granulyarlik", "Dekoder (MiB)", "Talab", "Enkoder (MiB)", "Talab"],
          [["per-operator", "16.0", "sig'adi", "16.0", "0.95x"],
           ["per-layer", "64.0", "3.81x", "48.0", "2.86x"],
           ["butun model", "1536.0", "91.4x", "1152.0", "68.6x"]],
          good_rows=(1,))
    table(doc, "2-jadval. Uchta arxitekturada kaskad holatlarining chastotasi.",
          ["Holat", "Shart", "Amal", "Whisper enk.", "Whisper dek.", "Llama"],
          [["1", "FP32 sig'adi", "o'zgartirish yo'q", "96/144", "240/240", "faqat attn"],
           ["2", "INT8 yetarli", "past-rank yo'q", "96/144", "240/240", "q/k/v/o 0.58x"],
           ["3", "byudjetdan oshadi", "past-rank qo'shiladi", "48/144", "0/240",
            "FFN 1.57x, head 5.81x"]],
          good_rows=(2,))
    para(doc,
         "Dekoder uchun chiqarilgan qatlam talabi 3.81x, INT8 esa 4.00x beradi, "
         "shuning uchun kaskad kvantlashni tanlaydi. 3-holat dekoderda umuman yuzaga "
         "kelmaydi, enkoderning FFN operatorlarida yuzaga keladi va Llama da keng "
         "tarqalgan — bu bizning to'plamimizdagi past-rank shoxchasi haqiqatan talab "
         "qilinadigan birinchi model.")
    para(doc, "Byudjetning qamrovi va alfa konstantasi.", bold=True, size=10)
    para(doc,
         "Byudjet VAZN baytlarini hisoblaydi, faollashuvlarni emas. Bu "
         "soddalashtirish emas, balki masalaning tuzilishidan kelib chiqadi: "
         "kaskad qaror qabul qiladigan kattalik — vaznlar, va blokli GEMM da "
         "keshda turishi kerak bo'lgan kattalik ham aynan vazn plitkasi, "
         "faollashuv plitkalari esa uning yonidan oqib o'tadi. Ikkalasining "
         "kesh xatti-harakati bir xil emas: vazn turg'un va pozitsiyalar "
         "bo'ylab qayta ishlatiladi, faollashuv esa bir marta iste'mol "
         "qilinadi. alpha = 0.7 koeffitsiyenti aynan shu qolgan hamma narsani "
         "— faollashuvlar, oraliq tenzorlar, boshqa jarayonlar — qamrab "
         "oluvchi zaxira sifatida kiritilgan.")
    para(doc,
         "Yagona qo'lda tanlangan konstanta faqat u beradigan QARORLAR "
         "barqaror bo'lsa zararsiz, shuning uchun alpha bo'yicha sezgirlik "
         "o'lchandi. Holat 2 dan holat 3 ga o'tish chegarasi "
         "alpha* = hajm / (4 x L3) nuqtasida joylashadi:")
    table(doc, "3-jadval. Kaskad qarorining alfa konstantasiga sezgirligi "
               "(L3 = 24 MiB, INT8 4.00x beradi).",
          ["Granulyarlik", "alpha*", "|0.7 - alpha*|", "Barqarorlik"],
          [["dekoder, per-layer", "0.667", "0.033", "chegaraga yaqin"],
           ["enkoder, per-layer", "0.500", "0.200", "barqaror"],
           ["dekoder, eng katta operator", "0.167", "0.533", "barqaror"],
           ["enkoder, eng katta operator", "0.167", "0.533", "barqaror"]],
          good_rows=(1, 2, 3), bad_rows=(0,))
    para(doc,
         "To'rtta qarordan uchtasi alphaning har qanday maqbul qiymatida "
         "o'zgarmaydi. Dekoderning qatlam darajasidagi qarori esa chegaraga "
         "yaqin: alpha 0.667 dan pastga tushsa, u 3-holatga o'tadi va kaskad "
         "dekoderga past-rank qo'shishni buyurgan bo'lardi — 15-jadval esa "
         "bunday o'zgarish qimmatga tushishini ko'rsatadi (WER 0.3056, FP32 "
         "ning 0.1793 iga qarshi). Buni ochiq qayd etamiz.")
    para(doc,
         "Ayni paytda zaiflik CHIQARISHDA, XULOSADA emas. alpha = 0.7 da "
         "olingan qaror — 'dekoderga INT8 yetarli, to'xta' — mustaqil ravishda "
         "uchdan-uchgacha tasdiqlangan: INT8 dekoder FP32 dan statistik "
         "jihatdan farqlanmaydi (dWER = +0.0032, 95% IO [-0.0039, +0.0100], "
         "4.9-bo'lim). Ya'ni alpha ning qiymati taxmin qilingan emas, "
         "natijasi bo'yicha tasdiqlangan; chiqarishning aniqligi esa uning "
         "asoslanishidan yuqori va bu keyingi ishda alphani apparat "
         "hisoblagichlaridan chiqarishni talab qiladi.", italic=True, size=10)

    para(doc, "Kesh hajmi qaror o'zgaruvchisi sifatida.", bold=True, size=10)
    para(doc,
         "Alfa bo'yicha sezgirlik chiqarishning bitta KOEFFITSIYENTIGA "
         "tegishli edi; ikkinchi va muhimroq savol esa uning APPARATGA "
         "javob berishi haqida: L3 boshqacha bo'lsa, chiqarish boshqa "
         "qaror beradimi, va o'sha qaror to'g'rimi? To'g'ridan-to'g'ri "
         "tekshiruv ikkinchi platformada o'lchash bo'lardi (5.4-bo'limda "
         "cheklov sifatida qayd etilgan), ammo arifmetikaning o'zi "
         "birinchi yarmiga javob beradi, ikkinchi yarmiga esa bizda "
         "allaqachon o'lchovlar bor.")
    table(doc, "4-jadval. Kaskad qarori L3 hajmi bo'yicha (alpha = 0.7). "
               "Yacheykada — chiqarilgan talab va u yuboradigan holat.",
          ["Granulyarlik", "8 MiB", "12 MiB", "16 MiB", "24 MiB", "32 MiB", "48 MiB"],
          [["dekoder, per-layer (64 MiB)", "11.43x / 3", "7.62x / 3", "5.71x / 3",
            "3.81x / 2", "2.86x / 2", "1.90x / 2"],
           ["enkoder, per-layer (48 MiB)", "8.57x / 3", "5.71x / 3", "4.29x / 3",
            "2.86x / 2", "2.14x / 2", "1.43x / 2"],
           ["eng katta operator (16 MiB)", "2.86x / 2", "1.90x / 2", "1.43x / 2",
            "sig'adi / 1", "sig'adi / 1", "sig'adi / 1"]])
    para(doc,
         "Chiqarish qotib qolgan emas: L3 ning maqbul diapazonida u uchta "
         "holatning HAMMASINI beradi va qaror chegaralari analitik "
         "aniqlanadi — dekoder qatlami uchun 2-holat L3 >= 22.9 MiB da, "
         "operator uchun 1-holat L3 >= 22.9 MiB da boshlanadi. Bu shuni "
         "anglatadiki, kaskadning ushbu mashinadagi xatti-harakati "
         "modelning xossasi emas, model va kesh JUFTLIGINING xossasi.")
    para(doc,
         "Muhimi shundaki, jadvaldagi o'zgarishlarning oqibatlari taxmin "
         "emas, o'lchangan. Dekoder 3-holatga o'tsa (L3 <= 16 MiB) kaskad "
         "past-rank buyuradi — 15- va 19-jadvallarga ko'ra bu WER ni 0.1833 "
         "dan 0.6101 ga ko'taradi. Dekoder 2-holatda qolsa, INT8 FP32 dan "
         "farqlanmaydi (0.1793, dWER = +0.0032 ns). Enkoder 3-holatda esa "
         "strukturaviy qisqartirish TEKIN chiqadi (0.1833 va 0.1847). "
         "Ya'ni jadvalning har bir ustuni bo'yicha kaskad tavsiya "
         "qiladigan amal, o'lchov bo'yicha, o'sha ustunda to'g'ri amal.")
    para(doc,
         "Shu bilan birga 24 MiB ning dekoder chegarasi 22.9 MiB ga yaqin "
         "ekani 3-jadvaldagi alpha* = 0.667 bilan bir xil zaiflikning "
         "ikkinchi ko'rinishi — ular bitta nuqtaning ikki koordinatasi. "
         "Shuning uchun ishning markaziy da'vosi tekshirilgan shaklda "
         "qo'yiladi. Tezlikning xotira kanali MISS HAJMI orqali ishlaydi "
         "va o'lchov bilan tasdiqlangan (baytlar-vaqt korrelyatsiyasi "
         "r = +0.974, xotira to'xtashlari umumiy vaqtdan tez — 2.41x ga "
         "qarshi 1.91x — qisqaradi, 4.10-bo'lim); rad etilgani faqat "
         "REZIDENTLIK CHEGARASI mexanizmi — 'butun vazn alpha*L3 ga "
         "sig'ishi shart' degan keskin shart — sozlangan bloklangan "
         "yadrolarda (4.10a; sodda yadrolarda u ham amal qiladi, "
         "1.56-2.3x). Kesh HAJMI esa qaror o'zgaruvchisi: u qayerda "
         "to'xtashni aytadi va to'xtash nuqtalari o'lchovda to'g'ri "
         "chiqadi.", italic=True, size=10)

    h(doc, "4.2. Kvantlash masshtabining hissasi", 2)
    table(doc, "5-jadval. Masshtabni baholash usulining operator chiqish xatosiga "
               "ta'siri (E_loc, held-out).",
          ["Masshtab usuli", "Enkoder fc1", "Yaxshilanish", "Dekoder fc1", "Yaxshilanish"],
          [["Q1 min/max (kutubxona standarti)", "0.00685", "—", "0.00441", "—"],
           ["Q2 almashinuvchi minimizatsiya", "0.00651", "+5.0%", "0.00442", "-0.1%"],
           ["Q3 kalibrlangan, per-tensor", "0.00525", "+23.3%", "0.00360", "+18.4%"],
           ["Q4 kalibrlangan, per-channel", "0.00179", "+73.8%", "0.00151", "+65.9%"]],
          good_rows=(3,))
    para(doc,
         "Foyda aynan kalibrlash bosqichidan keladi; almashinuvchi minimizatsiya "
         "yolg'iz o'zi deyarli hech narsa qo'shmaydi. Per-channel masshtablar har "
         "operator uchun m ta qo'shimcha FP32 qiymat, ya'ni vazn baytlarining taxminan "
         "0.5% ini talab qiladi.")
    para(doc,
         "Qamrov chegarasi. 5-jadval OPERATOR darajasidagi o'lchov bo'lib, ASR "
         "traktida uchdan-uchgacha joylashtirilmagan; sababi 5-bo'limda "
         "eksport yo'li cheklovi sifatida keltirilgan. Bu natijaning "
         "ahamiyatini pasaytirmaydi, chunki 4.9a-bo'limda kaskadning kvantlash "
         "bosqichi GPTQ ga topshirilgan va shu sababli kalibrlangan masshtab "
         "ishning yakuniy da'vosi emas; u kvantlash granulyarligi haqidagi "
         "diagnostik natija sifatida keltiriladi va 4.4-bo'limdagi bog'liqlikni "
         "asoslaydi.", italic=True, size=10)

    h(doc, "4.3. FFN bloklaridagi strukturaviy ortiqchalik", 2)
    table(doc, "6-jadval. tau = 0.99 da olib tashlanadigan FFN kanallari (Whisper enkoderi).",
          ["Qatlam", "Olib tashlandi", "Ulush", "Qatlam", "Olib tashlandi", "Ulush"],
          [["L0", "1764", "43.1%", "L12", "26", "0.6%"],
           ["L1", "2152", "52.5%", "L13", "10", "0.2%"],
           ["L2", "2376", "58.0%", "L14", "4", "0.1%"],
           ["L3", "2334", "57.0%", "L15-L20", "0", "0.0%"],
           ["L4", "2078", "50.7%", "L21", "3", "0.1%"],
           ["L5", "1535", "37.5%", "L22", "4", "0.1%"],
           ["L6", "1155", "28.2%", "L23", "2", "0.0%"],
           ["L8", "1048", "25.6%", "o'rtacha", "—", "17.1%"]],
          good_rows=(2,))
    para(doc,
         "Yagona chegara qiymati qatlamga kuchli bog'liq qaror hosil qiladi: L2-L3 da "
         "kanallarning 58% i olib tashlanadi, L15 dan boshlab esa hech biri. Hech "
         "narsa qatlam bo'yicha qo'lda sozlanmagan; profil o'lchov orqali aniqlangan "
         "model xossasi.")
    para(doc, "Profil nima uchun shunday va u ko'chadimi.", bold=True, size=10)
    para(doc,
         "Audio enkoderda ortiqchalik chuqurlik bo'yicha monoton kamayadi. "
         "Buning tabiiy izohi kirish ma'lumotining tuzilishida: mel "
         "spektrogrammada qo'shni chastota binlari va vaqt kadrlari kuchli "
         "korrelyatsiyalangan, shuning uchun birinchi transformer qatlamlari "
         "hali past o'lchamli akustik ko'pxillikka yaqin ishlaydi va FFN ning "
         "4096 oraliq kanali ko'p hollarda bir yo'nalishda javob beradi. "
         "Chuqurlik ortgani sari tasvir fonetik va leksik farqlarga ajraladi, "
         "bu esa haqiqatan yuqori o'lchamli, ya'ni kanallar mustaqil ish "
         "bajara boshlaydi.")
    para(doc,
         "Ammo bu qonun emas. mBERT da xuddi shu o'lchov (tau = 0.9) U "
         "shaklidagi profil beradi: L0-L4 da 8-17%, o'rta qatlamlarda eng "
         "past (L6 da 1.3%), oxirida esa qaytadan ko'tariladi (L11 da 20.5%). "
         "Buning izohi ham vazifadan kelib chiqadi — kirish diskret token "
         "vektorlari bo'lib, allaqachon yuqori o'lchamli; o'rta qatlamlar "
         "sintaktik va semantik ishning asosini bajaradi va eng zich bo'ladi; "
         "oxirgi qatlam esa niqoblangan til modeli boshiga, ya'ni lug'at "
         "proyeksiyasiga ixtisoslashadi va shu sababli qisman ortiqcha bo'lib "
         "qoladi.")
    para(doc,
         "Bu farq usulning loyihaviy tanlovini bevosita oqlaydi. Agar qatlam "
         "bo'yicha qo'lda yozilgan jadval ishlatilganda — masalan 'boshlang'ich "
         "qatlamlarni ko'proq qisqartir' — u audio enkoderda to'g'ri, mBERT da "
         "esa TESKARI natija bergan bo'lardi: eng zich o'rta qatlamlarni "
         "qisqartirib, ortiqcha oxirgi qatlamga tegmasdi. Bitta global tau "
         "esa taqsimotni ma'lumotning o'ziga qoldiradi va shu sababli uchala "
         "arxitekturada uch xil profil beradi, hech narsa sozlanmasdan.",
         italic=True, size=10)
    table(doc, "7-jadval. Xato qisqartirilgan qatlamlar bo'ylab to'planmaydi (FP32, "
               "enkoder chiqish xatosi).",
          ["Qisqartirilgan qatlamlar", "1", "4", "8", "12", "19"],
          [["Enkoder chiqish xatosi", "0.0209", "0.0206", "0.0295", "0.0297", "0.0298"]],
          good_rows=(0,))
    para(doc,
         "19 ta qatlamni qisqartirish bittasini qisqartirishdan deyarli qimmatga "
         "tushmaydi — bu past-rankli yaqinlashtirishdan keskin farq qiladi, u yerda "
         "xato operatorlar bo'ylab to'planadi. Sababi kollinear kanal yaqinlashtirilmay, "
         "aniq almashtirilishida.")

    figure(doc, 4,
           "Uchta arxitekturada FFN ortiqchaligining qatlamlar bo'yicha profili. "
           "Ortiqchalik — umumiy emas, model xossasi.",
           "A scientific line chart, white background, publication style, no 3D. "
           "X-axis 'layer index (normalized depth 0-1)', y-axis 'removable channels "
           "(%)' from 0 to 60. Three series with distinct markers and line styles in "
           "grayscale: 'Whisper encoder' rising to a peak of 58% at normalized depth "
           "0.1 then decaying to 0 by depth 0.6; 'open_llama_3b' starting at 3.4% and "
           "decaying to 0 by depth 0.3; 'mBERT' flat near 0 throughout. A horizontal "
           "dashed line marks the Whisper mean of 17.1%. Legend inside the plot area, "
           "thin axis lines, sans-serif labels.")

    h(doc, "4.4. Kompensatsiya va kvantlash granulyarligi bog'liqligi", 2)
    table(doc, "8-jadval. Kompensatsiya vazn diapazonini kengaytiradi (fc2, 2-qatlam).",
          ["Holat", "max |w|", "Satr normasi (mediana)", "Satr normasi (maks)", "Tarqoqlik"],
          [["asl", "0.2472", "0.0786", "0.7569", "9.6x"],
           ["kompensatsiyadan keyin", "11.3875", "0.4083", "76.9402", "188.4x"]],
          bad_rows=(1,))
    table(doc, "9-jadval. Kvantlash granulyarligi uchun oqibat (Whisper enkoderi).",
          ["Variant", "Hajm (MiB)", "Siqish", "E_glob", "Natija"],
          [["qisqartirilgan, INT8 per-tensor", "266", "4.33x", "0.7420", "model buzildi"],
           ["qisqartirilgan, INT8 per-channel", "267", "4.32x", "0.2226", "saqlandi"]],
          bad_rows=(0,), good_rows=(1,))
    para(doc,
         "gamma_j W[:, j] ni vakil ustunga qo'shish massani to'playdi, shuning uchun "
         "bitta tenzor-keng masshtab chetlanishlarni qoplashi kerak bo'ladi va qolgan "
         "hamma joyda aniqlikni yo'qotadi. Bu — mustaqil loyihaviy tanlov emas, "
         "usulning ikki komponenti orasidagi bog'liqlik.")

    h(doc, "4.5. Past-rank konstruksiyalarining taqqoslamasi", 2)
    table(doc, "10-jadval. Teng parametr byudjetida chiqish xatosi (held-out, 135 o'lchov).",
          ["Siqish", "oddiy SVD", "faollashuvga sezgir SVD", "funksional CUR", "leverage CUR"],
          [["2.00x", "0.3700", "0.2379", "0.5806", "0.6978"],
           ["3.81x (kesh)", "0.5534", "0.3689", "0.7180", "0.8223"],
           ["8.00x", "0.7025", "0.4730", "0.8227", "0.8963"]],
          good_rows=(1,))
    table(doc, "11-jadval. Xuddi shu byudjetlarda vazn xatosi — Ekart-Yang tasdig'i.",
          ["Siqish", "oddiy SVD", "faollashuvga sezgir SVD", "funksional CUR", "leverage CUR"],
          [["2.00x", "0.4019", "0.4926", "0.7790", "0.7873"],
           ["3.81x", "0.6010", "0.7314", "0.8861", "0.9085"],
           ["8.00x", "0.7583", "0.8516", "0.9540", "0.9749"]])
    table(doc, "12-jadval. 135 ta operator darajasidagi taqqoslash natijalari.",
          ["Taqqoslash", "G'alabalar", "Talqin"],
          [["funksional CUR > leverage CUR", "134 / 135",
            "kalibrlashga asoslangan ustun tartibi ishlaydi"],
           ["sezgir SVD > oddiy SVD", "135 / 135",
            "chiqish-optimallik vazn-optimallikni yutadi"],
           ["funksional CUR > sezgir SVD", "0 / 135",
            "CUR yig'ilishi raqobatbardosh emas"]],
          good_rows=(0, 1), bad_rows=(2,))
    para(doc,
         "9- va 11-jadvallar birgalikda markaziy tamoyilni ajratadi: faollashuvga "
         "sezgir yechim vazn xatosida YOMONROQ, chiqish xatosida esa YAXSHIROQ — va "
         "bu har bir o'lchovda takrorlanadi. Ekart-Yang buzilmaydi, u shunchaki bu "
         "maqsad uchun noto'g'ri funksionalni optimallashtiradi. CUR yig'ilishi esa "
         "qo'shimcha r^2 blokini olib yuradi (r(m+n) + r^2 va r(m+n)) va teng "
         "byudjetda pastroq rank bilan cheklanadi; u bir xilda yutqazadi.")
    para(doc,
         "Bu yerda rad etilayotgan narsa CUR YIG'ILISHI, ya'ni C U R "
         "faktorizatsiyasi. Uni ushbu ishning strukturaviy bosqichi bilan "
         "aralashtirmaslik kerak: strukturaviy bosqich faktorizatsiya "
         "qurmaydi, u operatorning o'zini kichraytiradi va qo'shimcha blok "
         "keltirmaydi (4.17-bo'lim). CUR adabiyotidan olingan ustun TANLASH "
         "tamoyili esa saqlanadi va 12-jadvalning birinchi qatorida "
         "tasdiqlanadi.", italic=True, size=10)

    h(doc, "4.6. Kalibrlash hajmiga talab", 2)
    table(doc, "13-jadval. Qator/rank nisbati va ortiqcha moslashuv (enkoder fc1, rank 409).",
          ["Moslash qatorlari", "Qator/rank", "Moslash xatosi", "Held-out xatosi", "Bo'shliq"],
          [["256", "0.6", "0.00000", "0.04355", "1 540 784x"],
           ["512", "1.3", "0.00035", "0.04624", "131x"],
           ["2048", "5.0", "0.00637", "0.02835", "4.4x"],
           ["4096", "10.0", "0.01151", "0.02199", "1.9x"],
           ["8192", "20.0", "0.01364", "0.01900", "1.4x"]],
          bad_rows=(0, 1), good_rows=(4,))
    para(doc,
         "Moslash xatosining aynan nolga teng bo'lishi usul sifatini emas, kalibrlash "
         "to'plamining yodlab olinganini bildiradi. Biz qator/rank nisbatini kamida "
         "10, ma'lumot imkon bergan joyda esa 20 qilib olishni tavsiya qilamiz.")

    h(doc, "4.7. Xatoning tarqalishi", 2)
    para(doc,
         "Operatorlarni birma-bir buzib, tarmoq chiqishidagi xatoni o'lchash "
         "per-operator ta'sir koeffitsientlarini beradi: c_i = E_glob(faqat i) / "
         "E_loc(i). 48 ta enkoder operatori bo'ylab E_loc 160 barobar (0.0014 dan "
         "0.225 gacha), E_glob esa atigi 4 barobar (0.012 dan 0.047 gacha) o'zgaradi. "
         "Koeffitsientlar kengaytiruvchi proyeksiya uchun 0.58-5.12, toraytiruvchi "
         "uchun 0.13-0.68. Assimetriya residual oqimdan kelib chiqadi: y = x + f(x) "
         "uchun f dagi nisbiy xato ||f|| / ||x + f|| koeffitsienti bilan suyultiriladi, "
         "kengaytiruvchi proyeksiyadagi xato esa unga yetib borgunicha nochiziqlik "
         "orqali kuchayadi.")

    figure(doc, 5,
           "O'lchangan xato yutilishi. Lokal operator xatosi ikki tartibga o'zgaradi, "
           "tarmoq chiqish xatosi esa tor doirada qoladi.",
           "A scientific scatter plot, white background, publication style. X-axis "
           "'per-operator error E_loc' on a logarithmic scale from 0.001 to 0.3, "
           "y-axis 'network output error E_glob' on a linear scale from 0 to 0.06. "
           "48 points using two distinct marker shapes: circles labeled 'expanding "
           "projection (fc1)' clustered at lower E_loc, triangles labeled 'contracting "
           "projection (fc2)' clustered at higher E_loc. Despite the 160x horizontal "
           "spread, all points lie within a narrow horizontal band. A shaded "
           "horizontal band marks the 4x E_glob range. An inset diagram in a corner "
           "shows a residual connection y = x + f(x) with the annotation 'relative "
           "error diluted by ||f||/||x+f||'. Grayscale, thin axis lines.")

    h(doc, "4.8. Rank taqsimoti", 2)
    table(doc, "14-jadval. Teng byudjetda bir xil va byudjet-optimal taqsimot "
               "(Whisper enkoderi, ikkala artefakt ham 203 MB).",
          ["Sxema", "Parametrlar", "Yig'indi xato", "WER (TEST, 300)",
           "WER (val, 80)"],
          [["bir xil rank", "100 515 840", "3.2682", "0.3513", "0.1719"],
           ["sezgirlikka asoslangan", "100 505 600", "2.8495", "0.3056", "0.0729"]],
          good_rows=(1,))
    para(doc,
         "Mustaqil test splitida farq dWER = -0.0457, 95% juftlik bootstrap "
         "oralig'i [-0.0885, -0.0138] — STATISTIK JIHATDAN AHAMIYATLI. Bu "
         "ishdagi kam sonli ahamiyatli yaxshilanishlardan biri va u eng kuchli "
         "protokolda olingan. Validation to'plamida effekt kattaroq ko'ringan "
         "edi (0.1719 -> 0.0729), ammo o'sha o'lchovning oraliqlari ancha keng "
         "bo'lgan; test splitida effekt kichrayadi, lekin ahamiyatliligi "
         "o'rnatiladi. Yig'indi maqsad esa atigi 12.8% ga yaxshilanadi — "
         "maqsad va WER o'rtasidagi bu nomutanosiblik 4.7-bo'limdagi "
         "nochiziqli tarqalish bilan izohlanadi.")
    para(doc,
         "Ikkala variant ham FP32 dan (0.1793) ancha yomon ekanini alohida "
         "ta'kidlaymiz: bu yerda past-rank enkoderning barcha 48 ta FFN "
         "operatoriga majburan qo'llanadi, ya'ni kaskadning 3-holati eng keng "
         "ko'lamda ishlatilgan. Jadvalning maqsadi past-rank shoxchasini FP32 "
         "ga qarshi oqlash emas — u TENG BYUDJETDA rankni QANDAY taqsimlash "
         "kerakligini ko'rsatadi. Shoxchaning o'zi qachon o'rinli ekanini "
         "4.1-bo'limdagi holat mezoni belgilaydi.")
    para(doc,
         "Qo'shimcha ravishda taqsimotni bevosita o'lchangan global zarar egri "
         "chiziqlariga qarshi sinadik (144 ta tarmoq yurishi); u bir xil WER va "
         "biroz yomonroq CER berdi, ya'ni arzon lokal proksi operatorlarni "
         "SARALASH uchun yetarli, garchi uchdan-uchgacha zararni BASHORAT "
         "QILISH uchun yaroqsiz bo'lsa ham.")

    h(doc, "4.9. Uchdan-uchgacha sifat", 2)
    table(doc, "15-jadval. Enkoder variantlari, Common Voice uz TEST splitining "
               "300 namunasi, 95% juftlik bootstrap IO. Kvantlagich — "
               "yaxlitlash (round-to-nearest), granulyarlik ustunda "
               "ko'rsatilgan.",
          ["Variant", "WER", "95% IO", "CER", "dWER (FP32 ga)", "Ahamiyatlilik"],
          [["FP32", "0.1793", "[0.1506, 0.2091]", "0.0522", "—", "—"],
           ["INT8 per-tensor (majburiy bosqich)", "0.1815", "[0.1522, 0.2122]",
            "0.0549", "+0.0023 [-0.0065, +0.0117]", "sezilarli emas"],
           ["strukturaviy qisqartirish + INT8 per-channel", "0.1943",
            "[0.1627, 0.2260]", "0.0589", "+0.0150 [+0.0013, +0.0303]", "sezilarli"],
           ["INT8 + taqsimlangan past-rank", "0.3056", "[0.2691, 0.3409]", "0.1018",
            "+0.1263 [+0.0999, +0.1539]", "sezilarli"]],
          bad_rows=(2, 3))
    para(doc,
         "Uchinchi qator kutilmagan: strukturaviy qisqartirish yaxlitlash bilan "
         "birga qo'llanganda FP32 dan SEZILARLI yomonlashuv beradi, kvantlashning "
         "o'zi esa bermaydi. Bu 4.4-bo'limdagi bog'liqlikning davomi. U yerda "
         "kompensatsiya satrlar bo'yicha diapazonni 9.6x dan 188.4x ga "
         "kengaytirgani va per-tensor yaxlitlashni butunlay yiqitgani "
         "ko'rsatilgan edi; per-channel granulyarlik modelni saqlaydi, ammo "
         "yuqoridagi raqamlar shuni qo'shadiki, u qoldiq zararni to'liq yo'q "
         "qilmaydi. Savol shunda: bu zarar strukturaviy o'qning o'z narximi, "
         "yoki yaxlitlashning cheklovimi? Javobni 18-jadvaldagi to'liq "
         "taqqoslash beradi.")
    para(doc,
         "Baholash to'plami haqida. Dastlab bu taqqoslash validation splitining "
         "80 ta namunasida o'tkazilgan edi va teskari tartib bergan (qisqartirilgan "
         "variant FP32 dan farqlanmaydi, kvantlashning o'zi esa sezilarli). "
         "Kalibrlash — funksional guruhlash qarori chiqariladigan ma'lumot — ham "
         "o'sha splitdan olinadi, shuning uchun namunalar kesishmasa ham taqsimot "
         "bir xil edi va kalibrlashga tayanadigan bosqich sun'iy ravishda ustun "
         "ko'ringan. Yuqoridagi raqamlar mustaqil test splitida va 3.75 barobar "
         "kattaroq to'plamda qayta o'lchandi. Bu 4.6-bo'limdagi kalibrlash hajmi "
         "talabini to'ldiradi: kalibrlashga asoslangan usullar uchun baholash "
         "to'plami faqat kesishmasligi emas, BOSHQA taqsimotdan bo'lishi kerak.",
         italic=True, size=10)
    para(doc,
         "Dekoder o'qi kaskadning qaroriga bevosita sinov bo'ladi. 4.1-bo'limda "
         "dekoder uchun chiqarilgan talab 3.81x, INT8 esa 4.00x beradi, shuning "
         "uchun kaskad 'kvantla va to'xta' deydi. Xuddi shu 300 namunada: FP32 "
         "0.1761, INT8 0.1793 (dWER +0.0032, IO [-0.0039, +0.0100]), per-channel "
         "INT8 0.1764 (+0.0003, [-0.0063, +0.0063]) — ikkalasi ham FP32 dan "
         "farqlanmaydi. Ya'ni kaskad qo'shimcha strukturaviy o'zgarishni rad "
         "etganda haq edi: u yerda olinadigan foyda yo'q, chunki kvantlashning "
         "o'zi maqsadga yetadi va sifatni buzmaydi.")

    h(doc, "4.9a. Nashr etilgan post-training usullari bilan taqqoslash", 2)
    para(doc,
         "Yuqoridagi taqqoslashlar bizning kvantlash sxemamizni kutubxona standarti "
         "(min/max yaxlitlash) bilan solishtirdi. Bu 2018-yilgi baza, shuning uchun "
         "xuddi shu sinfdagi — vazn-only, qayta o'qitishsiz — ikkita nashr etilgan "
         "usul qo'shildi: GPTQ [1] va AWQ [2]. Adolatlilik uchun barcha to'rt usul "
         "bir xil sharoitda ishlaydi: simmetrik INT8, per-output-channel masshtab, "
         "bir xil kalibrlash bo'linmasi, xato held-out qatorlarda. Ya'ni "
         "taqqoslanayotgan narsa algoritm, bit kengligi yoki granulyarlik emas.")
    para(doc,
         "MUHIM ESLATMA: auto-gptq va autoawq paketlari CUDA yadrolarini talab "
         "qiladi va CPU-only mashinada ishlamaydi, shuning uchun ikkala algoritm "
         "ham maqolalaridagi tavsif asosida qayta amalga oshirildi "
         "(nnopt/quantizer/baselines.py). Ular o'z invariantlari bo'yicha "
         "tekshirilgan (11 test): GPTQ anizotrop faollashuvlarda RTN ni yutadi va "
         "blok hajmiga sezgir emas, AWQ kuchliroq anizotropiyada kattaroq alpha "
         "tanlaydi. Natijalar rasmiy kod bilan aynan mos kelmasligi mumkin, chunki "
         "asl ishlar per-group asimmetrik kvantlash va boshqa bit kengliklaridan "
         "foydalanadi.", italic=True)
    table(doc, "16a-jadval. Operator darajasidagi taqqoslash (held-out E_loc, "
               "INT8 per-channel, har biri 30 operator).",
          ["Usul", "Encoder o'rtacha", "RTN ga nisbatan", "Decoder o'rtacha",
           "RTN ga nisbatan"],
          [["RTN (kalibrlashsiz)", "0.00873", "—", "0.00926", "—"],
           ["GPTQ (qayta amalga oshirilgan)", "0.00399", "-54.3%", "0.00662", "-28.5%"],
           ["AWQ (qayta amalga oshirilgan)", "0.00730", "-16.4%", "0.00788", "-14.9%"],
           ["bizning kalibrlangan masshtab", "0.00761", "-12.9%", "0.00834", "-9.8%"]],
          good_rows=(1,))
    table(doc, "16b-jadval. Juftlik g'alabalar (60 operator).",
          ["Taqqoslash", "Natija"],
          [["bizniki > RTN", "60 / 60"],
           ["bizniki > AWQ", "34 / 60"],
           ["bizniki > GPTQ", "13 / 60"]],
          bad_rows=(2,))
    callout_txt = (
        "Salbiy natija va uning oqibati: GPTQ ning Hessian orqali xato "
        "kompensatsiyasi bizning kalibrlangan masshtabimizdan sezilarli darajada "
        "ustun (encoderda 54.3% va 12.9% xato kamayishi, 60 operatordan 47 tasida "
        "g'alaba). Sabab tuzilishda: GPTQ bir ustunni kvantlagach kelib chiqqan "
        "xatoni qolgan ustunlarga tarqatadi, bizning usul esa har masshtabni "
        "mustaqil optimallashtiradi. Shuning uchun kaskadning majburiy kvantlash "
        "bosqichida GPTQ tavsiya etiladi. Bu ishning qolgan hissalariga ta'sir "
        "qilmaydi, chunki GPTQ faqat kvantlash usuli: u kanal olib tashlamaydi, "
        "kesh byudjetini hisobga olmaydi va rank taqsimlamaydi.")
    para(doc, callout_txt, italic=True, size=10, color=WARN)

    para(doc, "GPTQ ning sharti: kalibrlashning yetarliligi.", bold=True,
         size=10)
    para(doc,
         "Yuqoridagi tavsiya bir shartga bog'liq bo'lib chiqdi, va shartni "
         "topish uchun avval bir farazni tekshirish kerak bo'ldi. GPTQ ning "
         "WikiText-2 dagi INT4 natijasi tushunarsiz edi — u 8.646 beradi, "
         "oddiy yaxlitlashning 8.583 idan YOMONROQ (48-jadval) — holbuki "
         "operator darajasida u eng aniqi. Bizning o'z usulimizda xuddi "
         "shunday manzara SILJISHDAN kelib chiqqandi (gain 0.9896, chuqurlik "
         "bo'ylab 0.44), shuning uchun tabiiy faraz GPTQ da ham gain 1 dan "
         "past degan edi.")
    para(doc,
         "Faraz O'LCHANDI VA RAD ETILDI: GPTQ ning gaini 1.0000, ya'ni u "
         "hech qanday ko'paytiruvchi susayish keltirmaydi va chuqurlik "
         "bo'ylab to'planadigan narsa yo'q (78 operatorga gain^78 = 0.999). "
         "Chiqish domenidagi qayta masshtab unga foydasiz, hatto ozgina "
         "zararli. Additiv siljish ham yo'q: kvantlash xatosining o'rtacha "
         "faollashuvga proyeksiyasi GPTQ da 0.0002, RTN da 0.0836, bizda "
         "0.0554 — ya'ni uning xato kompensatsiyasi o'rtachani ham o'zi "
         "yo'q qiladi.")
    para(doc,
         "Sabab boshqa joyda bo'lib chiqdi. GPTQ Hessianni kalibrlash "
         "satrlaridan quradi va xatoni o'sha satrlarda minimallashtiradi, "
         "ya'ni uning e'lon qilingan xatosi MOSLASH statistikasi. Xuddi shu "
         "operatorlarni HELD-OUT faollashuvlarda o'lchash manzarani "
         "o'zgartiradi:")
    table(doc, "17-jadval. Kalibrlash hajmining ta'siri (open_llama_3b FFN, "
               "INT8, o'rtacha nisbiy chiqish xatosi).",
          ["Usul", "Moslash", "Held-out (4096 satr)", "Held-out (16384 satr)",
           "Siljish"],
          [["RTN", "0.00617", "0.00619", "0.00626", "0.0042"],
           ["bizniki", "0.00572", "0.00585", "0.00591", "0.0025"],
           ["GPTQ", "0.00330", "0.00640", "0.00558", "0.0002"]],
          bad_rows=(2,))
    para(doc,
         "4096 satrda GPTQ moslashda ikki barobar aniq (0.0023 va 0.0062) va "
         "HELD-OUT DA UCHALASINING ENG YOMONI (0.00640). Kalibrlash to'rt "
         "barobar oshirilganda u eng yaxshiga aylanadi (0.00558), moslash va "
         "held-out orasidagi farq esa 2.78x dan 1.69x ga tushadi. RTN va "
         "bizning usul deyarli qimirlamaydi — birinchisi kalibrlashdan "
         "umuman foydalanmaydi, ikkinchisi esa faqat kanal masshtabini "
         "sozlaydi. Ya'ni GPTQ ning ustunligi haqiqiy, lekin u KALIBRLASH "
         "YETARLI BO'LGANDA namoyon bo'ladi; yetarli bo'lmaganda u "
         "Hessianga ortiqcha moslashadi.")
    para(doc,
         "Bu 4.9a-bo'limdagi tavsiyani bekor qilmaydi, chunki o'sha "
         "taqqoslash Whisper da va HELD-OUT satrlarda o'tkazilgan. Ammo u "
         "da'voning QAMROVINI aniqlaydi: GPTQ ni tanlash kalibrlashning "
         "operator kengligiga nisbatan yetarliligini talab qiladi, va bu "
         "shart tekshirilishi kerak. Shuningdek u 48-jadvaldagi INT4 "
         "natijasini tushuntiradi — u yerda kalibrlash yanada tor "
         "bo'lgan.", italic=True, size=10)
    para(doc,
         "Bitta bashorat esa tasdiqlanmadi. Hessian n x n bo'lgani uchun "
         "kengligi eng katta operator (down_proj, n = 8640) 4096 satrda rang "
         "yetishmasligiga uchraydi va eng ko'p yutishi kutilgandi. U 19% "
         "yaxshilandi, gate_proj (n = 3200) esa 16% — deyarli bir xil. "
         "Demak sabab kalibrlash hajmi UMUMAN, aynan rang yetishmasligi "
         "emas.", size=9.5)
    para(doc, "Kvantlash uchun bias tuzatish.", bold=True, size=10)
    para(doc,
         "Jadvaldagi oxirgi ustun mustaqil bir imkoniyatni ochadi. Kvantlash "
         "xatosi nolga teng o'rtachaga ega bo'lishi shart emas, va uning "
         "doimiy qismi chiqishda qat'iy siljish beradi. Uni operatorning "
         "MAVJUD bias vektoriga qo'shish mumkin:")
    eq(doc, "b <- b + (W - W_kvant) mean(X) ,", 16)
    para(doc,
         "bu strukturaviy bosqichda ishlatiladigan ayniyatning aynan o'zi, "
         "faqat boshqa xato manbasiga qo'llangan. Hech bir butun sonli kod, "
         "bit kengligi yoki xotira formati o'zgarmaydi. Held-out xatosi "
         "bo'yicha foyda: INT4 da RTN uchun 4.8%, bizning usul uchun 6.0%; "
         "INT8 da mos ravishda 3.8% va 1.5%. GPTQ uchun esa 0.0%, chunki "
         "tuzatadigan siljish yo'q. Ya'ni bu masshtabga asoslangan "
         "kvantlagichlarga tekin qo'shimcha, GPTQ ga esa keraksiz.")
    para(doc,
         "Tuzatish uchdan-uchgacha ham o'lchandi. RTN per-channel "
         "enkoderning 144 operatoridan bias tashuvchi 120 tasiga (16) "
         "qo'llanib, butun sonli kodlar va formatga tegilmasdan TEST "
         "splitining 300 namunasida baholandi: WER 0.1858 dan 0.1798 ga "
         "tushadi — nuqtaviy bahoda RTN ning FP32 gacha bo'lgan "
         "farqining deyarli hammasi qaytadi (FP32 = 0.1793) — ammo "
         "juftlik farqi -0.0060 [-0.0142, +0.0023] nolni qamraydi, "
         "ya'ni ahamiyatlilik bu namunada O'RNATILMAYDI. Halol bayon: "
         "yo'nalish ijobiy va bepul, tasdiq esa kattaroq namuna talab "
         "qiladi; kaskadning GPTQ yo'lida esa tuzatish keraksizligicha "
         "qoladi.", italic=True, size=10)
    para(doc, "Bias tuzatishning qo'llanish sohasi: argmax va ehtimollik "
              "metrikalari.", bold=True, size=10)
    para(doc,
         "To'liq o'z-stack sinovi (kalibrlangan masshtab + chiqish-domen "
         "+ bias tuzatish, INT4) esa kutilmagan va MUHIM natija berdi: "
         "operator xatosini 6% ga kamaytiradigan tuzatish WikiText-2 "
         "perplexity ni 8.246 dan 8.268 ga YOMONLASHTIRDI. Bu ikki xil "
         "amalni ajratishga majbur qiladi, va ajratish ABLATSIYA bilan "
         "o'lchandi. STRUKTURAVIY buklash — olib tashlangan "
         "deyarli-doimiy kanallar o'rtachasini biasga yig'ish — tuzatish "
         "emas, SIGNAL: mBERT da bir xil kanallar bilan buklashni olib "
         "tashlash aniqlikni ham (-0.0386 [-0.0560, -0.0202], "
         "sezilarli), pseudo-perplexity ni ham (119 dan 230 ga) keskin "
         "buzadi — u ikkala metrika uchun ham usulning zaruriy qismi. "
         "KVANTLASH-QOLDIQ tuzatishi esa — (16) dagi kichik siljish — "
         "metrika turiga qarab ajraladi: argmax metrikasida foyda "
         "yo'nalishi (Whisper WER 0.1858 -> 0.1798), ehtimollik "
         "metrikasida zarar (Llama INT4 PPL +0.022). Amaliy qoida: "
         "strukturaviy buklash HAR DOIM qo'llanadi; kvantlash-qoldiq "
         "tuzatishi qaror-metrikali vazifalarga (ASR, klassifikatsiya) "
         "tavsiya etiladi, ehtimollik-metrikali vazifalarda (til "
         "modellash) qo'llanmaydi.", italic=True, size=10)

    h(doc, "4.9b. Strukturaviy qisqartirish kvantlash ustiga tekinga "
           "qo'shiladimi?", 2)
    para(doc,
         "15-jadval ochiq savol qoldirdi: qisqartirish yaxlitlash bilan birga "
         "sifatni yomonlashtirdi. Bu strukturaviy o'qning o'z narximi, yoki "
         "yaxlitlashning kompensatsiyadan keyingi qoldiqni ko'tara olmasligimi? "
         "Savolni ajratish uchun to'liq 2x2 taqqoslash o'tkazildi: qisqartirish "
         "qarori, kalibrlash ma'lumoti, eksport yo'li va baholash to'plami "
         "o'zgarmaydi, faqat kvantlagich almashadi. Granulyarlik hamma joyda "
         "per-output-channel, chunki 4.4-bo'limda per-tensor kompensatsiyadan "
         "keyin butunlay yiqilishi ko'rsatilgan.")
    table(doc, "18-jadval. Kvantlagich x strukturaviy qisqartirish, to'liq 2x2 "
               "(Whisper encoder, Common Voice uz TEST splitining 300 namunasi; "
               "kalibrlash VALIDATION splitidan, ya'ni boshqa taqsimotdan).",
          ["Variant", "Hajm (MiB)", "Latency (ms)", "WER", "dWER (FP32 ga)", "CER"],
          [["FP32", "1152", "11246.1", "0.1793", "—", "0.0522"],
           ["C: yaxlitlash (RTN) yolg'iz", "300", "—", "0.1858",
            "+0.0066 [-0.0024, +0.0158]", "0.0560"],
           ["D: qisqartirish + yaxlitlash", "267", "—", "0.1943",
            "+0.0150 [+0.0013, +0.0303]", "0.0589"],
           ["A: GPTQ yolg'iz", "300", "6944.9", "0.1847",
            "+0.0054 [-0.0021, +0.0138]", "0.0538"],
           ["B: qisqartirish + GPTQ", "267", "6740.1", "0.1833",
            "+0.0040 [-0.0056, +0.0147]", "0.0549"]],
          good_rows=(4,), bad_rows=(2,))
    para(doc,
         "Qisqartirishning narxi kvantlagichga bog'liq bo'lib chiqdi. GPTQ bilan "
         "u B - A = -0.0014, 95% IO [-0.0111, +0.0096] — ajratilmaydi va nuqtaviy "
         "baho hatto qisqartirilgan variant foydasiga; yaxlitlash bilan esa "
         "D - C = +0.0084, [-0.0038, +0.0223]. Yakuniy natijada farq aniq "
         "ko'rinadi: B FP32 dan farqlanmaydi, D esa SEZILARLI yomon. Muhimi, "
         "qisqartirishsiz ikkala kvantlagich deyarli teng (A - C = -0.0011, "
         "[-0.0116, +0.0090]), ya'ni GPTQ ning afzalligi o'z-o'zidan emas, "
         "AYNAN qisqartirish qo'llanganda namoyon bo'ladi.")
    para(doc,
         "Mexanizm 4.4-bo'limda o'lchangan. Kompensatsiya massani vakil "
         "ustunlarga to'playdi va satrlar bo'yicha diapazonni 9.6x dan 188.4x ga "
         "kengaytiradi. Per-tensor yaxlitlash buni umuman ko'tara olmaydi; "
         "per-channel granulyarlik modelni saqlaydi, ammo har bir satrda kattaroq "
         "kvantlash xatosi qoldiradi. GPTQ o'sha qoldiqni Hessian orqali hali "
         "kvantlanmagan ustunlarga tarqatadi, ya'ni kompensatsiya keltirgan "
         "zararni to'g'ridan-to'g'ri yutadi. Shu sababli ikki o'qning "
         "ortogonalligi shartsiz emas: u XATONI KOMPENSATSIYA QILADIGAN "
         "kvantlagich bilan amal qiladi, oddiy yaxlitlash bilan esa yo'q. Bu "
         "kaskadning kvantlash bosqichi uchun aniq tavsiya beradi va 4.9a dagi "
         "vosita tanlovini ikkinchi, mustaqil sabab bilan asoslaydi.")
    para(doc, "Bosqichlar TARTIBINING bevosita A/B sinovi.", bold=True,
         size=10)
    para(doc,
         "Kaskad kesishni kvantlashdan OLDIN qo'yadi va buni uch tayanch "
         "asoslaydi. Birinchisi strukturaviy: kompensatsiya butun sonli "
         "kodlarga nisbatan yopiq emas — kvantlangan qiymatlarning "
         "chiziqli kombinatsiyasi kvantlangan bo'lmaydi, ya'ni teskari "
         "tartibda joylashtiriladigan artefakt baribir QAYTA kvantlashni "
         "talab qiladi va ikkinchi kvantlash birinchisi ko'rmagan "
         "vaznlarga tushadi. Ikkinchisi o'lchangan zanjir (188x, "
         "4.4-bo'lim). Uchinchisi — bevosita A/B: bir xil kanallar va "
         "gammalar bilan ikki tartib qurilib solishtirildi. mBERT da — "
         "kollinearlik deyarli nol, gammalar massa tashimaydi — farq "
         "yo'q (+0.0009 [-0.0092, +0.0119]); Whisper da — 188x rejimi — "
         "teskari tartib +0.0046 [-0.0052, +0.0146] yomonroq. Yo'nalish "
         "mexanizmga mos va effekt kompensatsiya massasi bilan "
         "tartiblangan (5 barobar farq), ammo n = 300 da "
         "sertifikatlanmaydi; shuning uchun tartib-da'vosi birinchi "
         "ikki tayanchda turadi, A/B esa ularga zid kelmaydigan, "
         "yo'nalishi mos qo'shimcha dalil sifatida keltiriladi.",
         italic=True, size=10)
    para(doc,
         "Statistik ehtiyotkorlik. O'zaro ta'sirning o'zi — qisqartirilgan "
         "holatda GPTQ va yaxlitlash farqi (-0.0110, 95% IO [-0.0277, +0.0063]) — "
         "n = 300 da statistik jihatdan tasdiqlanmagan. Tasdiqlangani "
         "quyidagicha: qisqartirish + yaxlitlash FP32 dan sezilarli yomon, "
         "qisqartirish + GPTQ esa emas. Nuqtaviy baholar o'zaro ta'sir "
         "yo'nalishini ko'rsatadi (qisqartirish yo'q bo'lganda farq -0.0011, bor "
         "bo'lganda -0.0110, ya'ni 10 barobar katta), lekin uni qat'iy o'rnatish "
         "uchun kattaroq baholash to'plami kerak.", italic=True, size=10,
         color=WARN)

    h(doc, "4.9c. Butun model: kaskad qarori bir xil siyosatlarga qarshi", 2)
    para(doc,
         "Yuqoridagi barcha o'lchovlar bitta o'qni o'zgartirib, ikkinchisini "
         "qat'iy ushlab turdi. Bu operator o'zgarishining xavfsizligini "
         "ko'rsatadi, ammo kaskadning o'zi arziydimi degan savolga javob "
         "bermaydi: uning mahsuloti — har operator uchun QAROR, va bu modelda "
         "u ikkita qarama-qarshi qaror chiqaradi (dekoderda 'to'xta', enkoder "
         "FFN sida 'davom et'). Qaror faqat IKKALA yo'nalishda ham uni bekor "
         "qilish zarar keltirsa qimmatlidir, shuning uchun uni e'tiborsiz "
         "qoldiradigan ikki bir xil siyosat bilan taqqoslaymiz.")
    table(doc, "19-jadval. Butun model (enkoder + dekoder), Common Voice uz "
               "TEST splitining 300 namunasi.",
          ["Variant", "Hajm (MiB)", "Siqish", "WER", "dWER (FP32 ga)", "CER"],
          [["A: FP32", "2915", "1.00x", "0.1761", "—", "0.0504"],
           ["B: bir xil yumshoq (hamma joyda INT8)", "738", "3.95x", "0.1847",
            "+0.0086 [-0.0004, +0.0186]", "0.0538"],
           ["C: kaskad", "705", "4.14x", "0.1833",
            "+0.0072 [-0.0028, +0.0187]", "0.0549"],
           ["D: bir xil agressiv (hamma joyda past-rank)", "546", "5.34x",
            "0.6101", "+0.4340 [+0.3379, +0.5607]", "0.2219"]],
          good_rows=(2,), bad_rows=(3,))
    para(doc,
         "Ikki tomon keskin assimetrik. Kaskadni YUMSHOQ tomonga bekor qilish "
         "arzon: B ga nisbatan C 33 MiB (4.5%) tejaydi va aniqlik "
         "o'zgarmaydi (dWER = -0.0014, 95% IO [-0.0111, +0.0101]). AGRESSIV "
         "tomonga bekor qilish esa halokatli: D 159 MiB qo'shimcha tejaydi, "
         "lekin WER 0.1833 dan 0.6101 ga ko'tariladi (dWER = +0.4268, "
         "[+0.3263, +0.5618]) — model amalda ishlamay qoladi.")
    para(doc,
         "D varianti 5.34x siqish beradi. Aynan shu raqam 1-bo'limdagi "
         "tanqidni empirik tasdiqlaydi: qo'lda tanlanadigan darajalar odatda "
         "4x va 8x atrofida bo'ladi, 5.34x esa mutlaqo oqilona ko'rinadi — va "
         "u modelni yo'q qiladi. Kaskad 4.14x da to'xtaydi, chunki dekoder "
         "uchun chiqarilgan talab 3.81x va INT8 uni 4.00x bilan qondiradi. Bu "
         "qaror endi son bilan oqlanadi: uni bekor qilish 0.43 WER turadi.")
    para(doc,
         "Shu sababli usulning qiymatini to'g'ri ta'riflash muhim. Kaskad "
         "ko'proq siqmaydi — u eng sodda bazadan atigi 4.5% xotira yutadi. "
         "Uning qiymati QAYERDA TO'XTASHNI bilishida: dekoderda vaznlar har "
         "token uchun bir marta ishlatiladi, shuning uchun u yerda past-rank "
         "ning FLOP tejashi vaqtga aylanmaydi va faqat xato qo'shadi. "
         "Chiqarilgan maqsad buni oldindan aytadi, o'lchov esa tasdiqlaydi.")

    h(doc, "4.9d. Strukturaviy mezon: nashr etilgan bazalar va ablation", 2)
    para(doc,
         "Ushbu ishning o'z hissasi strukturaviy o'q bo'lgani uchun uni "
         "nashr etilgan qisqartirish usullariga qarshi qo'yish majburiy. "
         "Bazalar sifatida eng kichik ||W2[:, j]|| bo'yicha o'chirish "
         "(magnitude) va Wanda [27] ning ballash qoidasi olindi. Ular har "
         "qatlamda AYNAN biz olib tashlagan sonda kanal olib tashlaydi va bir "
         "xil GPTQ o'tishi bilan kvantlanadi, ya'ni yagona o'zgaruvchi — "
         "qaysi kanallar ketishi. Ikkalasi ham kompensatsiya qilmaydi; bu "
         "ularning ta'rifi, sun'iy cheklov emas.")
    para(doc, "Wanda ikki shaklda: nima aynan taqqoslanmoqda.", bold=True,
         size=10)
    para(doc,
         "Bu yerda aniqlik zarur, chunki Wanda [27] o'z asl shaklida "
         "STRUKTURASIZ usul: u S_ij = |W_ij| x ||X_j||_2 ballari bo'yicha "
         "alohida vaznlarni nolga aylantiradi va tanlovni har bir CHIQISH "
         "SATRI ichida o'tkazadi, tenzor shakli esa o'zgarmaydi. Yuqoridagi "
         "baza esa uning ballash qoidasini kanal darajasiga agregatsiya "
         "qiladi va butun kanallarni olib tashlaydi — bu USHBU ISHNING "
         "moslashtirishi bo'lib, asl usulga nisbat berilmaydi. Moslashtirish "
         "ataylab qilingan: strukturaviy olib tashlashni strukturasiz nolga "
         "aylantirishga qarshi qo'yish tanlash mezonini granulyarlik bilan "
         "chalkashtirgan bo'lar edi.")
    para(doc,
         "Ikkala shakl ham o'lchanadi, shuning uchun ikki omil ajraladi: "
         "MEZON (taklif etilgan ortiqchalik mezoni va Wanda ballari, "
         "granulyarlik teng) hamda GRANULYARLIK (kanal darajasidagi olib "
         "tashlash va strukturasiz nolga aylantirish, mezon teng). "
         "Strukturasiz variantning siyrakligi bizning parametr ulushimizga "
         "tenglashtirilgan, ya'ni nominal siqish bir xil.")
    table(doc, "20-jadval. Wanda ikki shaklda, ikkita byudjetda "
               "(TEST splitining 300 namunasi; FP32 = 0.1793).",
          ["Byudjet", "Arm", "WER", "MiB", "GMAC", "FP32 ga"],
          [["tau=0.99", "bizniki (ortiqchalik + kompensatsiya)", "0.1833",
            "267", "403.9", "1.022x"],
           ["(nominal 21.5%)", "Wanda mezoni, kanal (ushbu ish)", "0.1850",
            "267", "403.9", "1.032x"],
           ["", "Wanda [27] asl, strukturasiz", "0.1951", "300", "455.3",
            "1.088x"],
           ["tau=0.95", "bizniki (ortiqchalik + kompensatsiya)", "0.2006",
            "254", "384.2", "1.119x"],
           ["(nominal 23.6%)", "Wanda mezoni, kanal (ushbu ish)", "0.2202",
            "254", "384.2", "1.228x"],
           ["", "Wanda [27] asl, strukturasiz", "0.2696", "300", "455.3",
            "1.504x"]],
          good_rows=(0, 3), bad_rows=(2, 5))
    para(doc,
         "Ikki omilga ajratish natijasi quyidagicha. Mezon hissasi tau = 0.99 "
         "da +0.0017, tau = 0.95 da +0.0196; granulyarlik hissasi mos ravishda "
         "+0.0101 va +0.0495. Alohida olganda hech biri ahamiyatli emas, ammo "
         "ular TO'PLANADI: agressiv byudjetda umumiy farq +0.0690 bo'lib, 95% "
         "oralig'i [+0.0052, +0.1672] — statistik jihatdan ahamiyatli.")
    para(doc,
         "Ikkita xulosa muhim. Birinchidan, granulyarlik hissasi mezon "
         "hissasidan uch-besh barobar katta, ya'ni QANDAY olib tashlash "
         "nimani tanlashdan ustunroq. Ikkinchidan, faqat granulyarlik "
         "xotirani va arifmetikani ham kamaytiradi: strukturasiz variant "
         "ikkala byudjetda ham 300 MiB va 455.3 GMAC da qoladi, chunki nollar "
         "saqlanadi va zich yadrolarda baribir ko'paytiriladi. Bu kuzatuv "
         "yangi emas — Wanda ning o'zi shu sababdan 2:4 strukturaviy variantni "
         "taklif qiladi va [25], [26] ham shu bilan asoslanadi; bu yerda u "
         "ushbu ishning trakti va apparatida son bilan o'lchanadi.",
         italic=True, size=10)
    table(doc, "21-jadval. Strukturaviy mezonlar, ikki byudjetda "
               "(TEST splitining 300 namunasi, har bir guruhda kanal soni va "
               "kvantlagich bir xil).",
          ["Byudjet", "Mezon", "Kompensatsiya", "WER", "Bizdan farq"],
          [["tau=0.99, 267 MiB", "bizniki", "bor", "0.1833", "—"],
           ["(o'rt. 17.1%)", "magnitude", "yo'q", "0.1837",
            "+0.0004 [-0.0139, +0.0131]"],
           ["", "Wanda", "yo'q", "0.1850", "+0.0017 [-0.0101, +0.0119]"],
           ["tau=0.95, 254 MiB", "bizniki", "bor", "0.2006", "—"],
           ["(qatlamlarda 73% gacha)", "Wanda", "yo'q", "0.2202",
            "+0.0196 [-0.0151, +0.0696]"],
           ["", "bizniki, kompensatsiyasiz", "yo'q", "1.3393",
            "+1.1387 [+0.7811, +1.6176]"],
           ["", "magnitude", "yo'q", "2.7378",
            "+2.5372 [+2.1446, +2.9541]"]],
          good_rows=(0, 3), bad_rows=(5, 6))
    para(doc,
         "Mo'tadil byudjetda uchala mezon ham ajralmaydi. Buni mezonlar teng "
         "degan xulosa sifatida o'qish mumkin emas: o'sha byudjetda enkoder "
         "o'rtacha atigi 17.1% kanal yo'qotadi va 6-jadvalga ko'ra unda "
         "shunchalik ortiqchalik borki, har qanday oqilona mezon olib "
         "tashlanadigan kanallarni topadi. Mezonni sinash uchun zaxira "
         "tugagan nuqta kerak.")
    para(doc,
         "Agressiv byudjetda manzara keskin ajraladi va u ikki qismga "
         "bo'linadi. Faollashuvni hisobga oladigan ikkala yo'l ham modelni "
         "saqlaydi (bizniki 0.2006, Wanda 0.2202, o'zaro farq ajratilmaydi), "
         "faqat vaznga qaraydigan magnitude esa modelni buzadi (2.7378). "
         "Demak agressiv byudjetda hal qiluvchi omil — KALIBRLASH "
         "MA'LUMOTINING ishlatilishi. Ayni paytda taklif etilgan guruhlash "
         "Wanda dan ustun ekani KO'RSATILMADI; Wanda ancha sodda va bu yerda "
         "unga teng.")
    para(doc, "Mezonlar to'rt byudjetda: farq emas, SHAKL.", bold=True, size=10)
    para(doc,
         "Ikki byudjetdagi natijalar mezon farqi siqish darajasi bilan "
         "kengayayotgandek ko'rsatdi (tau = 0.99 da +0.0017, tau = 0.95 da "
         "+0.0196). Ikki nuqta trendni o'rnatmaydi, shuning uchun uchala "
         "mezon to'rtta byudjetda kuzatildi. Natija gipotezani RAD ETADI: "
         "farqlar ketma-ketligi +0.0017, -0.0066, +0.0196, +0.0216 bo'lib, "
         "monoton emas — tau = 0.97 da magnitude bizdan yaxshiroq — va har "
         "bir juftlik oralig'i nolni qamraydi. Ya'ni taklif etilgan mezon "
         "Wanda dan ustun ekani bu ma'lumotda o'rnatilmaydi.")
    table(doc, "22-jadval. Uchta strukturaviy mezon to'rtta teng byudjetda "
               "(TEST splitining 300 namunasi; har bir qatorda kanal soni va "
               "kvantlagich bir xil, FP32 = 0.1793).",
          ["tau", "MiB", "bizniki", "magnitude", "Wanda", "Wanda - bizniki"],
          [["0.99", "267", "0.1833", "0.1837", "0.1850",
            "+0.0017 [-0.0097, +0.0126]"],
           ["0.97", "261", "0.1916", "0.1832", "0.1851",
            "-0.0066 [-0.0213, +0.0068]"],
           ["0.95", "254", "0.2006", "2.7378", "0.2202",
            "+0.0196 [-0.0155, +0.0750]"],
           ["0.93", "248", "0.2179", "5.7990", "0.2395",
            "+0.0216 [-0.0214, +0.0935]"]],
          bad_rows=(2, 3))
    para(doc,
         "Ajratuvchi kattalik boshqa joyda: degradatsiyaning SHAKLIDA. "
         "Byudjet qadamlaridagi eng katta bir qadamli yomonlashish taklif "
         "etilgan mezonda +0.0173, Wanda da +0.0351, magnitude da esa "
         "+3.0612 ni tashkil qiladi. Kalibrlashga tayanadigan ikkala yo'l ham "
         "silliq (1.022x, 1.069x, 1.119x, 1.216x), faqat vaznga qaraydigan "
         "yo'l esa ogohlantirishsiz qulaydi.")
    para(doc, "Qulashning mexanizmi.", bold=True, size=10)
    table(doc, "23-jadval. Vazn normasi va faollik energiyasining "
               "bog'liqligi chuqurlik bo'ylab (fc2, har qatlamda 24% "
               "olib tashlash byudjeti).",
          ["Qatlam", "Spearman(||w||, ||h||)", "Mezon bilan kesishma",
           "Magnitude yo'qotgan eng katta hissa (medianaga)"],
          [["L0", "+0.879", "88.2%", "0.3x"],
           ["L2", "+0.927", "94.5%", "0.6x"],
           ["L5", "+0.804", "95.9%", "0.0x"],
           ["L8", "+0.575", "98.6%", "0.0x"],
           ["L16", "-0.431", "12.4% (tasodif 24%)", "10.3x"],
           ["L23", "-0.276", "28.4%", "27.4x"]],
          bad_rows=(4, 5))
    para(doc,
         "'Faollik-ko'rlik' nomni beradi, mexanizmni emas — mexanizm "
         "o'lchandi va u chuqurlik bilan ISHORA ALMASHINISHIDA. Sayoz "
         "qatlamlarda vazn normasi kanalning funksional hissasi "
         "(||w||В·||h||) bilan kuchli musbat bog'langan va magnitude "
         "mezon bilan deyarli bir xil kanallarni tanlaydi. Chuqur "
         "qatlamlarda tarmoq kichik normani katta faollik bilan "
         "muvozanatlaydi, korrelyatsiya manfiyga o'tadi va magnitude "
         "aynan ENG KATTA hissali kanallarni olib tashlaydi — L23 da "
         "medianadan 27 barobar katta hissa, mezon esa hech qachon 0.6x "
         "dan yuqorisini olmaydi. Agressiv byudjet aynan chuqur "
         "qatlamlarga yetgani uchun qulash aynan tau <= 0.95 da va "
         "ogohlantirishsiz yuz beradi. Dastlabki gipoteza "
         "('korrelyatsiya shunchaki zaif') o'lchov bilan "
         "ANIQLASHTIRILDI: bog'liqlik zaif emas, chuqurlikda teskari.")
    para(doc,
         "Bu yerdagi magnitude raqamlari BIZNING qatlam taqsimotimiz bilan "
         "olingan, ya'ni bazaga u tanlamagan taqsimot berilgan. Uning o'z "
         "global tartiblashi bilan qulash ancha yumshoqroq — 254 MiB da "
         "2.7378 emas, 0.4192 (29-jadval). Sifat xulosasi o'zgarmaydi, "
         "ammo bu qatordagi kattaliklarni magnitude usulining o'z xossasi "
         "sifatida o'qish mumkin emas.", italic=True, size=10)
    para(doc,
         "Bu 4.9e-bo'limdagi aniqlik-byudjetli tanlash uchun zaruriy shart. "
         "Byudjetni boshqarish mumkin bo'lishi uchun egri chiziq monoton va "
         "silliq bo'lishi kerak: magnitude bilan bir qadam byudjetni "
         "kichraytirish modelni ish holatidan chiqarib yuboradi, ya'ni hech "
         "qanday tolerans amalda ushlab turilmaydi. Shu sababli usulning "
         "amaliy qiymati eng past xatolikda emas, EGRI CHIZIQNING SHAKLIDA: "
         "ish nuqtasini tanlash uchun keyingi qadam qayerga tushishini "
         "bilish kerak.", italic=True, size=10)

    para(doc, "FLAP: mezon va kompensatsiya MOS bo'lishi kerak.", bold=True,
         size=10)
    para(doc,
         "Yuqoridagi bazalar bizning ikkala tarkibiy qismimizdan birortasidan "
         "mahrum edi. FLAP [28] esa ikkalasiga ham ega — kalibrlashdan "
         "foydalanadi va olib tashlangan kanalni kompensatsiya qiladi — "
         "shuning uchun u mexanizm haqidagi da'voni sinaydigan yagona baza. "
         "Farq kompensatsiya USULIDA, va har bir usulning mezoni o'z "
         "kompensatsiyasidan kelib chiqadi: FLAP olib tashlangan kanalni "
         "uning o'rtacha hissasi bilan almashtirib, uni chiqish biasiga "
         "qo'shadi, ya'ni faqat DOIMIY qism saqlanadi — shuning uchun u "
         "fluktuatsiyasi kichik kanallarni qidiradi. Bizda esa kanal vakilga "
         "qo'shiladi va O'ZGARUVCHI qism saqlanadi — shuning uchun biz "
         "kollinearlikni qidiramiz.")
    table(doc, "24-jadval. Kompensatsiya strategiyalari, ikkita byudjetda "
               "(TEST splitining 300 namunasi; FP32 = 0.1793).",
          ["Byudjet", "Usul", "Nima saqlanadi", "WER", "Bizdan farq"],
          [["B1 = 267 MiB", "bizniki (vakilga qo'shish)", "o'zgaruvchi qism",
            "0.1833", "—"],
           ["", "gibrid (vakilga + bias)", "ikkalasi", "0.1842",
            "+0.0009 [-0.0088, +0.0089]"],
           ["", "FLAP [28] (bias)", "doimiy qism", "0.1859",
            "+0.0027 [-0.0125, +0.0150]"],
           ["B2 = 254 MiB", "bizniki (vakilga qo'shish)", "o'zgaruvchi qism",
            "0.2006", "—"],
           ["", "gibrid (vakilga + bias)", "ikkalasi", "0.1967",
            "-0.0039 [-0.0135, +0.0036]"],
           ["", "FLAP [28] (bias)", "doimiy qism", "0.1925",
            "-0.0081 [-0.0266, +0.0101]"]],
          good_rows=(4, 5))
    para(doc,
         "Uchala strategiya ham bir-biridan ajratilmaydi, va bu mexanizm "
         "haqidagi da'voni TASDIQLAYDI: FLAP bizning ikkala qismimizga ega "
         "bo'lgani uchun biz bilan bir joyga tushadi, kompensatsiyasiz "
         "bazalar esa (18-jadval) o'sha byudjetda 1.3393 va 2.7378 beradi. "
         "Ayni paytda taklif etilgan usul FLAP dan ustun ekani ham "
         "KO'RSATILMAYDI — agressiv byudjetda FLAP nuqtaviy bahoda oldinda.")
    para(doc, "Ikkala kompensatsiyani birlashtirish.", bold=True, size=10)
    para(doc,
         "Taqqoslash o'z usulimizdagi bo'shliqni ko'rsatdi. Vakilga "
         "proyeksiya qoldiq r_j = h_j - gamma_j h_p ni qoldiradi; u qurilishi "
         "bo'yicha h_p ga ortogonal, ammo ORTOGONAL NOL O'RTACHA degani emas, "
         "va o'sha doimiy qism shunchaki tashlanadi. Uni chiqish biasiga "
         "qo'shish yetarli:")
    eq(doc, "c_j = mean(h_j - gamma_j h_p),    b <- b + c_j W2[:, j]", 17)
    para(doc,
         "Bu qat'iy qo'shimcha: vaznlar, kanal tanlovi va kvantlagich "
         "o'zgarmaydi, bias esa allaqachon mavjud vektor, ya'ni qo'shimcha "
         "xotira ham, arifmetika ham talab qilinmaydi. Tuzatmaning kattaligi "
         "byudjet bilan o'sadi — qoldiq o'rtachasining normasi tau = 0.99 da "
         "0.001-0.008, tau = 0.95 da esa 0.01-0.44 — va WER ta'siri xuddi "
         "shu yo'nalishda: B1 da o'zgarish yo'q (+0.0009), B2 da yaxshilanish "
         "(-0.0039). Ya'ni bashorat o'lchangan kattalikdan qilingan va ikkala "
         "byudjetda ham yo'nalish bo'yicha tasdiqlangan, garchi effekt "
         "n = 300 da ahamiyatlilikka yetmasa ham. Gibrid FLAP gacha bo'lgan "
         "masofaning taxminan yarmini yopadi, demak FLAP ning qolgan "
         "ustunligi bias hadida emas, uning fluktuatsiya mezonida.",
         italic=True, size=10)

    para(doc, "Ikkinchi olib tashlash mexanizmi: kaskadga yangi zina.",
         bold=True, size=10)
    para(doc,
         "Kichikroq byudjetga yetishning hozirgacha yagona vositasi bor edi — "
         "tau ni tushirish, ya'ni ortiqchalik mezonining O'ZINI bo'shatish. "
         "Bu ishlaydi, lekin sifat izchil yomonlashadi (tau 0.99, 0.97, 0.95, "
         "0.93 uchun 0.1833, 0.1916, 0.2006, 0.2179), chunki bo'shatilgan "
         "kosinus haqiqatan takrorlanmagan kanallarni ham qabul qila "
         "boshlaydi.")
    para(doc,
         "FLAP bilan taqqoslash boshqa vositani ko'rsatdi. Kanal olib "
         "tashlanishi mumkin bo'lgan IKKI mustaqil sabab bor va har biri o'z "
         "muomalasini talab qiladi: ortiqcha kanal vakilga qo'shiladi, "
         "deyarli doimiy kanal esa o'rtachasi bilan biasga o'tadi. Bizning "
         "mezon faqat birinchisini ko'radi. Shuning uchun mezonni bo'shatish "
         "o'rniga u tau = 0.99 da QAT'IY qoldirilib, byudjetning qolgan qismi "
         "omon qolganlar orasidan fluktuatsiyasi eng kichiklarini olib "
         "tashlash bilan yopiladi:")
    eq(doc, "score_p = ||W2_comp[:, p]||^2 * Var(h_p)", 18)
    para(doc, "To'liq ta'rif.", bold=True, size=10)
    para(doc,
         "Qatlamning oraliq kanallari to'plami N = {1, ..., n} bo'lsin. "
         "Birinchi bosqich (3)-(5) formulalari bo'yicha ishlaydi va K_1 "
         "vakillar to'plamini qoldiradi, har bir olib tashlangan j esa "
         "gamma_j koeffitsienti bilan o'z vakili p(j) ga qo'shiladi. "
         "Ikkinchi bosqich K_1 ning ICHIDA (18) balli bo'yicha eng kichik "
         "b - |N \\ K_1| tasini olib tashlaydi, bunda b — kesh-bog'langan "
         "maqsad talab qilgan umumiy olib tashlash soni:")
    eq(doc, "K_2 = K_1 \\ argmin_{|S| = b - |N \\ K_1|, S subset K_1} "
            "sum_{p in S} score_p .", 19)
    para(doc,
         "Ikkala bosqich ham o'zi tashlagan qismning O'RTACHASINI biasga "
         "topshiradi — birinchisi qoldiq r_j = h_j - gamma_j h_p ning "
         "o'rtachasini (17), ikkinchisi esa butun h_p ning o'rtachasini. "
         "Ularni alohida hisoblash shart emas: mu = mean(h) bo'lganda ikkala "
         "hissa ham bitta ayirmada jamlanadi,")
    eq(doc, "b <- b + W2 mu - W2_final mu[K_2] ,", 20)
    para(doc,
         "chunki W2_final ustunlari birinchi bosqichning kompensatsiyasini "
         "allaqachon o'z ichiga oladi. Bu ayniyat guruh tayinlanishini talab "
         "qilmaydi va amalga oshirishda aynan shu shaklda ishlatiladi. "
         "Natijada operator (1024 x n) dan (1024 x |K_2|) ga qisqaradi, bias "
         "esa uzunligini saqlaydi — ya'ni qo'shimcha xotira ham, arifmetika "
         "ham talab qilinmaydi.")
    table(doc, "25-jadval. 254 MiB ga ikki xil yo'l: mezonni bo'shatish yoki "
               "ikkinchi mexanizm qo'shish (TEST splitining 300 namunasi).",
          ["Yo'l", "Kosinus mezoni", "WER", "Bo'shatishga nisbatan"],
          [["tau ni 0.99 dan 0.95 ga bo'shatish", "bo'shatilgan", "0.2006",
            "—"],
           ["tau = 0.99 qat'iy + fluktuatsiya bosqichi", "saqlangan", "0.1921",
            "-0.0085 [-0.0252, +0.0080]"],
           ["FLAP [28] (mos yozuvlar)", "—", "0.1925", "-0.0081"]],
          good_rows=(1,))
    para(doc,
         "Teng byudjetda ikkinchi mexanizm qo'shish mezonni bo'shatishdan "
         "yaxshiroq va u FLAP darajasiga chiqadi, ayni paytda ortiqchalik "
         "mezoni o'z qat'iy holida qoladi. Ikkinchi bosqich boshlang'ich "
         "qatlamlarda byudjetning katta qismini oladi (L0 da 942 kanal, bias "
         "tuzatmasining normasi 1.38), oxirgi qatlamlarda esa deyarli "
         "ishlamaydi (L19 da 4 kanal, norma 0.0008) — ya'ni u ham 6-jadvaldagi "
         "profil kabi ma'lumot boshqaradigan qaror.")
    para(doc,
         "Bu kaskadning 'eng yumshoq yetarli o'zgarish' zinapoyasiga yangi "
         "pog'ona qo'shadi: ortiqcha kanallarni vakil ko'taradi (tekin), "
         "deyarli doimiylarini bias ko'taradi (arzon), va faqat ikkalasi ham "
         "tugagach past-rank kerak bo'ladi — 15-jadvalga ko'ra u haqiqiy "
         "sifat narxiga ega (0.3056). Yig'ma yaxshilanish 0.2006 dan 0.1921 "
         "gacha, ya'ni 0.0085; alohida qadamlarning hech biri n = 300 da "
         "ahamiyatlilikka yetmaydi, ammo uchalasi ham mexanizm bashorat "
         "qilgan yo'nalishda harakat qiladi.")
    para(doc,
         "Buni yangi algoritm sifatida emas, KOMPOZITSIYA sifatida taqdim "
         "etamiz: birinchi bosqich ushbu ishnikidir, ikkinchisi FLAP ning "
         "mexanizmi bo'lib, birinchisi qoldirgan kanallarga qo'llanadi. Yangi "
         "bo'lgan narsa — bu ishning natijalari majbur qilgan tashkiliy "
         "printsip: MUOMALA kanal nima uchun olib tashlanayotganiga MOS "
         "bo'lishi kerak.", italic=True, size=10)

    figure(doc, 6,
           "Ikki mexanizmli strukturaviy qisqartirish. Kanal olib tashlanishi "
           "mumkin bo'lgan ikki mustaqil sabab bor va har biri o'z "
           "kompensatsiyasini talab qiladi: ortiqcha kanal vakilga "
           "qo'shiladi, deyarli doimiy kanal esa o'rtachasi bilan biasga "
           "o'tadi.",
           "A clean schematic diagram of a transformer feed-forward block, "
           "white background, publication style, thin lines, grayscale with "
           "two accent colors. Left: a vertical stack of small squares "
           "labelled 'FFN intermediate channels h_1 ... h_n'. The stack is "
           "visually partitioned into three bands. Top band, in accent color "
           "one: three squares bracketed together as a group, with two of "
           "them drawn with dashed outlines and curved arrows pointing into "
           "the third solid one, annotated 'cos(h_j,h_p) >= tau  ->  "
           "W[:,p] += gamma_j W[:,j]' and labelled 'STAGE 1: redundant'. "
           "Middle band, in accent color two: two squares drawn with dashed "
           "outlines, each showing a nearly flat horizontal waveform inside "
           "to suggest low variance, with straight arrows pointing down to a "
           "single wide rectangle labelled 'bias b', annotated "
           "'small Var(h_p)  ->  b += mean(h_p) W2[:,p]' and labelled "
           "'STAGE 2: nearly constant'. Bottom band in plain gray: solid "
           "squares each showing a strongly varying waveform, labelled "
           "'kept'. To the right of the stack, the surviving squares feed a "
           "narrower matrix block labelled 'W2 (1024 x |K_2|)' whose reduced "
           "width is visibly smaller than an outlined ghost rectangle behind "
           "it labelled 'original W2 (1024 x n)'. Below the whole figure, a "
           "horizontal three-rung ladder labelled 'cascade: softest "
           "sufficient change', rungs reading '1 redundant - free', "
           "'2 nearly constant - cheap', '3 low-rank - real accuracy cost', "
           "with the third rung drawn faded. Sans-serif labels, no shadows, "
           "no 3D effects.")

    para(doc, "Taqqoslashni tau dan mustaqil qilish.", bold=True, size=10)
    para(doc,
         "Yuqoridagi barcha strukturaviy taqqoslashlar bazalarga BIZNING "
         "qatlam taqsimotimizni majburlagan. Bu mezonni toza ajratadi, ammo "
         "har bir bazaga u tanlamagan taqsimotni beradi — va bitta natija "
         "aynan shunga tayanadi: magnitude 0-4 qatlamlardan 66-73% olib "
         "tashlashga majbur qilinganda qulagan, holbuki uning o'z ballari "
         "bunday taqsimotni bermasligi mumkin edi. Shu sababli faqat UMUMIY "
         "byudjet qat'iy qoldirilib, taqsimot har bir usulning o'z qoidasiga "
         "topshirildi: bizda tau boshqaradi va qatlam sonlari guruhlashdan "
         "chiqadi, magnitude barcha qatlamlarning barcha kanallarini bitta "
         "global ro'yxatda tartiblaydi, Wanda nashr etilgan shakli bo'yicha "
         "har qatlamda bir xil ulush oladi, FLAP esa qatlam bo'yicha "
         "normallashtirilgan ball bilan o'z moslashuvchan taqsimotini quradi.")
    table(doc, "26-jadval. Tau dan mustaqil taqqoslash: umumiy byudjet qat'iy, "
               "taqsimot har bir usulning o'ziniki (TEST splitining 300 "
               "namunasi; FP32 = 0.1793).",
          ["Byudjet", "Usul", "Taqsimot", "WER", "Asl variantimizdan farq"],
          [["267 MiB", "bizniki: asl (kosinus + vakil)", "tau", "0.1833", "—"],
           ["(17.1% kanal)", "bizniki: + bias tuzatmasi", "tau", "0.1842",
            "+0.0009 [-0.0088, +0.0089]"],
           ["", "FLAP [28]", "moslashuvchan", "0.1838",
            "+0.0006 [-0.0084, +0.0105]"],
           ["", "magnitude", "global tartiblash", "0.1876",
            "+0.0044 [-0.0105, +0.0171]"],
           ["", "Wanda [27]", "uniform", "0.2140",
            "+0.0307 [+0.0129, +0.0503]"],
           ["254 MiB", "bizniki: asl (kosinus + vakil)", "tau", "0.2006", "—"],
           ["(23.6% kanal)", "bizniki: + bias tuzatmasi", "tau", "0.1967",
            "-0.0039 [-0.0135, +0.0036]"],
           ["", "bizniki: ikki bosqichli (kompozitsiya)", "tau", "0.1921",
            "-0.0085 [-0.0252, +0.0080]"],
           ["", "FLAP [28]", "moslashuvchan", "0.1984",
            "-0.0022 [-0.0135, +0.0086]"],
           ["", "Wanda [27]", "uniform", "0.2378",
            "+0.0372 [+0.0186, +0.0566]"],
           ["", "magnitude", "global tartiblash", "0.4192",
            "+0.2186 [+0.0761, +0.4397]"]],
          good_rows=(7,), bad_rows=(4, 9, 10))
    para(doc,
         "Ikki xulosa chiqadi va ikkalasi ham yuqoridagi taqdimotni "
         "tuzatadi.")
    para(doc,
         "Birinchidan, TAQSIMOT MEZONDAN MUHIMROQ, va bu ishning yagona "
         "statistik jihatdan ahamiyatli ustunligi aynan shundan keladi. "
         "Wanda ikkala byudjetda ham sezilarli darajada ortda, ammo sabab "
         "uning mezonida emas: xuddi shu Wanda ballari bizning taqsimotimiz "
         "bilan 0.1850 va 0.2202 bergan (21-jadval), o'zining uniform "
         "taqsimoti bilan esa 0.2140 va 0.2378. Ya'ni bir xil mezon, faqat "
         "boshqa taqsimot — farq 0.029 va 0.018. Kesh-bog'langan kaskad "
         "ortiqchalik o'lchovi orqali yaxshi taqsimotni qo'shimcha xarajatsiz "
         "beradi, uniform siyraklik esa 5-jadvaldagi qatlam profilini "
         "e'tiborsiz qoldirib buning narxini to'laydi. Buni mustaqil "
         "tasdiqlash ham bor: taqsimoti profilga ergashadigan uchala usul "
         "(bizniki, FLAP, magnitude) 267 MiB da 0.183-0.188 oralig'ida "
         "to'planadi.")
    para(doc,
         "Ikkinchidan, magnitude haqidagi da'vo BO'RTTIRILGAN edi va shu "
         "yerda tuzatiladi. O'z global tartiblashi bilan u 254 MiB da 0.4192 "
         "beradi, bizning taqsimotimiz bilan esa 2.7378 bergan. Sifat "
         "xulosasi saqlanadi — faqat vaznga qaraydigan tanlash agressiv "
         "byudjetda sezilarli darajada buziladi (+0.2186) — ammo miqdoriy "
         "da'vo 13.6 barobar emas, 2.1 barobar bo'lishi kerak. Qulashning "
         "katta qismi bazaga begona taqsimot berilganidan kelib chiqqan.",
         italic=True, size=10)
    para(doc,
         "Uchinchi qator ehtiyotkorlik talab qiladi: ikki bosqichli variant "
         "FLAP ning mexanizmini o'z ichiga oladi, shuning uchun uning FLAP "
         "dan oldinda turishi (0.1921 va 0.1984) ustunlik da'vosi sifatida "
         "o'qilmasligi kerak. U faqat shuni qo'llab-quvvatlaydi: ikkala "
         "mexanizm birgalikda har biridan alohida yaxshiroq ishlaydi.")

    para(doc, "Kompensatsiya ablationi.", bold=True, size=10)
    para(doc,
         "Bizning usul ikkita mustaqil g'oyani birlashtiradi: qaysi kanallar "
         "ketishi (saqlanadigan kanal bilan kollinear bo'lganlari) va ketgach "
         "nima bo'lishi (ular vakilga qo'shiladi). Wanda kompensatsiyasiz "
         "omon qolgani birinchi qarashda kompensatsiya keraksizdek "
         "ko'rsatadi, ammo bu boshqa usul haqidagi xulosa bo'lgani uchun "
         "to'g'ridan-to'g'ri o'lchandi: XUDDI SHU kanallar, kompensatsiya "
         "o'chirilgan holda. Natija — WER 0.2006 dan 1.3393 ga ko'tariladi, "
         "ya'ni 6.7 barobar yomonlashish.")
    para(doc,
         "Bu ikki usulning MEXANIZMI har xil ekanini ko'rsatadi. Wanda "
         "AHAMIYATSIZ kanallarni tanlaydi — ular kam hissa qo'shgani uchun "
         "ularni oddiy o'chirish xavfsiz, kompensatsiya qiladigan narsa yo'q. "
         "Bizning mezon esa ORTIQCHA kanallarni tanlaydi, va bunday kanal "
         "katta bo'lishi mumkin: u haqiqiy signal tashiydi, faqat o'sha "
         "signal boshqa kanalda takrorlanadi. Uni oddiy o'chirish signalni "
         "yo'qotadi, vakilga qo'shish esa saqlaydi. Shu sababli kompensatsiya "
         "usulning ixtiyoriy yaxshilanishi emas, TARKIBIY qismi: "
         "kompensatsiyasiz bu tanlash mezoni ichki ziddiyatli bo'lib qoladi.")
    para(doc,
         "Bu 4.4-bo'limdagi bog'liqlikni ham oqlaydi. Kompensatsiya satrlar "
         "bo'yicha vazn diapazonini 9.6x dan 188.4x ga kengaytirib, "
         "per-channel kvantlashni majburiy qiladi; endi ko'rinadiki, bu narx "
         "haqiqiy foyda evaziga to'lanadi.", italic=True, size=10)

    h(doc, "4.9e. Aniqlik byudjeti bo'yicha tau tanlash", 2)
    para(doc,
         "Shu paytgacha tavsiflangan kaskad bir tomonlama edi. Kesh byudjeti "
         "qancha siqish TALAB QILINISHINI aytadi, ammo hech narsa qancha "
         "siqish RUXSAT ETILISHINI aytmaydi: tau = 0.99 qat'iy giperparametr "
         "bo'lib, na chiqarilgan, na sifat bilan chegaralangan. 4.9d-bo'lim "
         "bu bo'shliqni ko'rsatdi — tau = 0.95 ancha ko'p olib tashlaydi va "
         "usulda uni to'xtatadigan tamoyil yo'q.")
    para(doc,
         "Aniqlik uchun: ushbu ishning barcha natijalarida ishlatilgan "
         "tau = 0.99 quyida tavsiflanadigan protsedura bilan TANLANMAGAN. U "
         "boshidanoq qat'iy belgilangan qiymat bo'lib, protsedura esa uni "
         "qanday chiqarish mumkinligini ko'rsatadi va o'z cheklovlari bilan "
         "birga keltiriladi.")
    para(doc,
         "Yopilishi kerak bo'lgan ikkinchi tomon quyidagicha qo'yiladi: kesh "
         "byudjeti minimal zarur siqishni, aniqlik byudjeti esa maksimal "
         "ruxsat etilgan siqishni beradi, tau shu ikkisi orasida qidiriladi. "
         "Byudjet MUTLAQ dWER sifatida emas, FP32 ga NISBATAN beriladi, "
         "chunki mutlaq tolerans ko'chmaydi: 0.02 WER modeli 0.05 da bo'lgan "
         "tizim uchun 40% yomonlashish, 0.50 da bo'lgani uchun esa 4%. "
         "Foydalanuvchi o'z FP32 modelining sifatini biladi, shuning uchun "
         "spetsifikatsiya tabiiy ravishda quyidagicha:")
    eq(doc, "WER_ruxsat = WER_FP32 x (1 + eps)", 21)
    para(doc,
         "Tanlash signali sifatida vazifa mezoni ishlatiladi. Bu ishning o'z "
         "natijalari shuni majbur qiladi: 4.7-bo'limda operator xatosi 160 "
         "barobar o'zgarganda tarmoq xatosi atigi 4 barobar o'zgargan, "
         "4.9a-bo'limda esa GPTQ ning 57% operator ustunligi WER ga umuman "
         "o'tmagan. Ya'ni E_loc va E_glob bu yerda sifat o'rnini bosa "
         "olmaydi. Bundan tashqari tau TEST splitida tanlanmaydi: "
         "4.9-bo'limda tanlash va baholash bir taqsimotdan bo'lganda "
         "variantlar tartibi teskari bo'lgani ko'rsatilgan, shuning uchun "
         "skanerlash kalibrlashdan keyingi VALIDATION namunalarida, hisobot "
         "esa TEST da beriladi.")
    table(doc, "27-jadval. tau egri chizig'i, 100 validation namunasi "
               "(tanlov uchun; hisobot 300 namunali TEST da).",
          ["Variant", "MiB", "WER", "FP32 ga", "Kvantlanganga nisbatan dWER"],
          [["FP32 (mos yozuvlar)", "1172", "0.0931", "1.000x",
            "-0.0080 [-0.0160, -0.0018]"],
           ["GPTQ yolg'iz (qisqartirishsiz)", "300", "0.1011", "1.086x", "—"],
           ["tau = 0.99", "267", "0.1013", "1.089x", "+0.0003 [-0.0068, +0.0081]"],
           ["tau = 0.97", "261", "0.0967", "1.039x", "-0.0043 [-0.0243, +0.0123]"],
           ["tau = 0.95", "254", "0.1004", "1.079x", "-0.0006 [-0.0219, +0.0180]"],
           ["tau = 0.93", "248", "0.1147", "1.232x", "+0.0136 [-0.0002, +0.0291]"],
           ["tau = 0.90", "237", "0.1125", "1.209x", "+0.0114 [-0.0089, +0.0285]"]],
          good_rows=(3,))
    para(doc,
         "Mexanizm ishlaydi va ikki xil foydali javob beradi. eps = 5% uchun "
         "u tau = 0.97 ni tanlaydi (261 MiB, 1.039x). eps = 1% va eps = 2% "
         "uchun esa HECH BIR tau mos kelmaydi, chunki kvantlashning o'zi "
         "1.086x beradi — bu strukturaviy o'qning ayibi emas, byudjet kesh "
         "talab qilgan majburiy bosqich uchun ham juda qattiq. Amaliy "
         "jihatdan bu qimmatli: foydalanuvchi tau ni behuda sozlamaydi, "
         "balki byudjetni yoki kvantlash bosqichini qayta ko'rib chiqadi.")
    para(doc,
         "Egri chiziq monoton emas: tau = 0.97 ko'proq olib tashlashiga "
         "qaramay tau = 0.99 dan yaxshiroq (0.0967 ga qarshi 0.1013), va "
         "ikkalasi ham qisqartirishsiz kvantlangan bazadan past. Bu "
         "4.9b-bo'limdagi kuzatuv bilan izchil (0.1833 ga qarshi 0.1847) va "
         "strukturaviy olib tashlash yumshoq regulyarizator sifatida ishlashi "
         "mumkinligini ko'rsatadi.")
    para(doc, "Ikkinchi bosqich tanlovni rad etadi.", bold=True, size=10)
    para(doc,
         "Skanerlashda tanlangan tau = 0.97 mustaqil test splitida "
         "tasdiqlandi va byudjetni QONDIRMADI: u yerda FP32 ning 1.069 "
         "karrasini beradi, ya'ni eps = 5% dan chetda. Xuddi shu to'plamda "
         "tau = 0.99 esa 1.022 karra beradi va byudjetga sig'adi. Ya'ni "
         "to'g'ri javob konservativ qiymat edi.")
    table(doc, "28-jadval. Tanlov ko'chmadi: validation skanerlashi va TEST "
               "splitidagi tasdiq (FP32 ga nisbatan).",
          ["Variant", "MiB", "Validation (n=100)", "TEST (n=300)"],
          [["FP32", "1172", "1.000x", "1.000x"],
           ["tau = 0.99", "267", "1.089x", "1.022x"],
           ["tau = 0.97 (skanerlash tanlovi)", "261", "1.039x", "1.069x"],
           ["tau = 0.95", "254", "1.079x", "1.119x"]],
          good_rows=(1,), bad_rows=(2,))
    para(doc,
         "Sabab ikki qismli. Birinchidan, n = 100 da juftlik oralig'i "
         "taxminan +-0.019, tau = 0.99, 0.97 va 0.95 orasidagi tarqoqlik esa "
         "atigi 0.0046 — ya'ni uchlik ichidagi tartib shovqin bilan "
         "belgilangan. Ikkinchidan, test splitida egri chiziq monoton "
         "(1.022x, 1.069x, 1.119x), validation dagi monoton emaslik esa o'sha "
         "shovqinning ko'rinishi edi. Skanerlash bu uchlikni tau = 0.93 va "
         "0.90 dan ishonchli ajratadi, ammo uchlik ICHIDAN tanlashga "
         "yetmaydi.")
    para(doc, "Tuzatilgan tanlov qoidasi.", bold=True, size=10)
    para(doc,
         "Xatoning manbai skanerlashning noto'g'ri raqam berishi emas — unga "
         "javob bera olmaydigan savol berilgani. Nuqtaviy bahoni o'qiydigan "
         "qoida o'z noaniqligini ko'ra olmaydi va eng omadli nomzodni "
         "ishonch bilan qaytaradi. Oraliqni o'qish buni tuzatadi: eps "
         "byudjeti uchun tau faqat quyidagi shart bajarilganda maqbul "
         "hisoblanadi:")
    eq(doc, "yuqori chegara( WER_tau - WER_FP32 )  <=  eps x WER_FP32", 22)
    table(doc, "29-jadval. Nuqtaviy baho va ishonch chegarasi bo'yicha "
               "tanlovlarning taqqoslanishi (skanerlash n = 100, natija "
               "300 namunali TEST splitida tekshirilgan).",
          ["eps", "Nuqtaviy baho tanlovi", "Testdagi natijasi",
           "Ishonch chegarasi tanlovi"],
          [["1%", "yo'q", "—", "yo'q"],
           ["2%", "yo'q", "—", "yo'q"],
           ["5%", "tau = 0.97, 261 MiB", "1.069x — BYUDJET BUZILDI", "yo'q"],
           ["10%", "tau = 0.95, 254 MiB", "1.119x — BYUDJET BUZILDI", "yo'q"]],
          bad_rows=(2, 3))
    para(doc,
         "Nuqtaviy qoida ikki holatda ham tanlov qildi va ikkalasi ham "
         "byudjetni buzdi. Ishonch chegarasiga asoslangan qoida esa ikkala "
         "holatda ham TANLASHDAN BOSH TORTDI — bu to'g'ri xulq, chunki "
         "n = 100 shu byudjetlarni sertifikatlashga yetmaydi. Skanerlashning "
         "o'lchangan ruxsati shu bilan aniqlanadi: tau = 0.99 uchun yuqori "
         "chegara +0.0182 ni tashkil qiladi, ya'ni bu hajmda faqat eps >= 20% "
         "sertifikatlanishi mumkin.")
    para(doc, "Nisbiy byudjetning baholash to'plamiga sezgirligi.", bold=True,
         size=10)
    para(doc,
         "Muammoning asosiy qismi namuna soni emas. Mutlaq zarar test "
         "splitida IKKI BAROBAR KICHIK (tau = 0.99 uchun +0.0040, validation "
         "da +0.0082), ammo nisbat aksincha kattaroq ko'rinadi, chunki maxraj "
         "— FP32 ning o'z xatoligi — validation da ikki barobar past "
         "(0.0931 ga qarshi 0.1793). Ya'ni bir xil mutlaq buzilish OSON "
         "to'plamda nisbatan yomonroq ko'rinadi va hech qanday namuna soni "
         "buni tuzatmaydi.")
    para(doc,
         "Bundan spetsifikatsiyaga uchinchi shart qo'shiladi. FP32 ga "
         "nisbatan berilgan byudjet modeldan modelga ko'chadi va shu sababli "
         "mutlaq tolerans dan afzal, ammo u baholash to'plamining QIYINLIGIGA "
         "bog'liq. Shuning uchun skanerlash to'plami kalibrlash va test bilan "
         "kesishmasligi yetarli emas — u joylashtirish taqsimotini vakillik "
         "qilishi kerak. Ushbu ishda validation bo'linmasi korpusning "
         "osonroq qismi bo'lib chiqdi va tau ni haddan tashqari konservativ "
         "baholadi.", italic=True, size=10)
    para(doc,
         "Umumiy xulosa protokolning o'zini oqlaydi: ikkinchi bosqich formal "
         "emas. U bu yerda haqiqiy xatoni ushladi, va agar tanlov faqat "
         "skanerlashga tayanganda, byudjetni buzadigan konfiguratsiya e'lon "
         "qilingan bo'lardi.")
    para(doc, "Protokolning tuzatilgan shakli (taklif; sinalmagan).",
         bold=True, size=10)
    para(doc,
         "Yuqoridagi muvaffaqiyatsizlikdan ikkita aniq tuzatish kelib "
         "chiqadi, va ular birgalikda protokolni ishlaydigan holga "
         "keltiradi.")
    para(doc,
         "Birinchisi — SKANERLASH TO'PLAMI. Muammoning asosiy qismi namuna "
         "soni emas edi: validation da tau = 0.99 nisbiy 8.9% turadi, test da "
         "esa 2.2%, ya'ni nuqtaviy baholarning o'zi kelishmaydi va cheksiz "
         "namuna ham buni tuzatmaydi. Sabab validation bo'linmasining "
         "osonroq bo'lishida (FP32 xatoligi 0.0931 ga qarshi 0.1793) va "
         "nisbiy byudjetning maxrajga sezgirligida. Shu sababli skanerlash "
         "joylashtirish taqsimotining O'ZIDAN olinishi kerak. Amalda bu test "
         "bo'linmasini ikkiga ajratishni anglatadi: bir yarmi tau ni "
         "tanlashga, ikkinchi yarmi hisobotga. Tanlov va hisobot bir "
         "taqsimotdan bo'ladi, ammo kesishmaydi.")
    para(doc,
         "Ikkinchisi — QIDIRUV YO'NALISHI. Maqsad byudjetga sig'adigan ENG "
         "KICHIK modelni topish bo'lgani uchun qidiruvni eng agressiv "
         "tomondan boshlash tabiiy: tau = 0.90 dan 0.99 tomon ko'tarilib, "
         "ishonch chegarasi byudjetni qanoatlantirgan birinchi qiymatda "
         "to'xtash. O'sha birinchi qoniqarli nuqta izlanayotgan javobning "
         "o'zi bo'ladi. Teskari yo'nalish esa avval qoniqarsiz nuqtani topib, "
         "so'ng bir qadam orqaga qaytishni talab qiladi, ya'ni bitta ortiqcha "
         "baholash.")
    para(doc,
         "Bu shakl ushbu ishda O'LCHANMAGAN va shu sababli taklif sifatida "
         "keltiriladi: mavjud natijalar tuzatishlarning zarurligini "
         "ko'rsatadi, ammo tuzatilgan protokolning byudjetga sig'adigan tau "
         "ni haqiqatan qaytarishini tasdiqlamaydi. Buni tekshirish uchun "
         "qo'shimcha o'lchov talab qilinadi va u keyingi ishga qoldiriladi.",
         italic=True, size=10)

    figure(doc, 7,
           "Butun model uchun siqish-sifat egri chizig'i va kaskadning "
           "to'xtash nuqtasi. Qo'lda tanlanadigan oqilona ko'rinuvchi daraja "
           "(5.34x) modelni yo'q qiladi.",
           "A scientific line-and-marker plot, white background, publication "
           "style, mostly grayscale with one highlighted marker. X-axis "
           "'whole-model compression factor' from 1 to 6, linear. Y-axis 'word "
           "error rate' from 0.15 to 0.65. Four labelled points connected by a "
           "thin line in x order: 'FP32' at (1.00, 0.1761); 'uniform INT8' at "
           "(3.95, 0.1847); 'cascade' at (4.14, 0.1833) drawn as a large "
           "filled highlighted marker; 'uniform aggressive' at (5.34, 0.6101) "
           "drawn in a warning color. The line is nearly flat from 1.0 to 4.14 "
           "and then rises almost vertically to the last point, forming a "
           "cliff. A horizontal dashed reference line at the FP32 WER 0.1761 "
           "spanning the plot. A light shaded vertical band between x=4.0 and "
           "x=4.3 labelled 'derived target: decoder needs 3.81x, INT8 supplies "
           "4.00x'. A short annotation with a leader line pointing at the last "
           "point reading 'hand-picked 5.34x: model destroyed'. Small vertical "
           "95% CI bars on each marker. Thin axis lines, sans-serif labels, "
           "legend omitted since points are labelled directly.")

    h(doc, "4.9f. Freymvork turli kesh hajmlarida: mashinalar bo'yicha baholash", 2)
    para(doc,
         "Shu paytgacha kaskad bitta mashinada, L3 = 24 MiB da baholandi. "
         "Freymvorkning da'vosi esa kengroq: u modelni, kesh hajmini va "
         "aniqlik byudjetini olib, o'sha apparat uchun konfiguratsiya "
         "tanlaydi. Bu bo'lim shu da'voni Whisper CPU da yurgiziladigan "
         "mashinalar diapazonida sinaydi va uni ikkita KO'R-KO'RONA "
         "amaliyot bilan taqqoslaydi.")
    para(doc, "Nima uchun bu hisoblab bo'ladigan tajriba.", bold=True, size=10)
    para(doc,
         "Har bir konfiguratsiyani har bir mashina uchun qayta baholash bir "
         "necha kun olardi. Bunga hojat yo'q, chunki WER — ARTEFAKTNING "
         "xossasi, mashinaning emas. Kesh hajmi qaysi artefakt tanlanishini "
         "va uning xotira trafigi qanchaligini o'zgartiradi; berilgan "
         "artefakt qanchalik yaxshi transkripsiya qilishini emas. Shuning "
         "uchun konfiguratsiyalar kutubxonasi BIR MARTA o'lchandi (o'sha 300 "
         "TEST namunasi, o'sha INT8 dekoder), barcha kesh hajmlari esa shu "
         "jadval ustidagi arifmetika bilan javob oldi.")
    table(doc, "30-jadval. Konfiguratsiyalar kutubxonasi (300 TEST "
               "namunasi, INT8 dekoder; kechikish NAVBATLASHGAN dizaynda, "
               "7 raund mediana, bitta oqim).",
          ["Konfiguratsiya", "Enk. MiB", "ms", "INT8 ga", "WER", "CER"],
          [["FP32 enkoder (nazorat)", "1172", "11550", "0.60x", "0.1793", "0.0522"],
           ["ko'r-ko'rona INT8", "300", "6981", "1.00x", "0.1847", "0.0538"],
           ["bizniki tau=0.99 (17%)", "267", "6602", "1.06x", "0.1833", "0.0549"],
           ["bizniki tau=0.97 (20%)", "261", "6467", "1.08x", "0.1916", "0.0575"],
           ["bizniki tau=0.95 (24%)", "254", "6598", "1.06x", "0.2006", "0.0621"],
           ["bizniki tau=0.93 (27%)", "248", "6477", "1.08x", "0.2179", "0.0710"],
           ["ko'r-ko'rona magnitude 30%", "242", "6498", "1.07x", "0.6294", "0.3626"],
           ["bizniki tau=0.90 (33%)", "237", "6368", "1.10x", "0.2365", "0.0847"],
           ["o'q gibridi (kanal L0-5, rank L6+)", "213", "6129", "1.14x", "0.3026", "0.0968"],
           ["kesh-majburiy 45% (tau panjarasi + trim)", "213", "6210", "1.12x", "0.3393", "0.1416"],
           ["kesh-majburiy 45% (uzluksiz tau bisektsiyasi)", "213", "6167", "1.13x", "0.3400", "0.1166"],
           ["ko'r-ko'rona magnitude 50%", "203", "6127", "1.14x", "0.7913", "0.4243"]],
          good_rows=(2,), bad_rows=(6, 11))
    para(doc, "Kechikish raqamlari qayta o'lchandi: blokli dizayndagi "
              "bitta yomon qator.", bold=True, size=10)
    para(doc,
         "Ushbu jadvalning kechikish ustuni dastlab BLOKLI tartibda "
         "o'lchangan edi — har konfiguratsiya alohida, natijalar "
         "keshlangan holda — ya'ni 4.10a-bo'limda ishonchsiz deb "
         "topilgan dizaynning aynan o'zi. Butun kutubxona navbatlashgan "
         "protokolda (A B C ... A B C, 7 raund, o'sha audio kirishi) "
         "qayta o'lchanganda o'n ikki qatordan O'N BIRTASI 0.4-6.3% "
         "ichida takrorlandi (shovqin darajasi), ko'r-ko'rona INT8 "
         "qatori esa 8658 dan 6981 ms ga, ya'ni 19.4% ga o'zgardi. "
         "Protokollar tizimli farq qilganda barcha qatorlar birga "
         "siljigan bo'lardi; faqat bittasining siljishi o'sha qator "
         "mashina sekin holatda o'lchanganini va blokli+keshlangan "
         "dizayn buni tuzatmasdan qoldirganini ko'rsatadi. Jadvaldagi "
         "raqamlar navbatlashgan o'lchovdan olingan; eski qiymatlar "
         "yozib qoldirilmadi, chunki ular bir xil artefaktning uch "
         "boshqa o'lchovi (38-jadvaldagi 7081 va 7417 ms) bilan ham "
         "ziddiyatda edi.", italic=True, size=10)
    para(doc,
         "Tuzatishning oqibati katta: tau = 0.99 ning ko'r-ko'rona INT8 "
         "ga nisbatan tezlanishi 1.29x emas, 1.06x. Bu esa 39-jadvaldagi "
         "teng-byudjet xulosasi bilan endi to'liq mos keladi — o'sha "
         "yerda usullar teng hajmda xotira xatti-harakati bo'yicha "
         "ajralmagan edi va 'tezlik baytlar soniga bog'liq, algoritmga "
         "emas' deb yozilgandi. Yangi ustun aynan shuni ko'rsatadi: "
         "1.06-1.14x oralig'idagi barcha qiymatlar hajm bilan tartibda, "
         "mezondan qat'i nazar (ko'r-ko'rona magnitude 30% ham 1.07x "
         "beradi). Eski 1.29x qiymati ishning O'Z xulosasiga zid edi va "
         "aynan qulay bo'lgani uchun e'tibordan chetda qolgan.")
    para(doc,
         "Ko'r-ko'ronalik uch o'qda o'lchanadi: nisbat qo'lda tanlanadi, "
         "taqsimot qatlamlar bo'ylab bir xil, mezon esa magnitude — "
         "strukturaviy pruning vositalarining standarti. Bu bazalarga "
         "bizning KOMPENSATSIYAMIZ berilgan; amalda kutubxona bilan "
         "ishlaydigan foydalanuvchi kompensatsiya qilmaydi va model butunlay "
         "quladi (o'xshash byudjetda 1.3393, 24-jadval). Bazaga yaxshiroq "
         "ishlovni berish uni yengishni QIYINLASHTIRADI, ya'ni omon qolgan "
         "har qanday ustunlik sodda yondashuvning eng kuchli shakliga "
         "qarshi qo'lga kiritilgan bo'ladi.")
    para(doc,
         "Natija keskin, ammo uni to'g'ri o'qish kerak. TEZLIK bo'yicha "
         "usullar deyarli ajralmaydi: hamma qisqartirilgan variant "
         "1.06-1.14x oralig'ida va tartib hajm bilan belgilanadi — "
         "ko'r-ko'rona magnitude 30% (242 MiB) 1.07x, bizning tau = 0.99 "
         "(267 MiB) 1.06x. Ajratuvchi o'lchov SIFAT: o'sha 1.07x uchun "
         "magnitude WER ni 0.6294 ga chiqaradi, 1.14x uchun esa 0.7913 "
         "ga; tau = 0.99 esa 11% kichikroq bo'lib, aniqligi ko'r-ko'rona "
         "INT8 dan farqlanmaydi (dWER = -0.0014, 95% IO [-0.0111, "
         "+0.0096], 15-jadval). Ya'ni tezlikni qisqartirishdan olish "
         "OSON — har qanday nisbat beradi — qiyin narsa uni SIFATNI "
         "yo'qotmasdan olish, va qidiruv aynan shuni to'laydi.")
    para(doc, "Mashina bo'yicha tanlov.", bold=True, size=10)
    para(doc,
         "Freymvorkning qoidasi 'L3 ga sig'dirish' EMAS. Kesh — maqsad "
         "funksiyasi, darvoza emas: sig'adigan konfiguratsiyalar orasida "
         "aniqlik byudjetini qanoatlantiradigani bo'lmasa, javob eng "
         "kichigi emas, byudjetni qanoatlantiradiganlar ichida MISS I ENG "
         "KICHIGI bo'ladi.")
    table(doc, "31-jadval. Mashina bo'yicha tanlov va ko'r-ko'rona INT8 ga "
               "nisbatan miss yutug'i (mutlaq byudjet, WER <= 0.2261).",
          ["Mashina (protsessor)", "L3 (MiB)", "Byudjet", "Tanlov", "Miss",
           "Yutuq", "Sig'adimi"],
          [["Raspberry Pi 5 (BCM2712, 4 yadro)", "2", "1.4M", "tau=0.93",
            "306961M", "1.26x", "yo'q"],
           ["Intel N100 (Alder Lake-N, 4 yadro)", "6", "4.2M", "tau=0.93",
            "206228M", "1.38x", "yo'q"],
           ["Core i5-1235U (2P + 8E)", "12", "8.4M", "tau=0.93", "73879M",
            "1.81x", "yo'q"],
           ["Tiger Lake H (bizniki, 16 yadro)", "24", "16.8M", "tau=0.93",
            "624M", "1.08x", "ha"],
           ["Ryzen 7 5800X (8 yadro, bitta CCD)", "32", "22.4M", "tau=0.93",
            "624M", "1.08x", "ha"],
           ["EPYC 7773X (Milan-X, CCD boshiga)", "96", "67.2M", "tau=0.93",
            "624M", "1.08x", "ha"]],
          good_rows=(2,))
    para(doc,
         "Uchta rejim ko'rinadi va ular 4.1-bo'limdagi taqsimot "
         "arifmetikasidan kelib chiqadi. L3 <= 6 MiB da hech bir qatlam "
         "byudjetga sig'maydi, miss faqat umumiy hajmga bog'liq va yutuq "
         "1.26-1.38x. L3 = 12 MiB da yutuq eng katta (1.81x), va uning "
         "manbai kutilmagan: tau-oilasi ENG KATTA qatlamni deyarli "
         "kichraytirmaydi (12.08 MiB, INT8 ning 12.10 MiB iga qarshi), "
         "chunki bizning mezonimiz ortiqchalikni erta qatlamlarda topadi. "
         "Yutuq esa aynan o'sha erta qatlamlar byudjetdan PASTGA "
         "tushishidan keladi — tau = 0.93 da L0 dan 72% kanal olinadi va "
         "qatlam 6.35 MiB ga tushib, R - 1 = 1499 ko'paytuvchisidan "
         "qutuladi. L3 >= 24 MiB da hammasi sig'adi va yutuq faqat hajm "
         "nisbatiga aylanadi (1.08x).")
    para(doc,
         "OGOHLANTIRISH. Ushbu jadvaldagi 1.26-1.81x lik yutuqlar miss "
         "ifodasining OVERFLOW HADIGA tayanadi, va o'sha had keyinchalik "
         "to'g'ridan-to'g'ri o'lchanganda tasdiqlanmadi: byudjetni kesib "
         "o'tishning narxi 2% dan oshmaydi (4.10a-bo'lim, 41-jadval). "
         "Shuning uchun L3 = 2, 6 va 12 MiB qatorlari TASDIQLANMAGAN model "
         "natijalari sifatida o'qilishi kerak. 24 MiB va undan yuqoridagi "
         "qatorlar overflow hadidan foydalanmaydi — u yerda miss oddiy bayt "
         "hisobiga teng — shuning uchun ular ta'sirlanmaydi.",
         italic=True, size=10)
    para(doc,
         "Diqqat: L3 = 2, 6 va 12 MiB da freymvork sig'dirishga URINMAYDI, "
         "chunki sig'adigan konfiguratsiyalar aniqlik byudjetidan chiqadi "
         "(kesh-majburiy 45% arm 0.3393 beradi). Bu muvaffaqiyatsizlik "
         "emas, yumshoq maqsad mavjud bo'lgan holatning o'zi.")
    para(doc, "Qaysi o'q qaysi qatlamda.", bold=True, size=10)
    para(doc,
         "Kesh maqsadini majburlash tau ni chuqurlik bo'ylab qulatadi — "
         "L3 = 12 MiB uchun qurilgan armda L0-L4 da 0.95-0.99, L19-L23 da "
         "esa 0.30. Bunday tau da guruhlash endi kollinearlikni "
         "topmayapti, faqat qaysi kollinear BO'LMAGAN kanalni qurbon "
         "qilishni saralayapti. Bu bitta turdagi ortiqchalik haqidagi "
         "bayonot; past-rank esa boshqasini — spektral kamayishni — "
         "ishlatadi. Teng parametr byudjetida, held-out faollashuvlarda "
         "o'lchandi:")
    table(doc, "32-jadval. Qaysi o'q kam xato beradi (fc2, teng parametr, "
               "held-out E_loc, k = 2253 kanalga rank 451 mos keladi).",
          ["Qatlamlar", "Kerakli tau", "E_kanal", "E_rank", "Yutuvchi"],
          [["L0-L4", "0.95-0.99", "0.0001-0.0048", "0.019-0.049", "kanal (100-400x)"],
           ["L5", "0.95", "0.0404", "0.0651", "kanal (1.6x)"],
           ["L6-L14", "0.80-0.90", "0.066-0.264", "0.055-0.174", "rank (1.1-1.7x)"],
           ["L15-L18", "0.60", "0.212-0.289", "0.195-0.221", "rank (1.1-1.3x)"],
           ["L19-L23", "0.30", "0.212-0.250", "0.166-0.246", "rank (1.0-1.4x)"]],
          good_rows=(0,))
    para(doc,
         "Rank 24 qatlamdan 18 tasida yutdi va mezon tugagan qatlamlarning "
         "(tau <= 0.60) 9/9 tasida yutdi. Qator/rank nisbati 18.2, ya'ni "
         "12-jadvaldagi 10-20 talabining ichida — past-rank armi "
         "kalibrlashni yodlab olmayapti, aks holda bu taqqoslash "
         "ishonchsiz bo'lardi.")
    para(doc,
         "Shu o'lchov asosida gibrid enkoder qurildi: L0-L5 da kanal, "
         "L6+ da rank, xuddi shu 213 MiB byudjetda. Uchdan-uchgacha natija "
         "esa ehtiyotkorlikni talab qiladi.")
    table(doc, "33-jadval. L3 = 12 MiB uchun uchta siyosat "
               "(300 TEST namunasi, INT8 dekoder).",
          ["Siyosat", "Enk. MiB", "WER", "95% IO", "Yumshoqqa nisbatan"],
          [["yumshoq: tau=0.90, mezon hurmat qilinadi", "237", "0.2365",
            "[0.2026, 0.2725]", "—"],
           ["o'q gibridi (kanal L0-5, rank L6+)", "213", "0.3026",
            "[0.2658, 0.3399]", "+0.0661 [+0.0382, +0.0935] SEZILARLI"],
           ["qat'iy: 45%/qatlam (tau panjarasi)", "213", "0.3393",
            "[0.2673, 0.4317]", "+0.1028 [+0.0430, +0.1902] SEZILARLI"],
           ["qat'iy: 45%/qatlam (uzluksiz tau)", "213", "0.3400",
            "[0.2929, 0.3958]", "panjaradan farqlanmaydi"]],
          good_rows=(0,), bad_rows=(2, 3))
    para(doc,
         "Markaziy xulosa shu jadvalda: kesh maqsadini MAJBURLASH atigi "
         "24 MiB tejash uchun aniqlikni sezilarli qurbon qiladi. Ikkala "
         "qat'iy variant ham yumshoq siyosatdan statistik jihatdan yomon. "
         "Bu 'sig'ishga intilmaslik' qoidasini fikr emas, o'lchov qiladi.")
    para(doc, "Tau ni uzluksiz qidirish.", bold=True, size=10)
    para(doc,
         "Yuqoridagi qat'iy arm tau ni oldindan tanlangan yetti nuqtadan "
         "iborat panjarada qidirgan va byudjetdan oshgan qismni "
         "qaytarish bilan tuzatgan. Tau esa haqiqiy son, panjara faqat "
         "hisob tejash edi. Guruhlash bosqichi tezlashgandan keyin uni "
         "10^-3 aniqlikda BISEKTSIYA bilan qidirish arzon bo'ldi, va "
         "natija sifat jihatidan boshqacha: tau hech qachon 0.536 dan "
         "pastga tushmaydi (panjarada 0.30 gacha tushgan edi) va byudjet "
         "uchun qaytariladigan kanallar soni 320-939 dan 9-256 ga "
         "kamayadi. Ya'ni har bir birlashma ancha yaxshi asoslangan.")
    para(doc,
         "Uchdan-uchgacha esa farq YO'Q: teng 213 MiB da "
         "dWER = +0.0007 [-0.0874, +0.0720]. Bu foydali salbiy natija — u "
         "qaysi kanallar olib tashlanishi tau ning ANIQ qiymatiga sezgir "
         "emasligini ko'rsatadi, ya'ni freymvork qo'pol panjara bilan "
         "ishlayverishi mumkin va bu qurilishni taxminan besh barobar "
         "tezlashtiradi. Diqqatga sazovori shundaki, CER bu yerda ham "
         "yaxshilanadi (0.1416 dan 0.1166 ga, 18%) — xuddi o'q gibridida "
         "bo'lgani kabi. Ikkala 'birlashmalarni yaxshiroq asoslash' "
         "o'zgarishi ham belgi darajasidagi buzilishni kamaytiradi, lekin "
         "so'z chegarasini tiklashga yetmaydi: bitta noto'g'ri belgi butun "
         "so'zni xato qiladi.")
    para(doc,
         "Shu bilan o'q tanlovining uchdan-uchgacha holati aniqlashadi va "
         "u CHEGARADA turadi. Gibrid panjara armidan farqlanmaydi "
         "(dWER = -0.0367 [-0.1151, +0.0191]), bisektsiya armidan esa "
         "sezilarli yaxshi (dWER = -0.0374 [-0.0839, -0.0028]) — nuqtaviy "
         "bahoning deyarli o'zi, ammo bisektsiya armi barqarorroq "
         "bo'lgani uchun oraliq torroq. Uchta juftlik taqqoslashidan "
         "bittasi sezilarli, chegara esa nolga juda yaqin, shuning uchun "
         "biz buni TASDIQ deb hisoblamaymiz. To'g'ri bayon: operator "
         "darajasidagi ustunlik mustahkam (18/24, mezon tugagan joyda "
         "9/9), uning WER ga ko'chishi esa chegarada va ko'proq namuna "
         "talab qiladi. Bu 50-jadvaldagi yutish qonuni bilan mos — "
         "operator xatosi 160 barobar o'zgarganda tarmoq xatosi atigi "
         "4 barobar o'zgaradi, ya'ni chuqur qatlamlardagi 1.0-1.7x lik "
         "ustunlikning yuvilishi kutilgan hol.", italic=True, size=10)
    para(doc, "Miss maqsadi nimani tasdiqladi va nimani yo'q.", bold=True, size=10)
    para(doc,
         "Maqsad funksiyasi chiqarilgan, kuzatilgan emas, shuning uchun u "
         "o'lchangan vaqt bilan bog'lanishi tekshirildi. 30-jadvaldagi "
         "kvantlangan armlar bo'yicha bayt va vaqt korrelyatsiyasi "
         "r = +0.974 (navbatlashgan o'lchov bo'yicha; blokli ustunda u "
         "+0.803 edi — bitta yomon nuqta korrelyatsiyani ham "
         "zaiflashtirgan). Tartib uchta juftlikda buziladi va uchalasi "
         "ham 2.0% dan kichik, ya'ni o'lchov shovqini darajasida. "
         "Ya'ni maqsad funksiyasining birinchi hadi — oqiziladigan bayt "
         "hajmi — o'lchangan vaqtni deyarli to'liq tartiblaydi.")
    para(doc,
         "Bu mashinada miss nisbati 1.05-1.08x, navbatlashgan qayta "
         "o'lchovdan keyingi tezlanish esa 1.06-1.14x — ya'ni ikkalasi "
         "BIR DIAPAZONDA va miss modeli kattaligi bo'yicha ham to'g'ri "
         "chiqadi (blokli o'lchovdagi 1.29x bilan bu moslik yo'q edi). "
         "Shunga qaramay tasdiqlangan narsa 'kam bayt -> kam vaqt' "
         "bog'lanishi, kesh REZIDENTLIGI mexanizmi emas: bu yerda hech "
         "bir qatlam byudjetdan chiqmaydi va tezlanishning katta qismi "
         "FLOP kamayishidan keladi (Memory Bound 9.7-18.2%, 39-jadval), "
         "bayt va FLOP esa olib tashlangan kanallar bilan birga "
         "kamayadi. 31-jadvaldagi 1.81x kabi bashoratlar overflow hadiga "
         "tayanadi va keshi to'lib ketmaydigan apparatda hosil qilib "
         "bo'lmaydi: ular ikkinchi mashinasiz tekshirilmagan bo'lib "
         "qoladi (5.4-bo'lim).", italic=True, size=10)
    para(doc, "Kesh hajmlari va chiplet ogohlantirishi.", bold=True, size=10)
    para(doc,
         "Jadvaldagi hajmlar ishlab chiqaruvchi hujjatlari bilan "
         "tekshirilgan; faqat 24 MiB qatori shu mashinada o'lchangan. "
         "Ulardan biri esa oddiy son emas. EPYC 7773X ning L3 i GLOBAL "
         "BO'LISHILMAYDI: sakkizta CCD ning har birida o'z 96 MiB i bor "
         "(32 MiB va ustiga taxlangan 64 MiB V-Cache), e'lon qilingan "
         "768 MiB esa ularning YIG'INDISI bo'lib, istalgan yadro "
         "foydalana oladigan umumiy hovuz emas. Ushbu usul bir guruh "
         "yadro KAFOLATLANGAN tarzda bo'lishadigan keshga bog'langani "
         "uchun 96 MiB faqat ishchi oqimlar bitta CCD da bo'lganda "
         "o'rinli. Bu chiplet arxitekturalari uchun umumiy holat: u yerda "
         "'L3' yagona narsa emas, va 2.2-bo'limdagi 'kafolatlangan umumiy "
         "kesh' ta'rifi aynan shuning uchun qo'yilgan. Taqqoslash uchun: "
         "Ryzen 7 5800X bitta CCD dan iborat, ya'ni uning 32 MiB i "
         "haqiqatan sakkizta yadroning hammasiga tegishli, bizning "
         "mashinamiz esa monolit.", size=9.5)

    h(doc, "4.9g. Freymvorkni uchdan-uchgacha yurgizish", 2)
    para(doc,
         "3.6-bo'lim freymvorkni tasvirlaydi; ushbu bo'lim uni YURGIZADI. "
         "Farq muhim, chunki tasvirlangan protsedura hech qachon "
         "bajarilmagunicha uning to'xtash qoidasi ishlashi tekshirilmagan "
         "bo'lib qoladi. Yurgizish uchta haqiqiy nuqsonni ochdi, va ularning "
         "hech biri kodni o'qib topilmagan bo'lardi.")
    table(doc, "34-jadval. Zinapoya bo'ylab yurish: L3 = 24 MiB, mutlaq "
               "byudjet 0.03, kalibrlash validation[100:106], tanlov "
               "validation[0:100].",
          ["#", "Pog'ona", "Vazn", "WER", "dWER (tayanchga)", "Qaror"],
          [["0", "o'zgartirishsiz (FP32)", "2915 MiB", "0.0961",
            "tayanch, zaxira +0.0300", "qabul"],
           ["1", "enkoder INT8", "2042 MiB", "0.1045",
            "+0.0083 [+0.0000, +0.0200]", "qabul"],
           ["2", "dekoder INT8", "737 MiB", "0.1036",
            "+0.0075 [+0.0017, +0.0148]", "qabul"],
           ["3", "enkoder qisqartirish tau=0.99", "704 MiB", "0.1039",
            "+0.0077 [+0.0013, +0.0163]", "qabul"],
           ["4", "enkoder qisqartirish tau=0.97", "699 MiB", "0.1060",
            "+0.0098 [-0.0054, +0.0236]", "QABUL, tanlandi"],
           ["5", "enkoder qisqartirish tau=0.95", "692 MiB", "0.1097",
            "+0.0135 [-0.0027, +0.0301]", "rad, to'xtadi"]],
          good_rows=(4,), bad_rows=(5,))
    para(doc,
         "Tanlangan konfiguratsiya — tau = 0.97 enkoder va INT8 dekoder. "
         "To'xtash haqiqiy aniqlik chegarasida yuz beradi, lekin uni ochiq "
         "qayd etish kerak: 0.0301 va ruxsat etilgan 0.0300 — chegara "
         "shovqin ichida, ya'ni boshqa bootstrap urug'i qarorni "
         "ag'darishi mumkin. Amalda bu tau = 0.97 va tau = 0.95 orasidagi "
         "tanlov shu byudjetda AJRATILMAYDI degani.")
    para(doc, "Birinchi nuqson: to'xtash qoidasi ishlamas edi.", bold=True,
         size=10)
    para(doc,
         "Dastlabki qoida har bir pog'onaning MUTLAQ WER ishonch chegarasini "
         "tayanchning NUQTAVIY bahosidan olingan shift bilan taqqoslardi. "
         "Birinchi yurgizishda o'zgartirilmagan modelning o'zi o'z "
         "byudjetidan chiqib ketdi: WER = 0.0961, byudjet 0.1009, yuqori "
         "chegara esa 0.1334. Sabab arifmetik — 100 namunada oraliqning "
         "yarim kengligi taxminan 0.035, 5% lik zaxira esa 0.0048, ya'ni "
         "test hech qachon o'ta olmaydi. Bu bir-biriga o'xshamagan "
         "kattaliklarni taqqoslash edi. To'g'ri savol — 'bu konfiguratsiya "
         "tayanchdan eps dan ko'proq yomonmi' — JUFTLIK taqqoslash bo'lib, "
         "namuna qiyinligi qisqaradi va farq ustidagi oraliq ancha tor "
         "bo'ladi; bu ishning qolgan hamma joyida ishlatiladigan mezonning "
         "aynan o'zi (4.9-bo'lim).")
    para(doc, "Ikkinchi nuqson: tayanch boshqa splitdan olinardi.",
         bold=True, size=10)
    para(doc,
         "Nisbiy byudjet konstanta sifatida berilgan tayanchdan hisoblanardi, "
         "u esa TEST splitida o'lchangan (0.1761). Tanlov esa validation da "
         "bo'ladi, u yerda FP32 tayanchi 0.0961 — deyarli ikki barobar "
         "farq, ya'ni bir xil 'eps = 0.05' ikki splitda butunlay boshqa "
         "chegarani anglatardi. Zinapoyaning birinchi pog'onasi "
         "o'zgartirilmagan modelning O'ZI bo'lgani uchun tayanch endi o'sha "
         "yerda o'lchanadi va byudjet tanlov splitiga nisbatan aniqlanadi.")
    para(doc, "Uchinchi nuqson: zinapoya ustunlik bilan yengilgan oilani "
              "sanardi.", bold=True, size=10)
    para(doc,
         "Rejalashtiruvchi strukturaviy pog'onalarni BIR XIL keep nisbati "
         "bilan sanardi, ya'ni har qatlamdan teng ulush olinardi. Bizning "
         "mezonimiz esa ortiqchalik qayerda bo'lsa, o'sha yerdan oladi. "
         "Ikkalasi bir xil splitda, bir xil tayanchga nisbatan o'lchandi:")
    table(doc, "35-jadval. Bir xil nisbat va mezon nuqtasi, "
               "validation[0:100], FP32 tayanchi 0.0961.",
          ["Konfiguratsiya", "Enkoder", "dWER (tayanchga)"],
          [["bir xil 10% qisqartirish", "281 MiB", "+0.0186 [+0.0055, +0.0342]"],
           ["bizniki tau = 0.99", "267 MiB", "+0.0077 [+0.0013, +0.0163]"]],
          good_rows=(1,))
    para(doc,
         "Farq -0.0108 [-0.0239, -0.0003], ya'ni SEZILARLI: mezon nuqtasi "
         "kichikroq va aniqroq, ikkala o'qda ham ustun. Bu 4.9d-bo'limdagi "
         "'taqsimot mezondan muhimroq' natijasining takrori, endi "
         "freymvorkning o'z zinapoyasiga qarshi. Nuqsonning aniq narxi ham "
         "o'lchandi: tuzatishdan oldin vosita 737 MiB tanlagan, keyin esa "
         "704 MiB — teng aniqlikda 33 MiB kichikroq. U buni o'z "
         "zinapoyasida bunday nomzod YO'QLIGI uchun o'tkazib yuborgan edi. "
         "Zinapoya endi mezonning o'lchangan ish nuqtalarini sanaydi "
         "(tau = 0.99 dan 0.90 gacha); bir xil nisbatlar yo'li mezoni "
         "tavsiflanmagan modellar uchun saqlanadi va kodda ochiq 'o'lchov "
         "bo'yicha yomonroq' deb belgilangan.")
    para(doc, "To'rtinchi nuqson: ulanmagan ishlov boshqasining yo'lini "
              "to'sardi.", bold=True, size=10)
    para(doc,
         "Zinapoya enkoder va dekoder qadamlarini navbatlashtirar edi, "
         "shuning uchun dekoderning birinchi strukturaviy qadami "
         "enkoderning IKKINCHISIDAN oldin turardi. Dekoderning kanal o'qi "
         "esa ulanmagan (uning FFN si bu kodda past-rank yoyilma bilan "
         "kichrayadi), va yalqov yurish o'sha yerda to'xtardi — natijada "
         "tau = 0.97 va undan keyingilari umuman sinalmagan, holbuki "
         "tau = 0.99 zaxiraning atigi yarmini ishlatgandi (+0.0163 va "
         "ruxsat +0.0300).")
    para(doc,
         "Yalqov to'xtash ANIQLIK muvaffaqiyatsizligi uchun asosli, chunki "
         "zinapoya monoton: buzilgan pog'onadan yuqoridagilar yo'qotilgan "
         "aniqlikni tiklay olmaydi. Ulanmaganlik uchun esa asossiz — "
         "implementatsiyaning yo'qligi aniqlik haqida hech narsa demaydi. "
         "Oddiy 'o'tkazib yuborish' ham yechim emas, chunki pog'onalar "
         "TO'PLANUVCHI holat: keyingi pog'ona {enkoder: tau=0.97, dekoder: "
         "qisqartirilgan} bo'lib, ulanmagan ishlov undan keyingi hamma "
         "pog'onada qoladi. Shuning uchun tuzatish rejalashtiruvchida: u "
         "qura olmaydigan konfiguratsiyani umuman SANAMAYDI. Bu belgi "
         "byudjet hisobidan alohida saqlanadi, chunki bajarilish tekshiruvi "
         "dekoder uchun 95% talabini baribir aytishi kerak.")
    para(doc,
         "Uchala tuzatishning to'plangan ta'siri o'lchandi: bir xil "
         "aniqlik byudjetida vosita 737 MiB dan 704 MiB ga, so'ngra "
         "699 MiB ga tushdi.")

    para(doc, "Ikkinchi model: freymvork mustaqil ravishda mBERT ga.",
         bold=True, size=10)
    para(doc,
         "Yuqoridagi hamma narsa Whisper da yurgizildi, va vosita uni uch "
         "joyda jimgina nazarda tutardi: modelda enkoder va dekoder bor, "
         "variantlarni aynan o'sha quruvchilar yasaydi, va sifat mezoni "
         "KICHIK yaxshi bo'ladigan kattalik. Uchtasi ham usulning xossasi "
         "emas, shuning uchun ular MODEL PROFILIGA ajratildi: profil "
         "qismlarni, quruvchilarni, baholovchini va metrikaning YO'NALISHINI "
         "beradi.")
    table(doc, "36-jadval. Xuddi shu freymvork mBERT ustida "
               "(L3 = 24 MiB, mutlaq byudjet 0.02, 400 matn).",
          ["#", "Pog'ona", "Vazn", "Yomonlashuv (juftlik)", "Qaror"],
          [["0", "o'zgartirishsiz (FP32)", "324 MiB",
            "tayanch, zaxira +0.0200", "qabul"],
           ["1", "enkoder INT8", "81 MiB",
            "+0.0092 [-0.0009, +0.0193]", "QABUL, tanlandi"],
           ["2", "enkoder qisqartirish 10% kanal", "254 MiB",
            "+0.0230 [+0.0101, +0.0358]", "rad, to'xtadi"]],
          good_rows=(1,), bad_rows=(2,))
    para(doc,
         "Vosita INT8 ni tanlaydi va qisqartirishni rad etadi — 43-jadvalda "
         "diagnostika asosida chiqarilgan hukmning aynan o'zi, endi esa "
         "qidiruv orqali, o'lchov bilan qayta topilgan. mBERT bitta qismdan "
         "iborat, uning metrikasi KATTA yaxshi bo'ladigan aniqlik, va ikkala "
         "farq ham profilda e'lon qilinadi, kodda emas.")
    para(doc,
         "Metrikaning yo'nalishi shu yerda alohida ahamiyat kasb etadi. "
         "Chegara 'yomonlashuv' tomoniga siljishi kerak, ya'ni so'z "
         "xatoligida YUQORIGA, aniqlikda esa PASTGA. Bu shartni byudjetni "
         "hisoblashning bir yo'lida hisobga olib, ikkinchisida "
         "(o'lchanadigan tayanch yo'lida) tashlab ketish mumkin — va shunday "
         "bo'lgan ham edi: chegara o'zgartirilmagan modeldan yuqorida "
         "turgani uchun ruxsat etilgan zaxira MANFIY chiqdi, har bir "
         "pog'ona rad etildi va vosita siqilmagan modelni qaytardi. Hech "
         "qanday xato xabari bo'lmadi, faqat noto'g'ri javob; yo'nalish "
         "endi profildan olinadi va ikkala yo'lda ham bir joydan "
         "kelib chiqadi.", italic=True, size=10)

    para(doc, "Uchinchi model: open_llama_3b, ONNX grafigisiz.",
         bold=True, size=10)
    para(doc,
         "Llama profili freymvorkning yana bir chegarasini sinaydi: bu "
         "modelning ONNX eksporti yo'q, qatlam o'lchamlari arxitektura "
         "konfiguratsiyasidan aniq hisoblanadi (Llama bloki uchun "
         "shakllar konfiguratsiyaning o'zi), qurish va baholash esa "
         "PyTorch orqali. L3 = 24 MiB, nisbiy byudjet 5%, baholash "
         "WikiText-2 ning 24 segmentida: tayanch 7.5466 (nashr etilgan "
         "FP32 qiymati 7.547 bilan mos — trakt tasdiqlangan), INT8 "
         "pog'onasi 7.5491 bilan qabul qilinadi (juftlik farqi +0.0003 "
         "[+0.0001, +0.0006] — zaxiraning 0.08% i). Mezon-asosli tau "
         "pog'onalari esa amalda hech narsa olmaydi: tau = 0.99 da 26 "
         "qatlamning jami BITTA kanali (8640 tadan) olib tashlanadi — "
         "4.15-bo'limdagi geometrik diagnostika bilan to'liq mos. Ya'ni "
         "freymvork bu modelda kesishni o'zi rad etmaydi — MEZON rad "
         "etadi, freymvork esa buni majburlamasdan hurmat qiladi. "
         "Yurish yakunigacha yetkazildi va yana BITTA — beshinchi — "
         "nuqsonni ochdi: juftlik farqi segment NLL fazosida, zaxira esa "
         "perplexity fazosida hisoblanardi, ya'ni 5% byudjet 8 barobar "
         "bo'shashgan (0.3773 va ln(1.05) = 0.0488) va yurish o'z "
         "chegarasidan ko'rinib turib chiqqan pog'onani (8.0687 > "
         "7.9240) qabul qilgan. Score chiziqli bo'lgan metrikalarda "
         "(WER, aniqlik) bu nomuvofiqlik yo'q; exp(o'rtacha NLL) esa "
         "nochiziqli va profil endi konvertatsiyani o'zi beradi. "
         "Tuzatilgan qoida bilan yakuniy hukm: tau pog'onalari qabul "
         "(+0.0003, +0.0021, +0.0071), majburiy 10% rad (+0.0669 > "
         "+0.0488), TANLANDI — tau = 0.90 (perplexity 7.6001). Ya'ni "
         "mezon topgan ozgina ortiqchalik (0.07%) byudjet ichida tekin, "
         "undan bir qadam nari esa chegaradan chiqadi; majburiy egri "
         "chiziq to'liq o'lchandi (10% - 8.0687, 20% - 8.9060, ikkinchisi "
         "mustaqil o'lchov bilan aynan mos). Uch arxitekturada uch xil "
         "yakun: Whisper 'kvantla va mezon nuqtalarigacha kes', mBERT "
         "'kvantla va to'xta', Llama 'kvantla va mezon ruxsat "
         "berganchagina kes' — hammasi bitta freymvork, uch profil "
         "bilan.")
    para(doc, "TEST splitida yakuniy tekshiruv: sertifikatlash "
              "chegarasi.", bold=True, size=10)
    para(doc,
         "Whisper yurishi TEST splitining 300 namunasida, 5% nisbiy "
         "byudjet bilan qayta yurgizildi. Tayanch 0.1761 (ma'lum TEST "
         "qiymati bilan aynan mos), ruxsat +0.0088. Birinchi pog'ona "
         "(enkoder INT8, WER 0.1838) nuqtaviy bahoda sig'adi (+0.0077), "
         "yuqori chegarada esa yo'q (+0.0162 > +0.0088) — yurish uni rad "
         "etdi va O'ZGARTIRILMAGAN modelni qaytardi. Bu 29-jadvalning "
         "mustaqil takrori: n = 300 da juftlik oralig'i 5% byudjetni "
         "sertifikatlashga yetmaydi va vosita buni yashirmasdan aytadi. "
         "Nuqtaviy o'quvchi bu yerda INT8 ni 'tasdiqlangan' deb "
         "qaytargan bo'lardi; ishonch-chegarali qoida esa to'g'ri "
         "savolni qaytaradi — kattaroq byudjet yoki kattaroq namuna "
         "keltiring. Shu bilan 'byudjet namunaga mos bo'lishi kerak' "
         "xossasi ikki mustaqil yo'lda o'lchangan bo'ldi.",
         italic=True, size=10)
    para(doc, "Monotonlik auditi: yalqov to'xtash qachon xavfsiz.",
         bold=True, size=10)
    para(doc,
         "Zinapoyaning yalqov to'xtashi bitta farazga tayanadi: buzilgan "
         "pog'onadan yuqoridagilar yo'qotilgan aniqlikni tiklay olmaydi. "
         "Bu faraz NUQTAVIY baholarda buziladi va buni o'z ma'lumotimiz "
         "ko'rsatadi: validation skanerlashida tau = 0.97 tau = 0.99 dan "
         "yaxshi chiqqan (0.0967 va 0.1013, 27-jadval) — zinapoya "
         "bo'ylab nuqtaviy egri monoton emas. Yalqov to'xtashni "
         "baribir xavfsiz qiladigan narsa TO'XTASH QOIDASINING o'zi: u "
         "nuqtaviy baho yomonlashganda emas, juftlik farqining YUQORI "
         "CHEGARASI byudjetdan chiqqanda to'xtaydi. Kuzatilgan "
         "monoton-emaslik amplitudasi (0.0046) juftlik oralig'idan "
         "(+-0.019) ancha kichik, ya'ni u shovqin ko'rinishi va "
         "chegarani buza olmaydi. Bajarilgan to'rtta yurishning "
         "(Whisper validation, Whisper TEST, mBERT, Llama) birortasida "
         "sezilarli buzilishdan KEYIN byudjetga sig'adigan pog'ona "
         "uchramagan; sezilarli buzilishlar esa hamma yurishda keskin "
         "va katta (masalan Llama da +0.0669 chegara +0.0488 ga qarshi, "
         "mBERT da +0.0230). Demak faraz aniqlashtiriladi: zinapoya "
         "SEZILARLI-farq fazosida monoton bo'lishi kifoya, va bu "
         "kuchsizroq shart o'lchovlarda buzilmagan. Buzilsa nima "
         "bo'lardi ham aniq: to'xtash bitta pog'ona kechroq sodir "
         "bo'lib, tanlov o'zgarmasdi — qoida konservativ tomonga "
         "xato qiladi.", italic=True, size=10)
    para(doc, "Foydalanish xossasi: byudjet namunaga mos bo'lishi kerak.",
         bold=True, size=10)
    para(doc,
         "Yuqoridagi birinchi nuqsondan umumiy qoida kelib chiqadi. Zaxira "
         "namunada ajratib bo'ladigan farqdan kichik bo'lsa, hech bir "
         "pog'ona byudjet ichida ekani ISBOTLANMAYDI va vosita "
         "o'zgartirilmagan modelni qaytaradi. Bu qo'yilgan savolga to'g'ri "
         "javob, lekin deyarli hech qachon nazarda tutilgan savol emas, "
         "shuning uchun vosita buni ochiq ogohlantirish bilan aytadi. "
         "Ogohlantirish JUFTLIK farqining oralig'idan hisoblanishi kerak: "
         "birinchi variant mutlaq WER tarqoqligidan foydalangan, u esa bir "
         "necha barobar keng, va natijada vosita 'hech narsa isbotlanmaydi' "
         "deb ogohlantirgan holda pog'onalarni qabul qilishda davom etgan. "
         "Ajratish chegarasi taxminan 1/sqrt(n) kabi kamayadi: shu ishdagi "
         "100 namunada juftlik oralig'ining yarim kengligi 0.011 atrofida, "
         "ya'ni 0.0048 lik zaxira yetarli emas, 0.0300 esa yetarli.")

    h(doc, "4.10. Tezlik va apparat hisoblagichlari", 2)
    para(doc, "Kesh misslari va tezlik: o'lchangan bog'lanish zanjiri.",
         bold=True, size=10)
    para(doc,
         "Kirishda qo'yilgan xotira-devori mantig'i — miss hajmi "
         "kamaysa, vaqt kamayadi — bu bo'limda OLTI mustaqil o'lchov "
         "bilan tasdiqlanadi, va ularni oldindan bir joyga yig'ib "
         "qo'yamiz. (1) Ko'chiriladigan baytlar va o'lchangan vaqt "
         "orasidagi korrelyatsiya kvantlangan konfiguratsiyalar bo'ylab "
         "r = +0.974, tartib faqat uchta 2% dan kichik juftlikda buzilgan, ya'ni to'liq"
         "mos (4.9f). (2) Kaskaddan keyin xotira to'xtashlari umumiy "
         "vaqtdan TEZROQ qisqaradi — 2.41x ga qarshi 1.91x — ya'ni "
         "model kamroq xotira bilan cheklangan holga o'tadi, bu miss "
         "kamayishining apparat hisoblagichlaridagi bevosita izi. "
         "(3) L3 bosimi 2.4% dan 1.0% ga tushadi. (4) Dekoder INT8 ga "
         "o'tganda ish to'plami DRAM dan L3 ga KO'CHADI (DRAM Bound "
         "9.9% dan 6.6-7.1% ga, L3 Bound 2.6% dan ko'tariladi) — "
         "kichrayРіР°РЅ vazn chuqurroq kesh darajasida yashay boshlaydi. "
         "(5) Qayta ishlatishi past dekoder qayta ishlatishi yuqori "
         "enkoderdan 1.9 barobar ko'proq xotira bilan cheklangan "
         "(18.2% va 9.7%) — miss narxi aynan trafik ko'p joyda katta. "
         "(6) Vazn matritsasini plitkalamaydigan yadroda byudjetdan "
         "chiqish 1.56-2.3x jarima beradi (4.10a) — rezidentlik "
         "effektining o'zi ham real, faqat sozlangan GEMM uni plitka "
         "darajasida yashiradi. Bulardan birinchi beshtasi miss "
         "HAJMI orqali, oltinchisi rezidentlik orqali ishlaydi; "
         "ikkalasi ham bitta printsipning ko'rinishi: protsessor "
         "hisoblashni kutmaydi, MA'LUMOTNI kutadi.")
    table(doc, "37-jadval. INT8 ustiga past-rank qo'shishning latencyga ta'siri "
               "(enkoder fc1, 1500 pozitsiya).",
          ["Variant", "Vaznlar (MiB)", "Latency (ms)", "FP32 ga", "INT8 ga", "E_loc"],
          [["zich FP32", "16.00", "117.9", "1.00x", "0.26x", "0"],
           ["zich INT8", "4.00", "30.3", "3.89x", "1.00x", "0.0082"],
           ["INT8 + SVD r=200", "0.98", "8.8", "13.40x", "3.44x", "0.0201"],
           ["INT8 + SVD r=128", "0.62", "6.7", "17.70x", "4.56x", "0.0305"]],
          good_rows=(3,))
    para(doc,
         "Bu jadvaldagi ranklar SIQISH KOEFFITSIYENTI bo'yicha qo'lda "
         "tanlangan (r = 409 aynan ikki barobar parametr siqishga to'g'ri "
         "keladi), chunki maqsad rankni chiqarish emas, rank pasayishining "
         "latencyga ta'sirini o'lchash edi. Chiqarilgan ranklar 15-jadvaldagi "
         "byudjet-optimal taqsimot artefaktlarida ishlatiladi.", italic=True,
         size=10)
    table(doc, "38-jadval. VTune apparat hisoblagichlari (enkoder fc1).",
          ["Variant", "ms/iter", "Memory bound", "L2", "L3", "DRAM", "CPI"],
          [["zich FP32", "121.85", "8.8%", "2.5%", "2.7%", "2.9%", "0.64"],
           ["zich INT8", "33.96", "12.7%", "1.9%", "2.4%", "6.5%", "0.67"],
           ["INT8 + SVD r=128", "7.62", "18.3%", "4.1%", "1.0%", "9.2%", "0.62"]],
          good_rows=(2,))
    para(doc, "Butun model darajasidagi hisoblagichlar.", bold=True, size=10)
    para(doc,
         "Kaskadning assimetrik qarori qayta ishlatish argumentiga tayanadi: "
         "enkoderda har vazn bir o'tishda 1500 marta, dekoderda esa har token "
         "uchun bir marta ishlatiladi va keyingi 23 qatlam uni siqib "
         "chiqaradi. Shu paytgacha bu ANALITIK dalil edi. 39-jadval uni "
         "bevosita o'lchaydi: ikkala model ham o'z tabiiy rejimida, bitta "
         "oqimda, 38-jadval bilan bir xil protokolda profillanadi.")
    table(doc, "39-jadval. Butun model apparat hisoblagichlari (bitta oqim, "
               "uarch-exploration). Kvadrat qavsdagi son — hajm MiB da; bir xil "
               "byudjetli variantlar yonma-yon turadi.",
          ["Variant", "MiB", "ms/iter", "Memory bound", "L3", "DRAM", "CPI"],
          [["enkoder FP32", "1172", "12120", "12.0%", "1.6%", "4.9%", "0.649"],
           ["enkoder INT8 per-tensor", "299", "7062", "9.8%", "1.3%", "3.6%", "0.469"],
           ["enkoder RTN per-channel", "300", "7090", "9.5%", "1.6%", "3.7%", "0.469"],
           ["enkoder GPTQ", "300", "7081", "9.8%", "1.0%", "4.1%", "0.468"],
           ["enkoder qisqartirish + RTN", "267", "6758", "9.2%", "1.2%", "3.5%", "0.462"],
           ["enkoder kaskad (qisqartirish + GPTQ)", "267", "6728", "9.6%", "1.6%",
            "3.4%", "0.460"],
           ["enkoder past-rank, bir xil taqsimot", "203", "6295", "8.9%", "1.6%",
            "3.6%", "0.450"],
           ["enkoder past-rank, optimal taqsimot", "203", "6301", "8.6%", "1.3%",
            "3.0%", "0.451"],
           ["dekoder FP32", "1743", "1620", "18.8%", "2.6%", "9.9%", "0.665"],
           ["dekoder INT8 per-tensor (kaskad tanlovi)", "438", "480.4", "17.7%",
            "8.2%", "6.6%", "0.627"],
           ["dekoder INT8 per-channel", "439", "469.2", "19.3%", "5.0%", "8.1%",
            "0.622"],
           ["dekoder INT8 + past-rank (rad etilgan)", "343", "463.9", "18.1%",
            "5.3%", "6.8%", "0.614"]],
          good_rows=(5, 9), bad_rows=(11,))
    para(doc,
         "Birinchidan, qayta ishlatish argumenti tasdiqlanadi. Siqilgan "
         "variantlar bo'yicha o'rtacha Memory Bound enkoderda 9.7%, dekoderda "
         "18.2% — dekoder taxminan 1.9 barobar ko'proq xotira bilan cheklangan, "
         "DRAM Bound ham shu yo'nalishda (3.0-4.1% ga qarshi 6.6-8.1%).")
    para(doc,
         "Ikkinchidan va hal qiluvchi tarzda, past-rank dekoderda VAQT "
         "BERMAYDI: 438 dan 343 MiB ga (-22% xotira) o'tish 480.4 dan 463.9 ms "
         "ga olib keladi, ya'ni farq o'lchov aniqligi darajasida. Xuddi shu "
         "o'zgarish enkoderda 6728 dan 6301 ms ga (-6.3%) tushiradi.")
    para(doc, "Teng byudjetda usullar ajraladimi?", bold=True, size=10)
    para(doc,
         "Siqilgan modelni faqat FP32 ga qarshi qo'yish noto'g'ri savolga javob "
         "beradi: 4 barobar kam bayt ko'chirish xotira to'xtashini albatta "
         "kamaytiradi. Usullar orasida tanlovni hal qiladigan savol boshqacha — "
         "XUDDI SHU hajmda ulardan biri xotira ierarxiyasida yaxshiroq "
         "ishlaydimi? 39-jadvalda uchta enkoder guruhi (299-300, 267 va 203 "
         "MiB) va bitta dekoder juftligi (438-439 MiB) aynan shu maqsadda "
         "byudjet bo'yicha moslashtirilgan.")
    para(doc,
         "Javob: yo'q. Guruh ichidagi xotira to'xtashi tarqoqligi mos ravishda "
         "3.0%, 3.9%, 3.4% va 6.5% ni tashkil qiladi. Bu qiymatlarni "
         "baholash uchun o'lchovning o'z aniqligi zarur, shuning uchun beshta "
         "model ikki marta profillandi: yugurishlararo farq 1.9% dan 7.2% "
         "gacha. Ya'ni guruh ichidagi BARCHA farqlar o'lchov "
         "o'zgaruvchanligi darajasida yoki undan past va usulga "
         "bog'lanmaydi.")
    para(doc,
         "Bundan aniq xulosa chiqadi: xotira xatti-harakati o'sha baytlarni "
         "qaysi algoritm hosil qilganiga emas, BAYTLAR SONIGA bog'liq "
         "(300 MiB -> ~690 ms, 267 MiB -> ~630 ms, 203 MiB -> ~550 ms). "
         "Shuning uchun kaskadning raqiblar oldidagi xotira ustunligi aynan "
         "uning hajm ustunligi (267 ga qarshi 300 MiB) va undan ortiq emas. "
         "Buni ochiq aytish muhim, chunki teskari da'vo — 'taklif etilgan usul "
         "kesh bilan yaxshiroq ishlaydi' — o'lchov bilan qo'llab-quvvatlanmaydi. "
         "Enkoder va dekoder orasidagi farq esa (9.7% ga qarshi 18.2%, ya'ni "
         "1.9 barobar) shovqindan ancha yuqori va o'z kuchida qoladi.",
         italic=True, size=10)
    para(doc,
         "Shu bilan kaskadning dekoderdagi rad etishi IKKALA o'qda ham "
         "oqlanadi: past-rank u yerda 22% xotira tejaydi, ammo vaqtdan hech "
         "narsa bermaydi va 19-jadvalga ko'ra 0.43 WER turadi — ya'ni qat'iy "
         "yutqazuvchi variant. Chiqarilgan maqsad buni operatorlarni ishga "
         "tushirmasdan oldin aytgan edi; hisoblagichlar esa sababini "
         "ko'rsatadi.")
    para(doc,
         "Memory Bound ulushini yolg'iz o'qish mumkin emas, chunki u "
         "maxraji o'zgaruvchan NISBAT: variant xotira to'xtashini mutlaq "
         "ravishda kamaytirib ham, umumiy vaqt undan tezroq qisqargani uchun "
         "yuqoriroq foiz ko'rsatishi mumkin. Dekoderda aynan shu yuz beradi — "
         "ulush 18.8% dan 19.4% ga chiqadi, mutlaq xotira to'xtashi esa 304.5 "
         "dan 87.5 ms ga, ya'ni 3.48 barobar tushadi. Butun model uchun ikkala "
         "ko'rinish quyidagicha:")
    table(doc, "40-jadval. Xotira to'xtashi: ulush va mutlaq vaqt "
               "(enkoder + dekoder, bitta oqim).",
          ["Konfiguratsiya", "Umumiy (ms)", "Xotira to'xtashi (ms)", "Ulush"],
          [["FP32", "13740", "1759", "12.8%"],
           ["kaskad", "7209", "731", "10.1%"],
           ["o'zgarish", "1.91x kamaydi", "2.41x kamaydi", "-2.7 p.p."]],
          good_rows=(1,))
    para(doc,
         "Xotira to'xtashi umumiy vaqtdan TEZROQ qisqaradi (2.41x ga qarshi "
         "1.91x), ya'ni kaskaddan keyingi model nisbatan kamroq xotira bilan "
         "cheklangan holga o'tadi. Bu kesh-bog'langan maqsadning bevosita "
         "kutilgan natijasi.")
    para(doc,
         "Dekoderdagi ichki taqsimot ish to'plamining ierarxiya bo'ylab "
         "ko'chishini ko'rsatadi: FP32 dan INT8 ga o'tganda DRAM Bound 9.9% dan "
         "tushadi, L3 Bound esa 2.6% dan ko'tariladi, ya'ni vaznlar DRAM dan "
         "L3 ga ko'chadi. Yo'nalish ikkala mustaqil yugurishda ham bir xil "
         "(DRAM 7.1% va 6.6%; L3 5.0% va 8.2%), ammo KATTALIK barqaror emas — "
         "L3 uchun ikki o'lchov 5.0% va 8.2% beradi. Shu sababli bu ko'chishni "
         "sifat jihatdan qayd etamiz, miqdoriy da'vo sifatida emas. Enkoderda "
         "bunday ko'chish kuzatilmaydi (DRAM 4.9% -> 3.4-4.1%, L3 1.6% -> "
         "1.0-1.6%), chunki u allaqachon hisoblash bilan cheklangan.")

    para(doc,
         "L3 bosimi 2.4% dan 1.0% ga tushadi, ya'ni kesh foydasi real va bilvosita "
         "emas, bevosita kuzatilgan. Ammo xotira bilan bog'liq to'xtashlar konveyer "
         "slotlarining atigi 9-18% ini tashkil qiladi, shuning uchun 4.46x "
         "tezlanishning asosiy mexanizmi keshda saqlanish emas, arifmetikaning 6.40x "
         "kamayishi. Shu sababli effektni compute-bound rejimda arifmetik hajmni "
         "strukturaviy kamaytirish deb ta'riflaymiz.")

    h(doc, "4.10a. Kesh-sig'imi mexanizmini to'g'ridan-to'g'ri sinash: "
           "salbiy natija", 2)
    para(doc,
         "Ishning chiqarish qismi bitta farazga tayanadi: operator vazni "
         "alpha*L3 ga sig'masa, xotira jarimasi paydo bo'ladi. Shu paytgacha "
         "bu faraz TEKSHIRILMAGAN edi va tekshirib bo'lmasdi ham — Whisper "
         "enkoderining eng katta qatlami INT8 dan keyin 12.1 MiB, byudjet "
         "esa 16.8 MiB, ya'ni model hech qachon o'sha rejimga kirmaydi. "
         "Ammo rejim OPERATORNING xossasi, faqat mashinaning emas: byudjetdan "
         "oshadigan operator kichikroq kesh Whisper ni qo'yadigan holatning "
         "aynan o'zida turadi. Shuning uchun byudjetni ikki tomondan qamrab "
         "oluvchi kvadrat operatorlar to'plami o'lchandi.")
    table(doc, "41-jadval. MAC boshiga vaqt vazn hajmiga qarab "
               "(1500 pozitsiya, byudjet 16.8 MiB).",
          ["Vazn (MiB)", "fp32, 1 oqim (ns/MAC)", "INT8, navbatma-navbat (ns/MAC)",
           "Byudjet"],
          [["1.00", "0.0185", "—", "ichida"],
           ["2.25", "0.0183", "—", "ichida"],
           ["4.00", "0.0187", "0.0050", "ichida"],
           ["9.00", "0.0189", "—", "ichida"],
           ["16.00", "0.0177", "0.0049", "ichida"],
           ["25.00", "0.0176", "0.0050", "TASHQARIDA"],
           ["36.00", "0.0177", "0.0051", "TASHQARIDA"],
           ["64.00", "0.0180", "—", "TASHQARIDA"]])
    para(doc,
         "Chiziq tekis. Vazn byudjetdan to'rt barobar oshib ketganda ham MAC "
         "boshiga vaqt o'zgarmaydi: fp32 da byudjetdan tashqaridagi mediana "
         "ichkaridagidan 0.95x, ya'ni hatto arzimas darajada TEZROQ. INT8 "
         "yadrosida — modellar aslida ishlatadigan yadroda — navbatma-navbat "
         "o'lchov 1.02x [10-90%: 1.00-1.05] va 1.02x [1.01-1.05] beradi, "
         "ya'ni jarima bor, lekin u 2%. Miss ifodasi esa o'sha baytlarga "
         "qayta ishlatish koeffitsiyentini (bu yerda 1499) ko'paytiradi.")
    para(doc,
         "Sabab tushunarli: bloklangan GEMM vazn matritsasini plitkalarga "
         "bo'ladi va keshda turishi kerak bo'lgan narsa PLITKA, butun "
         "matritsa emas. 'Vazn alpha*L3 ga sig'ishi kerak' degan shart "
         "yaxshi yozilgan GEMM uchun operatorning umumiy hajmiga bog'liq "
         "emas.")
    para(doc, "Metodologik ogohlantirish.", bold=True, size=10)
    para(doc,
         "Dastlabki o'lchov buning aksini ko'rsatgandi. Sakkiz oqimda "
         "hajmlar o'sish tartibida skanerlanganda byudjet kesishuvida "
         "1.74x va 2.18x lik keskin sakrash chiqdi — ishonarli tizza. "
         "Takrorlash uni tasdiqlamadi: o'sha 25 MiB operator 0.0055, keyin "
         "0.0032 ns/MAC berdi va ikkinchi yugurishdagi HAR BIR nuqta bir "
         "xilda tezroq edi. Bu mashina yuki, kesh emas. Tartiblangan "
         "skanerlash hajm effektini drift dan ajrata olmaydi, chunki "
         "ikkalasi birga o'sadi; navbatma-navbat o'lchov (A B C A B C) esa "
         "ajratadi va u tizza topmaydi (0.98x va 1.00x). Sakkiz oqimda shu "
         "mashinada raundlar orasidagi tarqoqlik 68-96% ni tashkil qiladi, "
         "ya'ni 1.3x dan kichik effektni aniqlab bo'lmaydi; ishonchli dalil "
         "bitta oqimdan va INT8 yadrosidan keladi, ularda tarqoqlik mos "
         "ravishda 15% va 2-15%.", italic=True, size=10)
    para(doc, "Sabab isboti: tizza yadroning bloklashiga bog'liq.",
         bold=True, size=10)
    para(doc,
         "Yuqoridagi 'bloklangan GEMM' izohi atributsiya edi, isbot emas "
         "— va u rad etilishi mumkin bashorat beradi: xuddi shu "
         "vazn-o'lchov qatorini vazn matritsasini PLITKALAMAYDIGAN "
         "yadrodan o'tkazsak, tizza PAYDO BO'LISHI kerak. Sodda yadro "
         "sifatida har chaqiruvda butun matritsani oqizadigan "
         "bo'lakli-matmul olindi (8 pozitsiyalik bo'laklar — vazn "
         "chaqiriqlar orasida keshda faqat SIG'SAGINA issiq qoladi), "
         "o'lchov o'sha interleaved intizomda:")
    table(doc, "42-jadval. Sodda (bloklanmagan) yadroda tizza paydo "
               "bo'ladi (MAC boshiga ns, mediana, 5 raund interleaved).",
          ["Vazn (MiB)", "4", "9", "16", "25", "36", "64"],
          [["sodda yadro", "0.0211", "0.0149", "0.0169", "0.0170",
            "0.0264", "0.0382"],
           ["byudjetga nisbatan", "ichida", "ichida", "ichida",
            "TASHQARIDA", "TASHQARIDA", "TASHQARIDA"]],
          bad_rows=())
    para(doc,
         "Bashorat tasdiqlandi: sodda yadroda tashqari/ichkari nisbati "
         "1.56x, 64 MiB da esa 2.3x gacha — gradatsiyali, sig'im "
         "misslari bilan mos o'sish; bloklangan yadroda o'sha qatorda "
         "1.02x. Shu bilan salbiy natija CHEGARALANGAN QONUNGA "
         "aylanadi: vazn-rezidentlik jarimasi apparatning emas, YADRO "
         "BLOKLASH STRATEGIYASINING xossasi. Miss ifodasining overflow "
         "hadi sodda yadrolar sohasida (masalan, bloklash "
         "optimallashuvisiz maxsus yoki embedded yadrolar) amal qiladi "
         "va sozlangan BLAS/ONNX Runtime sohasida amal qilmaydi. "
         "Freymvork joylashtirish muhitining yadro sinfini bilsa, "
         "overflow hadini yoqish yoki o'chirishni shu o'lchov "
         "asosida tanlashi mumkin.")
    para(doc, "Bu nimani rad etadi va nima kuchida qoladi.", bold=True, size=10)
    para(doc,
         "RAD ETILADI: miss ifodasining overflow hadi va uning ortidagi "
         "mexanizm. Butun vazn matritsasining keshda turishi shu apparatda, "
         "ikkala yadroda ham o'tkazuvchanlikni belgilamaydi. 31-jadvaldagi "
         "1.26-1.81x lik bashoratlar aynan shu hadga tayanadi, ya'ni ular "
         "TASDIQLANMAGAN model natijalari bo'lib qoladi.")
    para(doc,
         "KUCHIDA QOLADI: kaskadning qarorlari, chunki ular mustaqil "
         "ravishda uchdan-uchgacha tasdiqlangan — tau = 0.99 "
         "ko'r-ko'rona INT8 dan 11% kichik, 1.06x tez va aniqligi undan "
         "farqlanmaydi (30-jadval), kesh maqsadini majburlash esa "
         "0.1028 WER turadi (33-jadval). "
         "Tezlanishning xotira kanali esa REZIDENTLIK emas, MISS HAJMI "
         "orqali ishlashda davom etadi: baytlar kamayishi bilan miss "
         "hajmi kamayadi va vaqt shunga ergashadi (r = +0.974), xotira "
         "to'xtashlari 2.41x qisqaradi. Ushbu bo'lim aniqlashtirgan "
         "narsa — bu kanalning QAYERDAN o'tishi: bloklangan yadrolarda "
         "u faqat trafik hajmi orqali, sodda yadrolarda esa trafik + "
         "rezidentlik orqali.")
    para(doc,
         "Shuning uchun markaziy da'vo endi to'liq o'lchovga tayanadi: "
         "kesh MISSLARINI kamaytirish tezlikka bevosita ta'sir qiladi "
         "(olti dalilli zanjir, 4.10-bo'lim boshi), kesh HAJMI esa "
         "qaror o'zgaruvchisi — qayerda to'xtashni aytadi va to'xtash "
         "nuqtalari to'g'ri chiqadi. Rad etilgani ikkalasi ham emas, "
         "faqat bloklangan yadrolardagi keskin rezidentlik chegarasi.")
    para(doc, "Sinovning qamrovi.", bold=True, size=10)
    para(doc,
         "O'lchov sintetik kvadrat operatorlarda, ONNX Runtime da va bitta "
         "mashinada o'tkazildi. U mexanizm haqidagi umumiy da'voni sinaydi, "
         "31-jadvaldagi aniq sonlarni emas. Boshqa runtime, boshqa "
         "bloklash strategiyasi yoki ko'p jarayonli raqobat sharoitida "
         "natija boshqacha bo'lishi mumkin, va bu ikkinchi platformaga "
         "bo'lgan ehtiyojni kamaytirmaydi — aksincha, kuchaytiradi.",
         size=9.5)

    h(doc, "4.11. Arxitekturalararo ko'chish", 2)
    table(doc, "43-jadval. Uchta arxitekturada ortiqchalik va kesh hukmi.",
          ["Model", "Ortiqchalik (tau=0.99)", "Cho'qqi", "3-holatdagi operator"],
          [["Whisper enkoder", "17.1%", "58.0%", "yo'q (FFN INT8 dan keyin sig'adi)"],
           ["open_llama_3b", "0.6%", "3.4%", "gate/up/down 1.57x, head 5.81x"],
           ["mBERT", "0.1%", "0.7%", "lug'at matritsasi 5.21x"]])
    para(doc, "mBERT uchun vazifa metrikasi.", bold=True, size=10)
    para(doc,
         "43-jadval mBERT ni faqat DIAGNOSTIKA orqali tavsiflaydi: mezon "
         "unda deyarli ortiqchalik topmaydi (tau = 0.99 da 0.1%, cho'qqida "
         "0.7%). Whisper so'z xatoligi bilan, open_llama_3b esa perplexity "
         "bilan baholanadi; mBERT esa ikkalasi bilan ham emas, ya'ni "
         "'kaskadning HUKMI ko'chadi' degan da'vo shu model uchun sifat "
         "o'lchoviga tayanmagan edi. Bu bo'shliq quyidagicha yopiladi.")
    para(doc,
         "Metrika — held-out o'zbek matnida niqoblangan tokenni bashorat "
         "qilish aniqligi va o'sha bashoratlarning pseudo-perplexity si. U "
         "belgilangan ma'lumot ham, moslashtirilgan bosh ham talab qilmaydi "
         "va Llama uchun ishlatilgan perplexityning bevosita analogi. Kaskad "
         "mBERT uchun aniq bashorat qiladi — strukturaviy o'q hech narsa "
         "bermaydi, demak kvantlash tekin bo'lishi va kanal olib tashlash "
         "qimmatga tushishi kerak. Bashoratning ikkala yarmi ham "
         "tekshiriladi; niqoblash urug'i qat'iy, ya'ni har uch variant "
         "AYNAN bir xil pozitsiyalarda baholanadi va taqqoslash juftlik "
         "bo'ladi.")
    table(doc, "44-jadval. mBERT masked-LM, 2218 niqoblangan pozitsiya "
               "(kalibrlash matnidan ajratilgan o'zbek matni).",
          ["Variant", "MiB", "Aniqlik", "pseudo-PPL",
           "FP32 ga nisbatan (juftlik)"],
          [["FP32", "1029", "0.2656", "93.22", "—"],
           ["INT8", "259", "0.2633", "95.24",
            "-0.0023 [-0.0090, +0.0050] farqlanmaydi"],
           ["20% kanal, kosinus + INT8", "248", "0.2376", "103.22",
            "-0.0280 [-0.0397, -0.0167] SEZILARLI"],
           ["20% kanal, fluktuatsiya + INT8", "248", "0.2471", "109.91",
            "-0.0185 [-0.0293, -0.0081] SEZILARLI"]],
          good_rows=(1,), bad_rows=(2, 3))
    para(doc,
         "Ikkala yarmi ham tasdiqlandi. INT8 to'rt barobar siqishni beradi "
         "va aniqlikda FP32 dan FARQLANMAYDI. Kanal olib tashlash esa "
         "INT8 ustiga qo'shilganda aniqlikni sezilarli tushiradi "
         "(-0.0257 [-0.0365, -0.0153]) va buning evaziga atigi 11 MiB "
         "tejaydi — ya'ni o'n bir megabayt uchun 2.6 punkt. Whisper "
         "enkoderida xuddi shu strukturaviy o'q TEKIN edi (15-jadval); farq "
         "modelda, usulda emas, va kaskad bu farqni qurishdan oldin, faqat "
         "ortiqchalik diagnostikasidan bilgan.")
    para(doc,
         "Bu natija kaskadning RAD ETISHI ham o'lchanadigan qiymatga ega "
         "ekanini uchinchi arxitekturada takrorlaydi — 19-jadvaldagi "
         "dekoder qarori bilan bir xil shakl.", italic=True, size=10)
    para(doc,
         "Yo'l-yo'lakay bu tajriba 3.6-bo'limdagi tuzilmaviy blok "
         "topuvchini ham sinadi: mBERT o'zining oldinga uzatuvchi qismini "
         "intermediate.dense va output.dense deb ataydi — Whisper ning "
         "fc1/fc2 siga umuman o'xshamaydi — va topuvchi barcha 12 juftlikni "
         "nomga qaramasdan topdi.", size=9.5)
    para(doc, "Qaysi mezon bu modelga mos keladi.", bold=True, size=10)
    para(doc,
         "Kosinus mezonini byudjetga yetguncha majburlash (mBERT da tau "
         "0.70 gacha) yagona yo'l emas, va bu modelda u eng kam ehtimolli "
         "yo'l: mezon deyarli kollinear kanal topmaydi, ya'ni kesim uning "
         "O'ZI ortiqcha deb hisoblamaydigan kanallar orasidan qilinadi. "
         "4.9d-bo'limdagi ikki bosqichli usul aynan shunday holat uchun "
         "qurilgan: 1-bosqich kosinusni qat'iy qoldiradi, 2-bosqich esa "
         "byudjetning qolganini omon qolganlar orasidan fluktuatsiyasi eng "
         "kichiklariga sarflaydi va ikkala bosqich tashlagan qismning "
         "o'rtachasini chiqish biasiga yig'adi. To'rt mezon teng byudjetda, "
         "held-out faollashuvlarda taqqoslandi:")
    table(doc, "45-jadval. Olib tashlash mezonlari mBERT da (12 juftlik, "
               "20% teng byudjet, held-out E_loc).",
          ["Mezon", "O'rtacha E_loc", "G'alaba", "Izoh"],
          [["fluktuatsiya (+ bias)", "0.1169", "12/12", "har bir qatlamda"],
           ["ikki bosqichli", "0.1252", "8/12",
            "1-bosqich o'rtacha 2 kanal (0.07%)"],
           ["kosinus majburiy (tau=0.70)", "0.1476", "0/12", "1.26x yomon"],
           ["past-rank (teng parametr)", "0.1527", "0/12", "eng yomoni"]],
          good_rows=(0,), bad_rows=(2, 3))
    para(doc,
         "Fluktuatsiya mezoni O'N IKKI qatlamning hammasida yutadi. Ikki "
         "bosqichli arm esa amalda unga teng, chunki 1-bosqich o'rtacha "
         "atigi 2 kanal (0.07%) olib tashlaydi — ya'ni bu modelda 'ikki "
         "bosqichli' nomi chalg'ituvchi, ishlaydigan qism ikkinchisi. "
         "Muhimi shundaki, 1-bosqich haqiqatan ishlagan uchta qatlamda u "
         "ZARAR keltiradi: L0 da 16 kanalni kosinus bo'yicha olib tashlash "
         "xatoni 0.0129 dan 0.0443 ga ko'taradi. Kosinus mezoni "
         "kollinearlikni qidiradi, mBERT da esa kollinearlik yo'q; uni "
         "majburlash kollinearlik mavjudligini talab qilmaydigan mezondan "
         "yomonroq.")
    para(doc,
         "Uchdan-uchgacha esa ustunlik TASDIQLANMAYDI. 44-jadvalda "
         "fluktuatsiya armi aniqlikda oldinda (0.2471 va 0.2376), ammo "
         "juftlik farqi +0.0095 [-0.0018, +0.0203] — oraliq nolni qamrab "
         "oladi. Bundan tashqari pseudo-perplexity TESKARI tomonga ketadi "
         "(103.22 dan 109.91 ga). Dastlab buni bias tuzatishning logit "
         "o'rtachalarini siljitishi bilan izohlagan edik; ABLATSIYA bu "
         "izohni RAD ETDI. Bir xil kanallar bilan buklash olib "
         "tashlanganda aniqlik ham (0.2268 dan 0.1882 ga, farq -0.0386 "
         "[-0.0560, -0.0202], sezilarli), pseudo-perplexity ham (119 dan "
         "230 ga) keskin yomonlashadi — ya'ni deyarli-doimiy kanallar "
         "o'rtachasi tuzatish emas, HAQIQIY SIGNAL: uni buklash "
         "usulning tarkibiy qismi va ikkala metrikaga ham foyda. "
         "Fluktuatsiya-kosinus orasidagi perplexity farqining sababi "
         "demak bias emas, KANAL TANLOVI, va u ochiq savol bo'lib "
         "qoladi. Ikkala mezon bir yo'nalishga ketmagani uchun "
         "'fluktuatsiya yaxshiroq' degan xulosa chiqarilmaydi.",
         italic=True, size=10)
    para(doc,
         "Bu naqsh 4.9f-bo'limdagi o'q gibridi bilan aynan bir xil: "
         "operator darajasida qat'iy ustunlik, uchdan-uchgacha esa "
         "chegarada. Ikkalasi ham 50-jadvaldagi yutish qonunining "
         "bashoratiga mos. Va asosiy xulosa o'zgarmaydi: QAYSI mezon "
         "bilan bo'lmasin, 20% qisqartirish mBERT da INT8 dan sezilarli "
         "yomon (fluktuatsiya uchun -0.0162 [-0.0261, -0.0068]) va evaziga "
         "11 MiB beradi. Kaskadning hukmi — kvantla, qisqartirma — "
         "kuchida qoladi.")
    table(doc, "46-jadval. Past-rank shoxchasining birinchi uchdan-uchgacha tasdig'i "
               "(open_llama_3b, held-out o'zbek matnida perplexity).",
          ["Variant", "Perplexity", "FP32 ga"],
          [["FP32", "230.122", "1.000x"],
           ["per-channel INT8 (majburiy)", "230.900", "1.003x"],
           ["kesh-bog'langan past-rank + INT8", "231.642", "1.007x"]],
          good_rows=(2,))
    para(doc,
         "Chiqarilgan rank 1487 (to'liq rankning 46%) da o'rtacha per-operator xato "
         "0.148, maksimumi 0.308 bo'lishiga qaramay perplexity atigi 0.7% ga "
         "yomonlashadi. Bu 4.7-bo'limdagi yutish qonunini dekoder-only arxitekturada "
         "tasdiqlaydi va bizning tadqiqotimizda past-rank shoxchasining birinchi "
         "uchdan-uchgacha tekshiruvi hisoblanadi, chunki na Whisper, na mBERT o'zining "
         "FFN operatorlari uchun 3-holatga kirmagan.")

    h(doc, "4.12. Kvantlash rejimi: vazn-only va faollashuv kvantlashi", 2)
    para(doc,
         "46-jadvaldagi natijalar VAZN-ONLY kvantlash sxemasiga tegishli: vaznlar INT8 "
         "ga o'tkaziladi, faollashuvlar esa FP32 da qoladi. Bu farq Llama sinfidagi "
         "modellarda hal qiluvchi ahamiyatga ega, shuning uchun uni alohida "
         "o'lchadik. Xuddi shu FFN operatorlarini PyTorch ning dinamik kvantlashi "
         "bilan — u faollashuvlarni ham ish vaqtida kvantlaydi — qayta ishlaganda:")
    table(doc, "47-jadval. Kvantlash rejimining ta'siri (open_llama_3b, xuddi shu FFN "
               "operatorlari, real INT8 yadrolari).",
          ["Sxema", "Vaznlar", "Faollashuvlar", "Perplexity", "FP32 ga"],
          [["vazn-only (taklif etilgan)", "INT8 per-channel", "FP32", "230.900", "1.003x"],
           ["dinamik kvantlash", "INT8 per-channel", "INT8 dinamik", "3521.487", "15.303x"]],
          good_rows=(0,), bad_rows=(1,))
    para(doc,
         "Yagona farq — faollashuv kvantlashi. Sabab FFN oraliq tenzorining tuzilishida: "
         "gated FFN da h = SiLU(W_gate x) * (W_up x) ko'paytmasi ekstremal "
         "chetlanishlarga ega bo'ladi va per-tensor dinamik kvantlash ularni buzadi. "
         "Bu LLM.int8() [4] va SmoothQuant [3] hal qilishga qaratilgan hodisaning aynan "
         "o'zi. Attribusiya tahlili aybdorni ham aniqladi: faqat lm_head ni kvantlash "
         "perplexity ni atigi 1.023x yomonlashtiradi (235.385), ya'ni chiqish "
         "proyeksiyasi emas, FFN oraliq tenzori muammoli.")
    callout_note = ("Amaliy ahamiyati: taklif etilgan usul vazn-only kvantlashni "
                    "nazarda tutadi, bu esa amaliy INT8 LLM yechimlarida (llama.cpp "
                    "Q8_0, LLM.int8()) qo'llaniladigan standart rejim. Faollashuv "
                    "kvantlashi alohida muammo bo'lib, SmoothQuant uslubidagi "
                    "qayta taqsimlashni talab qiladi va bu ish doirasiga kirmaydi.")
    para(doc, callout_note, italic=True, size=10)

    para(doc, "Real vazn-only yadrolarining mavjudligi.", bold=True, size=10)
    para(doc,
         "Whisper eksperimentlarida tezlik ONNX Runtime ning haqiqiy INT8 yadrolarida "
         "o'lchangan (15-38-jadvallar bo'ylab). Llama uchun bunday o'lchov bu platformada "
         "bajarilmadi va sababi o'lchab aniqlandi, taxmin qilinmadi:")
    table(doc, "48-jadval. Vazn-only INT8 yadrolarining x86 CPU dagi holati "
               "(open_llama_3b FFN shakllari, bitta oqim).",
          ["Yo'l", "Tezlik", "Sifat", "To'siq"],
          [["ONNX Runtime (Whisper da ishlatilgan)", "real o'lchandi", "saqlanadi",
            "3B model 2 GiB protobuf chegarasidan oshadi"],
           ["torch.ao quantize_dynamic", "3.67x tezroq", "PPL 15.3x yomon",
            "faollashuvlarni ham kvantlaydi"],
           ["torch._weight_int8pack_mm", "0.01x (71x sekin)", "to'g'ri",
            "x86 uchun optimallashtirilmagan"]],
          bad_rows=(1, 2))
    para(doc,
         "Mikro-benchmark: gate/up_proj (8640 x 3200) shaklida FP32 Linear 73.22 ms, "
         "_weight_int8pack_mm esa 5234.30 ms; down_proj (3200 x 8640) uchun mos "
         "ravishda 93.75 ms va 5490.55 ms. Demak Llama uchun vazn-only INT8 tezligini "
         "o'lchash maxsus ish vaqti muhitini (llama.cpp, ONNX Runtime MatMulNBits yoki "
         "kompilyatsiya qilingan torchao yadrolari) talab qiladi. Bu ushbu ishning "
         "chegarasi bo'lib, usulning emas, mavjud CPU vositalarining cheklovi.")

    figure(doc, 8,
           "Enkoder uchun siqish-sifat ish nuqtalari, 95% ishonch oraliqlari bilan.",
           "A scientific scatter plot with error bars, white background, publication "
           "style, grayscale with one highlighted marker. X-axis 'compression factor' "
           "logarithmic from 1 to 8, y-axis 'WER' from 0.09 to 0.15. Four points with "
           "vertical 95% CI bars: 'FP32' at (1.0, 0.1007); 'structural removal + "
           "per-channel INT8' at (4.32, 0.1107) drawn as a filled highlighted marker "
           "and labeled 'proposed'; 'INT8' at (4.00, 0.1146); 'INT8 + low-rank' at "
           "(6.00, 0.1335). A horizontal dashed line at the FP32 WER. Marker area "
           "proportional to measured speedup. Legend inside the plot, sans-serif "
           "labels, thin axis lines.")

    # ============ 4.13 WIKITEXT-2 VA SILJISH ============
    h(doc, "4.13. Standart benchmark va chuqurlik bo'ylab siljishning "
           "to'planishi", 2)
    para(doc,
         "Yuqoridagi til-modeli natijalari o'zbekcha transkriptlarda o'lchangan "
         "bo'lib, ASR qismi bilan bir sohada qoladi, lekin adabiyot bilan "
         "bevosita taqqoslanmaydi. Shu sababli GPTQ, AWQ, SVD-LLM va SliceGPT "
         "hisobot beradigan standart protokol qo'shildi: WikiText-2 test "
         "to'plami birlashtiriladi, bir marta tokenlanadi va kesishmaydigan "
         "2048-tokenli segmentlarda baholanadi; kalibrlash faqat train "
         "to'plamidan olinadi. FP32 perplexity 7.547 chiqdi — nashr etilgan "
         "open_llama_3b qiymatlari (~7.6-7.9) bilan mos, ya'ni o'lchov trakti "
         "tashqi mos yozuvlar bilan tasdiqlangan.")
    para(doc,
         "Birinchi kuzatuv metodologik ahamiyatga ega: INT8 da usullar "
         "ajralmaydi. Yaxlitlash (RTN) 7.550, taklif etilgan masshtab 7.549, "
         "uning chiqish-domen varianti 7.548 — uchalasi ham FP32 ning 1.000 "
         "karrasi. Aynan shu sabab GPTQ va AWQ INT4/INT3 da hisobot beradi, "
         "hamda 16a-jadvaldagi operator darajasidagi farq (GPTQ ning E_loc i "
         "54% past) nima uchun uchdan-uchgacha yo'qolganini izohlaydi: "
         "4.7-bo'limdagi yutish qonuni shu kattalikdagi farqni to'liq "
         "singdiradi. Shuning uchun usullarni ajratuvchi taqqoslash INT4 da "
         "o'tkazildi.")
    table(doc, "49-jadval. WikiText-2 perplexity, open_llama_3b ning 78 ta FFN "
               "operatori, 24 x 2048 = 49 152 token. GPTQ va AWQ - qayta "
               "amalga oshirilgan (rasmiy paketlar CUDA talab qiladi).",
          ["Usul", "Bit", "Perplexity", "FP32 ga"],
          [["FP32 (baza)", "32", "7.547", "1.000x"],
           ["RTN", "8", "7.550", "1.000x"],
           ["taklif etilgan, vazn-domen masshtab", "8", "7.549", "1.000x"],
           ["taklif etilgan + chiqish-domen (LS)", "8", "7.548", "1.000x"],
           ["RTN", "4", "8.583", "1.137x"],
           ["GPTQ [1]", "4", "8.646", "1.146x"],
           ["AWQ [2]", "4", "8.222", "1.089x"],
           ["taklif etilgan, vazn-domen masshtab", "4", "12.799", "1.696x"],
           ["taklif etilgan + chiqish-domen (LS)", "4", "8.246", "1.093x"],
           ["taklif etilgan + siljimagan masshtab", "4", "8.258", "1.094x"]],
          good_rows=(2, 3, 8, 9), bad_rows=(7,))
    para(doc, "Manfiy natija va uning sababi.", bold=True, size=10)
    para(doc,
         "Kalibrlangan masshtabimiz INT4 da eng oxirgi o'rinni egalladi "
         "(1.696x), garchi INT8 da u min/max ga nisbatan operator xatosini "
         "65-74% ga kamaytirgan bo'lsa ham (6-jadval). Bu ziddiyat xulosa "
         "sifatida emas, tekshirilishi kerak bo'lgan gipoteza sifatida "
         "qaraldi. Kutilgan sabab — kesish (clipping) — o'lchov bilan rad "
         "etildi: INT4 da bizning usul har bir operatorda RTN dan HAM vazn, "
         "HAM chiqish xatosi bo'yicha yaxshiroq (o'rtacha chiqish xatosi "
         "0.084 va 0.114, ya'ni 26% past), kesilgan koeffitsiyentlar ulushi "
         "esa atigi 0.7-0.9%.")
    para(doc,
         "Haqiqiy sabab xatoning kattaligida emas, YO'NALISHIDA. Har bir "
         "chiqish kanali uchun kuchaytirish koeffitsiyenti o'lchandi:")
    eq(doc, "g_i = <y_i, y_i^> / <y_i, y_i> ,   y_i = X w_i", 23)
    table(doc, "50-jadval. Operator xatosi, kuchaytirish va uning chuqurlik "
               "bo'ylab to'planishi (9 ta o'lchangan operator, ikki bit "
               "kengligida).",
          ["Usul", "Bit", "O'rt. chiqish xatosi", "O'rt. gain g", "g^78"],
          [["RTN", "8", "0.00619", "1.0000", "1.002"],
           ["taklif etilgan, vazn-domen", "8", "0.00567", "0.9999", "0.995"],
           ["taklif etilgan + LS", "8", "0.00561", "1.0000", "0.997"],
           ["RTN", "4", "0.11389", "1.0005", "1.041"],
           ["taklif etilgan, vazn-domen", "4", "0.08405", "0.9896", "0.444"],
           ["taklif etilgan + LS", "4", "0.07743", "0.9926", "0.560"]],
          good_rows=(1, 2, 5), bad_rows=(4,))
    para(doc,
         "Siljish bit kengligiga keskin bog'liq. INT8 da taklif etilgan "
         "masshtab min/max ning 0.96-0.99 iga tushadi, gain esa 0.9999 bo'lib "
         "qoladi: butun tarmoq bo'ylab yig'ilgan susayish atigi 0.5%. INT4 da "
         "esa xuddi shu protsedura masshtabni 0.60-0.67 ga tushiradi va gain "
         "0.9896 ga pasayadi, ya'ni susayish 56% ni tashkil qiladi — 100 "
         "barobar katta. Shu sababli 49-jadvalda INT8 qatorlari tenglashadi, "
         "INT4 qatorlari esa keskin ajraladi.")
    para(doc,
         "Bu INT8 o'lchovi zarur edi. Uni \"RTN allaqachon FP32 ni "
         "takrorlagani uchun hech bir usul undan yaxshi bo'la olmaydi\" degan "
         "mulohaza bilan tashlab yuborish mumkin edi, ammo bu mulohaza "
         "usullarni faqat YUQORIDAN chegaralaydi; INT4 dagi buzilish esa aynan "
         "pastga og'ish edi. Shuning uchun avval arzon operator darajasidagi "
         "gain o'lchandi (50-jadvalning yuqori yarmi), u tenglikni bashorat "
         "qildi, so'ng bashorat to'liq tarmoqda tasdiqlandi.")
    para(doc,
         "Masshtabni vaznlarga moslash, s = <w,q>/<q,q>, taqsimotning asosiy "
         "qismiga to'g'rilanadi va shu bilan har bir kanalni ~1% "
         "KICHRAYTIRADI. Bitta operator uchun bu kichikroq xato. Ammo 78 ta "
         "operatorning hammasi bir xil yo'nalishda kichraytirgani uchun "
         "susayish ko'payadi, o'zaro so'nmaydi: 0.9896^78 = 0.44. RTN ning "
         "yaxlitlash xatosi esa afzal yo'nalishga ega emas va tasodifiy sayr "
         "kabi to'planadi. Bu 4.7-bo'limdagi yutish qonunini ham "
         "aniqlashtiradi — yutish faqat xatolar mustaqil bo'lganda amal "
         "qiladi.")
    para(doc, "Tuzatish.", bold=True, size=10)
    para(doc,
         "Butun sonli kodlar q o'zgarmaydi — faqat kanal bo'yicha bitta "
         "ko'paytuvchi qayta tanlanadi, ya'ni xotira formati, bit kengligi va "
         "siqish darajasi aynan o'sha qoladi. G = X^T X bo'lganda:")
    eq(doc, "s_LS = (q^T G w) / (q^T G q)", 24)
    eq(doc, "s_siljimagan = (w^T G w) / (q^T G w)", 25)
    para(doc,
         "(24) chiqish xatosini minimallashtiradi; qoldiq y^ ga ortogonal "
         "bo'lgani uchun gain = ||y^||^2/||y||^2 <= 1, ya'ni susayish "
         "kamayadi, lekin butunlay yo'qolmaydi. (25) esa <y, y^> = <y, y> "
         "shartini qo'yib gain ni aynan 1 ga tenglaydi, buning evaziga "
         "xatoni oshiradi. Qaysi biri yutishi — siljish/dispersiya savoli "
         "bo'lib, uni faqat to'liq tarmoq o'lchovi hal qila oladi.")
    para(doc,
         "49-jadval javobni beradi: perplexity 12.799 dan 8.246 ga tushdi va "
         "ikkala rejim amalda farqlanmaydi (8.246 va 8.258). Demak siljishning "
         "katta qismini olib tashlash hal qiluvchi, qolgan nozik farq esa "
         "shovqin darajasida. Tuzatilgan variant RTN va GPTQ dan yaxshiroq, "
         "AWQ ga (8.222) esa juda yaqin, lekin undan ozgina past. To'g'ri "
         "da'vo shunga mos ravishda cheklanadi: kalibrlangan masshtabimiz "
         "eng yaxshi post-training kvantlash usuli emas, balki u bilan "
         "raqobatbardosh — bu 4.9a-bo'limdagi pozitsiyani, ya'ni hissamiz "
         "kvantlashda emas, ortogonal strukturaviy o'qda ekanini "
         "mustahkamlaydi.")
    para(doc,
         "Ta'kidlash kerakki, bu cheklov ushbu ishning ish nuqtasiga taalluqli "
         "emas. Taklif etilgan kaskad INT8 da ishlaydi, u yerda esa masshtabimiz "
         "siljimagan bo'lib qoladi va uchdan-uchgacha RTN dan ozgina ustun "
         "turadi (7.549 va 7.548, RTN ning 7.550 iga qarshi, FP32 esa 7.547). "
         "Ya'ni 6-jadvaldagi operator darajasidagi ustunlik INT8 da "
         "saqlanadi. Buzilish faqat 4 bitga o'tilganda, masshtab min/max ning "
         "0.60-0.67 iga tushganda yuzaga keladi. Shuning uchun hissa endi "
         "o'lchangan chegara bilan qo'yiladi: kalibrlangan per-channel "
         "masshtab INT8 uchun asoslangan, INT4 ga esa faqat chiqish-domen "
         "tuzatishi bilan ko'chiriladi.")
    para(doc, "Cheklovlar.", bold=True, size=10)
    para(doc,
         "down_proj operatori uchun Hessian rank-kamchil (n = 8640 kirish, "
         "B = 4096 kalibrlash qatori), nashr etilgan GPTQ esa 262 144 qator "
         "ishlatadi. Shu sababli GPTQ ning bu yerdagi INT4 natijasi (8.646, "
         "RTN dan past) GPTQ algoritmi haqidagi xulosa emas, balki shu "
         "kalibrlash byudjetidagi xulosadir. Barcha usullar bir xil "
         "ma'lumot, bir xil granularlik (per-output-channel, simmetrik) va "
         "bir xil bit kengligida ishlagani uchun taqqoslash ichki jihatdan "
         "adolatli, lekin mutlaq qiymatlarni maqolalardagi raqamlar bilan "
         "bevosita solishtirish mumkin emas.")

    figure(doc, 9,
           "INT4 da operator darajasidagi aniqlik va tarmoq darajasidagi "
           "sifat o'rtasidagi uzilish, hamda uni siljish orqali izohlash.",
           "A two-panel scientific figure, white background, publication "
           "style, grayscale with one highlighted color. Left panel: grouped "
           "bar chart, x-axis four methods 'RTN', 'GPTQ', 'AWQ', 'proposed', "
           "two bars each — 'per-operator output error' (values 0.114, 0.095, "
           "0.099, 0.084) and 'WikiText-2 perplexity ratio to FP32' (values "
           "1.137, 1.146, 1.089, 1.696). The 'proposed' pair is highlighted to "
           "show it has the LOWEST operator error but the HIGHEST perplexity, "
           "an inversion. Right panel: line plot, x-axis 'number of stacked "
           "operators' from 0 to 78, y-axis 'cumulative signal gain' from 0.4 "
           "to 1.1 . Three curves labelled with their per-operator gain: "
           "'RTN g=1.0005' staying flat near 1.04, 'proposed g=0.9896' "
           "decaying to 0.44, 'proposed + output-domain scale g=0.9926' "
           "decaying only to 0.56. A horizontal dashed reference line at 1.0. "
           "Thin axis lines, sans-serif labels, legend inside each panel.")

    h(doc, "4.14. Strukturaviy kesish rad etilishining uchdan-uchgacha "
           "tekshiruvi (open_llama_3b)", 2)
    para(doc,
         "4.1-bo'limdagi kesh-bog'langan diagnostika open_llama_3b uchun "
         "kanalli kesishni rad etadi, chunki tau=0.99 da qoldiq ortiqchalik "
         "atigi 0.6% ni tashkil qiladi — mBERT ning 3.5% idan ham kam. Ammo "
         "bu xulosa faqat ARGUMENTLANGAN edi, mBERT uchun 4.11-bo'limda "
         "qilingani kabi O'LCHANMAGAN. Ikkalasi bir xil emas: mBERT da "
         "qoldiq ortiqchalik kam bo'lsa-da, majburiy kesish baribir "
         "sinaldi va xarajati raqamlashtirildi (40-jadval). Shu bo'shliqni "
         "yopish uchun 20% kanal 46-jadvaldagidek FFN dan (gate_proj, "
         "up_proj, down_proj — barchasi bir vaqtda) majburan olib "
         "tashlandi, ikkita mezon bilan: bizning kosinus guruhlash "
         "(byudjetgacha majburlangan) va fluktuatsiya balli (o'rtachasi "
         "bias ga qo'shilgan holda; bu arxitekturada tabiiy bias yo'q, "
         "shuning uchun 26 qatlamning har biriga 3200 o'lchamli bias "
         "vektori — jami 333 KiB — qo'shildi). Sinov 49-jadval bilan bir "
         "xil WikiText-2 protokolida o'tkazildi, shunday qilib natija "
         "kvantlash xarajati bilan bevosita solishtiriladi.")
    table(doc, "51-jadval. Strukturaviy kanal kesishning uchdan-uchgacha "
               "xarajati (open_llama_3b, WikiText-2 perplexity, 20% FFN "
               "kesish).",
          ["Variant", "Perplexity", "FP32 ga"],
          [["FP32", "7.547", "1.000x"],
           ["INT8 vazn-only (46-jadval)", "7.550", "1.000x"],
           ["20% kanal, kosinus (majburiy)", "9.490", "1.258x"],
           ["20% kanal, fluktuatsiya (+ bias)", "8.906", "1.180x"]],
          good_rows=(1,), bad_rows=(2, 3))
    para(doc,
         "Natija argumentni tasdiqlaydi, lekin kutilganidan ham qattiqroq. "
         "Fluktuatsiya kosinusdan yaxshiroq (8.906 va 9.490 ga — bias "
         "tuzatishi mBERT dagi kabi bu yerda ham foyda beradi), ammo "
         "ikkalasi ham INT8 ning 1.000x iga nisbatan 18-26% yomonlashadi, "
         "mBERT dagi 20% kesishning aniqlikka ta'siridan (40-jadval, "
         "-0.0162, nisbiy ~6%) SEZILARLI KATTAROQ. Buning sababi "
         "arxitekturaviy: mBERT da bitta kanal FFN ning ikkita "
         "matritsasiga (intermediate.dense, output.dense) tegadi, Llama "
         "da esa uchtasiga (gate_proj, up_proj, down_proj) — gated "
         "faollashuv orqali xato ko'payadi, chunki h = SiLU(g)*u "
         "ko'paytmasida ikkala omil ham kesilgan kanaldan ta'sirlanadi. "
         "Demak 4.1-bo'limdagi diagnostikaning yo'nalishi (kesish kerak "
         "emas) to'g'ri chiqdi, ammo uni faqat 0.6% raqami bilan "
         "asoslash xarajatning kattaligini kamsitgan bo'lar edi — "
         "o'lchov buni 44-47-jadvallardagi kvantlash xarajatidan (0.3% "
         "dan kam) o'n baravar farqli ekanini ko'rsatadi.")
    para(doc,
         "Kaskadning uch modeldagi hukmi endi bir xil dalilga tayanadi: "
         "Whisper (17.1% ortiqchalik, past-rank shoxchasi orqali 0.7% "
         "xarajat bilan ishlaydi, 46-jadval, yuqorida), mBERT (3.5%, "
         "kesish majburlansa 6% xarajat, 40-jadval) va Llama (0.6%, "
         "kesish majburlansa 18-26% xarajat, yuqorida) — ortiqchalik "
         "kamayishi bilan majburiy kesish xarajati monoton o'sadi, bu "
         "kaskadning kvantlashni birinchi, kesishni oxirgi qo'yishini "
         "(3.6-bo'lim) uchala arxitekturada ham mustaqil ravishda "
         "asoslaydi.")

    h(doc, "4.15. Nima uchun Llama da kollinear kanal yo'q: faollashuv "
           "geometriyasi", 2)
    para(doc,
         "Yuqoridagi 0.6% raqami tabiiy savol tug'diradi: mezon haqiqatan "
         "ortiqchalik topmayaptimi, yoki uni chaqirishda xatolik bormi? "
         "Savol o'rinli, chunki xuddi shu kod Whisper enkoderida 17.1% "
         "topadi, va nol natija ko'pincha nuqson belgisi bo'ladi. Shu "
         "sababli javob guruhlash kodiga TAYANMASDAN olindi: mezon "
         "cos(h_i, h_j) >= tau bo'lgani uchun juftlik kosinuslarini "
         "to'g'ridan-to'g'ri sanash qancha juftlik birlashishi MUMKIN "
         "ekanini aytadi, va algoritm natijasini shunga solishtirish "
         "mumkin.")
    table(doc, "52-jadval. open_llama_3b, 8-qatlam down_proj kirishi: "
               "8640 kanal, 2048 qator; barcha juftlik kosinuslari "
               "to'liq sanab chiqilgan.",
          ["O'lchov", "Qiymat"],
          [["eng yaqin qo'shni, mediana", "0.2330"],
           ["eng yaqin qo'shni, 99.9-protsentil", "0.6594"],
           ["BUTUN matritsadagi eng katta kosinus", "0.7681"],
           ["tau >= 0.90 juftliklar", "0"],
           ["tau >= 0.70 juftliklar", "1"],
           ["tau >= 0.50 juftliklar", "134"]])
    para(doc,
         "37 million juftlik orasida eng kollinear ikkitasining kosinusi "
         "0.7681, ya'ni tau = 0.90 ga yetadigan bironta ham juftlik yo'q. "
         "Algoritm esa aynan shuni qaytaradi. Ikkinchi tekshiruv sifatida "
         "guruhlash eps chegarasi cheksiz qilib ham chaqirildi — natija "
         "har bir tau da bir xil (0.00% / 0.01% / 0.82%), demak kanallarni "
         "rad etayotgan ta'sir mezoni emas, yo'nalish mezoni. Qisqasi, "
         "0.6% modelning xossasi, kodning emas.")
    para(doc, "Sabab: faollashuvning ishorasi.", bold=True, size=10)
    para(doc,
         "Ikki arxitektura qisqartiriladigan o'qqa TURLI taqsimotdagi "
         "faollashuv beradi. Whisper da fc2 operatori GELU chiqishini "
         "o'qiydi, u esa deyarli butunlay manfiy emas; shuning uchun kanal "
         "javob vektorlari musbat konusda yotadi va ikkitasining kosinusi "
         "tuzilishiga ko'ra yuqori bo'ladi. Llama da down_proj gated "
         "ko'paytmani, SiLU(W_gate x) * (W_up x) ni o'qiydi, uning ishorasi "
         "erkin o'zgaradi — o'lchandi: musbat qiymatlar ulushi 50.1%. "
         "Vektorlar butun sferaga tarqaladi va deyarli-kollinear juftlik "
         "amalda uchramaydi.")
    table(doc, "53-jadval. Faollashuv geometriyasi va u qaysi "
               "shoxchani ochishi.",
          ["Model", "Qisqartiriladigan o'q kirishi", "Musbat ulush",
           "Juftlik ortiqchaligi"],
          [["Whisper enkoder", "GELU(fc1 x + b)", "~100%", "17.1%"],
           ["open_llama_3b", "SiLU(W_gate x) * (W_up x)", "50.1%", "0%"]])
    para(doc,
         "Bu kuzatuv ikkita ilgari bog'lanmagan natijani birlashtiradi. "
         "Kanalli kesish Llama da qimmat (4.14-bo'lim, +18-26%), past-rank "
         "esa arzon (46-jadval, +0.7%) — va endi nima uchun ekani "
         "ko'rinadi: 2048 o'lchamli javob fazosida 8640 ta vektor "
         "MAJBURAN chiziqli bog'liq, ammo bu bog'liqlik JUFTLIKDA emas, "
         "taqsimlangan holda mavjud. Hech qaysi ikki kanal bir-biriga "
         "o'xshamaydi, biroq hammasi birgalikda rank jihatdan kamchil. "
         "Kosinus mezoni juftlik ortiqchaligini qidiradi, past-rank yoyilma "
         "esa taqsimlanganini.")
    para(doc,
         "Shuning uchun to'g'ri formulirovka 'Llama da ortiqchalik yo'q' "
         "emas, balki 'ortiqchalikning TURI boshqa' bo'ladi — va bu "
         "kaskadning ikki strukturaviy shoxchaga (kanal olib tashlash va "
         "past-rank yoyilma) ega bo'lishini asoslaydi. Bitta shoxchali "
         "usul bu modelda yo hech narsa topmagan, yo 18-26% xarajat bilan "
         "majburlagan bo'lardi.")
    para(doc,
         "Operator darajasidagi to'liq o'lchov shu xulosani raqamlaydi "
         "(7 qatlam, held-out faollashuvlarda):", size=10)
    table(doc, "54-jadval. open_llama_3b FFN: mezon tanlagan va majburiy "
               "ish nuqtalarining operator xatosi (7 qatlam o'rtachasi, "
               "held-out).",
          ["Ish nuqtasi", "Olindi", "E_loc", "Bias tuzatishi bilan"],
          [["tau = 0.99", "0.00%", "0.0000", "0.0000"],
           ["tau = 0.95", "0.04%", "0.0045", "0.0044"],
           ["tau = 0.90", "0.11%", "0.0183", "0.0179"],
           ["10% majburiy kosinus", "10%", "0.2307", "0.2261"],
           ["10% fluktuatsiya", "10%", "0.1801", "0.1763"],
           ["20% majburiy kosinus", "20%", "0.3059", "0.3017"],
           ["20% fluktuatsiya", "20%", "0.2672", "0.2604"],
           ["30% majburiy kosinus", "30%", "0.3379", "0.3342"],
           ["30% fluktuatsiya", "30%", "0.3418", "0.3326"],
           ["INT8 (kesishsiz)", "0%", "0.0079", "0.0078"]],
          good_rows=(9,), bad_rows=(3, 5, 7))
    para(doc,
         "Uch narsa ko'rinadi. Birinchidan, INT8 raqobatdan tashqarida: "
         "uning operator xatosi eng arzon kesishnikidan 23 barobar kichik, "
         "bu esa kaskadning 'avval kvantla' qoidasini operator darajasida "
         "ham tasdiqlaydi. Ikkinchidan, fluktuatsiya mezoni YUMSHOQ "
         "rejimda qat'iy yutadi (10% da 22%, 20% da 13% yaxshiroq), 30% da "
         "esa ustunlik yo'qoladi va teskarisiga o'tadi (0.3418 va 0.3379) "
         "— tanlash uchun joy qolmaganda qaysi mezon bilan tanlash ham "
         "ahamiyatini yo'qotadi. Uchinchidan, kesish qoldig'iga bias "
         "tuzatishini qo'llash foydasi kichik: qoldiq energiyasining "
         "atigi 2-6% i o'zgarmas komponentga to'g'ri keladi, xato esa "
         "1.5-4% ga kamayadi. Bu tuzatish shu sababli kaskadga "
         "kiritilmadi — u bepul bo'lsa-da, o'lchov uni asoslamadi.")

    para(doc, "Kuzatuvdan qonunga: oldindan e'lon qilingan bashorat "
              "sinovi.", bold=True, size=10)
    para(doc,
         "Uch model — kuzatuv; qonun da'vosi uchun ikki YANGI model "
         "faqat faollashuv turi bo'yicha tanlab olindi va bashoratlar "
         "o'lchovdan OLDIN qayd etildi: SwiGLU li Qwen2.5-0.5B uchun "
         "~50% musbat faollik va tau >= 0.90 da deyarli nol; GELU li "
         "DistilBERT uchun kuchli bir ishoralilik va gated modellardan "
         "aniq yuqori kosinus poli. Ikkinchi bashorat ataylab TAQSIMOT "
         "haqida, olib tashlash ulushi haqida emas — mBERT ko'rsatadiki, "
         "GELU enkoder yuqori kosinus poli bilan ham ish nuqtasida kam "
         "ortiqchalik berishi mumkin.")
    table(doc, "55-jadval. Bashorat sinovi: eng yaqin qo'shni |cos| "
               "taqsimoti, har modelda uch chuqurlik.",
          ["Model / qatlam", "Musbat ulush", "Mediana", "Maks",
           "tau>=0.90"],
          [["Qwen2.5-0.5B L6 (SwiGLU)", "50.0%", "0.306", "0.952",
            "0.08%"],
           ["Qwen2.5-0.5B L12", "50.1%", "0.270", "0.945", "0.14%"],
           ["Qwen2.5-0.5B L18", "50.3%", "0.303", "0.904", "0.04%"],
           ["DistilBERT L1 (GELU)", "14.4%", "0.572", "0.997", "0.78%"],
           ["DistilBERT L3", "13.1%", "0.661", "0.979", "0.62%"],
           ["DistilBERT L5", "12.7%", "0.864", "0.998", "39.65%"],
           ["(ma'lumot) open_llama_3b L8", "50.1%", "0.233", "0.768",
            "0.00%"]],
          good_rows=(0, 1, 2, 3, 4, 5))
    para(doc,
         "Ikkala bashorat ham tasdiqlandi. Qwen ning musbat ulushi aynan "
         "50%, medianasi Llama diapazonida; halol nuance — tau >= 0.90 "
         "da qat'iy nol emas, 4864 kanaldan 2-7 tasi (Llama da qat'iy "
         "nol edi). DistilBERT esa 86-87% bir ishorali va kosinus poli "
         "keskin yuqori: mediana 0.57-0.86 (gated modellarda 0.23-0.31), "
         "L5 da kanallarning 39.65% i tau = 0.90 dan yuqori qo'shniga "
         "ega. Demak ishora tuzilishi kosinus taqsimotini modellar "
         "OILASI bo'ylab bashorat qiladi — bu 53-jadvaldagi mezonning "
         "ilgari ko'rilmagan modellardagi tasdig'i.")
    para(doc,
         "Taqsimot natijaga ham aylanadi. DistilBERT da tau = 0.90 "
         "mezoni o'rtacha 4.14% kanalni tasdiqlaydi (gated modellardagi "
         "~0.1% dan qirq barobar ko'p) va olib tashlash aynan taqsimot "
         "ko'rsatgan qatlamda to'planadi (L5: poli 39.65%, olib "
         "tashlash 22.3%). Uchdan-uchgacha (o'zbek matni, 1051 "
         "niqoblangan pozitsiya, juftlik): mezon armi FP32 dan "
         "farqlanmaydi (+0.0048 [-0.0038, +0.0133]); teng hajmdagi "
         "tasodifiy kesish yo'nalishda yomonroq (-0.0067), ammo bu "
         "hajmda ahamiyatlilik o'rnatilmaydi.", italic=True, size=10)
    para(doc, "Kollinearlik korpusga ham bog'liq.", bold=True, size=10)
    table(doc, "56-jadval. Bir xil model, bir xil qatlam, ikki korpus "
               "(open_llama_3b, tau = 0.99 da olib tashlanadigan ulush).",
          ["Qatlam", "O'zbek matni", "WikiText-2"],
          [["L0", "3.37%", "0.00%"],
           ["L4", "0.65%", "0.00%"],
           ["L8", "0.15%", "0.00%"],
           ["L20", "0.00%", "0.00%"]])
    para(doc,
         "O'xshashlik faollashuvlarda o'lchanadi, faollashuvlar esa "
         "matnning funksiyasi — demak qaysi kanallar birga ishlashi "
         "korpusga bog'liq. Bu usulning kamchiligi emas, ta'rifiy "
         "xossasining natijasi, ammo undan amaliy qoida chiqadi: "
         "ortiqchalik diagnostikasi JOYLASHTIRISH taqsimotidagi "
         "kalibrlashda o'tkazilishi kerak, va 'model X da Y% ortiqchalik "
         "bor' shaklidagi da'volar korpussiz to'liq emas.")

    h(doc, "4.16. Mezonning ikkita taklif etilgan kengaytmasi: biri qabul, "
           "biri rad etildi", 2)
    para(doc,
         "Yuqoridagi natijalar mezonning o'zini o'zgartirish bo'yicha ikkita "
         "taklifni tug'dirdi. Ikkalasi ham qoldiq nisbati orqali "
         "ifodalangan: R_jp = ||h_j - gamma_jp h_p|| / ||h_j||, hamda uning "
         "affin varianti, unda gamma bilan birga o'zgarmas c_j ham "
         "moslashtiriladi. Ikkalasi ham o'lchandi va natija turlicha "
         "chiqdi.")
    para(doc, "Birinchi shakl yangi mezon emas.", bold=True, size=10)
    para(doc,
         "Eng kichik kvadratlar bo'yicha optimal gamma da "
         "||h_j - gamma h_p||^2 = ||h_j||^2 (1 - cos^2), demak "
         "R_jp = sqrt(1 - cos^2) = |sin(theta)|, ya'ni R <= eps_R sharti "
         "|cos| >= sqrt(1 - eps_R^2) shartiga AYNAN teng. Bu 3.2-bo'limdagi "
         "kosinus chegarasining qayta parametrlanishi bo'lib, yangi qabul "
         "sohasi bermaydi; ushbu ishning guruhlash moduli uni allaqachon "
         "sin_theta metrikasi sifatida saqlaydi. Shu bilan birga taklifga "
         "ilova qilingan 'har xil masshtabdagi bog'liqliklarni qabul "
         "qiladi' izohi ham o'rinli emas: kosinus masshtabga invariant va "
         "gamma masshtabni aniq o'zlashtiradi, ya'ni bu jihatdan ikki "
         "mezon farq qilmaydi.")
    para(doc, "Ishorani tashlash esa haqiqiy kamchilikni ochdi.", bold=True,
         size=10)
    para(doc,
         "Guruhlash darvozasi cos >= tau shaklida ishorali edi. Ammo "
         "cos = -0.95 bo'lgan juftlik cos = +0.95 bilan BIR XIL darajada "
         "birlashtiriladi: gamma manfiy chiqadi, kompensatsiya esa gamma ni "
         "ishorasi bilan qo'llaydi (3.2-bo'lim). Ya'ni to'siq faqat "
         "darvozada edi, mexanizmda emas. Ikki modelda o'lchandi:")
    table(doc, "57-jadval. Ishorali va ishorasiz yo'nalish mezoni "
               "(chegaradan yuqori kanallar ulushi).",
          ["Model / qatlam", "Eng katta", "tau=0.99", "tau=0.90", "tau=0.70"],
          [["Llama L8, ishorali", "0.7681", "0.00%", "0.00%", "0.02%"],
           ["Llama L8, |cos|", "0.8488", "0.00%", "0.00%", "0.13%"],
           ["Whisper L8, ishorali", "1.0000", "26.46%", "56.76%", "93.70%"],
           ["Whisper L8, |cos|", "1.0000", "26.46%", "56.76%", "93.90%"]],
          good_rows=(1,))
    para(doc,
         "Whisper da ish nuqtasida (tau = 0.99) farq to'rtta tekshirilgan "
         "qatlamda ham AYNAN nol, Llama da esa tau = 0.70 da ishorasiz "
         "shakl 6.5 barobar ko'p kanal topadi. Sabab yana faollashuv "
         "geometriyasida: Whisper ning qisqartiriladigan o'qi bir ishorali "
         "(o'lchandi: qiymatlarning atigi 1.1-3.7% i musbat, ya'ni GELU "
         "ning 'o'chiq' rejimi ustunlik qiladi), shuning uchun u yerda "
         "anti-kollinear juftlik deyarli yo'q; Llama ning gated ko'paytmasi "
         "esa ikki ishorali (50.1% musbat) va bunday juftliklar mavjud. "
         "Shu sababli ishorasiz variant kiritildi, ammo standart holatda "
         "O'CHIQ qoldirildi: maqoladagi barcha ish nuqtalari ishorali "
         "darvoza bilan o'lchangan va tau = 0.95 da ikki shakl ozgina "
         "farqlanadi, ya'ni standartni almashtirish allaqachon baholangan "
         "artefaktlarni jimgina o'zgartirgan bo'lardi.")
    para(doc, "Affin variant esa teskari natija berdi.", bold=True, size=10)
    para(doc,
         "Xuddi shu hisob bilan affin shakl R^aff = sqrt(1 - corr^2) ga "
         "keltiriladi, ya'ni u kosinusni emas, KORRELYATSIYANI "
         "chegaralaydi. Nazariy jozibasi aniq: markazlashtirish barcha "
         "kanallar bo'lishadigan umumiy siljishni olib tashlaydi. O'lchov "
         "esa ikkala modelda ham uni rad etdi — korrelyatsiya kosinusdan "
         "KAMROQ ortiqchalik topadi (Llama L8, tau=0.50: 0.91% va 1.56%; "
         "Whisper L0, tau=0.99: 29.47% va 47.14%), ustiga har bir kanal "
         "uchun c_j ni saqlashni talab qiladi. Taklifning 'ko'proq "
         "funksional bog'liqlikni qabul qiladi' degan asosi shu tariqa "
         "o'lchov bilan tasdiqlanmadi.")
    para(doc,
         "Bu o'lchovning yon mahsuloti mustaqil ahamiyatga ega. Whisper "
         "dagi ortiqchalik umumiy siljish artefakti bo'lishi mumkin edi — "
         "faolliklar bir ishorali bo'lgani uchun barcha kanallar umumiy "
         "poydevorga ega va bu har qanday kosinusni sun'iy ko'taradi. "
         "Markazlashtirish shu ehtimolni sinaydi, va ish nuqtasida javob "
         "salbiy: tau = 0.99 da ortiqchalik L8, L16 va L23 da deyarli "
         "to'liq saqlanadi (26.46% -> 26.25%, 0.05% -> 0.05%, "
         "0.10% -> 0.10%). Faqat birinchi blokda uning uchdan bir qismi "
         "siljishdan kelib chiqadi (47.14% -> 29.47%). Demak 4.2-bo'limda "
         "da'vo qilingan funksional kollinearlik ish nuqtasida haqiqiy, "
         "bo'sh chegaralarda esa siljishning ulushi sezilarli — bu "
         "chegaralarda ishlanmagani uchun xulosaga ta'sir qilmaydi, ammo "
         "qayd etilishi kerak.")

    h(doc, "4.17. Ikki strukturaviy oila: ustun tanlash va past-rank "
           "yoyilma", 2)
    para(doc,
         "Yuqoridagi natijalar kaskadning ikki strukturaviy shoxchasini "
         "matritsa yaqinlashtirish adabiyotidagi ikki oila bilan bevosita "
         "bog'laydi, va bu bog'lash bir noaniqlikni ham yo'q qiladi. "
         "4.5-bo'limda CUR yig'ilishi faollashuvga sezgir SVD ga 135 "
         "o'lchovdan 135 tasida yutqazgan (12-jadval), shu bilan birga "
         "strukturaviy shoxcha Whisper enkoderida asosiy natijani beradi "
         "(20-jadval). Bu ziddiyat emas: 'CUR' nomi ostida IKKI boshqa "
         "narsa turibdi.")
    table(doc, "58-jadval. Uch konstruksiya va ularning bu ishdagi holati.",
          ["Konstruksiya", "Nima quriladi", "Operatorlar soni", "Holat"],
          [["CUR yig'ilishi (C U R)",
            "ustunlar, satrlar va r x r o'rta blok",
            "uch", "rad etildi (12-jadval, 0/135)"],
           ["ustun tanlash / ID",
            "operatorning O'ZI kichrayadi, gamma vakilga buklanadi",
            "bitta", "qabul (Whisper: 17.1%, dWER -0.0014)"],
           ["past-rank yoyilma (SVD)",
            "ikki kichik matritsa U V",
            "ikki", "qabul (Llama: rank 46%, +0.7%)"]],
          good_rows=(1, 2), bad_rows=(0,))
    para(doc,
         "Farq shakl darajasida emas, xarajat darajasida. CUR yig'ilishi "
         "teng byudjetda qo'shimcha r^2 blokini olib yuradi, ya'ni r(m+n) "
         "o'rniga r(m+n) + r^2 sarflaydi va shu sababli pastroq rankka "
         "majbur bo'ladi — 4.5-bo'limda o'lchangan yutqazishning sababi "
         "shu. Bizning strukturaviy shoxchamiz esa umuman faktorizatsiya "
         "EMAS: kanal olib tashlanadi va uning hissasi vakil ustunga "
         "buklanadi, natijada oraliq kenglik 4096 dan 3400 ga tushadi va "
         "qo'shimcha matmul paydo bo'lmaydi. Adabiyotda bunga eng yaqin "
         "konstruksiya interpolative decomposition [23], ya'ni o'rta "
         "bloksiz ustun tanlash.")
    para(doc,
         "CUR adabiyotidan olingan g'oya esa saqlanadi va tasdiqlanadi: "
         "ustunlarni QAYSI tartibda tanlash. Leverage skorlari vazn "
         "matritsasining o'z geometriyasidan kelib chiqadi, bizning "
         "funksional tartibimiz esa kalibrlashdagi javobdan; teng "
         "byudjetda ikkinchisi 135 operatordan 134 tasida yutadi "
         "(12-jadval). Ya'ni CUR ning TANLASH tamoyili ishlaydi, uning "
         "YIG'ILISHI esa ishlamaydi, va bu ikkisini ajratish ushbu ishning "
         "aniqlashtirishlaridan biridir.")
    para(doc, "Qaysi oila qachon yutadi.", bold=True, size=10)
    para(doc,
         "4.15-bo'limdagi geometrik kuzatuv bu tanlovga bashorat qiluvchi "
         "mezon beradi. Ustun tanlash faqat haqiqiy ustunlar bir-birini "
         "ifodalay olganda ishlaydi, ya'ni juftlik kollinearligi mavjud "
         "bo'lganda; past-rank yoyilma esa sun'iy bazis quradi va "
         "taqsimlangan bog'liqlikni ham oladi. Faollashuv geometriyasi "
         "esa qaysi holat ekanini oldindan aytadi:")
    table(doc, "59-jadval. Faollashuv geometriyasi qaysi oilani "
               "tanlashini bashorat qiladi.",
          ["Model", "Qisqartiriladigan o'q", "Bir ishorali",
           "Juftlik kollinearligi", "G'olib oila"],
          [["Whisper enkoder", "GELU chiqishi", "ha (97%)", "17.1%",
            "ustun tanlash"],
           ["mBERT", "GELU chiqishi", "ha", "3.5%",
            "hech biri (INT8 yetarli)"],
           ["open_llama_3b", "SiLU(g) * u", "yo'q (50.1%)", "0%",
            "past-rank yoyilma"]])
    para(doc,
         "Bir ishorali faollashuv kanal vektorlarini konusga joylashtiradi "
         "va juftlik kosinuslarini tuzilishiga ko'ra ko'taradi, ikki "
         "ishorali gated ko'paytma esa ularni sferaga tarqatadi. Shu "
         "sababli 4.1-bo'limdagi diagnostikaning natijasi tasodifiy emas "
         "va uni model ISHGA TUSHIRILMASDAN, faqat arxitektura tavsifidan "
         "taxmin qilish mumkin. Bu kaskadning ikki shoxchali bo'lishini "
         "ham asoslaydi: bitta oila bilan cheklangan usul bu uch modeldan "
         "birida albatta noto'g'ri javob bergan bo'lardi.")

    h(doc, "4.18. Yakuniy taqqoslash: kaskad va mavjud konveyerlar teng "
           "sharoitda", 2)
    para(doc,
         "Nashr etilgan usullar bilan taqqoslashlar yuqorida o'z "
         "bo'limlarida sochilgan; bu jadval ularni BITTA teng-sharoit "
         "ko'rinishga yig'adi. Barcha qatorlar Whisper enkoderi, TEST "
         "splitining o'sha 300 namunasi, INT8 dekoder bilan; strukturaviy "
         "qatorlar teng byudjetda (267 MiB, tau = 0.99 nuqtasi). Har "
         "raqam yuqoridagi qaysi jadvaldan kelgani ko'rsatilgan.",
         size=10)
    table(doc, "60-jadval. Kaskad va konveyerlar: bitta teng-sharoit "
               "ko'rinish (FP32 = 0.1793; strukturaviy qatorlar 267 MiB).",
          ["Konveyer", "MiB", "WER", "Manba bo'lim"],
          [["FP32 (baza)", "1172", "0.1793", "—"],
           ["ko'r-ko'rona INT8 (RTN)", "300", "0.1858", "4.9"],
           ["GPTQ yolg'iz", "300", "0.1847", "4.9b"],
           ["KASKAD: mezon-kesish + GPTQ", "267", "0.1833",
            "4.9b; FP32 dan farqlanmaydi"],
           ["magnitude-kesish + INT8", "267", "0.1837", "4.9d"],
           ["Wanda-kesish + INT8", "267", "0.1850", "4.9d"],
           ["FLAP + INT8", "267", "0.1859", "4.9d"],
           ["kesish KOMPENSATSIYASIZ + INT8", "254", "1.3393", "4.9d"],
           ["magnitude, agressiv (30%)", "242", "0.6294", "4.9f"],
           ["taqsimlangan past-rank + INT8", "203", "0.3056", "4.9"]],
          good_rows=(3,), bad_rows=(7, 8, 9))
    para(doc,
         "Uch xulosa bir qarashda ko'rinadi. Birinchidan, mo''tadil "
         "byudjetda kalibrlash + kompensatsiya ishlatgan BARCHA "
         "konveyerlar bir guruhda (0.1833-0.1859) va kaskad shu "
         "guruhning eng kichigi va eng yaxshisi — ustunlik da'vo "
         "qilinmaydi, teng-yoki-yaxshiroq o'lchanadi. Ikkinchidan, "
         "guruhdan chiqishning uch yo'li uch xil sababdan halokatli: "
         "kompensatsiyani tashlash (1.34), mezonni faollikka ko'r "
         "qilish (0.63), byudjetni mezondan nariga surish (0.31). "
         "Uchinchidan, kaskadning farqlovchi qiymati qatorlarning "
         "birortasida emas, ULARNING ORASIDA: qaysi qatordan qaysisiga "
         "o'tish xavfsizligini faqat u oldindan aytadi.", size=10)

    # ===================== 5. MUHOKAMA =====================
    doc.add_page_break()
    h(doc, "5. Muhokama", 1)

    para(doc, "Ikki strukturaviy vositaning yig'ma hukmi.", bold=True,
         size=10)
    para(doc,
         "Natijalar bo'ylab sochilgan xulosalarni bir joyga yig'amiz. "
         "USTUN TANLASH (ID oilasi) bir ishorali faollashuv konusi "
         "bo'lgan modellarda ishlaydi — Whisper enkoderida 17.1% tekin "
         "(dWER -0.0014) — va kollinearlik yo'q joyda majburlansa "
         "qimmatga tushadi (mBERT ~6%, Llama 18-26%). PAST-RANK YOYILMA "
         "(faollashuvga sezgir SVD) teskari sohada ishlaydi — gated, "
         "taqsimlangan-ortiqchali modellarda (Llama: rank 46% da +0.7%) "
         "— va uch shart bilan: qator/rank >= 10-20 (aks holda "
         "kalibrlashni yodlaydi), rank byudjet-optimal taqsimlanadi "
         "(bir xildan -0.0457 yaxshi), qayta ishlatishi past operatorga "
         "majburlanmaydi (dekoderda vaqt bermay 0.43 WER turadi). "
         "Qaysi vosita qachonligini faollashuv ishorasi modelni ishga "
         "tushirmasdan aytadi (4.15) — kaskadning ikki shoxchali "
         "bo'lishining o'zi shu ikki sohaning mavjudligidan kelib "
         "chiqadi.", size=10)

    h(doc, "5.1. Usul nima uchun CPU ga bog'langan va bu qayerda ahamiyatli", 2)
    para(doc,
         "Ushbu ishning barcha o'lchovlari CPU da o'tkazilgan. Buni cheklov "
         "sifatida emas, usulning TA'RIF SOHASI sifatida qarash kerak, chunki "
         "kaskadning kirish parametri faqat shu apparatda mavjud. Maqsad "
         "kafolatlangan umumiy keshning hajmidan chiqariladi va uning ma'nosi "
         "uch shartga tayanadi: kesh hajmi aniq va barcha yadrolar uchun "
         "umumiy; vaznlar turg'un, faollashuvlar esa oqim tarzida o'tadi; "
         "qayta ishlatish koeffitsiyenti operatorning hisoblash yoki xotira "
         "bilan cheklanganini belgilaydi.")
    para(doc,
         "Grafik protsessorlarda bu tuzilma mavjud emas. U yerda L2 keshi model "
         "hajmiga nisbatan juda kichik, har bir hisoblash bloki o'zining "
         "mahalliy xotirasiga ega, va hukmron cheklov kesh sig'imi emas, HBM "
         "o'tkazuvchanligi hisoblanadi. Shu sababli alpha x L3 byudjetining "
         "GPU da to'g'ridan-to'g'ri analogi yo'q va uni sun'iy ko'chirish "
         "usulning asosini yo'qotadi. Ya'ni taklif etilgan yondashuv GPU "
         "xizmatiga raqib emas — u GPU MAVJUD BO'LMAGAN muhitda siqish "
         "maqsadini qanday chiqarish kerakligini ko'rsatadi.")
    para(doc,
         "Bunday muhitlar ushbu ishda o'rganilgan uchala model uchun ham "
         "haqiqiy, ammo bir xil darajada emas. mBERT sinfidagi kodlovchilar "
         "klassifikatsiya, nomlangan obyektlarni ajratish va qidiruvni "
         "tartiblash uchun ishlatiladi; bu vazifalarda partiya kichik va "
         "kechikish talabi qattiq bo'lgani uchun xizmat asosan CPU da "
         "quriladi. Whisper sinfidagi ASR modellari noutbuk, avtomobil, "
         "tibbiy diktovka va tashkilot ichidagi call-center qurilmalarida "
         "ishlaydi; whisper.cpp kabi loyihalarning mavjudligi bu talabning "
         "bevosita ko'rsatkichi. Llama sinfidagi modellar uchun manzara "
         "murakkabroq: katta til modellari xizmati asosan GPU da quriladi, "
         "ammo llama.cpp va shunga o'xshash muhitlar butunlay CPU ga "
         "yo'naltirilgan va 3B atrofidagi o'lcham aynan shu maqsad uchun "
         "tanlanadi.")
    para(doc,
         "Ishning til konteksti bu yo'nalishni kuchaytiradi. Kam resursli "
         "tillar uchun nutqni tanish tizimlari ko'pincha markazlashtirilgan "
         "GPU infratuzilmasi bo'lmagan sharoitda joylashtiriladi, ya'ni "
         "apparat qat'iy va kamtar bo'ladi — bu esa maqsadni qo'lda tanlash "
         "emas, mavjud apparatdan chiqarish zarur bo'lgan holatning o'zi.")
    para(doc,
         "Ikkita halol qayd. Birinchidan, taqqoslanayotgan GPTQ va AWQ "
         "usullari dastlab GPU xizmati uchun ishlab chiqilgan; ular "
         "apparatdan mustaqil VAZN TRANSFORMATSIYALARI bo'lgani uchun bu yerda "
         "qo'llanishi o'rinli, farq faqat bajarish yadrolarida, ammo ularning "
         "nashr etilgan tezlik ko'rsatkichlari bu muhitga taalluqli emas. "
         "Ikkinchidan, 4.12-bo'limda ko'rsatilganidek, Llama sinfi uchun x86 "
         "da haqiqiy vazn-only INT8 yadrolari mavjud emas, shuning uchun CPU "
         "dagi LLM tezligi haqidagi da'volar bu ishda keltirilmaydi. Shu "
         "sababli usulning amaliy ahamiyati eng kuchli tarzda mBERT va Whisper "
         "sinflarida namoyon bo'ladi.", italic=True, size=10)

    h(doc, "5.2. Umumlashadigan natijalar", 2)
    para(doc,
         "Uchta natija o'rganilgan konkret modellardan tashqariga umumlashadi. "
         "Birinchidan, kesh-bog'langan maqsad ixtiyoriy giperparametrni chiqariladigan "
         "kattalikka aylantiradi va chiqarilgan qiymat uchala arxitekturada ham "
         "informativ bo'lib chiqdi — u audio enkoderda FFN operatorlarini, mBERT da "
         "lug'at matritsasini, Llama da esa ham FFN, ham head operatorlarini "
         "ko'rsatdi. Ikkinchidan, yutish qonuni qatlam darajasidagi bilvosita "
         "mezonlar nega chalg'itishini tushuntiradi: residual oqim nisbiy xatoni "
         "suyultirgani uchun katta lokal buzilish zararsiz bo'lishi, boshqacha "
         "taqsimlangan o'rtacha buzilish esa zararli bo'lishi mumkin. Uchinchidan, "
         "kompensatsiya va kvantlash granulyarligi orasidagi bog'liqlik — bu ikki "
         "komponent alohida baholanganda ko'rinmaydigan loyihaviy cheklov.")
    para(doc,
         "Kvantlash komponenti bo'yicha natija aniq: GPTQ ning Hessian orqali xato "
         "kompensatsiyasi bizning kalibrlangan masshtabimizdan ustun, shuning uchun "
         "kaskadning majburiy kvantlash bosqichida GPTQ tavsiya etiladi. Bu ishning "
         "pozitsiyasini kuchaytiradi: taklif etilgan usul kvantlash sxemasini "
         "takomillashtirishga da'vo qilmaydi, balki unga ORTOGONAL o'q qo'shadi. "
         "18-jadval buni empirik tasdiqlaydi — GPTQ ustiga qo'llangan strukturaviy "
         "qisqartirish qo'shimcha 11% xotira tejaydi va aniqlikni o'zgartirmaydi. "
         "Ammo o'sha jadval ortogonallikning SHARTLI ekanini ham ko'rsatadi: "
         "kvantlagich oddiy yaxlitlashga almashtirilsa, birikma FP32 dan sezilarli "
         "yomonlashadi, holbuki qisqartirishsiz ikkala kvantlagich teng. Demak "
         "ikki o'q resurs uchun raqobatlashmaydi, lekin ular BOG'LIQ: "
         "kompensatsiya qoldiq xato qoldiradi va uni yutish uchun xatoni qayta "
         "taqsimlaydigan kvantlagich kerak. Bu 4.4-bo'limdagi granulyarlik "
         "cheklovining tabiiy davomi — u yerda per-tensor yetarli emasligi, bu "
         "yerda esa per-channel ham o'z-o'zicha yetarli emasligi aniqlanadi.")
    para(doc,
         "To'rtinchi umumlashadigan natija 4.13-bo'limdan chiqadi va u yuqoridagi "
         "yutish qonunini muhim tarzda cheklaydi. Kvantlash masshtabini har bir "
         "operatorda alohida optimallashtirish standart amaliyot bo'lib, biz uni "
         "INT8 da foydali deb topdik; INT4 da esa xuddi shu optimallashtirish "
         "tarmoq sifatini deyarli ikki barobar yomonlashtirdi, garchi har bir "
         "operatorning xatosi 26% kamaygan bo'lsa ham. Sabab shundaki, vazn "
         "xatosini minimallashtiruvchi masshtab SILJIGAN bo'ladi — u har bir "
         "kanalni ozgina kichraytiradi — va bir xil yo'nalishdagi siljish "
         "chuqurlik bo'ylab ko'payadi, tasodifiy xato esa so'nadi. Bundan "
         "quyidagi umumiy tavsiya kelib chiqadi: chuqur tarmoqlarda operator "
         "darajasidagi maqsad funksiyasi xato KATTALIGINI emas, SILJIMAGANLIKNI "
         "muhofaza qilishi kerak. Buni ta'minlash uchun xotira formatini "
         "o'zgartirish shart emas — kanal bo'yicha bitta ko'paytuvchini chiqish "
         "domenida qayta tanlash perplexity ni 12.799 dan 8.246 ga tushirdi. "
         "Bu kuzatuv ushbu ishning kaskadidan tashqarida, ixtiyoriy post-training "
         "kvantlash sxemasiga tegishli.")
    para(doc,
         "Salbiy natijalar ham xuddi shunday informativ. Dastlabki loyihani "
         "asoslagan CUR konstruksiyasi past-rankli yaqinlashtiruvchi sifatida "
         "raqobatbardosh emas: u r^2 jarimasini olib yuradi va Ekart-Yang erishish "
         "mumkin bo'lgan vazn xatosini chegaralaydi. Shunga qaramay uning ustun "
         "tanlash mezoni 134/135 taqqoslashda leverage skorlaridan ustun turadi va "
         "asosidagi guruhlash mexanizmi yaqinlashtirishdan aniq strukturaviy olib "
         "tashlashga yo'naltirilganda muvaffaqiyat qozonadi. Shuning uchun biz "
         "kalibrlashga asoslangan ustun tanlashni yoyilma usuli emas, ortiqchalik "
         "detektori sifatida talqin qilishni tavsiya qilamiz.")
    para(doc,
         "Dalil qamrovi. 4.7-bo'limdagi yutish qonuni operator darajasidagi "
         "xato bilan vazifa mezoni 40 barobargacha ajralishi mumkinligini "
         "ko'rsatgani uchun, har bir da'vo uchun dalil QAYSI darajada "
         "to'planganini ochiq ko'rsatamiz. 61-jadval ushbu ishning barcha "
         "asosiy tasdiqlarini shu bo'yicha guruhlaydi.")
    table(doc, "61-jadval. Da'volar bo'yicha dalil darajasi va baholash to'plami.",
          ["Da'vo", "Daraja", "Mezon", "Baholash to'plami"],
          [["Kesh maqsadi — dekoder qarori", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Kesh maqsadi — enkoder qarori", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Butun model: kaskad vs bir xil siyosatlar", "uchdan-uchgacha",
            "WER", "TEST, 300 namuna"],
           ["Qayta ishlatish argumenti (enkoder vs dekoder)", "apparat",
            "VTune: Memory/DRAM bound, CPI", "butun model, bitta oqim"],
           ["Strukturaviy mezon vs magnitude/Wanda", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Kompensatsiyaning zarurligi (ablation)", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Funksional guruhlash (strukturaviy olib tashlash)",
            "uchdan-uchgacha", "WER", "TEST, 300 namuna"],
           ["Kvantlash bilan ortogonallik (2x2)", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Siljishning chuqurlik bo'ylab to'planishi", "uchdan-uchgacha",
            "perplexity", "WikiText-2"],
           ["Kompensatsiya per-channel ni talab qiladi", "uchdan-uchgacha",
            "WER", "validation, 80 namuna"],
           ["Byudjet-optimal rank taqsimoti", "uchdan-uchgacha", "WER",
            "TEST, 300 namuna"],
           ["Past-rank shoxchasi (Llama)", "uchdan-uchgacha", "perplexity",
            "o'zbek matni"],
           ["Kvantlash rejimi (vazn-only)", "uchdan-uchgacha", "perplexity",
            "o'zbek matni"],
           ["Kalibrlangan masshtab (5-jadval)", "operator", "E_loc", "—"],
           ["Kalibrlash hajmi qoidasi", "operator", "E_loc (fit/held-out)", "—"],
           ["CUR va SVD ustun tanlash", "operator", "E_loc", "—"],
           ["mBERT ga ko'chish", "operator", "ortiqchalik, E_loc", "—"]],
          good_rows=(0, 1, 2, 3, 4, 6))
    para(doc,
         "Jadvaldan ikkita chegara ko'rinadi. Birinchidan, oxirgi to'rt qator "
         "operator darajasida qoladi; ular ishning diagnostik qismi bo'lib, "
         "vazifa darajasidagi da'vo sifatida keltirilmaydi. Ikkinchidan, "
         "uchdan-uchgacha o'lchanganlarning bir qismi validation to'plamidan "
         "olingan, ya'ni 4.9-bo'limda tavsiflangan kalibrlashga qo'shni "
         "protokolda; markaziy da'volar esa mustaqil TEST splitiga "
         "ko'chirilgan. Bu taqsimot ataylab shunday: eng kuchli protokol "
         "kaskadning qarorlari va strukturaviy o'qqa, ya'ni ishning asosiy "
         "hissasiga yo'naltirilgan.", italic=True, size=10)

    para(doc,
         "Eksport yo'lining cheklovi. ASR tajribalarida joylashtirilgan model "
         "ONNX Runtime ning quantize_dynamic funksiyasi orqali hosil qilinadi, "
         "u esa kvantlash masshtabini o'ziga berilgan float vaznlardan QAYTA "
         "hisoblaydi. Sxemaning o'zi bizning ta'rifimiz bilan aynan mos "
         "tushadi — simmetrik int8, nol nuqtasi 0, chiqish kanali bo'yicha "
         "s = max|w|/127 (o'lchangan nisbat 1.0000) — ammo masshtabni tashqaridan "
         "belgilash imkoni yo'q. Oqibati aniq va o'lchandi: kodi В±127 ga yetgan "
         "kanal aynan saqlanadi, yetmagan kanal esa qayta yaxlitlanib yarim "
         "qadamgacha siljiydi. Haqiqiy enkoder vaznlarida kanallarning 78-93% i "
         "aynan o'tadi.")
    para(doc,
         "Bundan uchta natija kelib chiqadi. Birinchidan, strukturaviy o'q "
         "ta'sirlanmaydi: ONNX operator SHAKLLARINI aniq saqlaydi, shuning uchun "
         "14-, 15- va 18-jadvallardagi barcha da'volar kuchida qoladi. "
         "Ikkinchidan, GPTQ tarmog'i to'g'ri joylashtirilgan, chunki u vaznlarni "
         "eksportdan oldin o'z panjarasiga yaxlitlaydi; buni tekshirish uchun "
         "vaznlar joylashtirilgan ONNX fayllaridan qayta o'qib o'lchandi va "
         "GPTQ ning ustunligi 57.3% chiqdi, ya'ni xotiradagi 54.3% ga mos — "
         "eksport uni yemaydi. Uchinchidan, 5-jadvaldagi kalibrlangan masshtab "
         "ASR traktida uchdan-uchgacha joylashtirilmagan va bu natija operator "
         "darajasida qoladi. Qoldiq siljish (yarim kvantlash qadami) INT8 da "
         "tarmoq sezadigan chegaradan ancha past — buni 4.7-bo'limdagi yutish "
         "qonuni, 49-jadvaldagi INT8 natijasi (yaxlitlash FP32 ni uch xonagacha "
         "takrorlaydi) va GPTQ ning 57% operator ustunligining WER ga umuman "
         "o'tmasligi mustaqil ravishda ko'rsatadi. Shunga qaramay, kalibrlangan "
         "masshtabni joylashtirilgan artefaktda ko'rsatish uchun kvantlangan "
         "tenzorlarni bevosita yozish kerak bo'ladi; format buni to'liq "
         "qo'llab-quvvatlaydi va bu keyingi ishning aniq, tor vazifasi.")
    para(doc,
         "Cheklovlar. Asosiy taqqoslashlar 300 namunada o'tkazilgan va ishonch "
         "oraliqlari taxminan 0.01 WER kengligida qoladi, ya'ni undan kichik "
         "farqlar hamon ajratilmaydi; to'liq test splitida takrorlash ularni "
         "yanada toraytiradi. "
         "Kesh-bog'langan maqsad bitta kesh konfiguratsiyasida tekshirilgan va "
         "tahlilni turli L3 hajmli platformalarda takrorlash markaziy da'voning eng "
         "to'g'ridan-to'g'ri sinovi hisoblanadi. Til modeli tajribalarida sifat "
         "vazn-only INT8 sxemasida o'lchandi; bu platformada real vazn-only yadrolari "
         "mavjud emasligi 4.12-bo'limda o'lchab ko'rsatilgan (dinamik kvantlash "
         "faollashuvlarni ham kvantlab perplexity ni 15.3x yomonlashtiradi, "
         "_weight_int8pack_mm esa x86 da FP32 dan 71x sekin), shuning uchun Llama "
         "uchun tezlik raqamlari keltirilmaydi. WikiText-2 taqqoslashida "
         "kalibrlash byudjeti 4096 qator bilan cheklangan, down_proj uchun esa "
         "bu Hessian ni rank-kamchil qoldiradi (n = 8640); barcha usullar bir "
         "xil ma'lumotda ishlagani uchun taqqoslash ichki jihatdan adolatli, "
         "ammo GPTQ ning bu sharoitdagi natijasini uning nashr etilgan "
         "ko'rsatkichlari bilan tenglashtirib bo'lmaydi. Nihoyat, E_glob "
         "dan vazifa mezonlariga o'tish empirik bo'lib qolmoqda: WER ni operator "
         "darajasidagi xatodan bashorat qilish kaskadni to'liq bashoratli qilar edi, "
         "hozircha u bunday emas.")

    # ===================== 6. XULOSALAR =====================
    h(doc, "6. Xulosalar", 1)
    para(doc,
         "Biz siqish maqsadi kesh topologiyasidan chiqariladigan va o'zgartirish har "
         "operator uchun o'lchov asosida tanlanadigan kaskadni taqdim etdik. "
         "Kvantlash bosqichida o'lchov eng yaxshi mavjud usulni (GPTQ) tanlaydi, "
         "ishning o'z hissasi esa unga ortogonal strukturaviy o'qdir. "
         "Whisper-medium o'zbek ASR enkoderida funksional jihatdan ortiqcha FFN "
         "kanallarini olib tashlash GPTQ ustiga qo'shilganda xotirani 11% "
         "kamaytiradi va 3% tezlik beradi, so'z xatoligini esa o'zgartirmaydi "
         "(dWER = -0.0014, IO [-0.0111, +0.0096]); ikkala variant ham FP32 dan "
         "farqlanmaydi. Bu Common Voice TEST splitining 300 namunasida, "
         "kalibrlashdan boshqa taqsimotda tasdiqlangan. Butun model "
         "darajasida kaskadning asosiy qiymati esa siqish darajasida emas, "
         "TO'XTASH nuqtasini to'g'ri aniqlashda: u 4.14x da to'xtaydi va "
         "sifatni saqlaydi, xuddi shu retseptni 5.34x gacha davom ettirish "
         "esa WER ni 0.1833 dan 0.6101 ga ko'taradi. Bu qaror apparat "
         "hisoblagichlari bilan ham oqlanadi: dekoder o'z tabiiy rejimida "
         "enkoderdan 1.9 barobar ko'proq xotira bilan cheklangan (Memory Bound "
         "18.2% ga qarshi 9.7%), shuning uchun u yerda past-rank 22% xotira "
         "tejab, vaqtdan hech narsa bermaydi. Butun model bo'yicha kaskad "
         "xotira to'xtashini 2.41 barobar, umumiy vaqtni esa 1.91 barobar "
         "qisqartiradi, ya'ni model nisbatan kamroq xotira bilan cheklangan "
         "holga o'tadi. Ayni paytda teng byudjetda usullar xotira "
         "xatti-harakati bo'yicha ajralmaydi — farqlar o'lchov "
         "o'zgaruvchanligi darajasida — ya'ni bu ustunlik hajm ustunligining "
         "aksi, algoritmning alohida xossasi emas. Ayni paytda bu xossa "
         "shartli: kvantlagich oddiy yaxlitlashga almashtirilganda birikma FP32 "
         "dan sezilarli yomonlashadi, ya'ni strukturaviy o'q xatoni kompensatsiya "
         "qiladigan kvantlagichni talab qiladi — bu kaskadning vosita tanlovi "
         "uchun aniq amaliy qoida. Usul qayta o'qitishni talab qilmaydi. "
         "Mustaqil ahamiyatga "
         "ega uchta yordamchi natija: kompensatsiya per-channel kvantlashni ixtiyoriy "
         "emas, majburiy qiladi; kalibrlashga asoslangan past-rank yoyilma qator/rank "
         "nisbati kamida 10-20 bo'lishini talab qiladi; transformer bloklari lokal "
         "buzilishlarni kuchli yutadi va bu qatlam darajasidagi bilvosita mezonlar "
         "nega vazifa sifatining ishonchsiz bashoratchisi ekanini tushuntiradi. "
         "Arxitekturalararo tekshiruv diagnostika ko'chishini, aniq vosita esa "
         "ko'chmasligini ko'rsatadi va bu qat'iy belgilangan siqish retseptlari "
         "o'rniga o'lchovga asoslangan yondashuv foydasiga dalil bo'ladi.")
    para(doc,
         "Standart WikiText-2 benchmarkida (FP32 perplexity 7.547, nashr etilgan "
         "qiymatlar bilan mos) to'rtinchi, kaskaddan mustaqil natija olindi. "
         "INT4 da kvantlash masshtabini vazn xatosi bo'yicha optimallashtirish "
         "har bir operatorda 26% aniqroq bo'lsa-da, tarmoq perplexity sini "
         "1.696x ga yomonlashtiradi, chunki bunday masshtab har bir kanalni bir "
         "xil yo'nalishda ~1% kichraytiradi va bu siljish 78 ta operator bo'ylab "
         "ko'payadi (0.44 karra susayish). Masshtabni chiqish domenida qayta "
         "tanlash — butun sonli kodlarni, bit kengligini va xotira formatini "
         "o'zgartirmasdan — perplexity ni 8.246 ga tushiradi va usulni AWQ bilan "
         "raqobatbardosh qiladi. Siljish bit kengligiga keskin bog'liq: ish "
         "nuqtasi bo'lgan INT8 da u 100 barobar kichik (gain 0.9999) va usul "
         "uchdan-uchgacha RTN dan ozgina ustun turadi (7.549 va 7.550, FP32 "
         "esa 7.547), ya'ni kaskadning asosiy natijalari o'zgarmaydi. Amaliy "
         "xulosa: chuqur tarmoqlar uchun post-training kvantlashda operator "
         "darajasidagi maqsad xato kattaligini emas, siljimaganlikni muhofaza "
         "qilishi kerak.")

    # ===================== YAKUNIY BO'LIMLAR =====================
    h(doc, "Mualliflar hissasi", 1)
    para(doc, "Konseptualizatsiya, X.Y.; metodologiya, X.Y.; dasturiy ta'minot, X.Y.; "
              "validatsiya, X.Y. va Z.W.; formal tahlil, X.Y.; tadqiqot, X.Y.; "
              "ma'lumotlarni tayyorlash, X.Y.; qo'lyozmani yozish, X.Y.; ko'rib chiqish "
              "va tahrirlash, Z.W.; vizualizatsiya, X.Y.; rahbarlik, Z.W. Barcha "
              "mualliflar qo'lyozmaning nashr etilgan variantini o'qib chiqdi va "
              "roziligini bildirdi.", size=9.5)
    h(doc, "Moliyalashtirish", 1)
    para(doc, "Ushbu tadqiqot tashqi moliyalashtirishsiz bajarilgan.", size=9.5)
    h(doc, "Ma'lumotlar mavjudligi", 1)
    para(doc, "Common Voice o'zbek korpusi ochiq foydalanishda. Dastur kodi, "
              "eksperiment skriptlari va barcha o'lchangan natija fayllari "
              "mualliflardan so'rov asosida taqdim etiladi.", size=9.5)
    h(doc, "Manfaatlar to'qnashuvi", 1)
    para(doc, "Mualliflar manfaatlar to'qnashuvi yo'qligini bildiradi.", size=9.5)

    h(doc, "Adabiyotlar", 1)
    para(doc, "Eslatma muallifga: quyidagi manbalarning mualliflari, sarlavhalari, "
              "nashr joyi va yili tekshirilgan; arXiv identifikatorlari keltirilgan. "
              "Jurnal maqolalari uchun DOI va sahifa diapazonlarini yuborishdan oldin "
              "asl nashrdan tasdiqlang.", italic=True, size=8.5, color=CRIT)

    para(doc, "O'qitilgandan keyingi kvantlash", bold=True, size=9.5)
    mono(doc,
         "1.  Frantar, E.; Ashkboos, S.; Hoefler, T.; Alistarh, D. GPTQ: Accurate\n"
         "    Post-Training Quantization for Generative Pre-trained Transformers. In\n"
         "    Proceedings of the International Conference on Learning Representations\n"
         "    (ICLR), Kigali, Rwanda, 1-5 May 2023. arXiv:2210.17323.\n"
         "2.  Lin, J.; Tang, J.; Tang, H.; Yang, S.; Dang, X.; Gan, C.; Han, S. AWQ:\n"
         "    Activation-Aware Weight Quantization for On-Device LLM Compression and\n"
         "    Acceleration. In Proceedings of Machine Learning and Systems (MLSys),\n"
         "    Santa Clara, CA, USA, 13-16 May 2024. arXiv:2306.00978.\n"
         "3.  Xiao, G.; Lin, J.; Seznec, M.; Wu, H.; Demouth, J.; Han, S. SmoothQuant:\n"
         "    Accurate and Efficient Post-Training Quantization for Large Language\n"
         "    Models. In Proceedings of the International Conference on Machine\n"
         "    Learning (ICML), Honolulu, HI, USA, 23-29 July 2023. arXiv:2211.10438.\n"
         "4.  Dettmers, T.; Lewis, M.; Belkada, Y.; Zettlemoyer, L. LLM.int8():\n"
         "    8-bit Matrix Multiplication for Transformers at Scale. In Advances in\n"
         "    Neural Information Processing Systems (NeurIPS), New Orleans, LA, USA,\n"
         "    28 November - 9 December 2022. arXiv:2208.07339.\n"
         "5.  Yao, Z.; Yazdani Aminabadi, R.; Zhang, M.; Wu, X.; Li, C.; He, Y.\n"
         "    ZeroQuant: Efficient and Affordable Post-Training Quantization for\n"
         "    Large-Scale Transformers. In Advances in Neural Information Processing\n"
         "    Systems (NeurIPS), 2022. arXiv:2206.01861.\n"
         "6.  Frantar, E.; Alistarh, D. Optimal Brain Compression: A Framework for\n"
         "    Accurate Post-Training Quantization and Pruning. In Advances in Neural\n"
         "    Information Processing Systems (NeurIPS), 2022. arXiv:2208.11580.\n"
         "7.  Nagel, M.; Amjad, R.A.; van Baalen, M.; Louizos, C.; Blankevoort, T. Up\n"
         "    or Down? Adaptive Rounding for Post-Training Quantization. In\n"
         "    Proceedings of the International Conference on Machine Learning (ICML),\n"
         "    2020. arXiv:2004.10568.\n"
         "8.  Li, Y.; Gong, R.; Tan, X.; Yang, Y.; Hu, P.; Zhang, Q.; Yu, F.; Wang, W.;\n"
         "    Gu, S. BRECQ: Pushing the Limit of Post-Training Quantization by Block\n"
         "    Reconstruction. In Proceedings of the International Conference on\n"
         "    Learning Representations (ICLR), 2021. arXiv:2102.05426.\n"
         "9.  Jacob, B.; Kligys, S.; Chen, B.; Zhu, M.; Tang, M.; Howard, A.; Adam, H.;\n"
         "    Kalenichenko, D. Quantization and Training of Neural Networks for\n"
         "    Efficient Integer-Arithmetic-Only Inference. In Proceedings of the IEEE\n"
         "    Conference on Computer Vision and Pattern Recognition (CVPR), Salt Lake\n"
         "    City, UT, USA, 18-23 June 2018; pp. 2704-2713.\n"
         "10. Nagel, M.; Fournarakis, M.; Amjad, R.A.; Bondarenko, Y.; van Baalen, M.;\n"
         "    Blankevoort, T. A White Paper on Neural Network Quantization. arXiv 2021,\n"
         "    arXiv:2106.08295.\n"
         "11. Gholami, A.; Kim, S.; Dong, Z.; Yao, Z.; Mahoney, M.W.; Keutzer, K. A\n"
         "    Survey of Quantization Methods for Efficient Neural Network Inference.\n"
         "    arXiv 2021, arXiv:2103.13630.\n"
         "12. Ashkboos, S.; Mohtashami, A.; Croci, M.L.; Li, B.; Jaggi, M.;\n"
         "    Alistarh, D.; Hoefler, T.; Hensman, J. QuaRot: Outlier-Free 4-Bit\n"
         "    Inference in Rotated LLMs. In Advances in Neural Information Processing\n"
         "    Systems (NeurIPS), 2024. arXiv:2404.00456.", 8)

    para(doc, "Past-rankli yoyilma va matritsa yaqinlashtirish", bold=True, size=9.5)
    mono(doc,
         "13. Eckart, C.; Young, G. The Approximation of One Matrix by Another of\n"
         "    Lower Rank. Psychometrika 1936, 1, 211-218.\n"
         "14. Hsu, Y.-C.; Hua, T.; Chang, S.; Lou, Q.; Shen, Y.; Jin, H. Language Model\n"
         "    Compression with Weighted Low-Rank Factorization. In Proceedings of the\n"
         "    International Conference on Learning Representations (ICLR), 2022.\n"
         "    arXiv:2207.00112.\n"
         "15. Yuan, Z.; Shang, Y.; Song, Y.; Wu, Q.; Yan, Y.; Sun, G. ASVD:\n"
         "    Activation-Aware Singular Value Decomposition for Compressing Large\n"
         "    Language Models. arXiv 2023, arXiv:2312.05821.\n"
         "16. Wang, X.; Zheng, Y.; Wan, Z.; Zhang, M. SVD-LLM: Truncation-Aware\n"
         "    Singular Value Decomposition for Large Language Model Compression.\n"
         "    arXiv 2024, arXiv:2403.07378.\n"
         "17. Kaushal, A.; Vaidhya, T.; Rish, I. LORD: Low Rank Decomposition of\n"
         "    Monolingual Code LLMs for One-Shot Compression. arXiv 2023,\n"
         "    arXiv:2309.14021.\n"
         "18. Denton, E.; Zaremba, W.; Bruna, J.; LeCun, Y.; Fergus, R. Exploiting\n"
         "    Linear Structure Within Convolutional Networks for Efficient Evaluation.\n"
         "    In Advances in Neural Information Processing Systems (NeurIPS),\n"
         "    Montreal, QC, Canada, 8-13 December 2014.\n"
         "19. Sainath, T.N.; Kingsbury, B.; Sindhwani, V.; Arisoy, E.; Ramabhadran, B.\n"
         "    Low-Rank Matrix Factorization for Deep Neural Network Training with\n"
         "    High-Dimensional Output Targets. In Proceedings of the IEEE\n"
         "    International Conference on Acoustics, Speech and Signal Processing\n"
         "    (ICASSP), Vancouver, BC, Canada, 26-31 May 2013; pp. 6655-6659.\n"
         "20. Lan, Z.; Chen, M.; Goodman, S.; Gimpel, K.; Sharma, P.; Soricut, R.\n"
         "    ALBERT: A Lite BERT for Self-Supervised Learning of Language\n"
         "    Representations. In Proceedings of the International Conference on\n"
         "    Learning Representations (ICLR), 2020. arXiv:1909.11942.\n"
         "21. Mahoney, M.W.; Drineas, P. CUR Matrix Decompositions for Improved Data\n"
         "    Analysis. Proc. Natl. Acad. Sci. USA 2009, 106, 697-702.\n"
         "22. Drineas, P.; Mahoney, M.W.; Muthukrishnan, S. Relative-Error CUR Matrix\n"
         "    Decompositions. SIAM J. Matrix Anal. Appl. 2008, 30, 844-881.\n"
         "23. Halko, N.; Martinsson, P.-G.; Tropp, J.A. Finding Structure with\n"
         "    Randomness: Probabilistic Algorithms for Constructing Approximate Matrix\n"
         "    Decompositions. SIAM Rev. 2011, 53, 217-288.", 8)

    para(doc, "Strukturaviy va strukturaviy bo'lmagan qisqartirish", bold=True, size=9.5)
    mono(doc,
         "24. Ma, X.; Fang, G.; Wang, X. LLM-Pruner: On the Structural Pruning of Large\n"
         "    Language Models. In Advances in Neural Information Processing Systems\n"
         "    (NeurIPS), 2023. arXiv:2305.11627.\n"
         "25. Ashkboos, S.; Croci, M.L.; Nascimento, M.G.; Hoefler, T.; Hensman, J.\n"
         "    SliceGPT: Compress Large Language Models by Deleting Rows and Columns.\n"
         "    In Proceedings of the International Conference on Learning\n"
         "    Representations (ICLR), 2024. arXiv:2401.15024.\n"
         "26. Xia, M.; Gao, T.; Zeng, Z.; Chen, D. Sheared LLaMA: Accelerating Language\n"
         "    Model Pre-Training via Structured Pruning. In Proceedings of the\n"
         "    International Conference on Learning Representations (ICLR), 2024.\n"
         "    arXiv:2310.06694.\n"
         "27. Sun, M.; Liu, Z.; Bair, A.; Kolter, J.Z. A Simple and Effective Pruning\n"
         "    Approach for Large Language Models. In Proceedings of the International\n"
         "    Conference on Learning Representations (ICLR), 2024. arXiv:2306.11695.\n"
         "28. An, Y.; Zhao, X.; Yu, T.; Tang, M.; Wang, J. Fluctuation-Based Adaptive\n"
         "    Structured Pruning for Large Language Models. In Proceedings of the AAAI\n"
         "    Conference on Artificial Intelligence, 2024. arXiv:2312.11983.\n"
         "29. Michel, P.; Levy, O.; Neubig, G. Are Sixteen Heads Really Better than\n"
         "    One? In Advances in Neural Information Processing Systems (NeurIPS),\n"
         "    Vancouver, BC, Canada, 8-14 December 2019. arXiv:1905.10650.\n"
         "30. Voita, E.; Talbot, D.; Moiseev, F.; Sennrich, R.; Titov, I. Analyzing\n"
         "    Multi-Head Self-Attention: Specialized Heads Do the Heavy Lifting, the\n"
         "    Rest Can Be Pruned. In Proceedings of the Annual Meeting of the\n"
         "    Association for Computational Linguistics (ACL), Florence, Italy,\n"
         "    28 July - 2 August 2019. arXiv:1905.09418.\n"
         "31. Han, S.; Pool, J.; Tran, J.; Dally, W.J. Learning Both Weights and\n"
         "    Connections for Efficient Neural Networks. In Advances in Neural\n"
         "    Information Processing Systems (NeurIPS), Montreal, QC, Canada,\n"
         "    7-12 December 2015. arXiv:1506.02626.\n"
         "32. LeCun, Y.; Denker, J.S.; Solla, S.A. Optimal Brain Damage. In Advances in\n"
         "    Neural Information Processing Systems (NeurIPS), Denver, CO, USA, 1989.", 8)

    para(doc, "Modellar, arxitekturalar va korpuslar", bold=True, size=9.5)
    mono(doc,
         "33. Vaswani, A.; Shazeer, N.; Parmar, N.; Uszkoreit, J.; Jones, L.;\n"
         "    Gomez, A.N.; Kaiser, L.; Polosukhin, I. Attention Is All You Need. In\n"
         "    Advances in Neural Information Processing Systems (NeurIPS), Long Beach,\n"
         "    CA, USA, 4-9 December 2017. arXiv:1706.03762.\n"
         "34. Radford, A.; Kim, J.W.; Xu, T.; Brockman, G.; McLeavey, C.; Sutskever, I.\n"
         "    Robust Speech Recognition via Large-Scale Weak Supervision. In\n"
         "    Proceedings of the International Conference on Machine Learning (ICML),\n"
         "    Honolulu, HI, USA, 23-29 July 2023. arXiv:2212.04356.\n"
         "35. Devlin, J.; Chang, M.-W.; Lee, K.; Toutanova, K. BERT: Pre-Training of\n"
         "    Deep Bidirectional Transformers for Language Understanding. In\n"
         "    Proceedings of NAACL-HLT, Minneapolis, MN, USA, 2-7 June 2019;\n"
         "    pp. 4171-4186.\n"
         "36. Conneau, A.; Khandelwal, K.; Goyal, N.; Chaudhary, V.; Wenzek, G.;\n"
         "    Guzman, F.; Grave, E.; Ott, M.; Zettlemoyer, L.; Stoyanov, V.\n"
         "    Unsupervised Cross-Lingual Representation Learning at Scale. In\n"
         "    Proceedings of the Annual Meeting of the Association for Computational\n"
         "    Linguistics (ACL), 2020. arXiv:1911.02116.\n"
         "37. Touvron, H.; Lavril, T.; Izacard, G.; Martinet, X.; Lachaux, M.-A.;\n"
         "    Lacroix, T.; Roziere, B.; Goyal, N.; Hambro, E.; Azhar, F.; et al.\n"
         "    LLaMA: Open and Efficient Foundation Language Models. arXiv 2023,\n"
         "    arXiv:2302.13971.\n"
         "38. Touvron, H.; Martin, L.; Stone, K.; Albert, P.; Almahairi, A.; Babaei, Y.;\n"
         "    Bashlykov, N.; Batra, S.; Bhargava, P.; Bhosale, S.; et al. Llama 2:\n"
         "    Open Foundation and Fine-Tuned Chat Models. arXiv 2023, arXiv:2307.09288.\n"
         "39. Geng, X.; Liu, H. OpenLLaMA: An Open Reproduction of LLaMA. 2023.\n"
         "    Available online: https://github.com/openlm-research/open_llama\n"
         "    (accessed on [DATE]).\n"
         "40. Shazeer, N. GLU Variants Improve Transformer. arXiv 2020,\n"
         "    arXiv:2002.05202.\n"
         "41. Ardila, R.; Branson, M.; Davis, K.; Henretty, M.; Kohler, M.; Meyer, J.;\n"
         "    Morais, R.; Saunders, L.; Tyers, F.M.; Weber, G. Common Voice: A\n"
         "    Massively-Multilingual Speech Corpus. In Proceedings of the Language\n"
         "    Resources and Evaluation Conference (LREC), Marseille, France,\n"
         "    11-16 May 2020; pp. 4218-4222.", 8)

    para(doc, "Apparat, ishlash unumdorligi va samarali inferens", bold=True, size=9.5)
    mono(doc,
         "42. Williams, S.; Waterman, A.; Patterson, D. Roofline: An Insightful Visual\n"
         "    Performance Model for Multicore Architectures. Commun. ACM 2009, 52,\n"
         "    65-76.\n"
         "43. Goto, K.; van de Geijn, R.A. Anatomy of High-Performance Matrix\n"
         "    Multiplication. ACM Trans. Math. Softw. 2008, 34, 1-25.\n"
         "44. Dao, T.; Fu, D.Y.; Ermon, S.; Rudra, A.; Re, C. FlashAttention: Fast and\n"
         "    Memory-Efficient Exact Attention with IO-Awareness. In Advances in Neural\n"
         "    Information Processing Systems (NeurIPS), 2022. arXiv:2205.14135.\n"
         "45. Hinton, G.; Vinyals, O.; Dean, J. Distilling the Knowledge in a Neural\n"
         "    Network. arXiv 2015, arXiv:1503.02531.\n"
         "46. Sanh, V.; Debut, L.; Chaumond, J.; Wolf, T. DistilBERT, a Distilled\n"
         "    Version of BERT: Smaller, Faster, Cheaper and Lighter. arXiv 2019,\n"
         "    arXiv:1910.01108.\n"
         "47. Gandhi, S.; von Platen, P.; Rush, A.M. Distil-Whisper: Robust Knowledge\n"
         "    Distillation via Large-Scale Pseudo Labelling. arXiv 2023,\n"
         "    arXiv:2311.00430.", 8)

    para(doc, "Statistik baholash", bold=True, size=9.5)
    mono(doc,
         "48. Efron, B.; Tibshirani, R.J. An Introduction to the Bootstrap; Chapman &\n"
         "    Hall/CRC: Boca Raton, FL, USA, 1993.\n"
         "49. Bisani, M.; Ney, H. Bootstrap Estimates for Confidence Intervals in ASR\n"
         "    Performance Evaluation. In Proceedings of the IEEE International\n"
         "    Conference on Acoustics, Speech, and Signal Processing (ICASSP),\n"
         "    Montreal, QC, Canada, 17-21 May 2004; pp. 409-412.\n"
         "50. Koehn, P. Statistical Significance Tests for Machine Translation\n"
         "    Evaluation. In Proceedings of the Conference on Empirical Methods in\n"
         "    Natural Language Processing (EMNLP), Barcelona, Spain, 25-26 July 2004;\n"
         "    pp. 388-395.\n"
         "51. Optimum Intel / OpenVINO. Joint Pruning, Quantization and\n"
         "    Distillation for Efficient Inference of Transformers, 2023.\n"
         "52. Qu, X. va b. Automatic Joint Structured Pruning and Quantization for\n"
         "    Efficient Neural Network Training and Compression (GETA). CVPR 2025.\n"
         "    arXiv:2502.16638.\n"
         "53. Joint Structural Pruning and Mixed-Precision Quantization for LLM\n"
         "    Compression. arXiv:2606.07819, 2026.", 8)

    doc.save(OUT)
    print(f"saqlandi: {OUT}")


if __name__ == "__main__":
    main()


