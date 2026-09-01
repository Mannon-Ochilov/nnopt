"""Build the experimental-results report as a .docx the author can paste
straight into the dissertation chapters.

Tables are real Word tables (not images, not preformatted text), so they
stay editable and restyleable in Word.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "../DSc_eksperimental_natijalar.docx"

ACCENT = RGBColor(0x0F, 0x64, 0x70)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def callout(doc, title, text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    p2 = doc.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True
    p2.paragraph_format.left_indent = Pt(18)
    return p2


def table(doc, caption, headers, rows, highlight_rows=(), fail_rows=()):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri in highlight_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in fail_rows:
                run.font.color.rgb = CRIT
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    h(doc, "DSc — Eksperimental natijalar hisoboti", 0)
    para(doc, "Neyron tarmoq operatorlarini optimallashtirish: o'lchangan natijalar va ular nimani ko'rsatadi",
         italic=True, size=12)
    para(doc, "Model: Kotib/uzbek_stt_v1 (Whisper-medium)  |  Decoder: 240 operator, 402 653 184 parametr  |  "
              "Mashina: 16 yadro, L3 = 24 MiB  |  Sana: 2026-08-11", size=9)
    para(doc, "Barcha raqamlar real o'lchov: real o'zbek nutqi kalibratsiyasi (Mozilla Common Voice), "
              "real ONNX Runtime (intra_op_threads=1), held-out namunalar.", size=9, italic=True)

    # ---- Key findings ----
    h(doc, "Asosiy topilmalar", 1)
    for tag, text in [
        ("1. INT8 bu model uchun mutlaqo bepul.",
         "4x siqish, 2.80x tezlanish, WER umuman o'zgarmaydi (0.0417 -> 0.0417). "
         "Transkripsiya belgi-ba-belgi bir xil."),
        ("2. E_loc / E_glob chalg'ituvchi mezon.",
         "E_glob = 0.23 -> hech qanday zarar yo'q; E_glob = 1.46 -> to'liq halokat. "
         "Mezon monoton emas; oraliqda keskin chegara bor va uni faqat WER ko'rsatadi."),
        ("3. Tezisning konseptual yadrosi tasdiqlandi.",
         "Vazn-optimallik chiqish-optimallik emas: 135/135 o'lchovda kalibrlashga asoslangan "
         "yaqinlashtirish vazn xatosida yomonroq, lekin chiqish xatosida yaxshiroq."),
        ("4. Ammo bu argumentni CUR emas, activation-aware SVD yutyapti.",
         "CUR act-aware SVD ga qarshi 0/135. Funksional klasterlash esa leverage-CUR ni "
         "134/135 yutadi — hissa haqiqiy, lekin u yaxshilayotgan ramka o'zi zaifroq."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(tag)
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(" " + text)
        r2.font.size = Pt(10)

    # ---- 1 ----
    h(doc, "1. Kesh hajmi — kaskadning birlamchi kirish parametri", 1)
    para(doc, "Kaskadning siqish maqsadi qo'lda tanlanmaydi; u kesh byudjetidan chiqariladi. "
              "Kafolatlangan umumiy kesh — barcha 16 yadro baham ko'radigan L3 = 24 MiB, "
              "samarali byudjet alpha*L3 = 0.7 x 24 = 16.8 MiB.")
    table(doc, "1-jadval. Kesh-bog'langan talab, granulyarlik bo'yicha",
          ["Granulyarlik", "Hajm (MiB)", "Talab", "Baho"],
          [["DECODER — per-operator", "16.0", "—", "allaqachon sig'adi"],
           ["DECODER — per-layer", "64.0", "3.81x", "HAQIQIY MAQSAD"],
           ["DECODER — whole-model", "1536.0", "91.43x", "imkonsiz"],
           ["ENCODER — per-operator", "16.0", "0.95x", "sig'adi"],
           ["ENCODER — per-layer", "48.0", "2.86x", "erishiladigan"],
           ["ENCODER — whole-model", "1152.0", "68.57x", "imkonsiz"]],
          highlight_rows=(1,))
    callout(doc, "Nega bu muhim:",
            "Decoder uchun kesh-bog'langan maqsad 3.81x, va INT8 aynan 4.00x beradi — ya'ni talabni "
            "qoplaydigan ENG YUMSHOQ usul. Kaskadning 'eng yumshoq yetarli o'zgarish' tamoyili aynan "
            "shuni tanlashi kerak edi, va tanladi. Ilgari 4x ixtiyoriy raqam edi; endi u kesh hajmidan "
            "chiqarilgan.", GOOD)

    # ---- 2 ----
    h(doc, "2. Kesh-sig'imi qachon ma'noga ega: qayta ishlatish tahlili", 1)
    table(doc, "2-jadval. Ish rejimi va cheklovchi resurs",
          ["Qism", "Pozitsiya/o'tish", "Vazn qayta ishlatilishi", "Cheklovchi", "Nima yordam beradi"],
          [["Encoder", "1500", "1500x", "hisoblash", "FLOPs kamaytirish (CUR/SVD)"],
           ["Decoder (batch=1)", "1", "1x", "xotira", "bayt kamaytirish (kvantlash)"]],
          highlight_rows=(0,))
    para(doc, "Decoderda vazn har token uchun bir marta o'qiladi, va layer 0 ga qaytguncha qolgan 23 qatlam "
              "uni keshdan siqib chiqaradi — demak kesh-sig'imi qayta ishlatish bermaydi, faqat umumiy bayt "
              "hajmi muhim. O'lchovda tasdiqlangan: INT8 2.80x, CUR qo'shilganda atigi 2.97x, FLOPs "
              "kamayishiga qaramay.")

    # ---- 3 ----
    h(doc, "3. WER / CER — yakuniy hukm", 1)
    para(doc, "Butun tadqiqotdagi eng muhim jadval. Barcha oldingi mezonlar (E_loc, E_glob) proxy edi; "
              "transkripsiya o'zgaradimi-yo'qmi degan savolga faqat WER javob beradi.")
    table(doc, "3-jadval. Butun decoder: siqish x aniqlik x tezlik x WER/CER (8 ta held-out namuna, greedy dekodlash)",
          ["Usul", "Vazn (MiB)", "Siqish", "E_glob", "WER", "CER", "Tezlanish"],
          [["FP32 (asl)", "1536", "1.00x", "0.0000", "0.0417", "0.0030", "1.00x"],
           ["INT8 per-tensor", "384", "4.00x", "0.2275", "0.0417", "0.0030", "2.80x"],
           ["INT8 per-channel", "386", "3.98x", "0.1997", "0.0417", "0.0030", "2.79x"],
           ["fc1'ga CUR majburlangan", "336", "4.57x", "1.4559", "0.9185", "0.6993", "2.97x"]],
          highlight_rows=(1, 2), fail_rows=(3,))

    para(doc, "Namuna transkripsiya (1-namuna):", bold=True, size=10)
    for line in [
        "REF               : shu fransiyaga ikkita o'yinda ham mag'lub bo'lmagan turkiyamidi",
        "FP32              : shu fransiyaga ikkita o'yinda ham mag'lub bo'lmagan turkiyamidi",
        "INT8 per-tensor   : shu fransiyaga ikkita o'yinda ham mag'lub bo'lmagan turkiyamidi",
        "INT8 per-channel  : shu fransiyaga ikkita o'yinda ham mag'lub bo'lmagan turkiyamidi",
        "CUR majburlangan  : (buzilgan)",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)

    callout(doc, "Metodologik xulosa — barcha bo'limlarga taalluqli:",
            "E_glob = 0.2275 logits xatosi HECH QANDAY transkripsiya zararini bermadi, E_glob = 1.4559 esa "
            "modelni butunlay yaroqsiz qildi. Demak E_loc/E_glob monoton emas va ular asosida 'usul yaxshi/yomon' "
            "degan xulosa chiqarish xavfli. Har qanday yakuniy da'vo WER/CER bilan tasdiqlanishi shart.", CRIT)

    # ---- 4 ----
    h(doc, "4. Kvantlash granulyarligi: per-tensor vs per-channel", 1)
    table(doc, "4-jadval. Masshtab granulyarligining ta'siri",
          ["Usul", "Vazn (MiB)", "Siqish", "E_glob", "WER", "Latency (ms)"],
          [["FP32", "1536", "1.00x", "0.0000", "0.0417", "1742.0"],
           ["INT8 per-tensor", "384", "4.00x", "0.2275", "0.0417", "621.4"],
           ["INT8 per-channel", "386", "3.98x", "0.1997", "0.0417", "624.5"]])
    para(doc, "Per-channel logits xatosini 12.2% yaxshiladi (0.2275 -> 0.1997), lekin WER o'zgarmadi — chunki u "
              "allaqachon FP32 darajasida edi. Ya'ni bu modelda per-tensor INT8 yetarli; granulyarlikni oshirish "
              "zaxira aniqlik beradi, amaliy foyda emas. Kutilgan 'eng katta yutuq' gipotezasi tasdiqlanmadi.")

    # ---- 5 ----
    h(doc, "5. SVD vs CUR — hal qiluvchi taqqoslash", 1)
    para(doc, "Teng PARAMETR byudjetida (teng rankda emas — CUR r^2 qo'shimcha parametr talab qiladi), barcha "
              "usullar x_fit da moslashtiriladi va KO'RILMAGAN x_eval da baholanadi. 45 operator x 3 nuqta = 135 o'lchov.")
    table(doc, "5-jadval. Chiqish xatosi E_loc (held-out). 3.81x — kesh-bog'langan maqsad",
          ["Siqish", "plain SVD", "act-aware SVD", "CUR (funksional)", "CUR (leverage)"],
          [["2.00x", "0.3700", "0.2379", "0.5806", "0.6978"],
           ["3.81x (kesh)", "0.5534", "0.3689", "0.7180", "0.8223"],
           ["8.00x", "0.7025", "0.4730", "0.8227", "0.8963"]])
    table(doc, "6-jadval. Vazn xatosi (Frobenius). Ekart-Yang-Mirski: plain SVD bu ustunda doim optimal",
          ["Siqish", "plain SVD", "act-aware SVD", "CUR (funksional)", "CUR (leverage)"],
          [["2.00x", "0.4019", "0.4926", "0.7790", "0.7873"],
           ["3.81x", "0.6010", "0.7314", "0.8861", "0.9085"],
           ["8.00x", "0.7583", "0.8516", "0.9540", "0.9749"]])
    callout(doc, "Tezisning konseptual yadrosi tasdiqlandi:",
            "5- va 6-jadvallarni solishtiring: act-aware SVD ning VAZN xatosi yomonroq (0.4926 vs 0.4019), lekin "
            "CHIQISH xatosi yaxshiroq (0.2379 vs 0.3700) — va bu 135/135 o'lchovda takrorlanadi. Ya'ni "
            "'vazn-optimallik chiqish-optimallik emas; kalibrlash orqali chiqishni optimallashtirish kerak' degan "
            "g'oya to'g'ri. Ekart-Yang teoremasi ham buzilmadi — ikkalasi bir vaqtda to'g'ri.", GOOD)
    table(doc, "7-jadval. Operator-darajasidagi g'alabalar (135 o'lchov)",
          ["Taqqoslash", "Natija", "Xulosa"],
          [["funksional CUR > leverage CUR", "134 / 135", "hissangiz mustahkam"],
           ["act-aware SVD > plain SVD", "135 / 135", "kalibrlash g'oyasi to'g'ri"],
           ["funksional CUR > plain SVD", "3 / 135", "CUR ramkasi zaif"],
           ["funksional CUR > act-aware SVD", "0 / 135", "CUR ramkasi zaif"]],
          highlight_rows=(0, 1), fail_rows=(2, 3))
    table(doc, "8-jadval. Overfitting bo'shlig'i (o'quv -> held-out)",
          ["Siqish", "aa-SVD (fit)", "aa-SVD (eval)", "CUR (fit)", "CUR (eval)"],
          [["2.00x", "0.0065", "0.2379", "0.5836", "0.5806"],
           ["3.81x", "0.0596", "0.3689", "0.7237", "0.7180"],
           ["8.00x", "0.2296", "0.4730", "0.8264", "0.8227"]])
    para(doc, "CUR deyarli umuman overfitting qilmaydi (fit ~ eval), aa-SVD esa juda kuchli overfitting qiladi — "
              "lekin baribir yutadi. Bu CUR kalibrlashdan juda oz ma'lumot olishini ko'rsatadi (faqat ustunlarni "
              "tartiblash uchun). Ya'ni CUR'ning muammosi umumlashtirishda emas, IFODA KUCHIDA.")

    h(doc, "CUR ning oxirgi gipotezasi: kvantlashga chidamlilik", 2)
    para(doc, "Amalda past-rank hech qachon yolg'iz qo'llanmaydi — u doim INT8 bilan birga ishlatiladi. Bu CUR uchun "
              "qolgan yagona strukturaviy argumentni ochadi: C va R — asl matritsaning haqiqiy ustun/satrlari, demak "
              "ular W ning qiymat taqsimotini meros qilib oladi. SVD faktorlari (U*S, V^T) esa yo'q — singulyar "
              "qiymatlar tartiblar bo'yicha farq qiladi, ya'ni faktor elementlarining dinamik diapazoni ancha keng, "
              "va aynan shuni qat'iy-nuqtali panjara yomon ko'taradi. Agar bu muhim bo'lsa, CUR kvantlashdan kamroq "
              "zarar ko'rishi kerak.")
    table(doc, "11-jadval. Past-rank + INT8: kvantlashga chidamlilik (27 operator x 3 nuqta = 81 o'lchov, held-out)",
          ["Siqish", "SVD fp32", "SVD int8", "CUR fp32", "CUR int8", "SVD zarar", "CUR zarar"],
          [["2.00x", "0.2356", "0.2361", "0.5798", "0.5808", "0.0005", "0.0010"],
           ["3.81x", "0.3634", "0.3635", "0.7144", "0.7150", "0.0002", "0.0006"],
           ["8.00x", "0.4611", "0.4612", "0.8166", "0.8169", "0.0001", "0.0004"]])
    table(doc, "12-jadval. Gipoteza sinovi",
          ["Sinov", "Natija", "Xulosa"],
          [["CUR zarari < SVD zarari", "18 / 81", "gipoteza rad etildi"],
           ["CUR+int8 mutlaq g'alaba", "0 / 81", "gipoteza rad etildi"]],
          fail_rows=(0, 1))
    callout(doc, "Gipoteza rad etildi — ikki tomonlama:",
            "Birinchidan, kvantlash zarari IKKALA usulda ham amalda nolga teng (0.0001–0.0010), ya'ni past-rank "
            "faktorlarni INT8 ga o'tkazish deyarli bepul va bu yerda ajratuvchi farq umuman yo'q. Ikkinchidan, farq "
            "bo'lgan joyda ham SVD chidamliroq (63/81 da), CUR emas. Demak 'C, R asl taqsimotni saqlaydi, shuning "
            "uchun yaxshiroq kvantlanadi' degan taxmin amalda ishlamaydi. Bu CUR uchun qolgan oxirgi TEXNIK "
            "himoyani yopadi — unga qoladigan yagona afzallik interpretatsiya qilinuvchanlik (C va R real ustun/satrlar).",
            CRIT)

    # ---- 6 ----
    h(doc, "6. Per-operator taqqoslash: qaysi qatlamlar CUR'ga mos", 1)
    table(doc, "9-jadval. Operator turi bo'yicha E_loc (5 qatlam x 10 tur = 50 operator; INT8 = 4x, qolganlari 8x)",
          ["Operator turi", "#", "INT8 (4x)", "INT4 (8x)", "CUR bizniki (8x)", "CUR leverage (8x)"],
          [["fc1", "5", "0.0107", "0.0585", "0.3314", "0.4584"],
           ["enc_attn/k_proj", "5", "0.0371", "0.1686", "0.5684", "0.7641"],
           ["enc_attn/out_proj", "5", "0.0420", "0.1915", "0.5791", "0.6893"],
           ["enc_attn/q_proj", "5", "0.0289", "0.1584", "0.6224", "0.7333"],
           ["fc2", "5", "0.0301", "0.1287", "0.6337", "0.7190"],
           ["enc_attn/v_proj", "5", "0.0233", "0.1505", "0.6380", "0.8314"],
           ["self_attn/q_proj", "5", "0.0258", "0.1546", "0.7161", "0.7836"],
           ["self_attn/k_proj", "5", "0.0313", "0.1604", "0.7230", "0.7796"],
           ["self_attn/out_proj", "5", "0.0342", "0.1775", "0.7481", "0.8344"],
           ["self_attn/v_proj", "5", "0.0234", "0.1505", "0.8019", "0.8610"],
           ["O'RTACHA", "50", "0.0287", "0.1499", "0.6362", "0.7454"]],
          highlight_rows=(0,))
    para(doc, "Aniq tanlash mezoni: fc1 (0.33) va enc_attn/* (0.57–0.64) past-rankli tuzilishga ega; self_attn/* "
              "(0.72–0.80) esa deyarli to'liq rankli. Ya'ni CUR barcha qatlamlarga emas, FFN kengaytiruvchi va "
              "cross-attention operatorlariga qo'llanishi mantiqiy. Shuningdek, bir xil 8x da INT4 CUR'ni 50/50 yutadi.")

    # ---- 7 ----
    h(doc, "7. Xato qatlamlar bo'ylab qanday to'planadi", 1)
    table(doc, "10-jadval. Per-operator xatodan uchdan-uchgacha xatoga",
          ["Usul", "Per-operator o'rtacha", "Uchdan-uchgacha (E_glob)", "Kuchayish", "WER ta'siri"],
          [["INT8", "0.0287", "0.2275", "~8x", "yo'q"],
           ["fc1 CUR (8x)", "~0.35", "1.4559", "~4x", "halokatli"]],
          fail_rows=(1,))
    para(doc, "Per-operator xato uchdan-uchgacha 4–8 barobar kuchayadi. Bu 9-jadval kabi per-operator jadvallarga "
              "tayanib yakuniy xulosa chiqarish mumkin emasligini ko'rsatadi.")

    # ---- 8 ----
    h(doc, "8. Yakuniy yo'nalish", 1)
    h(doc, "Decoder uchun tugallangan javob", 2)
    para(doc, "Kesh-bog'langan maqsad 3.81x; INT8 4.00x beradi, WER o'zgarmaydi, 2.80x tezlanish. Kaskad to'g'ri "
              "javobni topadi va CUR bu yerda kerak emas. Kaskadning CUR'ni rad etishi ehtiyotkorlik emas — "
              "zaruriyat (WER 0.042 -> 0.919).")
    h(doc, "Ochiq yo'nalishlar", 2)
    for txt in [
        "Encoder — vazn 1500 marta qayta ishlatiladi, compute-bound. Past-rank usullar real tezlik berishi "
        "mumkin bo'lgan yagona joy. Hali sinalmagan.",
        "CUR ning kvantlashga chidamliligi — SINALDI VA RAD ETILDI (11–12-jadval). Bu yo'nalish yopiq.",
        "WER'ni asosiy mezon sifatida barcha keyingi tajribalarga kiritish.",
        "Sezgirlikka asoslangan rank taqsimoti — bir xil 8x o'rniga har qatlamga spektriga qarab turlicha rank.",
    ]:
        p = doc.add_paragraph(txt, style="List Number")
        for r in p.runs:
            r.font.size = Pt(10)

    callout(doc, "Dissertatsiya uchun halol ramka:",
            "Ikkita mustahkam, himoya qilinadigan da'vo bor: (1) vazn-optimallik chiqish-optimallik emas — 135/135 "
            "tasdiq; (2) funksional klasterlash generik leverage-score tanlovidan yaxshiroq — 134/135 tasdiq. "
            "Uchinchi da'vo — 'CUR eng yaxshi past-rank usuli' — ma'lumotlar bilan qo'llab-quvvatlanmaydi: aniqlikda "
            "0/135, kvantlashga chidamlilikda 0/81. CUR ga qoladigan yagona afzallik — interpretatsiya qilinuvchanlik. "
            "Buni oldindan tan olish himoyada kuchli pozitsiya beradi, chunki retsenzent SVD haqida baribir so'raydi.",
            ACCENT)

    doc.add_paragraph()
    para(doc, "Kesh-miss hisoblagichlari (PMC) bu mashinada mavjud emas — kesh tahlili topologiya va hajm hisobiga "
              "asoslangan, apparat hisoblagichlariga emas. Skriptlar: experiments/, natijalar: "
              "experiments/results_*.json", size=8.5, italic=True)

    doc.save(OUT)
    print(f"saqlandi: {OUT}")


if __name__ == "__main__":
    main()
