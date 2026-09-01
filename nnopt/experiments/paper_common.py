"""Shared DOCX helpers for the three split papers (A, B, C).

Extracted verbatim from build_q1_paper_uz.py so the three generators and the
original produce identically styled documents. The original keeps its own
copies for now -- it remains the single full manuscript until the split is
approved -- so edits made here for the split papers do not silently restyle it.

Table and figure captions in the split papers use the placeholder labels
(NEW-jadval) and are numbered by scratchpad/autonumber.py per paper, each
paper numbering from 1.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x0F, 0x64, 0x70)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
GOOD = RGBColor(0x1F, 0x7A, 0x4D)
CRIT = RGBColor(0xA3, 0x2F, 0x2F)
WARN = RGBColor(0x9A, 0x64, 0x10)
FIGURE_DIR = "figures"


def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Palatino Linotype"
    st.font.size = Pt(10.5)
    return doc


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2A, 0x2A, 0x2A)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def eq(doc, text, number=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if number:
        r2 = p.add_run(f"\t\t({number})")
        r2.font.size = Pt(10.5)
    return p


def mono(doc, text, size=8.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def table(doc, caption, headers, rows, good_rows=(), bad_rows=()):
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ht in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(ht)
        run.bold = True
        run.font.size = Pt(8)
        if i > 0:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            if ri in good_rows:
                run.bold = True
                run.font.color.rgb = GOOD
            elif ri in bad_rows:
                run.font.color.rgb = CRIT
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


def bullets(doc, items, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + it[1])
            r2.font.size = Pt(10.5)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10.5)


def figure(doc, num, caption, prompt, src=None):
    """`num` is the figure's number IN THIS PAPER; `src` names the rendered
    PNG in the shared figures/ directory, which keeps the full manuscript's
    numbering. Without `src` the old num-based lookup applies, so the main
    manuscript's behaviour is unchanged."""
    path = src if src else os.path.join(FIGURE_DIR, f"fig{num}.png")
    if os.path.exists(path):
        pimg = doc.add_paragraph()
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.add_run().add_picture(path, width=Inches(5.6))
    p = doc.add_paragraph()
    r = p.add_run(f"{num}-rasm. {caption}")
    r.bold = True
    r.font.size = Pt(9)
    if not os.path.exists(path):
        p2 = doc.add_paragraph()
        r2 = p2.add_run("[AI prompt — ingliz tilida; SXEMA uchun, "
                        "grafik uchun EMAS] ")
        r2.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = ACCENT
        r3 = p2.add_run(prompt)
        r3.font.size = Pt(8)
        r3.font.color.rgb = MUTED
        p2.paragraph_format.left_indent = Pt(14)
    doc.add_paragraph()


def todo(doc, text):
    """A visibly marked migration note: what to move here from the full
    manuscript. Every one of these must be gone before submission."""
    p = doc.add_paragraph()
    r = p.add_run("[KO'CHIRILADI] " + text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = CRIT
    return p
